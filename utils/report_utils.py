"""
CyberMind AI

Report Utilities

Enterprise Production Version
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

import json
import pandas as pd

from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph
)

from docx import Document

from openpyxl import Workbook

from core.logger import logger


class ReportUtils:
    """
    Enterprise Report Utility.

    Responsibilities

    • JSON Report

    • Excel Report

    • Word Report

    • PDF Report

    • Export Reports
    """

    def __init__(
        self
    ) -> None:

        logger.info(

            "Report Utils initialized."

        )

    def json_report(
        self,
        report: dict[str, Any],
        output: str
    ) -> str:
        """
        Export JSON report.
        """

        path = Path(output)

        path.parent.mkdir(

            parents=True,

            exist_ok=True

        )

        with path.open(

            "w",

            encoding="utf-8"

        ) as file:

            json.dump(

                report,

                file,

                indent=4,

                ensure_ascii=False

            )

        return str(path)

    def excel_report(
        self,
        report: dict[str, Any],
        output: str
    ) -> str:
        """
        Export Excel report.
        """

        workbook = Workbook()

        sheet = workbook.active

        sheet.title = "CyberMind Report"

        row = 1

        for key, value in report.items():

            sheet.cell(

                row=row,

                column=1

            ).value = key

            sheet.cell(

                row=row,

                column=2

            ).value = str(value)

            row += 1

        Path(output).parent.mkdir(

            parents=True,

            exist_ok=True

        )

        workbook.save(

            output

        )

        return output

    def word_report(
        self,
        report: dict[str, Any],
        output: str
    ) -> str:
        """
        Export Word report.
        """

        document = Document()

        document.add_heading(

            "CyberMind AI Report",

            level=1

        )

        for key, value in report.items():

            document.add_paragraph(

                f"{key}: {value}"

            )

        Path(output).parent.mkdir(

            parents=True,

            exist_ok=True

        )

        document.save(

            output

        )

        return output

    def pdf_report(
        self,
        report: dict[str, Any],
        output: str
    ) -> str:
        """
        Export PDF report.
        """

        styles = getSampleStyleSheet()

        story = []

        story.append(

            Paragraph(

                "<b>CyberMind AI Report</b>",

                styles["Heading1"]

            )

        )

        for key, value in report.items():

            story.append(

                Paragraph(

                    f"<b>{key}</b>: {value}",

                    styles["BodyText"]

                )

            )

        Path(output).parent.mkdir(

            parents=True,

            exist_ok=True

        )

        pdf = SimpleDocTemplate(

            output

        )

        pdf.build(

            story

        )

        return output

    def dataframe_report(
        self,
        dataframe: pd.DataFrame,
        output: str
    ) -> str:
        """
        Export DataFrame.
        """

        Path(output).parent.mkdir(

            parents=True,

            exist_ok=True

        )

        dataframe.to_excel(

            output,

            index=False

        )

        return output

    def supported_formats(
        self
    ) -> list[str]:
        """
        Supported report formats.
        """

        return [

            "JSON",

            "Excel",

            "Word",

            "PDF"

        ]

    def health_check(
        self
    ) -> dict[str, Any]:
        """
        Health check.
        """

        return {

            "service":

                "Report Utils",

            "status":

                "Healthy",

            "version":

                "2.0"

        }

    def __repr__(
        self
    ) -> str:

        return (

            "ReportUtils("

            "Enterprise Version)"

        )


report_utils = ReportUtils()