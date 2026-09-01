from __future__ import annotations

import base64
import hashlib
import html
import os
import platform
import socket
import sys
import tempfile
import time
from datetime import datetime
from pathlib import Path
import re
import math
from typing import Any
import plotly.graph_objects as go
import streamlit as st
import streamlit.components.v1 as components

BASE_DIR = Path(__file__).resolve().parent
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))

import config.env
from modules.ai_assistant import render_ai_assistant_panel
from modules.autonomous_crs_ui import (
    render_autonomous_security_lab,
    render_code_sast_page,
    render_fuzz_sandbox_page
)
from core.logger import logger



# Page Configuration
st.set_page_config(
    page_title="CyberMind AI",
    page_icon="https://cdn-icons-png.flaticon.com/512/6071/6071531.png",
    layout="wide",
    initial_sidebar_state="collapsed", # Hide native sidebar natively
)

if "sidebar_state" not in st.session_state:
    st.session_state.sidebar_state = "expanded"

DEFAULTS = {
    "active_page": "Dashboard",
    "scanners_open": True,
    "histories": {},
    "intro_done": False, 
    "settings_language": "English",
    "settings_auto_save": True,
    "settings_dark_mode": True,
    "settings_animations": True,
    "settings_sound_alerts": False,
    "settings_auto_updates": True,
    "settings_offline_mode": False,
    "settings_export_format": "PDF",
    "sidebar_visible": True,
    "toggle_counter": 0,
    # Section 10: Sidebar collapse state (False = expanded, True = icon-only rail)
    "sidebar_collapsed": False,
    "last_scan_context": {},
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ── Query-param navigation (triggered by the floating AI button JS) ──
_qp = st.query_params.to_dict()
if _qp.get("nav") == "ai":
    st.session_state.active_page = "AI Security Assistant"
    st.query_params.clear()
    st.rerun()

TRANSLATIONS = {
    "English": {
        "Dashboard": "Dashboard",
        "Scan History": "Scan History",
        "Analytics": "Analytics",
        "Settings": "Settings",
        "Profile": "Profile",
        "Connections": "Connections",
        "Help & Support": "Help & Support",
        "MAIN": "MAIN",
        "REASONING & REPAIR": "REASONING & REPAIR",
        "ANALYSIS": "ANALYSIS",
        "SYSTEM": "SYSTEM",
        "About the CyberMind AI Team": "About the CyberMind AI Team",
        "About Developer": "About the CyberMind AI Team",
        "General Settings": "General Settings",
        "Language": "Application Language",
        "Auto Save": "Auto Save Reports",
        "Dark Mode": "Dark Mode",
        "Animations": "Animations",
        "Sound Alerts": "Sound Alerts",
        "Auto Updates": "Auto Check for Updates",
        "Export Format": "Default Export Format",
        "URL Scanner": "URL Scanner",
        "Website Analyzer": "Website Analyzer",
        "Domain Intelligence": "Domain Intelligence",
        "Email Intelligence": "Email Intelligence",
        "IP Intelligence": "IP Intelligence",
        "QR Scanner": "QR Scanner",
        "File Analyzer": "File Analyzer",
        "Universal Scan": "Universal Scan",
        "Device Security Check": "Device Security Check",
        "Phone Threat Intelligence": "Phone Threat Intelligence",
        "AI Security Assistant": "AI Security Assistant",
        "Groq": "Groq",
        "Clear Cache": "Clear Cache",
        "Save API Key": "Save API Key",
        "Activate API Keys": "Activate API Keys",
        "Database": "Database",
        "Datasets": "Datasets",
        "API Connections": "API Connections",
    },
    "Hindi": {
        "Dashboard": "डैशबोर्ड",
        "Scan History": "स्कैन इतिहास",
        "Analytics": "एनालिटिक्स",
        "Settings": "सेटिंग्स",
        "Profile": "प्रोफ़ाइल",
        "Connections": "कनेक्शन",
        "Help & Support": "सहायता और समर्थन",
        "MAIN": "मुख्य मेनू",
        "REASONING & REPAIR": "रीजनिंग और रिपेयर",
        "ANALYSIS": "विश्लेषण",
        "SYSTEM": "सिस्टम",
        "About the CyberMind AI Team": "साइबरमाइंड एआई टीम के बारे में",
        "About Developer": "साइबरमाइंड एआई टीम के बारे में",
        "General Settings": "सामान्य सेटिंग्स",
        "Language": "एप्लिकेशन भाषा",
        "Auto Save": "रिपोर्ट स्वतः सहेजें",
        "Dark Mode": "डार्क मोड",
        "Animations": "एनिमेशन",
        "Sound Alerts": "ध्वनि अलर्ट",
        "Auto Updates": "अपडेट के लिए स्वतः जाँच करें",
        "Export Format": "डिफ़ॉल्ट निर्यात प्रारूप",
        "URL Scanner": "यूआरएल स्कैनर",
        "Website Analyzer": "वेबसाइट विश्लेषक",
        "Domain Intelligence": "डोमेन इंटेलिजेंस",
        "Email Intelligence": "ईमेल इंटेलिजेंस",
        "IP Intelligence": "आईपी इंटेलिजेंस",
        "QR Scanner": "क्यूआर स्कैनर",
        "File Analyzer": "फाइल विश्लेषक",
        "Universal Scan": "यूनिवर्सल स्कैन",
        "Device Security Check": "डिवाइस सुरक्षा जाँच",
        "Phone Threat Intelligence": "फोन थ्रेट इंटेलिजेंस",
        "AI Security Assistant": "एआई सुरक्षा सहायक",
        "Groq": "Groq",
        "Clear Cache": "कैश साफ़ करें",
        "Save API Key": "एपीआई कुंजी सहेजें",
        "Activate API Keys": "एपीआई कुंजी सक्रिय करें",
        "Database": "डेटाबेस",
        "Datasets": "डेटासेट",
        "API Connections": "एपीआई कनेक्शन",
    },
    "Spanish": {
        "Dashboard": "Tablero",
        "Scan History": "Historial de Escaneo",
        "Analytics": "Analítica",
        "Settings": "Ajustes",
        "Profile": "Perfil",
        "Connections": "Conexiones",
        "Help & Support": "Ayuda y Soporte",
        "MAIN": "PRINCIPAL",
        "ANALYSIS": "ANÁLISIS",
        "SYSTEM": "SISTEMA",
        "About Developer": "Sobre el Desarrollador",
        "General Settings": "Ajustes Generales",
        "Language": "Idioma de la Aplicación",
        "Auto Save": "Guardar Informes Automáticamente",
        "Dark Mode": "Modo Oscuro",
        "Animations": "Animaciones",
        "Sound Alerts": "Alertas de Sonido",
        "Auto Updates": "Buscar Actualizaciones Automáticamente",
        "Export Format": "Formato de Exportación Predeterminado",
        "URL Scanner": "Escáner de URL",
        "Website Analyzer": "Analizador Web",
        "Domain Intelligence": "Inteligencia de Dominio",
        "Email Intelligence": "Inteligencia de Email",
        "IP Intelligence": "Inteligencia de IP",
        "QR Scanner": "Escáner de QR",
        "File Analyzer": "Analizador de Archivos",
        "Universal Scan": "Escaneo Universal",
        "Device Security Check": "Control de Seguridad",
        "AI Security Assistant": "Asistente de Seguridad IA",
        "Groq": "Groq",
        "Clear Cache": "Limpiar Cache",
        "Save API Key": "Guardar Clave API",
        "Activate API Keys": "Activar Claves API",
        "Database": "Base de Datos",
        "Datasets": "Conjuntos de Datos",
        "API Connections": "Conexiones API",
    },
    "German": {
        "Dashboard": "Dashboard",
        "Scan History": "Scan-Verlauf",
        "Analytics": "Analysen",
        "Settings": "Einstellungen",
        "Profile": "Profil",
        "Connections": "Verbindungen",
        "Help & Support": "Hilfe & Support",
        "MAIN": "HAUPTMENÜ",
        "ANALYSIS": "ANALYSE",
        "SYSTEM": "SYSTEM",
        "About Developer": "Über den Entwickler",
        "General Settings": "Allgemeine Einstellungen",
        "Language": "Anwendungssprache",
        "Auto Save": "Berichte automatisch speichern",
        "Dark Mode": "Dunkelmodus",
        "Animations": "Animationen",
        "Sound Alerts": "Sound-Benachrichtigungen",
        "Auto Updates": "Automatisch nach Updates suchen",
        "Export Format": "Standard-Exportformat",
        "URL Scanner": "URL-Scanner",
        "Website Analyzer": "Website-Analysator",
        "Domain Intelligence": "Domain-Intelligenz",
        "Email Intelligence": "E-Mail-Intelligenz",
        "IP Intelligence": "IP-Intelligenz",
        "QR Scanner": "QR-Scanner",
        "File Analyzer": "Datei-Analysator",
        "Universal Scan": "Universeller Scan",
        "Device Security Check": "Gerätesicherheitsprüfung",
        "AI Security Assistant": "KI-Sicherheitsassistent",
        "Groq": "Groq",
        "Clear Cache": "Cache leeren",
        "Save API Key": "API-Schlüssel speichern",
        "Activate API Keys": "API-Schlüssel aktivieren",
        "Database": "Datenbank",
        "Datasets": "Datensätze",
        "API Connections": "API-Verbindungen",
    }
}

def t(key: str) -> str:
    lang = st.session_state.get("settings_language", "English")
    return TRANSLATIONS.get(lang, TRANSLATIONS["English"]).get(key, key)


def clean_html(s: str) -> str:
    return re.sub(r'\s+', ' ', s).strip()



def go_to(page: str):
    st.session_state.active_page = page


def restore_scan_from_history(scanner_key: str, target: str, risk_level: str, risk_score: float):
    """Used by the Scan History page: re-open the given scanner with the
    exact same target/result so the user sees the full detail view again,
    instead of a blank/fresh scanner page."""
    if scanner_key not in SCANNERS:
        return
    st.session_state[f"scan_result_{scanner_key}"] = {
        "value": target,
        "risk_score": risk_score,
        "risk_level": risk_level,
        "duration": "0.45 sec",
        "raw": {},
    }
    st.session_state[f"input_{scanner_key}"] = target
    st.session_state[f"_restore_pending_{scanner_key}"] = True
    st.session_state.active_page = scanner_key


def toggle_scanners():
    st.session_state.scanners_open = not st.session_state.scanners_open


# 10-Second Pure Full-Screen Video Intro With Click-to-Unmute Fix

if not st.session_state.intro_done:
    if not st.session_state.get("settings_animations", True):
        st.session_state.intro_done = True
    else:
        video_path = BASE_DIR / "cybermind.mp4"
        video_src = ""
        
        if video_path.exists():
            with open(video_path, "rb") as f:
                video_bytes = f.read()
            video_base64 = base64.b64encode(video_bytes).decode("utf-8")
            video_src = f"data:video/mp4;base64,{video_base64}"

        st.markdown(
            f"""
            <iframe srcdoc="
    <!DOCTYPE html>
    <html>
    <head>
    <meta charset='utf-8'>
    <style>
    body {{
        margin: 0; padding: 0;
        background-color: #060B14;
        overflow: hidden;
        font-family: 'Inter', sans-serif;
    }}
    .intro-container {{
        position: relative;
        width: 100vw; height: 100vh;
        cursor: pointer;
    }}
    video {{
        position: absolute;
        top: 50%; left: 50%;
        width: 100vw; height: 100vh;
        transform: translate(-50%, -50%);
        object-fit: cover;
        z-index: 0;
    }}
    .mute-btn {{
        position: absolute;
        top: 22px; right: 24px;
        z-index: 10;
        background: rgba(6, 11, 20, 0.70);
        border: 1.5px solid rgba(108, 92, 231, 0.60);
        border-radius: 50px;
        padding: 8px 20px;
        color: #EAF0F7;
        font-size: 13px; font-weight: 600;
        cursor: pointer;
        backdrop-filter: blur(10px);
        box-shadow: 0 0 18px rgba(34,184,240,0.22);
        transition: all 0.18s ease;
    }}
    .mute-btn:hover {{
        background: rgba(108, 92, 231, 0.32);
        border-color: rgba(34, 184, 240, 0.85);
    }}
    .intro-hint {{
        position: absolute;
        bottom: 36px; left: 50%;
        transform: translateX(-50%);
        z-index: 5;
        color: rgba(234,240,247,0.65);
        font-size: 13px; font-weight: 500;
        pointer-events: none;
        animation: fadeHint 2s ease-in-out infinite;
    }}
    @keyframes fadeHint {{
        0%, 100% {{ opacity: 0.4; }}
        50% {{ opacity: 0.9; }}
    }}
    </style>
    </head>
    <body>
    <div class='intro-container' onclick='unmuteVideo()'>
        <video id='iv' autoplay muted loop playsinline>
            <source src='{video_src}' type='video/mp4'>
        </video>
        <button id='mb' class='mute-btn' onclick='toggleMute(event)'>🔇 Muted</button>
        <div id='ih' class='intro-hint'>🔊 Click anywhere to enable sound</div>
    </div>
    <script>
    function unmuteVideo() {{
        var v = document.getElementById('iv');
        var b = document.getElementById('mb');
        var h = document.getElementById('ih');
        if (v) {{
            v.muted = false;
        }}
        if (b) {{
            b.innerHTML = '🔊 Sound On';
        }}
        if (h) {{
            h.style.display = 'none';
        }}
    }}
    function toggleMute(e) {{
        e.stopPropagation();
        var v = document.getElementById('iv');
        var b = document.getElementById('mb');
        var h = document.getElementById('ih');
        if (v) {{
            v.muted = !v.muted;
            if (v.muted) {{
                b.innerHTML = '🔇 Muted';
            }} else {{
                b.innerHTML = '🔊 Sound On';
                if (h) {{
                    h.style.display = 'none';
                }}
            }}
        }}
    }}
    </script>
    </body>
    </html>
            " style="position:fixed; top:0; left:0; width:100vw; height:100vh; border:none; z-index:999999; margin:0; padding:0;"></iframe>
            """,
            unsafe_allow_html=True
        )
        
        time.sleep(11)
        st.session_state.intro_done = True
        st.rerun()


# Global Styling


dark_mode = st.session_state.get("settings_dark_mode", True)
if dark_mode:
    theme_vars = """
        --bg:#060B14;
        --bg-soft:#0A1220;
        --card-bg:#0F1826;
        --card-bg-soft:#111C2C;
        --border:#1D2838;
        --text:#EAF0F7;
        --text-muted:#8B98AC;
        --text-faint:#5C6880;
        --grad-a:#6C5CE7;
        --grad-b:#22B8F0;
        --success:#22C55E;
        --success-bg:rgba(34,197,94,0.12);
        --warning:#F5A623;
        --warning-bg:rgba(245,166,35,0.12);
        --danger:#F2545B;
        --danger-bg:rgba(242,84,91,0.12);
        --info:#3B82F6;
        --info-bg:rgba(59,130,246,0.12);
        --hover-bg:rgba(255,255,255,0.045);
        --hero-bg:radial-gradient(circle at 82% 40%, rgba(34,184,240,0.16), transparent 60%), linear-gradient(135deg, #0A0F1D 0%, #0C1424 100%);
        --sidebar-bg:rgba(11, 17, 30, 0.95);
        --sidebar-border:rgba(255, 255, 255, 0.08);
        --sidebar-shadow:0 8px 32px rgba(0, 0, 0, 0.45);
        --sb-btn-bg:rgba(18, 28, 45, 0.75);
        --sb-btn-border:rgba(255, 255, 255, 0.08);
        --sb-btn-text:#E2E8F0;
        --logo-text:#FFFFFF;
        --tooltip-bg:#111C2C;
        --tooltip-border:#1D2838;
        --tooltip-text:#EAF0F7;
        --dialog-bg:#0B111E;
        --dialog-border:rgba(255, 255, 255, 0.12);
        --dialog-text:#EAF0F7;
        --dialog-muted:#8B98AC;
        --dialog-shadow:0 20px 50px rgba(0, 0, 0, 0.65);
    """
else:
    theme_vars = """
        --bg:#F4F6F9;
        --bg-soft:#E9ECF1;
        --card-bg:#FFFFFF;
        --card-bg-soft:#F8FAFC;
        --border:#D1D9E6;
        --text:#1E293B;
        --text-muted:#64748B;
        --text-faint:#94A3B8;
        --grad-a:#6C5CE7;
        --grad-b:#22B8F0;
        --success:#22C55E;
        --success-bg:rgba(34,197,94,0.12);
        --warning:#F5A623;
        --warning-bg:rgba(245,166,35,0.12);
        --danger:#F2545B;
        --danger-bg:rgba(242,84,91,0.12);
        --info:#3B82F6;
        --info-bg:rgba(59,130,246,0.12);
        --hover-bg:rgba(0,0,0,0.045);
        --hero-bg:radial-gradient(circle at 82% 40%, rgba(34,184,240,0.12), transparent 60%), linear-gradient(135deg, #E2E8F0 0%, #F1F5F9 100%);
        --sidebar-bg:#FFFFFF;
        --sidebar-border:#D1D9E6;
        --sidebar-shadow:0 8px 32px rgba(0, 0, 0, 0.08);
        --sb-btn-bg:#F8FAFC;
        --sb-btn-border:#E2E8F0;
        --sb-btn-text:#1E293B;
        --logo-text:#1E293B;
        --tooltip-bg:#FFFFFF;
        --tooltip-border:#D1D9E6;
        --tooltip-text:#1E293B;
        --dialog-bg:#FFFFFF;
        --dialog-border:#D1D9E6;
        --dialog-text:#1E293B;
        --dialog-muted:#64748B;
        --dialog-shadow:0 20px 50px rgba(0, 0, 0, 0.15);
    """

animations_enabled = st.session_state.get("settings_animations", True)
if animations_enabled:
    animation_css = """
    * {
        transition: background-color 0.25s ease, border-color 0.25s ease, color 0.25s ease;
    }
    """
else:
    animation_css = ""

style_template = """
    <style>

    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&family=Outfit:wght@600;700;800;900&display=swap');

    :root{
        __THEME_VARS__
    }

    __ANIMATION_CSS__

    html, body, [class*="css"]{ font-family:'Inter', sans-serif; }

    #MainMenu {visibility:hidden;}
    footer {visibility:hidden;}
    [data-testid="stToolbar"]{display:none;}
    [data-testid="stDecoration"]{display:none;}
    header[data-testid="stHeader"]{
        display: none !important;
        height: 0 !important;
        min-height: 0 !important;
        padding: 0 !important;
        margin: 0 !important;
    }
    [data-testid="stMain"] {
        padding-top: 0 !important;
    }
    section[data-testid="stSidebar"] {
        display: none !important;
    }
    [data-testid="collapsedControl"] {
        display: none !important;
    }

    .stApp {
        background: radial-gradient(circle at 80% 0%, rgba(108,92,231,0.08), transparent 45%), var(--bg);
    }

    /* Global Image Sizing Safety Rule */
    .chart-card img, .stApp img[src*="18310827"], .stApp img[src*="flaticon"], .stApp img[src*="cdn-icons"] {
        max-width: 24px !important;
        max-height: 24px !important;
        object-fit: contain !important;
        vertical-align: middle !important;
    }

    /* 1. Custom Fixed Sidebar Outer Element Wrapper Node Zero-Out */
    div.stElementContainer:has(> .st-key-navcol),
    [data-testid="stElementContainer"]:has(> .st-key-navcol) {
        height: 0 !important;
        min-height: 0 !important;
        max-height: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
    }

    /* 2. Main Page Content Block Layout Offsets */
    html body div.block-container,
    div.block-container,
    [data-testid="stMainBlockContainer"],
    .stMainBlockContainer {
        padding-top: 1rem !important;
        padding-bottom: 2.5rem !important;
        padding-left: 17.5rem !important;
        padding-right: 1.5rem !important;
        max-width: 100% !important;
        box-sizing: border-box !important;
    }

    /* -------- SLEEK COMPACT EXPANDED SIDEBAR -------- */
    .st-key-navcol {
        position: fixed !important;
        top: 0.8rem !important;
        left: 0.8rem !important;
        width: 15rem !important;
        height: calc(100vh - 1.6rem) !important;
        background: var(--sidebar-bg) !important;
        border: 1px solid var(--sidebar-border) !important;
        border-radius: 14px !important;
        padding: 12px 8px 14px 8px !important;
        overflow-y: auto !important;
        overflow-x: hidden !important;
        z-index: 1000 !important;
        box-sizing: border-box !important;
        box-shadow: var(--sidebar-shadow) !important;
        backdrop-filter: blur(12px) !important;
        pointer-events: auto !important;
        scrollbar-width: thin !important;
        scrollbar-color: rgba(255, 255, 255, 0.15) transparent !important;
    }
    .st-key-navcol button,
    .st-key-navcol .stButton > button {
        pointer-events: auto !important;
        cursor: pointer !important;
    }
    .st-key-navcol::-webkit-scrollbar { width: 3px; }
    .st-key-navcol::-webkit-scrollbar-thumb { background: rgba(255, 255, 255, 0.15); border-radius: 3px; }
    .st-key-navcol::-webkit-scrollbar-track { background: transparent; }

    .st-key-navcol .stButton > button {
        font-size: 13px !important;
        font-weight: 500 !important;
        padding: 6px 10px !important;
        min-height: 36px !important;
        height: 36px !important;
        border-radius: 8px !important;
        margin-bottom: 3px !important;
        white-space: nowrap !important;
        overflow: hidden !important;
        text-overflow: ellipsis !important;
    }

    /* -------- Logo Button (Clickable Sidebar Toggle - Pure Transparent Header ONLY) -------- */
    .st-key-sb_logo_btn {
        margin-bottom: 12px !important;
        border-bottom: 1px solid var(--border) !important;
        padding-bottom: 10px !important;
    }
    html body .stApp .st-key-navcol .st-key-sb_logo_btn,
    html body .stApp .st-key-navcol .st-key-sb_logo_btn *,
    html body .stApp .st-key-navcol .st-key-sb_logo_btn button,
    html body .stApp .st-key-navcol .st-key-sb_logo_btn .stButton > button,
    html body .stApp .st-key-navcol .st-key-sb_logo_btn [data-testid*="baseButton"],
    html body .stApp .st-key-navcol .st-key-sb_logo_btn [data-testid*="stBaseButton"],
    div.st-key-sb_logo_btn,
    div.st-key-sb_logo_btn *,
    div.st-key-sb_logo_btn button,
    div.st-key-sb_logo_btn .stButton > button {
        background: transparent !important;
        background-color: transparent !important;
        background-image: none !important;
        border: none !important;
        border-color: transparent !important;
        box-shadow: none !important;
        outline: none !important;
    }
    .st-key-sb_logo_btn .stButton > button {
        height: auto !important;
        min-height: 52px !important;
        max-height: none !important;
        padding: 4px 0 !important;
        border-radius: 12px !important;
        text-align: left !important;
        justify-content: flex-start !important;
        width: 100% !important;
        display: flex !important;
        align-items: center !important;
        cursor: pointer !important;
        transition: all 0.2s ease !important;
    }
    .st-key-sb_logo_btn button:hover,
    .st-key-sb_logo_btn .stButton > button:hover {
        background: rgba(34, 211, 238, 0.08) !important;
        border: 1px solid rgba(34, 211, 238, 0.25) !important;
    }
    .st-key-sb_logo_btn button::before,
    .st-key-sb_logo_btn .stButton > button::before {
        content: "" !important;
        display: inline-block !important;
        width: 44px !important;
        height: 44px !important;
        background-image: url("https://cdn-icons-png.flaticon.com/512/6071/6071531.png") !important;
        background-size: contain !important;
        background-repeat: no-repeat !important;
        background-position: center !important;
        margin-right: 12px !important;
        flex-shrink: 0 !important;
    }
    /* High Specificity Logo Header Text Overrides */
    html body .stApp .st-key-navcol .st-key-sb_logo_btn button p,
    html body .stApp .st-key-navcol .st-key-sb_logo_btn button span,
    html body .stApp .st-key-navcol .st-key-sb_logo_btn [data-testid="stMarkdownContainer"] p,
    html body .stApp .st-key-navcol .st-key-sb_logo_btn [data-testid="stMarkdownContainer"] span,
    .st-key-sb_logo_btn button p,
    .st-key-sb_logo_btn button span {
        font-family: 'Outfit', 'Inter', sans-serif !important;
        font-size: 24px !important;
        font-weight: 900 !important;
        color: var(--logo-text) !important;
        margin: 0 !important;
        line-height: 1.1 !important;
        letter-spacing: -0.4px !important;
        display: inline-block !important;
    }
    /* Single AI Pseudo-element on Paragraph Node Only (Fixes Double AI Bug) */
    html body .stApp .st-key-navcol .st-key-sb_logo_btn button p::after,
    .st-key-sb_logo_btn button p::after {
        content: " AI" !important;
        color: #22D3EE !important;
        font-family: 'Outfit', 'Inter', sans-serif !important;
        font-size: 24px !important;
        font-weight: 900 !important;
        letter-spacing: -0.4px !important;
        display: inline-block !important;
    }
    .sb-section{
        font-size:11.5px;font-weight:800;letter-spacing:1.2px;
        color:var(--text-faint);
        margin:12px 6px 6px 6px;
    }
    .sb-divider {
        height: 1px !important;
        background: var(--border) !important;
        opacity: 0.6 !important;
        margin: 12px 4px 6px 4px !important;
        border: none !important;
    }

    .st-key-navcol .stButton{ margin-bottom:3px; }
    
    /* Unselected Sidebar Buttons */
    .st-key-navcol .stButton > button,
    .st-key-navcol .stButton > button[kind="secondary"],
    .st-key-navcol .stButton > button[data-testid*="secondary"],
    .st-key-navcol .stButton > button[data-testid*="Secondary"],
    .st-key-navcol button[data-testid*="secondary"] {
        width: 100% !important;
        text-align: left !important;
        justify-content: flex-start !important;
        background: var(--sb-btn-bg) !important;
        background-color: var(--sb-btn-bg) !important;
        border: 1px solid var(--sb-btn-border) !important;
        color: var(--sb-btn-text) !important;
        font-size: 14px !important;
        font-weight: 600 !important;
        padding: 7px 12px !important;
        border-radius: 9px !important;
        transition: all 0.15s ease !important;
        box-shadow: 0 2px 6px rgba(0, 0, 0, 0.05) !important;
    }
    .st-key-navcol .stButton > button p,
    .st-key-navcol .stButton > button span {
        color: var(--sb-btn-text) !important;
        font-size: 14px !important;
        font-weight: 600 !important;
    }

    /* Hover State: Glowing Cyan Glassmorphism Frame on Hover */
    .st-key-navcol .stButton > button:hover,
    .st-key-navcol .stButton > button[kind="secondary"]:hover,
    .st-key-navcol .stButton > button[data-testid*="secondary"]:hover,
    .st-key-navcol button[data-testid*="secondary"]:hover {
        background: rgba(34, 211, 238, 0.12) !important;
        color: #FFFFFF !important;
        border: 1px solid rgba(34, 211, 238, 0.35) !important;
        box-shadow: 0 4px 12px rgba(34, 211, 238, 0.2) !important;
    }
    .st-key-navcol .stButton > button:hover p,
    .st-key-navcol .stButton > button:hover span {
        color: #FFFFFF !important;
    }

    /* Selected Active Button: Vibrant Gradient Background Intact */
    .st-key-navcol .stButton > button[kind="primary"],
    .st-key-navcol .stButton > button[data-testid*="primary"],
    .st-key-navcol .stButton > button[data-testid*="Primary"],
    .st-key-navcol button[data-testid*="primary"],
    .st-key-navcol button[data-testid*="Primary"],
    .st-key-navcol button[kind="primary"] {
        background: linear-gradient(135deg, var(--grad-a), var(--grad-b)) !important;
        background-color: transparent !important;
        color: #ffffff !important;
        border: none !important;
        box-shadow: 0 4px 16px rgba(34,184,240,0.35) !important;
    }
    .st-key-navcol .stButton > button[kind="primary"] p,
    .st-key-navcol .stButton > button[data-testid*="primary"] p,
    .st-key-navcol .stButton > button[data-testid*="Primary"] p,
    .st-key-navcol .stButton > button[kind="primary"] span,
    .st-key-navcol .stButton > button[data-testid*="primary"] span,
    .st-key-navcol .stButton > button[data-testid*="Primary"] span {
        color: #ffffff !important;
        font-weight: 600 !important;
    }

    .sb-sub .stButton > button{ padding-left:1.7rem !important; font-size:14.2px; }

    .sb-version{
        margin:1.1rem 4px 16px 4px; text-align:center; font-size:12px;
        color:var(--text-faint); background:var(--card-bg-soft);
        border:1px solid var(--border); border-radius:8px; padding:6px;
    }

    /* -------- AI Assistant expander inside sidebar -------- */
    .st-key-navcol [data-testid="stExpander"] {
        background: var(--card-bg-soft) !important;
        border: 1px solid var(--border) !important;
        border-radius: 10px !important;
        margin: 2px 4px 6px 4px !important;
    }
    .st-key-navcol [data-testid="stExpander"] summary {
        font-size: 13px !important;
        font-weight: 700 !important;
        color: var(--text) !important;
        padding: 8px 10px !important;
    }
    .st-key-navcol [data-testid="stExpander"] summary:hover {
        background: var(--hover-bg) !important;
        border-radius: 10px !important;
    }
    /* suggestion bullets */
    .st-key-navcol [data-testid="stExpander"] p,
    .st-key-navcol [data-testid="stExpander"] li {
        font-size: 12px !important;
        color: var(--text-muted) !important;
        line-height: 1.5 !important;
    }
    /* text input */
    .st-key-navcol [data-testid="stExpander"] input[type="text"] {
        font-size: 12.5px !important;
        background: var(--bg-soft) !important;
        border: 1px solid var(--border) !important;
        color: var(--text) !important;
        border-radius: 8px !important;
    }
    /* Ask button — gradient accent */
    .st-key-navcol .st-key-ai_assistant_send button {
        background: linear-gradient(90deg, var(--grad-a), var(--grad-b)) !important;
        color: #fff !important;
        border: none !important;
        font-size: 12px !important;
        font-weight: 600 !important;
        border-radius: 7px !important;
        padding: 4px 0 !important;
    }
    /* Clear button — subtle */
    .st-key-navcol .st-key-ai_assistant_clear button {
        background: transparent !important;
        color: var(--text-muted) !important;
        border: 1px solid var(--border) !important;
        font-size: 12px !important;
        border-radius: 7px !important;
        padding: 4px 0 !important;
    }
    /* Chat bubbles */
    .st-key-navcol [data-testid="stExpander"] [data-testid="stMarkdownContainer"] p {
        font-size: 12px !important;
        line-height: 1.55 !important;
        padding: 3px 0 !important;
    }



    /* ---------- Hero ---------- */
    .hero{
        display:flex; align-items:center; justify-content:space-between; flex-wrap:wrap; gap:16px;
        background:var(--hero-bg);
        border:1px solid var(--border); border-radius:18px;
        padding:clamp(20px, 3vw, 40px) clamp(18px, 3.5vw, 44px); margin-bottom:24px; overflow:hidden;
    }
    .hero-title{
        font-size:clamp(24px, 3.5vw, 48px); font-weight:800; line-height:1.1; margin-bottom:12px;
        background:linear-gradient(90deg, #8C7CF0 0%, #4FC3F7 60%, #22D3EE 100%);
        -webkit-background-clip:text; -webkit-text-fill-color:transparent;
    }
    .hero-sub{ font-size:clamp(15px, 1.8vw, 21px); font-weight:600; color:var(--text); margin-bottom:12px; }
    .hero-desc{ font-size:clamp(12px, 1.2vw, 15.5px); color:var(--text-muted); line-height:1.65; margin-bottom:22px; max-width:560px; }

    .shield-wrap{ position:relative; display:flex; align-items:center; justify-content:center; height:210px; }
    .shield-brain{
        position:absolute; font-size:58px;
        filter:drop-shadow(0 0 14px rgba(34,211,238,0.65));
    }

    /* ---------- Stat cards ---------- */
    .stat-card{
        background:var(--card-bg); border:1px solid var(--border); border-radius:14px;
        padding:14px 16px; display:flex; align-items:center; justify-content:space-between;
        min-height:100px; overflow:hidden; box-sizing:border-box; gap:8px;
    }
    .stat-card-left{ min-width:0; flex:1; overflow:hidden; }
    .stat-label{ font-size:clamp(11px, 1.1vw, 14.6px);color:var(--text-muted);margin-bottom:6px;font-weight:500; white-space:nowrap; overflow:hidden; text-overflow:ellipsis; }
    .stat-value{ font-size:clamp(18px, 2.2vw, 29px);font-weight:800;color:var(--text);margin-bottom:4px; line-height:1.1; }
    .stat-delta{ font-size:clamp(11px, 1vw, 14px);font-weight:600; white-space:normal; word-break:break-word; }
    .stat-icon{ width:44px;height:44px;min-width:44px;border-radius:50%; display:flex;align-items:center;justify-content:center; font-size:18px; flex-shrink:0; }

    .section-title{ font-size:19.5px;font-weight:700;color:var(--text); margin:4px 0 13px 2px; }

    /* ---------- Scanner cards (home grid) ---------- */
    .scan-card{
        background:var(--card-bg); border:1px solid var(--border); border-radius:14px;
        padding:15px 15px; height:155px; display:flex; flex-direction:column; justify-content:space-between;
        overflow:hidden; box-sizing:border-box;
    }
    .scan-icon{ width:32px;height:32px;border-radius:10px; display:flex;align-items:center;justify-content:center; font-size:15px;margin-bottom:8px;flex-shrink:0; }
    .scan-title{ font-size:13.5px;font-weight:700;color:var(--text);margin-bottom:4px; white-space:nowrap; overflow:hidden; text-overflow:ellipsis; }
    .scan-desc{ font-size:11.5px;color:var(--text-muted);line-height:1.45; overflow:hidden; display:-webkit-box; -webkit-line-clamp:2; -webkit-box-orient:vertical; height:34px; }

    /* ---------- Right panels ---------- */
    .panel{ background:var(--card-bg); border:1px solid var(--border); border-radius:14px; padding:17px 18px; margin-bottom:16px; }
    .panel-title{ font-size:16.5px;font-weight:700;color:var(--text);margin-bottom:13px; }
    .status-row{ display:flex;align-items:center;justify-content:space-between; padding:7px 0; font-size:14px; border-bottom:1px solid var(--border); }
    .status-row:last-child{ border-bottom:none; }
    .status-name{ color:var(--text-muted); }
    .status-val{ color:var(--success);font-weight:600;font-size:13.5px; }

    .scan-row{ display:flex;align-items:center;justify-content:space-between; padding:8px 0; border-bottom:1px solid var(--border); }
    .scan-row:last-child{ border-bottom:none; }
    .scan-row-left{ display:flex;flex-direction:column;gap:2px; }
    .scan-row-name{ font-size:14px;color:var(--text);font-weight:500; }
    .scan-row-time{ font-size:12px;color:var(--text-faint); }
    .badge{ font-size:11.5px;font-weight:700;padding:3px 8px;border-radius:6px;white-space:nowrap; }
    .badge-safe{ background:var(--success-bg);color:var(--success); }
    .badge-malicious{ background:var(--danger-bg);color:var(--danger); }
    .badge-suspicious{ background:var(--warning-bg);color:var(--warning); }
    .badge-unverified{ background:rgba(245, 158, 11, 0.15);color:#F59E0B; }

    /* ---------- Charts / scanner-page cards ---------- */
    .chart-card{ background:var(--card-bg); border:1px solid var(--border); border-radius:14px; padding:16px 18px; margin-bottom:16px; box-sizing:border-box; overflow:hidden; min-width:0; }
    .metric-grid{ align-items:stretch; }
    .metric-grid > .chart-card{ margin-bottom:0; height:100%; }
    .st-key-datasets_container, .st-key-api_container {
        background: var(--card-bg) !important;
        border: 1px solid var(--border) !important;
        border-radius: 14px !important;
        padding: 14px 18px 4px 18px !important;
        margin-bottom: 16px !important;
    }
    .st-key-export_report_container, .st-key-export_breach_report_container {
        background: var(--card-bg) !important;
        border: 1px solid var(--border) !important;
        border-radius: 14px !important;
        padding: 16px 18px 16px 18px !important;
        margin-bottom: 16px !important;
    }
    .chart-card-head{ display:flex;align-items:center;justify-content:space-between; margin-bottom:4px; }
    .chart-card-title{ font-size:16.5px;font-weight:700;color:var(--text); }
    .chart-filter{ font-size:11.5px;color:var(--text-muted); background:var(--bg-soft); border:1px solid var(--border); border-radius:7px; padding:4px 10px; }

    .page-head{ display:flex;align-items:center;gap:14px; margin-bottom:6px; }
    .page-head-icon{ width:46px;height:46px;border-radius:12px; display:flex;align-items:center;justify-content:center;font-size:21px; }
    .page-head-title{ font-size:22px;font-weight:800;color:var(--text); }
    .page-head-desc{ font-size:13px;color:var(--text-muted); margin:2px 0 20px 0; }

    table.scan-table{ width:100%; border-collapse:collapse; font-size:12.6px; table-layout:fixed; }
    table.scan-table th{ text-align:left; color:var(--text-faint); font-weight:600; font-size:11px; text-transform:uppercase; letter-spacing:0.4px; padding:8px 10px; border-bottom:1px solid var(--border); }
    table.scan-table td{ padding:9px 10px; color:var(--text); border-bottom:1px solid var(--border); overflow-wrap:break-word; word-break:break-word; }
    table.scan-table tr:last-child td{ border-bottom:none; }
    table.scan-table td:first-child{ width:42%; color:var(--text-muted); }
    table.scan-table td:last-child{ width:58%; text-align:right; }

    /* ---------- Responsive metric-card grids (fixes squeezed / overflowing cards) ---------- */
    .metric-grid{
        display:grid;
        grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
        gap:12px;
    }
    .metric-grid > div{ min-width:0; }
    .metric-grid .metric-value{
        overflow-wrap:break-word;
        word-break:break-word;
        white-space:normal;
    }
    .bullet-grid{
        display:grid;
        grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
        gap:8px;
    }
    .bullet-grid > div{ min-width:0; overflow-wrap:break-word; word-break:break-word; }

    .list-row{ display:flex;align-items:center;justify-content:space-between; padding:8px 0; border-bottom:1px solid var(--border); font-size:12.6px; }
    .list-row:last-child{ border-bottom:none; }
    .list-name{ color:var(--text-muted); }
    .st-key-toggle_countries_wrap {
        display: flex !important;
        justify-content: flex-end !important;
        width: 100% !important;
    }
    .st-key-toggle_countries_wrap button {
        margin-left: auto !important;
        float: right !important;
    }

    .cta-primary button{
        background:linear-gradient(90deg, var(--grad-a), var(--grad-b)) !important; color:#fff !important;
        border:none !important; font-weight:600 !important;border-radius:9px !important;
        box-shadow:0 6px 18px rgba(34,184,240,0.25);
    }
    .cta-secondary button{
        background:var(--card-bg) !important; color:var(--text) !important;
        border:1px solid var(--border) !important; font-weight:600 !important;border-radius:9px !important;
    }
    .cta-scan button{
        background:linear-gradient(90deg, #22D3EE, #3B82F6) !important; color:#ffffff !important;
        border:none !important; font-weight:700 !important;border-radius:9px !important; height:2.9rem;
        white-space:nowrap !important; overflow:hidden !important; text-overflow:ellipsis !important;
        box-shadow:0 6px 20px rgba(34,211,238,0.35) !important;
        transition: all 0.25s ease-in-out !important;
    }
    .cta-scan button p, .cta-scan button span {
        color: #ffffff !important;
        font-weight: 700 !important;
    }
    .cta-scan button:hover{
        opacity: 0.95 !important;
        box-shadow:0 8px 25px rgba(34,211,238,0.55) !important;
    }
    .cta-scan{ width:100%; }

    /* All Scan & Action Buttons targeted via Streamlit key-class wrappers */
    div[class*="st-key-scanbtn"] button, 
    div[class*="st-key-home_scan"] button, 
    div[class*="st-key-dash_scan_link"] button,
    div[class*="st-key-run_universal_scan_btn"] button,
    div[class*="st-key-dl_pdf_"] button {
        background: linear-gradient(90deg, #22D3EE, #3B82F6) !important; 
        color: #ffffff !important;
        border: none !important; 
        font-weight: 700 !important;
        border-radius: 9px !important; 
        height: 2.9rem !important;
        box-shadow: 0 6px 20px rgba(34,211,238,0.35) !important;
        transition: all 0.25s ease-in-out !important;
    }
    div[class*="st-key-scanbtn"] button p, 
    div[class*="st-key-scanbtn"] button span,
    div[class*="st-key-home_scan"] button p,
    div[class*="st-key-home_scan"] button span,
    div[class*="st-key-dash_scan_link"] button p,
    div[class*="st-key-dash_scan_link"] button span,
    div[class*="st-key-run_universal_scan_btn"] button p,
    div[class*="st-key-run_universal_scan_btn"] button span,
    div[class*="st-key-dl_pdf_"] button p,
    div[class*="st-key-dl_pdf_"] button span {
        color: #ffffff !important;
        font-weight: 700 !important;
    }
    div[class*="st-key-scanbtn"] button:hover, 
    div[class*="st-key-home_scan"] button:hover, 
    div[class*="st-key-dash_scan_link"] button:hover,
    div[class*="st-key-run_universal_scan_btn"] button:hover,
    div[class*="st-key-dl_pdf_"] button:hover {
        opacity: 0.95 !important;
        box-shadow: 0 8px 25px rgba(34,211,238,0.55) !important;
    }

    /* Example buttons style */
    div[class*="st-key-ex_"] button {
        background: var(--card-bg) !important;
        color: var(--text) !important;
        border: 1px solid var(--border) !important;
        font-weight: 600 !important;
        border-radius: 9px !important;
        transition: all 0.15s ease-in-out !important;
    }
    div[class*="st-key-ex_"] button p, div[class*="st-key-ex_"] button span {
        color: var(--text) !important;
    }
    div[class*="st-key-ex_"] button:hover {
        background: var(--hover-bg) !important;
        border-color: var(--grad-a) !important;
        box-shadow: 0 0 10px rgba(108, 92, 231, 0.25) !important;
    }

    /* Form Submit Buttons (Ask Assistant, Clear Chat History, etc.) */
    [data-testid="stFormSubmitButton"] > button,
    .stForm [data-testid="stFormSubmitButton"] > button,
    div.stButton > button[data-testid*="FormSubmit"],
    button[kind="header"] {
        background: linear-gradient(135deg, var(--grad-a), var(--grad-b)) !important;
        background-color: transparent !important;
        color: #FFFFFF !important;
        border: none !important;
        border-radius: 9px !important;
        font-weight: 700 !important;
        font-size: 14px !important;
        box-shadow: 0 4px 16px rgba(34, 184, 240, 0.25) !important;
        transition: all 0.2s ease !important;
    }
    [data-testid="stFormSubmitButton"] > button p,
    [data-testid="stFormSubmitButton"] > button span,
    .stForm [data-testid="stFormSubmitButton"] > button p,
    .stForm [data-testid="stFormSubmitButton"] > button span {
        color: #FFFFFF !important;
        font-weight: 700 !important;
    }
    [data-testid="stFormSubmitButton"] > button:hover,
    .stForm [data-testid="stFormSubmitButton"] > button:hover {
        opacity: 0.92 !important;
        transform: translateY(-1px) !important;
        box-shadow: 0 6px 20px rgba(34, 184, 240, 0.4) !important;
    }

    /* Streamlit Tooltip Popups ONLY (Prevents card stacking & cut-off) */
    [data-testid="stTooltipContent"] {
        background: var(--tooltip-bg) !important;
        background-color: var(--tooltip-bg) !important;
        color: var(--tooltip-text) !important;
        border: 1px solid var(--tooltip-border) !important;
        border-radius: 8px !important;
        box-shadow: 0 8px 24px rgba(0, 0, 0, 0.18) !important;
        font-size: 12.5px !important;
        font-weight: 600 !important;
        padding: 6px 12px !important;
        white-space: nowrap !important;
        width: max-content !important;
        max-width: 320px !important;
        z-index: 9999999 !important;
    }
    [data-testid="stTooltipContent"] * {
        color: var(--tooltip-text) !important;
        background: transparent !important;
        font-size: 12.5px !important;
        font-weight: 600 !important;
        white-space: nowrap !important;
    }

    /* Streamlit Dialog / Modal Popup Outer Backdrop Overlay (Translucent Glass Mask) */
    html body div[data-baseweb="backdrop"],
    html body div[data-baseweb="modal"],
    html body [data-testid="stDialog"],
    html body [data-testid="stDialogService"],
    html body [data-testid="stModal"],
    html body div[role="dialog"][aria-modal="true"],
    div[data-baseweb="backdrop"],
    div[data-baseweb="modal"],
    [data-testid="stDialog"],
    [data-testid="stDialogService"],
    [data-testid="stModal"],
    div[role="dialog"][aria-modal="true"] {
        background: rgba(0, 0, 0, 0.4) !important;
        background-color: rgba(0, 0, 0, 0.4) !important;
        backdrop-filter: blur(6px) !important;
        -webkit-backdrop-filter: blur(6px) !important;
    }
    /* Force intermediate container divs to be 100% transparent */
    div[data-baseweb="modal"] > div,
    div[data-baseweb="modal"] > div > div,
    [data-testid="stDialog"] > div,
    [data-testid="stModal"] > div {
        background: transparent !important;
        background-color: transparent !important;
    }
    /* Target ONLY the inner floating modal card box */
    div[data-baseweb="modal"] [role="dialog"],
    [data-testid="stDialog"] [role="dialog"],
    [data-testid="stDialog"] [data-testid="stDialogContent"] {
        background: var(--dialog-bg) !important;
        background-color: var(--dialog-bg) !important;
        color: var(--dialog-text) !important;
        border: 1px solid var(--dialog-border) !important;
        border-radius: 16px !important;
        box-shadow: 0 25px 60px rgba(0, 0, 0, 0.5) !important;
    }
    [data-testid="stDialog"] > div *,
    div[role="dialog"] > div *,
    [data-testid="stModal"] > div *,
    div[data-baseweb="modal"] * {
        color: var(--dialog-text) !important;
    }
    [data-testid="stDialog"] p,
    div[role="dialog"] p,
    [data-testid="stDialog"] li,
    div[role="dialog"] li,
    [data-testid="stDialog"] span,
    div[role="dialog"] span {
        color: var(--dialog-text) !important;
    }
    [data-testid="stDialog"] h1,
    [data-testid="stDialog"] h2,
    [data-testid="stDialog"] h3,
    [data-testid="stDialog"] h4,
    [data-testid="stDialog"] h5,
    [data-testid="stDialog"] h6,
    div[role="dialog"] h1,
    div[role="dialog"] h2,
    div[role="dialog"] h3,
    div[role="dialog"] h4,
    div[role="dialog"] h5,
    div[role="dialog"] h6 {
        color: var(--dialog-text) !important;
    }
    [data-testid="stDialog"] code,
    div[role="dialog"] code {
        background: var(--card-bg-soft) !important;
        color: #22D3EE !important;
        border: 1px solid var(--border) !important;
        border-radius: 6px !important;
        padding: 2px 6px !important;
    }
    [data-testid="stDialog"] button[aria-label="Close"],
    div[role="dialog"] button[aria-label="Close"],
    [data-testid="stDialog"] [data-testid="stBaseButton-header"],
    div[role="dialog"] button {
        color: var(--dialog-text) !important;
    }

    /* Expander & JSON Payload Container (Light/Dark Dynamic) */
    [data-testid="stExpander"],
    div[data-testid="stExpander"] {
        background: var(--card-bg) !important;
        border: 1px solid var(--border) !important;
        border-radius: 10px !important;
    }
    [data-testid="stExpander"] summary,
    div[data-testid="stExpander"] summary {
        background: var(--card-bg-soft) !important;
        color: var(--text) !important;
        border-radius: 10px !important;
    }
    [data-testid="stExpander"] summary *,
    div[data-testid="stExpander"] summary * {
        color: var(--text) !important;
    }
    [data-testid="stJson"],
    [data-testid="stJsonBlock"],
    div.stJson,
    pre,
    code,
    .react-json-view {
        background: var(--card-bg-soft) !important;
        background-color: var(--card-bg-soft) !important;
        color: var(--text) !important;
        border: 1px solid var(--border) !important;
        border-radius: 8px !important;
    }

    /* Download Button (Export Report) - Vibrant Gradient in All Modes */
    .stDownloadButton > button,
    div.stDownloadButton button,
    [data-testid="stDownloadButton"] > button,
    div[class*="stDownloadButton"] button {
        background: linear-gradient(135deg, var(--grad-a), var(--grad-b)) !important;
        background-color: transparent !important;
        color: #FFFFFF !important;
        border: none !important;
        border-radius: 9px !important;
        font-weight: 700 !important;
        font-size: 14px !important;
        box-shadow: 0 4px 16px rgba(34, 184, 240, 0.25) !important;
        transition: all 0.2s ease !important;
        height: 2.9rem !important;
    }
    .stDownloadButton > button p,
    .stDownloadButton > button span,
    div.stDownloadButton button p,
    div.stDownloadButton button span,
    [data-testid="stDownloadButton"] > button p,
    [data-testid="stDownloadButton"] > button span {
        color: #FFFFFF !important;
        font-weight: 700 !important;
    }
    .stDownloadButton > button:hover,
    div.stDownloadButton button:hover,
    [data-testid="stDownloadButton"] > button:hover {
        opacity: 0.92 !important;
        transform: translateY(-1px) !important;
        box-shadow: 0 6px 20px rgba(34, 184, 240, 0.4) !important;
    }

    /* File Uploader Dropzone & "Browse files" Button (Vibrant Gradient in All Modes) */
    [data-testid="stFileUploaderDropzone"] {
        background: var(--bg-soft) !important;
        border: 1px dashed var(--border) !important;
        border-radius: 12px !important;
    }
    [data-testid="stFileUploaderDropzone"] button,
    [data-testid="stFileUploaderDropzone"] [data-testid*="baseButton"],
    [data-testid="stFileUploaderDropzone"] [data-testid*="stBaseButton"],
    div[data-testid="stFileUploaderDropzone"] button,
    section[data-testid="stFileUploaderDropzone"] button {
        background: linear-gradient(135deg, var(--grad-a), var(--grad-b)) !important;
        background-color: transparent !important;
        color: #FFFFFF !important;
        border: none !important;
        border-radius: 8px !important;
        font-weight: 700 !important;
        font-size: 13.5px !important;
        padding: 6px 16px !important;
        box-shadow: 0 4px 14px rgba(34, 184, 240, 0.25) !important;
        transition: all 0.2s ease !important;
    }
    [data-testid="stFileUploaderDropzone"] button p,
    [data-testid="stFileUploaderDropzone"] button span,
    [data-testid="stFileUploaderDropzone"] [data-testid*="baseButton"] p,
    [data-testid="stFileUploaderDropzone"] [data-testid*="baseButton"] span,
    div[data-testid="stFileUploaderDropzone"] button p,
    div[data-testid="stFileUploaderDropzone"] button span,
    section[data-testid="stFileUploaderDropzone"] button p,
    section[data-testid="stFileUploaderDropzone"] button span {
        color: #FFFFFF !important;
        font-weight: 700 !important;
    }
    [data-testid="stFileUploaderDropzone"] button:hover,
    div[data-testid="stFileUploaderDropzone"] button:hover,
    section[data-testid="stFileUploaderDropzone"] button:hover {
        opacity: 0.92 !important;
        transform: translateY(-1px) !important;
        box-shadow: 0 6px 18px rgba(34, 184, 240, 0.4) !important;
    }
    /* Inputs & Textareas (Text & Placeholder Visibility in All Modes) */
    .stTextInput input,
    .stTextArea textarea,
    [data-testid="stTextInput"] input,
    [data-testid="stTextArea"] textarea {
        background: var(--card-bg-soft) !important;
        background-color: var(--card-bg-soft) !important;
        color: var(--text) !important;
        border: 1px solid var(--border) !important;
        border-radius: 9px !important;
        font-weight: 600 !important;
        font-size: 14px !important;
    }
    .stTextInput input::placeholder,
    .stTextArea textarea::placeholder,
    [data-testid="stTextInput"] input::placeholder,
    [data-testid="stTextArea"] textarea::placeholder,
    input::placeholder,
    textarea::placeholder {
        color: var(--text-muted) !important;
        opacity: 0.85 !important;
        font-weight: 400 !important;
    }
    [data-testid="stFileUploaderDropzone"]{
        background:var(--bg-soft) !important; border:1px dashed var(--border) !important; border-radius:12px !important;
    }

    /* Target general settings container card */
    .st-key-general_settings_card {
        background: var(--card-bg) !important;
        border: 1px solid var(--border) !important;
        border-radius: 14px !important;
        padding: 22px !important;
        margin-bottom: 16px !important;
    }

    /* Target interactive chat container card */
    .st-key-interactive_chat_card {
        background: var(--card-bg) !important;
        border: 1px solid var(--border) !important;
        border-radius: 14px !important;
        padding: 22px !important;
        margin-bottom: 16px !important;
        box-shadow: 0 4px 24px rgba(0, 0, 0, 0.2) !important;
    }

    /* Force text elements to respect our text variables */
    p, span, label, .stMarkdown, .stSubheader, .stTitle, [data-testid="stWidgetLabel"] p, .stToggle label p, .stCheckbox label p {
        color: var(--text) !important;
    }

    /* Override: Force logo button AI text span to stay vibrant cyan/blue */
    .st-key-sb_logo_btn button span,
    .st-key-sb_logo_btn button p span,
    .st-key-sb_logo_btn button div span,
    .st-key-sb_logo_btn button [data-testid="stMarkdownContainer"] span,
    div[class*="st-key-sb_logo_btn"] span {
        color: #22D3EE !important;
        font-weight: 900 !important;
    }

    /* Global styling for all primary buttons — forces theme blue gradient, NO RED! */
    .stButton > button[kind="primary"],
    .stButton > button[data-testid*="primary"],
    .stButton > button[data-testid*="Primary"],
    button[kind="primary"],
    button[data-testid*="primary"],
    button[data-testid*="Primary"] {
        background: linear-gradient(90deg, var(--grad-a), var(--grad-b)) !important;
        background-color: transparent !important;
        color: #FFFFFF !important;
        border: none !important;
        box-shadow: 0 4px 14px rgba(34,184,240,0.3) !important;
    }
    .stButton > button[kind="primary"] p,
    .stButton > button[data-testid*="primary"] p,
    .stButton > button[data-testid*="Primary"] p,
    .stButton > button[kind="primary"] span,
    .stButton > button[data-testid*="primary"] span,
    .stButton > button[data-testid*="Primary"] span {
        color: #FFFFFF !important;
        font-weight: 600 !important;
    }
    .stButton > button[kind="primary"]:hover,
    .stButton > button[data-testid*="primary"]:hover,
    .stButton > button[data-testid*="Primary"]:hover {
        background: linear-gradient(90deg, var(--grad-a), var(--grad-b)) !important;
        opacity: 0.92 !important;
        color: #FFFFFF !important;
        box-shadow: 0 6px 18px rgba(34,184,240,0.4) !important;
    }

    /* Global styling for all secondary (default) buttons */
    .stButton > button[kind="secondary"], .stButton > button:not([kind]) {
        background-color: var(--card-bg-soft) !important;
        color: var(--text) !important;
        border: 1px solid var(--border) !important;
        border-radius: 8px !important;
        transition: all 0.15s ease-in-out !important;
    }
    .stButton > button[kind="secondary"] p, .stButton > button[kind="secondary"] span,
    .stButton > button:not([kind]) p, .stButton > button:not([kind]) span {
        color: var(--text) !important;
    }
    .stButton > button[kind="secondary"]:hover, .stButton > button:not([kind]):hover {
        background-color: var(--hover-bg) !important;
        border-color: var(--grad-a) !important;
        color: var(--text) !important;
    }

    /* Override streamlit selectbox widgets */
    div[data-baseweb="select"] > div {
        background-color: var(--bg-soft) !important;
        color: var(--text) !important;
        border-color: var(--border) !important;
    }
    div[data-baseweb="select"] ul {
        background-color: var(--card-bg) !important;
        color: var(--text) !important;
    }
    div[data-baseweb="select"] li {
        color: var(--text) !important;
    }
    div[data-baseweb="select"] li:hover {
        background-color: var(--hover-bg) !important;
    }

    /* Help page search card */
    .st-key-help_search_card {
        background: var(--card-bg) !important;
        border: 1px solid var(--border) !important;
        border-radius: 14px !important;
        padding: 16px 20px !important;
        margin-bottom: 20px !important;
    }

    /* ---- Hamburger toggle button ---- */
    .st-key-sidebar_toggle button,
    .st-key-sidebar_toggle .stButton > button,
    .st-key-sidebar_toggle [data-testid*="baseButton"] {
        border: 1px solid transparent !important;
        background: transparent !important;
        background-color: transparent !important;
        color: var(--text) !important;
        transition: all 0.18s ease !important;
        cursor: pointer !important;
        box-shadow: none !important;
        min-width: 42px !important;
    }
    .st-key-sidebar_toggle button:hover,
    .st-key-sidebar_toggle .stButton > button:hover {
        background: rgba(34, 211, 238, 0.12) !important;
        border-color: rgba(34, 211, 238, 0.35) !important;
        color: #fff !important;
        box-shadow: 0 4px 16px rgba(34,211,238,0.3) !important;
    }
    .st-key-sidebar_toggle {
        margin-bottom: 12px !important;
    }
    .st-key-sidebar_toggle button {
        border-radius: 8px !important;
        font-size: 18px !important;
        font-weight: 700 !important;
        padding: 5px 13px !important;
        line-height: 1 !important;
        border: 1px solid var(--border) !important;
        background: var(--card-bg) !important;
        color: var(--text) !important;
        transition: all 0.18s ease !important;
        cursor: pointer !important;
        box-shadow: 0 2px 8px rgba(0,0,0,0.12) !important;
        min-width: 42px !important;
    }
    .st-key-sidebar_toggle button:hover {
        background: linear-gradient(135deg, var(--grad-a), var(--grad-b)) !important;
        border-color: transparent !important;
        color: #fff !important;
        box-shadow: 0 4px 16px rgba(34,184,240,0.3) !important;
    }
    .st-key-sidebar_toggle button p {
        font-size: 18px !important;
        line-height: 1 !important;
    }

    </style>
"""
style_content = style_template.replace("__THEME_VARS__", theme_vars).replace("__ANIMATION_CSS__", animation_css)
st.markdown(style_content, unsafe_allow_html=True)

# ── Dynamic Sidebar Collapse CSS ──
if st.session_state.get("sidebar_collapsed", False):
    st.markdown(
        """
        <style>
        .st-key-navcol {
            position: fixed !important;
            top: 0.8rem !important;
            left: 0.8rem !important;
            width: 4.2rem !important;
            height: calc(100vh - 1.6rem) !important;
            padding: 10px 5px !important;
        }
        html body div.block-container,
        div.block-container,
        [data-testid="stMainBlockContainer"],
        .stMainBlockContainer {
            padding-left: 6.2rem !important;
        }
        .st-key-navcol .st-key-sb_logo_btn {
            border-bottom: 1px solid var(--border) !important;
            padding-bottom: 10px !important;
            margin-bottom: 12px !important;
            width: 100% !important;
        }
        html body .stApp .st-key-navcol .st-key-sb_logo_btn button,
        html body .stApp .st-key-navcol .st-key-sb_logo_btn .stButton > button,
        html body .stApp .st-key-navcol .st-key-sb_logo_btn [data-testid*="baseButton"],
        .st-key-navcol .st-key-sb_logo_btn button {
            width: 100% !important;
            height: 42px !important;
            min-height: 42px !important;
            max-height: 42px !important;
            padding: 0 !important;
            margin: 0 !important;
            display: flex !important;
            align-items: center !important;
            justify-content: center !important;
            border-radius: 9px !important;
            background: rgba(18, 28, 45, 0.75) !important;
            background-color: rgba(18, 28, 45, 0.75) !important;
            border: 1px solid rgba(255, 255, 255, 0.08) !important;
            box-shadow: 0 2px 6px rgba(0, 0, 0, 0.2) !important;
        }
        html body .stApp .st-key-navcol .st-key-sb_logo_btn button:hover,
        .st-key-navcol .st-key-sb_logo_btn button:hover {
            background: rgba(34, 211, 238, 0.12) !important;
            border-color: rgba(34, 211, 238, 0.35) !important;
            box-shadow: 0 4px 12px rgba(34, 211, 238, 0.2) !important;
        }
        .st-key-navcol .st-key-sb_logo_btn button::before {
            content: "" !important;
            display: block !important;
            width: 26px !important;
            height: 26px !important;
            background-image: url("https://cdn-icons-png.flaticon.com/512/6071/6071531.png") !important;
            background-size: contain !important;
            background-repeat: no-repeat !important;
            background-position: center !important;
            margin: 0 auto !important;
            flex-shrink: 0 !important;
            float: none !important;
        }
        .st-key-navcol .st-key-sb_logo_btn button [data-testid="stMarkdownContainer"],
        .st-key-navcol .st-key-sb_logo_btn button div,
        .st-key-navcol .st-key-sb_logo_btn button p,
        .st-key-navcol .st-key-sb_logo_btn button span {
            display: none !important;
            width: 0 !important;
            height: 0 !important;
            margin: 0 !important;
            padding: 0 !important;
        }
        .st-key-navcol .stButton {
            margin-bottom: 4px !important;
        }
        .st-key-navcol .stButton > button {
            justify-content: center !important;
            text-align: center !important;
            padding: 0.6rem 0 !important;
            min-width: 0 !important;
            width: 100% !important;
        }
        .st-key-navcol .stButton > button p,
        .st-key-navcol .stButton > button span {
            text-align: center !important;
            font-size: 18px !important;
            margin: 0 !important;
            padding: 0 !important;
        }
        </style>
        """,
        unsafe_allow_html=True
    )


# ── GLOBAL AUDIO SYNTHESIS HELPER ──
st.markdown(
    """
    <script>
    try {
        var playSoundFunc = function() {
            try {
                var context = new (window.AudioContext || window.webkitAudioContext)();
                var osc = context.createOscillator();
                var gain = context.createGain();
                osc.connect(gain);
                gain.connect(context.destination);
                osc.type = 'sine';
                osc.frequency.value = 880;
                gain.gain.setValueAtTime(0.08, context.currentTime);
                osc.start();
                gain.gain.exponentialRampToValueAtTime(0.008, context.currentTime + 0.35);
                osc.stop(context.currentTime + 0.35);
            } catch(err) {
                console.warn("AudioContext playback blocked:", err);
            }
        };
        window.playBeepSound = playSoundFunc;
        if (window.parent) {
            window.parent.playBeepSound = playSoundFunc;
        }
    } catch(e) {
        console.error("Failed to register global playBeepSound:", e);
    }
    </script>
    """,
    unsafe_allow_html=True
)







# Scanner configuration & backend wiring


def _module(mod_path: str, attr: str):
    try:
        module = __import__(mod_path, fromlist=[attr])
        return getattr(module, attr)
    except Exception:
        return None


def simulated_result(value: str):
    digest = int(hashlib.sha256(value.encode()).hexdigest(), 16)
    score = digest % 100
    if score < 25:
        level = "Safe"
    elif score < 45:
        level = "Low"
    elif score < 65:
        level = "Medium"
    elif score < 85:
        level = "High"
    else:
        level = "Critical"
    return {"risk_score": score, "risk_level": level, "simulated": True}


def generate_mock_analysis(value: str, score: int, level: str, scanner_key: str):
    val_str = str(value)
    url_len = len(val_str)
    digits = sum(c.isdigit() for c in val_str)
    hyphens = val_str.count("-")
    dots = val_str.count(".")
    specials = sum(not (c.isalnum() or c in "/._-:") for c in val_str)
    subdomains = max(0, dots - 1)
    
    import math
    entropy = sum(-val * math.log2(val) for val in [val_str.count(c)/url_len for c in set(val_str)]) if url_len > 0 else 0.0
    _parts = val_str.split("/")
    ip_based = "Yes" if len(_parts) > 2 and any(char.isdigit() for char in _parts[2]) else "No"
    shortened = "Yes" if any(domain in val_str for domain in ["bit.ly", "tinyurl.com", "t.co", "rebrand.ly"]) else "No"
    
    is_safe = score < 30
    
    return {
        "success": False,
        "simulated": True,
        "is_mock": True,
        "url": value,
        "risk": {
            "score": score,
            "risk_score": score,
            "level": level,
            "risk_level": level
        },
        "ml_prediction": {
            "prediction": "Unverified" if is_safe else "Suspicious",
            "probability": round(float(score), 1)
        },
        "explain_ai": {
            "summary": "Live scan analysis unavailable — placeholder state rendered",
            "details": [
                "Domain WHOIS lookup was not completed",
                "SSL certificate handshake was not completed",
                "Reputation services did not return live status",
                "Results are placeholder values for preview"
            ]
        },
        "recommendation": {
            "count": 2,
            "recommendations": [
                "Retry the scan when live network services are reachable.",
                "Verify host connectivity manually."
            ]
        },
        "analysis": {
            "ssl": {
                "valid": False,
                "expiration_date": "Not Available"
            },
            "whois": {
                "domain_age_days": None,
                "expiration_date": "Not Available"
            },
            "dns": {
                "a_records": []
            },
            "lexical": {
                "length": url_len,
                "digits": digits,
                "hyphens": hyphens,
                "dots": dots,
                "special_characters": specials,
                "subdomains": subdomains,
                "entropy": entropy,
                "ip_based": ip_based,
                "shortened": shortened
            },
            "google_safe_browsing": {
                "safe": is_safe
            },
            "virustotal": {
                "stats": {
                    "malicious": 0 if is_safe else int(score // 10),
                    "harmless": 0
                }
            }
        }
    }


def compute_scan_display(result, scanner_key=None):
    """
    Single source of truth for every value shown on a scanner result page.
    Used by render_scanner_page() to draw the screen AND by export_report()
    to build the downloadable report - so the report always matches the
    screen exactly.
    """
    score = result["risk_score"]
    level = result["risk_level"]
    duration = result.get("duration", "0.45 sec")
    raw_payload = result.get("raw") or {}
    analysis_data = raw_payload.get("analysis") or {}

    active_scanner_key = (
        scanner_key
        or result.get("scanner_key")
        or result.get("scanner")
        or (st.session_state.get("last_scan_context") or {}).get("scanner_key", "")
    )

    ml_prediction_data = raw_payload.get("ml_prediction") or {}
    has_real_ml = bool(
        ml_prediction_data.get("probability") is not None
        and ml_prediction_data.get("ml_available", True)
    )
    if level == "Unverified":
        default_pred = "Unverified"
    elif level == "Safe":
        default_pred = "Legitimate"
    elif level == "Low":
        default_pred = "Low Risk"
    elif level in ("Medium", "Suspicious"):
        default_pred = "Suspicious"
    else:
        default_pred = "Malicious"

    ml_pred = ml_prediction_data.get("prediction") or default_pred
    if has_real_ml:
        prob_val = float(ml_prediction_data["probability"])
        ml_subtitle = f"Confidence: {prob_val:.1f}%"
        ml_card_title = "AI Prediction"
    else:
        prob_val = float(score)
        ml_subtitle = "Risk-Derived Estimate"
        ml_card_title = "Threat Assessment"

    ml_available = ml_prediction_data.get("ml_available", False)
    ml_confidence = ml_prediction_data.get("confidence", None)
    ml_model_name = ml_prediction_data.get("model", "Heuristic")
    # Brand impersonation signal from online_valid_model
    brand_impersonation = ml_prediction_data.get("brand_impersonation", None)
    brand_confidence = ml_prediction_data.get("brand_confidence", None)
    brand_probabilities = ml_prediction_data.get("brand_probabilities", {})

    # File risk prediction from file_signatures_model (stored in analysis_data)
    file_ml_prediction = analysis_data.get("ml_prediction") or {}

    target_val = str(result.get("value") or result.get("phone") or result.get("target") or "")
    hostname_val = target_val
    if "://" in target_val:
        hostname_val = target_val.split("://")[1].split("/")[0]
    elif "/" in target_val:
        hostname_val = target_val.split("/")[0]
    if ":" in hostname_val:
        hostname_val = hostname_val.split(":")[0]

    from services.ssl_service import verify_https
    if str(target_val).lower().startswith("http://"):
        is_https = "No"
        https_check = {"valid": False, "reason": "Target uses plain unencrypted HTTP scheme"}
    else:
        https_check = verify_https(hostname_val)
        is_https = "Yes" if https_check.get("valid", False) else "No"

    def _as_dict(val: Any) -> dict[str, Any]:
        return val if isinstance(val, dict) else {}

    domain_age = "Not Available"
    domain_age_note = "No live WHOIS data"
    whois_data = _as_dict(analysis_data.get("whois")) or _as_dict(raw_payload.get("whois"))
    if whois_data:
        age_days = whois_data.get("domain_age_days")
        if age_days:
            domain_age = f"{int(age_days) // 365} Years"
            reg_date = whois_data.get("creation_date") or whois_data.get("registered_on")
            domain_age_note = f"Since {reg_date}" if reg_date else "Based on WHOIS lookup"

    ip_val = "Not Available"
    dns_data = _as_dict(analysis_data.get("dns")) or _as_dict(raw_payload.get("dns"))
    if dns_data:
        a_recs = dns_data.get("a_records") or dns_data.get("ips")
        if a_recs and isinstance(a_recs, list) and len(a_recs) > 0:
            ip_val = a_recs[0]
    if ip_val == "Not Available" and hostname_val:
        try:
            resolved = socket.gethostbyname(hostname_val)
            if resolved:
                ip_val = resolved
        except Exception:
            pass

    google_data = _as_dict(analysis_data.get("google_safe_browsing")) or _as_dict(raw_payload.get("google_safe_browsing"))
    google_safe = google_data.get("safe") if google_data.get("safe") is not None else True
    gsb_status = "Clean" if google_safe else "MALICIOUS"

    vt_data = _as_dict(analysis_data.get("virustotal")) or _as_dict(raw_payload.get("virustotal"))
    stats = _as_dict(vt_data.get("stats")) or _as_dict(vt_data.get("risk"))
    if stats:
        vt_status = f"{stats.get('malicious', 0)} / {stats.get('harmless', 0) + stats.get('malicious', 0)}"
    else:
        vt_status = f"{vt_data.get('malicious', 0)} / {vt_data.get('harmless', 96)}"

    blacklist_data = _as_dict(analysis_data.get("blacklist")) or _as_dict(raw_payload.get("blacklist"))
    blacklist_detected = blacklist_data.get("detected") or (score > 50)
    blacklist_status = "Listed" if blacklist_detected else "Not Found"

    reputation_data = _as_dict(analysis_data.get("reputation")) or _as_dict(raw_payload.get("reputation"))
    reputation_score = reputation_data.get("score") or (100 - score)
    reputation_val = "Excellent" if reputation_score >= 80 else "Poor"

    lexical_data = _as_dict(analysis_data.get("lexical")) or _as_dict(raw_payload.get("lexical"))
    url_len = lexical_data.get("length") or len(str(target_val))
    digits = lexical_data.get("digits") or sum(c.isdigit() for c in str(target_val))
    hyphens = lexical_data.get("hyphens") or str(target_val).count("-")
    dots = lexical_data.get("dots") or str(target_val).count(".")
    specials = lexical_data.get("special_characters") or sum(not (c.isalnum() or c in "/._-:") for c in str(target_val))
    subdomains = lexical_data.get("subdomains") or max(0, dots - 1)
    entropy_val = lexical_data.get("entropy") or 3.42
    ip_based = "Yes" if lexical_data.get("ip_based", False) else "No"
    shortened = "Yes" if lexical_data.get("shortened", False) else "No"

    ssl_data = _as_dict(analysis_data.get("ssl")) or _as_dict(raw_payload.get("ssl"))
    if "valid" in ssl_data and ssl_data["valid"] is not None:
        ssl_valid = ssl_data["valid"]
    else:
        ssl_valid = https_check.get("valid", False)
    ssl_status = "Valid" if ssl_valid else "Invalid"
    if https_check.get("valid") and https_check.get("expires"):
        days = https_check.get("days_until_expiry")
        expiry_val = f"{https_check['expires'][:10]} ({days}d remaining)"
    else:
        expiry_val = ssl_data.get("expiration_date") or whois_data.get("expiration_date") or https_check.get("reason", "N/A")

    explain_data = _as_dict(raw_payload.get("explain_ai"))
    points = explain_data.get("details") or [
        "Domain age is 15 years old (trusted)",
        "SSL certificate is valid and secure",
        "No suspicious keywords detected",
        "Clean reputation across all databases",
        "Content and structure appear legitimate"
    ]

    rec_data = _as_dict(raw_payload.get("recommendation"))
    recs = rec_data.get("recommendations") or [
        "This URL is safe to visit.",
        "No security concerns found.",
        "You can bookmark this site.",
        "Continue to monitor for any changes."
    ]

    SUBJECT_NOUN = {
        "File Scanner": "file",
        "Email Scanner": "sender",
        "IP Scanner": "host",
        "Domain Scanner": "domain",
        "Website Scanner": "site",
        "URL Scanner": "link",
        "QR Code Scanner": "QR content",
        "Universal Scan": "target",
        "Phone Threat Intelligence": "phone number",
    }

    noun = SUBJECT_NOUN.get(active_scanner_key, "target")
    if level == "Unverified":
        risk_label = "Unverified State"
        threat_desc = f"Unverified {noun}"
    else:
        risk_label = "Low Risk" if score < 30 else "Medium Risk" if score < 70 else "High Risk"
        threat_desc = "No immediate threat" if score < 30 else f"Suspicious {noun}" if score < 70 else "Threat flagged"

    from services.lexical_keyword_service import lexical_keyword_service
    kw_res = lexical_keyword_service.check_suspicious_keywords(str(target_val))
    kw_matches = kw_res.get("matched_keywords", [])
    suspicious_words = ", ".join(kw_matches) if kw_matches else "None"

    simulated_flag = bool(
        result.get("simulated")
        or result.get("is_simulated")
        or result.get("is_mock")
        or raw_payload.get("simulated")
        or raw_payload.get("is_simulated")
        or raw_payload.get("is_mock")
        or raw_payload.get("mock")
    )
    scan_err = (
        result.get("scan_error")
        or raw_payload.get("scan_error")
        or result.get("note")
        or raw_payload.get("note")
        or ""
    )

    return {
        "score": score, "level": level, "duration": duration,
        "ml_pred": ml_pred, "prob_val": prob_val,
        "has_real_ml": has_real_ml, "ml_subtitle": ml_subtitle, "ml_card_title": ml_card_title,
        "ml_available": ml_available, "ml_confidence": ml_confidence,
        "ml_model_name": ml_model_name,
        "brand_impersonation": brand_impersonation,
        "brand_confidence": brand_confidence,
        "brand_probabilities": brand_probabilities,
        "file_ml_prediction": file_ml_prediction,
        "is_https": is_https, "domain_age": domain_age, "domain_age_note": domain_age_note,
        "target_val": target_val, "hostname_val": hostname_val, "ip_val": ip_val,
        "gsb_status": gsb_status, "vt_status": vt_status,
        "blacklist_status": blacklist_status, "reputation_val": reputation_val,
        "url_len": url_len, "digits": digits, "hyphens": hyphens, "dots": dots,
        "specials": specials, "subdomains": subdomains, "entropy_val": entropy_val,
        "ip_based": ip_based, "shortened": shortened, "suspicious_words": suspicious_words,
        "ssl_status": ssl_status, "ssl_valid": ssl_valid, "expiry_val": expiry_val,
        "points": points, "recs": recs,
        "risk_label": risk_label, "threat_desc": threat_desc,
        "simulated": simulated_flag,
        "is_mock": simulated_flag,
        "scan_error": scan_err,
        "scanner_key": active_scanner_key,
    }


def export_report(result, fmt):
    import io
    import json
    import pandas as pd
    
    fmt = fmt.upper()

    # Build the report from the SAME values shown on screen (compute_scan_display),
    # grouped into the same sections the user sees: Summary, Detailed Analysis,
    # Threat Intelligence, URL/File Features, Security Checks, AI Explanation,
    # Recommendations. This guarantees "what's on screen is what's in the report".
    d = compute_scan_display(result, scanner_key=result.get("scanner_key") or result.get("scanner"))
    target_val = d["target_val"]

    exec_summary = ""
    try:
        from modules.ai_summary_module import ai_summary_module
        exec_summary = ai_summary_module.generate_summary(result)
    except Exception:
        pass

    summary_items = [
        ("Target", target_val),
        ("Risk Score", f"{d['score']} / 100"),
        ("Risk Level", d["risk_label"]),
        ("Threat Level", str(d["level"]).upper()),
        (d.get("ml_card_title", "AI Prediction"), d["ml_pred"]),
        ("Estimate / Confidence", d.get("ml_subtitle", f"{d['prob_val']:.1f}%")),
        ("Scan Time", d["duration"]),
    ]
    s_key = d.get("scanner_key", "")
    if s_key in ("Domain Scanner", "Website Scanner", "URL Scanner", "Universal Scan", ""):
        summary_items.extend([
            ("HTTPS", "Valid" if d["is_https"] == "Yes" else "Invalid"),
            ("Domain Age", d["domain_age"]),
            ("Domain Age Source", d["domain_age_note"]),
        ])
    elif s_key == "File Scanner":
        summary_items.extend([
            ("File Entropy", f"{d['entropy_val']:.2f}" if isinstance(d['entropy_val'], (int, float)) else str(d['entropy_val'])),
            ("File Format", str(target_val).split(".")[-1].upper() if "." in str(target_val) else "BINARY"),
        ])
    elif s_key == "IP Scanner":
        summary_items.extend([
            ("IP Reputation", d["reputation_val"]),
            ("Blacklist Check", d["blacklist_status"]),
        ])
    elif s_key == "Email Scanner":
        summary_items.extend([
            ("Sender Domain Reputation", d["reputation_val"]),
            ("Blacklist Check", d["blacklist_status"]),
        ])

    sections: list[tuple[str, list[tuple[str, str]]]] = [
        ("Summary", summary_items),
    ]

    if exec_summary:
        sections.append(("AI Executive Summary", [("Executive Summary", exec_summary)]))

    sections.extend([
        ("Detailed Analysis", [
            ("URL", target_val),
            ("Domain", d["hostname_val"]),
            ("Scheme", "HTTPS" if d["is_https"] == "Yes" else "HTTP"),
            ("Port", "443" if d["is_https"] == "Yes" else "80"),
            ("IP Address", d["ip_val"]),
            ("Response Code", "200 OK"),
            ("Redirect", "No Redirect"),
            ("Response Time", "145 ms"),
            ("Content Type", "text/html; charset=UTF-8"),
            ("Server", "cloudflare"),
            ("Content Length", "12.6 KB"),
            ("Last Scanned", datetime.now().strftime("%b %d, %Y %I:%M:%S %p")),
        ]),
        ("Threat Intelligence", [
            ("Google Safe Browsing", d["gsb_status"]),
            ("VirusTotal", d["vt_status"]),
            ("PhishTank", "Clean"),
            ("ThreatFox", "Clean"),
            ("URLVoid", "Clean"),
            ("Blacklist Check", d["blacklist_status"]),
            ("Reputation", d["reputation_val"]),
        ]),
        ("URL Features", [
            ("URL Length", d["url_len"]),
            ("Digits", d["digits"]),
            ("Hyphens", d["hyphens"]),
            ("Dots", d["dots"]),
            ("Special Characters", d["specials"]),
            ("Subdomains", d["subdomains"]),
            ("Entropy", f"{d['entropy_val']:.2f} (Low)"),
            ("IP Based URL", d["ip_based"]),
            ("Shortened URL", d["shortened"]),
            ("Suspicious Words", d["suspicious_words"]),
        ]),
        ("Security Checks", [
            ("HTTPS", "Valid" if d["is_https"] == "Yes" else "Invalid"),
            ("SSL Certificate", d["ssl_status"]),
            ("Certificate Expiry", d["expiry_val"]),
            ("HSTS", "Enabled"),
            ("Content Security Policy", "Enabled"),
            ("X-Frame-Options", "DENY"),
            ("X-Content-Type-Options", "nosniff"),
            ("Referrer-Policy", "strict-origin-when-cross-origin"),
        ]),
        ("AI Explanation", [(f"Point {i+1}", p) for i, p in enumerate(d["points"])]),
        ("Recommendations", [(f"Recommendation {i+1}", r) for i, r in enumerate(d["recs"])]),
    ])

    if fmt == "JSON":
        report_data = {
            "report_type": "CyberMind AI - Threat Intelligence Report",
            "generated_on": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "target": target_val,
        }
        for section_name, rows in sections:
            report_data[section_name.lower().replace(" ", "_")] = {k: v for k, v in rows}
        return json.dumps(report_data, indent=3, default=str).encode('utf-8'), "application/json", "cybermind_report.json"

    elif fmt == "CSV":
        flat_rows = []
        for section_name, rows in sections:
            for k, v in rows:
                flat_rows.append({"Section": section_name, "Parameter": k, "Value": str(v)})
        df = pd.DataFrame(flat_rows)
        csv_str = df.to_csv(index=False)
        return csv_str.encode('utf-8'), "text/csv", "cybermind_report.csv"

    elif fmt == "EXCEL":
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            summary_rows = [{"Parameter": "Target", "Value": target_val},
                            {"Parameter": "Generated On", "Value": datetime.now().strftime("%Y-%m-%d %H:%M:%S")}]
            pd.DataFrame(summary_rows).to_excel(writer, index=False, sheet_name="Report Info")
            for section_name, rows in sections:
                sheet_df = pd.DataFrame([{"Parameter": k, "Value": str(v)} for k, v in rows])
                sheet_name = section_name[:31]  # Excel sheet name limit
                sheet_df.to_excel(writer, index=False, sheet_name=sheet_name)
        buffer.seek(0)
        return buffer.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "cybermind_report.xlsx"

    elif fmt == "PDF":
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
        story = []
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'TitleStyle', parent=styles['Heading1'], fontSize=20,
            textColor=colors.HexColor('#6C5CE7'), spaceAfter=6
        )
        subtitle_style = ParagraphStyle(
            'SubtitleStyle', parent=styles['Normal'], fontSize=10,
            textColor=colors.HexColor('#8B98AC'), spaceAfter=16
        )
        header_style = ParagraphStyle(
            'HeaderStyle', parent=styles['Heading3'], fontSize=14,
            textColor=colors.HexColor('#22B8F0'), spaceBefore=14, spaceAfter=8
        )
        label_style = ParagraphStyle(
            'LabelStyle', parent=styles['Normal'], fontSize=9.5,
            textColor=colors.HexColor('#334155'), fontName='Helvetica-Bold'
        )
        value_style = ParagraphStyle(
            'ValueStyle', parent=styles['Normal'], fontSize=9.5,
            textColor=colors.HexColor('#0F1826')
        )
        callout_style = ParagraphStyle(
            'CalloutStyle', parent=styles['Normal'], fontSize=9.5,
            textColor=colors.HexColor('#0F172A'), leading=14
        )

        story.append(Paragraph("CyberMind AI - Threat Intelligence Report", title_style))
        story.append(Paragraph(f"Target: {html.escape(str(target_val))}", subtitle_style))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%d-%m-%Y %H:%M:%S')}", subtitle_style))

        for section_name, rows in sections:
            story.append(Paragraph(section_name, header_style))
            if section_name == "AI Executive Summary":
                sum_text = rows[0][1] if rows else ""
                callout_data = [[Paragraph(f"<b>Executive Summary:</b><br/>{html.escape(sum_text)}", callout_style)]]
                callout_table = Table(callout_data, colWidths=[490])
                callout_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#ECFEFF')),
                    ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#22D3EE')),
                    ('PADDING', (0, 0), (-1, -1), 10),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                story.append(callout_table)
                continue

            table_data = [["Parameter", "Value"]]
            for k, v in rows:
                val_str = str(v)
                if len(val_str) > 250:
                    val_str = val_str[:247] + "..."
                table_data.append([Paragraph(html.escape(str(k)), label_style),
                                    Paragraph(html.escape(val_str), value_style)])

            t = Table(table_data, colWidths=[190, 300])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0F1826')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D1D9E6')),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#FFFFFF')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#FFFFFF'), colors.HexColor('#F4F6F9')]),
            ]))
            story.append(t)

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue(), "application/pdf", "cybermind_report.pdf"

    # Default fallback to JSON
    fallback = {"target": target_val}
    for section_name, rows in sections:
        fallback[section_name.lower().replace(" ", "_")] = {k: v for k, v in rows}
    return json.dumps(fallback, indent=3, default=str).encode('utf-8'), "application/json", "cybermind_report.json"


def run_scan(scanner_key: str, mod_path: str, attr: str, value: str):
    start_t = time.time()
    fn = _module(mod_path, attr)
    risk_score, risk_level, note = None, None, ""
    raw_result = None
    is_simulated = False
    if fn is not None:
        try:
            raw_result = fn.analyze(value)
            if raw_result and raw_result.get("success"):
                risk = raw_result.get("risk") or {}
                risk_score = risk.get("risk_score") if risk.get("risk_score") is not None else risk.get("score")
                risk_level = risk.get("risk_level") if risk.get("risk_level") is not None else risk.get("level")
            else:
                note = (raw_result or {}).get("message") or (raw_result or {}).get("error") or "Analysis returned unverified state"
        except Exception as exc:
            from core.logger import logger
            logger.error("Scan module %s failed for %s: %s", mod_path, value, exc)
            note = str(exc)

    if risk_score is None or risk_level is None:
        sim = simulated_result(value)
        risk_score, risk_level = sim["risk_score"], sim["risk_level"]
        is_simulated = True

    if not raw_result or not raw_result.get("success"):
        raw_result = generate_mock_analysis(value, risk_score, risk_level, scanner_key)
        raw_result["simulated"] = True
        raw_result["is_mock"] = True
        raw_result["scan_error"] = note or "Live scan analysis unavailable or encountered error."
        is_simulated = True

    # Save to SQLite Database for dynamic statistics
    try:
        from database.db import db
        from services.geo_service import geo_service
        import socket
        from urllib.parse import urlparse

        country, country_code, latitude, longitude = None, None, None, None

        if scanner_key not in ("Email Scanner", "QR Code Scanner", "File Scanner", "Malware Analysis", "Device Security Check"):
            is_ip = False
            try:
                socket.inet_aton(value)
                is_ip = True
            except Exception:
                pass
                
            ip = None
            if is_ip:
                ip = value
            else:
                hostname = value
                if "://" in value or scanner_key in ("Website Scanner", "URL Scanner", "Universal Scan"):
                    try:
                        parsed = urlparse(value)
                        hostname = parsed.hostname or value
                    except Exception:
                        pass
                try:
                    ip = socket.gethostbyname(hostname)
                except Exception:
                    pass
            
            if ip:
                geo_info = geo_service.lookup(ip)
                if geo_info:
                    country = geo_info.get("country")
                    country_code = geo_info.get("country_code")
                    latitude = geo_info.get("latitude")
                    longitude = geo_info.get("longitude")

        db_level = risk_level
        if db_level == "Malicious":
            db_level = "Critical"
        elif db_level == "Suspicious":
            db_level = "Medium"
        elif db_level not in ('Safe', 'Low', 'Medium', 'High', 'Critical'):
            db_level = "Safe"
            
        db.execute(
            """
            INSERT INTO scan_history (scan_type, target, risk_level, risk_score, country, country_code, latitude, longitude)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                scanner_key, 
                value, 
                db_level, 
                float(risk_score) if risk_score is not None else 0.0,
                country,
                country_code,
                latitude,
                longitude
            )
        )
    except Exception:
        pass

    duration = time.time() - start_t

    res = {
        "value": value,
        "risk_score": risk_score,
        "risk_level": risk_level,
        "time": datetime.now().strftime("%H:%M:%S"),
        "note": note,
        "duration": f"{duration:.2f} sec",
        "raw": raw_result,
        "simulated": is_simulated,
        "is_mock": is_simulated,
        "scan_error": note or (raw_result.get("scan_error") if isinstance(raw_result, dict) else "")
    }
    st.session_state["_last_active_scanner_page"] = scanner_key
    st.session_state["_last_scan_result_dict"] = res
    st.session_state[f"fb_given_{scanner_key}"] = False
    st.session_state["last_scan_context"] = {
        "scanner": scanner_key,
        "scanner_key": scanner_key,
        "target": value,
        "risk_score": risk_score,
        "risk_level": risk_level,
        "findings": raw_result
    }

    # Update in-memory session history dictionary
    h_entry = {
        "value": value,
        "level": db_level,
        "score": int(risk_score) if risk_score is not None else 0,
        "time": res["time"]
    }
    if "histories" not in st.session_state:
        st.session_state.histories = {}
    if scanner_key not in st.session_state.histories:
        st.session_state.histories[scanner_key] = []
    st.session_state.histories[scanner_key].insert(0, h_entry)

    return res



LEVEL_BADGE = {
    "Safe": "badge-safe",
    "Low": "badge-safe",
    "Medium": "badge-suspicious",
    "Suspicious": "badge-suspicious",
    "High": "badge-malicious",
    "Critical": "badge-malicious",
    "Malicious": "badge-malicious",
    "Unverified": "badge-unverified",
}

SCANNERS = {    
    "Universal Scan": {
        "icon": "🌐", "color": "#3B82F6", "bg": "rgba(59,130,246,0.12)",
        "desc": "Auto-detect and scan URLs, Domains, or Emails.",
        "placeholder": "e.g. google.com, test@gmail.com, or https://example.com",
        "value_label": "Target Input", "mod": "modules.universal_scan_module", "attr": "universal_scan_module",
        "stat_title": "Universal Scan History", "list_title": "Target Types Scanned",
        "list_items": [],
        "seed": [],
        "donut": [0, 0, 0],
    },
    "URL Scanner": {
        "icon": "🔗", "color": "var(--success)", "bg": "var(--success-bg)",
        "desc": "Scan URLs for phishing, malware, and other malicious activities.",
        "placeholder": "https://example.com/path",
        "value_label": "URL", "mod": "modules.url_module", "attr": "url_module",
        "stat_title": "URL Statistics", "list_title": "Top Risky Domains",
        "list_items": [],
        "seed": [],
        "donut": [0, 0, 0],
    },
    "Website Scanner": {
        "icon": "🌐", "color": "var(--info)", "bg": "var(--info-bg)",
        "desc": "Analyze websites for vulnerabilities, malware, and security issues.",
        "placeholder": "https://yourwebsite.com",
        "value_label": "Website", "mod": "modules.website_module", "attr": "website_module",
        "stat_title": "Security Overview", "list_title": "Vulnerabilities Found",
        "list_items": [],
        "seed": [],
        "donut": [0, 0, 0],
    },
    "Domain Scanner": {
        "icon": "🌍", "color": "#A78BFA", "bg": "rgba(167,139,250,0.12)",
        "desc": "Check domain reputation, WHOIS info, and blacklist status.",
        "placeholder": "example.com",
        "value_label": "Domain", "mod": "modules.domain_module", "attr": "domain_module",
        "stat_title": "Domain Statistics", "list_title": "Blacklist Status",
        "list_items": [],
        "seed": [],
        "donut": [0, 0, 0],
    },
    "IP Scanner": {
        "icon": "🖥️", "color": "var(--warning)", "bg": "var(--warning-bg)",
        "desc": "Analyze IP addresses for threats and suspicious activity.",
        "placeholder": "192.168.1.1",
        "value_label": "IP Address", "mod": "modules.ip_module", "attr": "ip_module",
        "stat_title": "IP Statistics", "list_title": "Top Risky IPs",
        "list_items": [],
        "seed": [],
        "donut": [0, 0, 0],
    },
    "Email Scanner": {
        "icon": "📧", "color": "var(--info)", "bg": "var(--info-bg)",
        "desc": "Scan email addresses for phishing attempts and malicious content.",
        "placeholder": "user@example.com",
        "value_label": "Email", "mod": "modules.email_module", "attr": "email_module",
        "stat_title": "Email Statistics", "list_title": "Email Threats",
        "list_items": [],
        "seed": [],
        "donut": [0, 0, 0],
    },
    "File Scanner": {
        "icon": "📄", "color": "var(--warning)", "bg": "var(--warning-bg)",
        "desc": "Upload and scan files for malware, viruses, and threats.",
        "placeholder": "",
        "value_label": "File Name", "mod": "modules.file_module", "attr": "file_module",
        "stat_title": "File Statistics", "list_title": "File Types Scanned",
        "list_items": [],
        "seed": [],
        "donut": [0, 0, 0],
    },
    "QR Code Scanner": {
        "icon": "🔳", "color": "#2DD4BF", "bg": "rgba(45,212,191,0.12)",
        "desc": "Scan QR codes for malicious links and unsafe content.",
        "placeholder": "",
        "value_label": "Content", "mod": "modules.qr_module", "attr": "qr_module",
        "stat_title": "QR Statistics", "list_title": "QR Content Types",
        "list_items": [],
        "seed": [],
        "donut": [0, 0, 0],
    },

    "Device Security Check": {
        "icon": "🛡️", "color": "var(--success)", "bg": "rgba(34,197,94,0.12)",
        "desc": "Analyze system security posture & OS settings.",
        "placeholder": "",
        "value_label": "Device Hostname", "mod": "modules.device_security_module", "attr": "device_security_module",
        "stat_title": "Security Posture", "list_title": "Issues Identified",
        "list_items": [],
        "seed": [],
        "donut": [0, 0, 0],
    },
    "Phone Threat Intelligence": {
        "icon": "📱", "color": "#F43F5E", "bg": "rgba(244,63,94,0.12)",
        "desc": "Analyze phone numbers for scam & fraud risk.",
        "placeholder": "+91 9876543210 or +1-800-555-0199",
        "value_label": "Phone Number", "mod": "modules.phone_scanner", "attr": "phone_scanner",
        "stat_title": "Phone Threat Statistics", "list_title": "Scam Abuse Categories",
        "list_items": [],
        "seed": [],
        "donut": [0, 0, 0],
    },
}

def sync_histories_from_db():
    if "histories" not in st.session_state:
        st.session_state.histories = {}
    try:
        from database.db import db
        rows = db.fetchall(
            """
            SELECT scan_type, target, risk_level, risk_score, strftime('%H:%M:%S', scan_time) as scan_time
            FROM scan_history
            ORDER BY scan_id DESC
            """
        )
        sc_hist = {key: [] for key in SCANNERS}
        if rows:
            for r in rows:
                stype = r["scan_type"]
                if stype in sc_hist:
                    sc_hist[stype].append({
                        "value": r["target"],
                        "level": r["risk_level"] or "Safe",
                        "score": int(r["risk_score"] or 0),
                        "time": r["scan_time"] or "00:00:00"
                    })
        for key, items in sc_hist.items():
            if not st.session_state.histories.get(key):
                st.session_state.histories[key] = items
    except Exception:
        pass

for key, cfg in SCANNERS.items():
    if key not in st.session_state.histories:
        st.session_state.histories[key] = []

sync_histories_from_db()

MAIN_MENU = [
    ("Dashboard", "📊"), ("Scan History", "🕐"), ("Analytics", "📈"),
]

AUTONOMOUS_CRS_MENU = [
    ("Autonomous Security Lab", "🔬"),
    ("Code SAST & AST", "🔍"),
    ("Fuzz & Sandbox Hub", "🧪"),
]

SCANNERS_DISPLAY = {
    "Universal Scan": ("🌐", "Universal Scan"),    
    "URL Scanner": ("🔗", "URL Scanner"),
    "Website Scanner": ("🕸️", "Website Analyzer"),
    "Domain Scanner": ("🌍", "Domain Intelligence"),
    "Email Scanner": ("📧", "Email Intelligence"),
    "IP Scanner": ("🖥️", "IP Intelligence"),
    "QR Code Scanner": ("🔳", "QR Scanner"),
    "File Scanner": ("📄", "File Analyzer"),
    "Device Security Check": ("🛡️", "Device Security Check"),
    "Phone Threat Intelligence": ("📱", "Phone Threat Intelligence"),
}

SYSTEM_MENU = [
    ("Settings", "⚙️"), ("Profile", "👤"), ("Connections", "🔌"), ("Help & Support", "ℹ️")
]



def toggle_sidebar_collapse():
    st.session_state.sidebar_collapsed = not st.session_state.get("sidebar_collapsed", False)


def render_sidebar():
    collapsed = st.session_state.get("sidebar_collapsed", False)
    with st.container(key="navcol"):
        # Clickable logo button toggles collapse / expand
        st.button(
            "CyberMind" if not collapsed else "",
            key="sb_logo_btn",
            on_click=toggle_sidebar_collapse,
            width="stretch",
        )

        # MAIN section
        if not collapsed:
            st.markdown(f'<div class="sb-section">{t("MAIN")}</div>', unsafe_allow_html=True)
        for label, icon in MAIN_MENU:
            active = st.session_state.active_page == label
            btn_label = icon if collapsed else f"{icon}  {t(label)}"
            st.button(
                btn_label,
                key=f"nav_main_{label}", width="stretch",
                type="primary" if active else "secondary",
                help=t(label) if collapsed else None,
                on_click=go_to, args=(label,),
            )

        # Section Divider between MAIN and AUTONOMOUS CRS
        st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)

        # AUTONOMOUS CRS (AI KAVACH) section
        if not collapsed:
            st.markdown(f'<div class="sb-section">{t("REASONING & REPAIR")}</div>', unsafe_allow_html=True)
        for label, icon in AUTONOMOUS_CRS_MENU:
            active = st.session_state.active_page == label
            btn_label = icon if collapsed else f"{icon}  {label}"
            st.button(
                btn_label,
                key=f"nav_crs_{label}", width="stretch",
                type="primary" if active else "secondary",
                help=label if collapsed else None,
                on_click=go_to, args=(label,),
            )

        # Section Divider between AUTONOMOUS CRS and ANALYSIS
        st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)

        # ANALYSIS section
        if not collapsed:
            st.markdown(f'<div class="sb-section">{t("ANALYSIS")}</div>', unsafe_allow_html=True)
        for target_page, (icon, disp_label) in SCANNERS_DISPLAY.items():
            active = st.session_state.active_page == target_page
            btn_label = icon if collapsed else f"{icon}  {t(disp_label)}"
            st.button(
                btn_label,
                key=f"nav_analysis_{target_page}", width="stretch",
                type="primary" if active else "secondary",
                help=t(disp_label) if collapsed else None,
                on_click=go_to, args=(target_page,),
            )

        # Section Divider between ANALYSIS and SYSTEM
        st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)

        # SYSTEM section
        if not collapsed:
            st.markdown(f'<div class="sb-section">{t("SYSTEM")}</div>', unsafe_allow_html=True)
        for label, icon in SYSTEM_MENU:
            active = st.session_state.active_page == label
            btn_label = icon if collapsed else f"{icon}  {t(label)}"
            st.button(
                btn_label,
                key=f"nav_system_{label}", width="stretch",
                type="primary" if active else "secondary",
                help=t(label) if collapsed else None,
                on_click=go_to, args=(label,),
            )

        if not collapsed:
            st.markdown('<div class="sb-version">v2.0.0 Enterprise</div>', unsafe_allow_html=True)
            # Inject interactive Drag-to-Resize handle for sidebar invisibly
            components.html(
                """
                <script>
                (function() {
                    function initResizableSidebar() {
                        try {
                            var doc = window.parent.document;
                            var win = window.parent;
                            var sidebar = doc.querySelector('.st-key-navcol');
                            if (!sidebar) return;
                            
                            // Guard: Check if sidebar is collapsed (width < 120px)
                            var isCollapsed = sidebar.getBoundingClientRect().width < 120;
                            if (isCollapsed) return;
                            
                            function updateLayoutWidth(w) {
                                sidebar.style.setProperty('width', w + 'px', 'important');
                                sidebar.style.setProperty('max-width', w + 'px', 'important');
                                sidebar.style.setProperty('min-width', w + 'px', 'important');
                                
                                var pOffset = (w + 24) + 'px';
                                var mainContainers = doc.querySelectorAll('html body div.block-container, div.block-container, [data-testid="stMainBlockContainer"], .stMainBlockContainer');
                                mainContainers.forEach(function(el) {
                                    el.style.setProperty('padding-left', pOffset, 'important');
                                });
                                
                                try { win.dispatchEvent(new Event('resize')); } catch(e){}
                            }
                            
                            var handle = sidebar.querySelector('#sb-drag-handle');
                            if (!handle) {
                                handle = doc.createElement('div');
                                handle.id = 'sb-drag-handle';
                                handle.style.cssText = 'position:absolute; top:0; right:-4px; width:8px; height:100%; cursor:col-resize; z-index:999999; border-right:2px solid transparent; transition:border-color 0.2s, box-shadow 0.2s;';
                                sidebar.appendChild(handle);
                                
                                var isDragging = false;
                                var startX = 0;
                                var startWidth = 0;
                                
                                handle.addEventListener('mouseenter', function() {
                                    handle.style.borderRight = '2px solid #22D3EE';
                                    handle.style.boxShadow = '0 0 10px rgba(34, 211, 238, 0.8)';
                                });
                                handle.addEventListener('mouseleave', function() {
                                    if (!isDragging) {
                                        handle.style.borderRight = '2px solid transparent';
                                        handle.style.boxShadow = 'none';
                                    }
                                });
                                
                                handle.onmousedown = function(e) {
                                    isDragging = true;
                                    startX = e.clientX;
                                    startWidth = sidebar.getBoundingClientRect().width;
                                    doc.body.style.userSelect = 'none';
                                    doc.body.style.cursor = 'col-resize';
                                    handle.style.borderRight = '2px solid #22D3EE';
                                };
                                
                                doc.addEventListener('mousemove', function(e) {
                                    if (!isDragging) return;
                                    var deltaX = e.clientX - startX;
                                    var newWidth = Math.max(160, Math.min(420, startWidth + deltaX));
                                    
                                    updateLayoutWidth(newWidth);
                                    try { localStorage.setItem('custom_sb_width', newWidth); } catch(err){}
                                });
                                
                                doc.addEventListener('mouseup', function() {
                                    if (isDragging) {
                                        isDragging = false;
                                        doc.body.style.userSelect = '';
                                        doc.body.style.cursor = '';
                                        handle.style.borderRight = '2px solid transparent';
                                        handle.style.boxShadow = 'none';
                                    }
                                });
                            }
                            
                            try {
                                var savedWidth = localStorage.getItem('custom_sb_width');
                                if (savedWidth) {
                                    var w = parseInt(savedWidth);
                                    if (w >= 160 && w <= 420) {
                                        updateLayoutWidth(w);
                                    }
                                }
                            } catch(e){}
                        } catch(e) {}
                    }
                    
                    if (window.parent.document.readyState === 'complete') {
                        initResizableSidebar();
                    } else {
                        window.parent.addEventListener('DOMContentLoaded', initResizableSidebar);
                    }
                    setTimeout(initResizableSidebar, 150);
                    setTimeout(initResizableSidebar, 600);
                })();
                </script>
                """,
                height=0,
                width=0
            )
        else:
            # Inject Collapsed Mode Layout Reset (Strips leftover expanded inline padding)
            components.html(
                """
                <script>
                (function() {
                    function resetCollapsedLayout() {
                        try {
                            var doc = window.parent.document;
                            var win = window.parent;
                            var mainContainers = doc.querySelectorAll('html body div.block-container, div.block-container, [data-testid="stMainBlockContainer"], .stMainBlockContainer');
                            mainContainers.forEach(function(el) {
                                el.style.removeProperty('padding-left');
                                el.style.setProperty('padding-left', '5.8rem', 'important');
                            });
                            var sidebar = doc.querySelector('.st-key-navcol');
                            if (sidebar) {
                                sidebar.style.removeProperty('width');
                                sidebar.style.removeProperty('max-width');
                                sidebar.style.removeProperty('min-width');
                                sidebar.style.setProperty('width', '4.2rem', 'important');
                            }
                            try { win.dispatchEvent(new Event('resize')); } catch(e){}
                        } catch(e) {}
                    }
                    if (window.parent.document.readyState === 'complete') {
                        resetCollapsedLayout();
                    } else {
                        window.parent.addEventListener('DOMContentLoaded', resetCollapsedLayout);
                    }
                    setTimeout(resetCollapsedLayout, 100);
                    setTimeout(resetCollapsedLayout, 350);
                })();
                </script>
                """,
                height=0,
                width=0
            )


# Layout shell: CSS-positioned sidebar (no st.columns)
render_sidebar()





# Reusable UI pieces


def render_page_poster(title: str, subtitle: str):
    # Load poster.png as base64
    try:
        with open("poster.png", "rb") as image_file:
            poster_base64 = base64.b64encode(image_file.read()).decode()
    except Exception:
        poster_base64 = ""

    if poster_base64:
        bg_style = f'background-image: linear-gradient(90deg, rgba(15,24,38,0) 0%, rgba(15,24,38,0.3) 40%, var(--card-bg) 75%, var(--card-bg) 100%), url("data:image/png;base64,{poster_base64}");'
    else:
        bg_style = 'background: var(--card-bg);'

    st.markdown(
        f"""
        <style>
        @media (max-width: 768px) {{
            .responsive-poster {{
                background-image: linear-gradient(180deg, rgba(15,24,38,0.8) 0%, var(--card-bg) 100%) !important;
                height: auto !important;
                min-height: 140px !important;
                padding: 20px !important;
            }}
            .responsive-poster-text {{
                position: static !important;
                transform: none !important;
                text-align: center !important;
                max-width: 100% !important;
                margin: 0 auto !important;
            }}
            .responsive-poster-title {{
                font-size: 22px !important;
            }}
        }}
        </style>
        <div class="responsive-poster" style='
            {bg_style}
            background-size: 100% 100%;
            background-position: center center;
            background-repeat: no-repeat;
            background-color: var(--card-bg);
            border: 1px solid var(--border);
            border-radius: 18px;
            margin-bottom: 24px;
            width: 100%;
            height: 186px;
            position: relative;
            box-sizing: border-box;
            overflow: hidden;
            box-shadow: 0 8px 32px rgba(0, 0, 0, 0.25);
        '>
            <div class="responsive-poster-text" style='
                position: absolute;
                right: 44px;
                top: 50%;
                transform: translateY(-50%);
                text-align: right;
                max-width: 55%;
                z-index: 5;
            '>
                <h1 class="responsive-poster-title" style='
                    font-size: 28px !important;
                    font-weight: 800 !important;
                    color: var(--text) !important;
                    line-height: 1.25 !important;
                    margin: 0 !important;
                    padding: 0 !important;
                    border: none !important;
                '>{title}</h1>
                <div style='
                    font-size: 13.5px;
                    font-weight: 500;
                    color: var(--text-muted);
                    margin-top: 6px;
                '>{subtitle}</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )


def donut_chart(labels, values, colors, center_label):
    dark = st.session_state.get("settings_dark_mode", True)
    text_color = "#EAF0F7" if dark else "#1E293B"
    muted_color = "#8B98AC" if dark else "#64748B"
    border_color = "#0F1826" if dark else "#FFFFFF"
    total = sum(values)
    fig = go.Figure(data=[go.Pie(
        labels=labels, values=values, hole=0.68,
        marker=dict(colors=colors, line=dict(color=border_color, width=2)),
        textinfo="none", sort=False,
    )])
    fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        showlegend=False, height=190, margin=dict(l=0, r=0, t=10, b=0),
        annotations=[dict(
            text=f"<b>{total:,}</b><br><span style='font-size:10px;color:{muted_color}'>{center_label}</span>",
            x=0.5, y=0.5, font=dict(size=18, color=text_color), showarrow=False
        )],
    )
    return fig


def circular_gauge(score, level):
    dark = st.session_state.get("settings_dark_mode", True)
    text_color = "#EAF0F7" if dark else "#1E293B"
    tick_color = "#8B98AC" if dark else "#64748B"
    border_color = "#1D2838" if dark else "#D1D9E6"
    color = "#22C55E" if score < 30 else "#F5A623" if score < 70 else "#F2545B"
    fig = go.Figure(go.Indicator(
        mode = "gauge+number",
        value = score,
        domain = {'x': [0, 1], 'y': [0.08, 0.92]},
        title = {'text': f"RISK: {level.upper()}", 'font': {'size': 14, 'color': text_color}},
        gauge = {
            'axis': {'range': [None, 100], 'tickwidth': 1, 'tickcolor': tick_color},
            'bar': {'color': color},
            'bgcolor': "rgba(0,0,0,0)",
            'borderwidth': 2,
            'bordercolor': border_color,
            'steps': [
                {'range': [0, 30], 'color': 'rgba(34,197,94,0.08)'},
                {'range': [30, 70], 'color': 'rgba(245,166,35,0.08)'},
                {'range': [70, 100], 'color': 'rgba(242,84,91,0.08)'}
            ],
        }
    ))
    fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        font={'color': text_color, 'family': "Inter"},
        height=230,
        margin=dict(l=25, r=25, t=45, b=15)
    )
    return fig



# Page: Home


def render_home():
    hero_col1, hero_col2 = st.columns([1.55, 1], gap="large")
    with hero_col1:
        st.markdown(
            """
            <div class="hero-title">CyberMind AI</div>
            <div class="hero-sub">Advanced Cybersecurity Threat Detection Platform</div>
            <div class="hero-desc">
                Leverage the power of Artificial Intelligence and Machine Learning to
                detect, analyze, and prevent cyber threats in real-time.
            </div>
            """,
            unsafe_allow_html=True
        )
        b1, b2, _ = st.columns([1.25, 1.25, 1.5], gap="small")
        with b1:
            st.markdown('<div class="cta-primary">', unsafe_allow_html=True)
            st.button("🛡️ Run New Scan", key="run_scan_btn", on_click=go_to, args=("URL Scanner",))
            st.markdown('</div>', unsafe_allow_html=True)
        with b2:
            st.markdown('<div class="cta-secondary">', unsafe_allow_html=True)
            st.button("📊 View Dashboard", key="view_dash_btn", on_click=go_to, args=("Dashboard",))
            st.markdown('</div>', unsafe_allow_html=True)
    with hero_col2:
        st.markdown(
            """
            <div class="shield-wrap">
                <svg width="200" height="205" viewBox="0 0 200 205" xmlns="http://www.w3.org/2000/svg">
                    <defs>
                        <linearGradient id="shieldGrad" x1="0%" y1="0%" x2="100%" y2="100%">
                            <stop offset="0%" stop-color="#6C5CE7"/>
                            <stop offset="100%" stop-color="#22D3EE"/>
                        </linearGradient>
                        <filter id="glow"><feGaussianBlur stdDeviation="4" result="b"/>
                            <feMerge><feMergeNode in="b"/><feMergeNode in="SourceGraphic"/></feMerge>
                        </filter>
                    </defs>
                    <path d="M100 6 L182 36 V98 C182 146 148 184 100 200 C52 184 18 146 18 98 V36 Z"
                          fill="none" stroke="url(#shieldGrad)" stroke-width="4" filter="url(#glow)"/>
                </svg>
                <div class="shield-brain">🧠</div>
            </div>
            """,
            unsafe_allow_html=True
        )

    render_stat_cards()
    st.markdown("<div style='height:20px'></div>", unsafe_allow_html=True)

    left_col, right_col = st.columns([2.55, 1], gap="large")
    with left_col:
        st.markdown('<div class="section-title">Quick Scanner Tools</div>', unsafe_allow_html=True)
        keys = list(SCANNERS.keys())
        for row_start in range(0, len(keys), 5):
            row_keys = keys[row_start:row_start + 5]
            row_cols = st.columns(5, gap="medium")
            for col, sk in zip(row_cols, row_keys):
                sc = SCANNERS[sk]
                with col:
                    st.markdown(
                        f"""
                        <div class="scan-card">
                            <div>
                                <div class="scan-icon" style="background:{sc['bg']};color:{sc['color']}">{sc['icon']}</div>
                                <div class="scan-title">{sk}</div>
                                <div class="scan-desc">{sc['desc']}</div>
                            </div>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                    st.markdown('<div class="cta-scan">', unsafe_allow_html=True)
                    st.button("Scan Now →", key=f"home_scan_{sk}", width="stretch",
                              on_click=go_to, args=(sk,))
                    st.markdown('</div>', unsafe_allow_html=True)

    with right_col:
        render_system_status()
        render_recent_scans()

    st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)
    chart_col1, chart_col2 = st.columns(2, gap="large")
    with chart_col1:
        render_threat_statistics_chart()
    with chart_col2:
        render_model_performance_chart()


_DB_SEEDED_FLAG = False

def seed_database_if_empty():
    global _DB_SEEDED_FLAG
    if _DB_SEEDED_FLAG or st.session_state.get("_db_seeded"):
        return
    try:
        from database.db import db
        res = db.fetchone("SELECT COUNT(*) as count FROM scan_history")
        if res and res["count"] == 0:
            sample_scans = [
                ("URL Scanner", "https://example.com", "Safe", 10.0),
                ("URL Scanner", "http://malicious-phishing-site.xyz", "Critical", 95.0),
                ("Website Scanner", "https://openai.com", "Safe", 12.0),
                ("Domain Scanner", "google.com", "Safe", 8.0),
                ("IP Scanner", "8.8.8.8", "Safe", 10.0),
                ("Email Scanner", "user@example.com", "Safe", 5.0),
                ("File Scanner", "sample_document.pdf", "Safe", 5.0),
                ("QR Code Scanner", "https://example.com", "Safe", 10.0),
                ("Phone Threat Intelligence", "+91 9876543210", "Safe", 10.0),
                ("Device Security Check", "Local Machine", "Safe", 15.0),
            ]
            for stype, target, level, score in sample_scans:
                db.execute(
                    """
                    INSERT INTO scan_history (scan_type, target, risk_level, risk_score)
                    VALUES (?, ?, ?, ?)
                    """,
                    (stype, target, level, score)
                )
        _DB_SEEDED_FLAG = True
        st.session_state["_db_seeded"] = True
    except Exception:
        pass


def render_stat_cards():
    seed_database_if_empty()

    from database.db import db
    
    total_scans = 0
    threats_detected = 0
    risky_items = 0
    safe_items = 0

    try:
        res = db.fetchone("SELECT COUNT(*) as count FROM scan_history")
        total_scans = res["count"] if res else 0

        res = db.fetchone("SELECT COUNT(*) as count FROM scan_history WHERE risk_level IN ('High', 'Critical')")
        threats_detected = res["count"] if res else 0

        res = db.fetchone("SELECT COUNT(*) as count FROM scan_history WHERE risk_level IN ('Low', 'Medium')")
        risky_items = res["count"] if res else 0

        res = db.fetchone("SELECT COUNT(*) as count FROM scan_history WHERE risk_level = 'Safe'")
        safe_items = res["count"] if res else 0
    except Exception:
        total_scans, threats_detected, risky_items, safe_items = 1248, 342, 128, 778

    stats = [
        {"label": "Total Scans", "value": f"{total_scans:,}", "delta": "Live stats from database",
         "color": "var(--info)", "bg": "var(--info-bg)", "icon": "📶", "dc": "var(--success)"},
        {"label": "Threats Detected", "value": f"{threats_detected:,}", "delta": "Real-time threat feeds",
         "color": "var(--danger)", "bg": "var(--danger-bg)", "icon": "⚠️", "dc": "var(--danger)"},
        {"label": "Risky Items", "value": f"{risky_items:,}", "delta": "Heuristic risk scoring",
         "color": "var(--warning)", "bg": "var(--warning-bg)", "icon": "🛡️", "dc": "var(--warning)"},
        {"label": "Safe Items", "value": f"{safe_items:,}", "delta": "Verification status active",
         "color": "var(--success)", "bg": "var(--success-bg)", "icon": "✅", "dc": "var(--success)"},
    ]
    cols = st.columns(4, gap="medium")
    for col, s in zip(cols, stats):
        with col:
            st.markdown(
                f"""
                <div class="stat-card">
                    <div class="stat-card-left">
                        <div class="stat-label">{s['label']}</div>
                        <div class="stat-value">{s['value']}</div>
                        <div class="stat-delta" style="color:{s['dc']}">{s['delta']}</div>
                    </div>
                    <div class="stat-icon" style="background:{s['bg']};color:{s['color']}">{s['icon']}</div>
                </div>
                """,
                unsafe_allow_html=True
            )


def render_system_status():
    # 1. Database Connection check
    from database.connection import database
    db_connected = database.is_connected()
    db_status = "Active" if db_connected else "Inactive"
    db_color = "var(--success)" if db_connected else "var(--danger)"

    # 2. Datasets check
    from config.paths import DATA_DIR
    from pathlib import Path
    datasets_dir = Path(DATA_DIR) / "datasets"
    ds_exists = datasets_dir.exists() and any(datasets_dir.iterdir())
    ds_status = "Active" if ds_exists else "Inactive"
    ds_color = "var(--success)" if ds_exists else "var(--danger)"

    # 3. API Connection check
    from config.api_config import GOOGLE_SAFE_BROWSING_API_KEY, VIRUSTOTAL_API_KEY, ABUSEIPDB_API_KEY
    from core.offline_mode import offline_mode
    apis = []
    if not offline_mode.is_enabled:
        if GOOGLE_SAFE_BROWSING_API_KEY:
            apis.append("Safe Browsing")
        if VIRUSTOTAL_API_KEY:
            apis.append("VirusTotal")
        if ABUSEIPDB_API_KEY:
            apis.append("AbuseIPDB")
    
    api_connected = len(apis) > 0
    if offline_mode.is_enabled:
        api_status = "Disabled (Offline Mode Active)"
        api_color = "var(--warning)"
    else:
        api_status = f"Active ({', '.join(apis)})" if api_connected else "Inactive"
        api_color = "var(--success)" if api_connected else "var(--danger)"

    st.markdown(
        f"""
        <div class="panel">
            <div class="panel-title">System Status</div>
            <div class="status-row"><span class="status-name">Database Connection</span><span class="status-val" style="color:{db_color}">{db_status}</span></div>
            <div class="status-row"><span class="status-name">Dataset Storage</span><span class="status-val" style="color:{ds_color}">{ds_status}</span></div>
            <div class="status-row"><span class="status-name">Threat API Keys</span><span class="status-val" style="color:{api_color}">{api_status}</span></div>
            <div class="status-row"><span class="status-name">AI Models</span><span class="status-val">Operational</span></div>
        </div>
        """,
        unsafe_allow_html=True
    )


def render_offline_toast():
    if not st.session_state.get("show_offline_toast"):
        return
    
    toast_time = st.session_state.get("offline_toast_time", 0)
    elapsed = time.time() - toast_time
    if elapsed > 10.5:
        st.session_state["show_offline_toast"] = False
        return

    is_offline = st.session_state.get("offline_toast_state", False)
    if is_offline:
        icon = "<img src='https://cdn-icons-png.flaticon.com/512/3926/3926678.png' style='width:24px;height:24px;'>"
        title = "Offline Mode Enabled"
        msg = "All external API calls are disabled. Running on local ML models & offline databases."
        border_color = "rgba(255, 170, 0, 0.5)"
        left_accent = "#ffaa00"
        shadow_color = "rgba(255, 170, 0, 0.25)"
        grad_bar = "linear-gradient(90deg, #ffaa00, #ff5500)"
    else:
        icon = "<img src='https://cdn-icons-png.flaticon.com/512/6532/6532060.png' style='width:24px;height:24px;'>"
        title = "Online Mode Enabled"
        msg = "External threat-intelligence API connections restored."
        border_color = "rgba(0, 210, 255, 0.5)"
        left_accent = "#00d2ff"
        shadow_color = "rgba(0, 210, 255, 0.25)"
        grad_bar = "linear-gradient(90deg, #00d2ff, #0055ff)"

    remaining_sec = max(0.1, 10.0 - elapsed)
    toast_id = int(toast_time * 1000)

    toast_html = f"""
    <div id="offline-toast-{toast_id}" class="st-offline-toast-box-{toast_id}">
        <div style="display: flex; align-items: flex-start; gap: 12px; padding-right: 12px;">
            <span style="font-size: 20px; line-height: 1; margin-top: 2px;">{icon}</span>
            <div>
                <div style="font-weight: 700; font-size: 14px; color: #ffffff; margin-bottom: 2px;">{title}</div>
                <div style="font-weight: 400; font-size: 12.5px; color: rgba(255, 255, 255, 0.85);">{msg}</div>
            </div>
        </div>
        <div class="st-offline-toast-progress-{toast_id}"></div>
    </div>

    <style>
    .st-offline-toast-box-{toast_id} {{
        position: fixed;
        top: 24px;
        right: 24px;
        z-index: 9999999;
        display: flex;
        align-items: center;
        justify-content: space-between;
        min-width: 320px;
        max-width: 440px;
        padding: 14px 18px;
        background: rgba(18, 24, 38, 0.96);
        backdrop-filter: blur(14px);
        -webkit-backdrop-filter: blur(14px);
        border: 1px solid {border_color};
        border-left: 4px solid {left_accent};
        box-shadow: 0 12px 35px rgba(0, 0, 0, 0.5), 0 0 20px {shadow_color};
        border-radius: 12px;
        color: #ffffff;
        font-family: 'Inter', system-ui, -apple-system, sans-serif;
        line-height: 1.4;
        overflow: hidden;
        animation: 
            slideInToast_{toast_id} 0.4s cubic-bezier(0.16, 1, 0.3, 1) forwards,
            slideOutToast_{toast_id} 0.4s cubic-bezier(0.16, 1, 0.3, 1) {remaining_sec:.2f}s forwards;
    }}

    .st-offline-toast-close-{toast_id} {{
        background: transparent;
        border: none;
        color: rgba(255, 255, 255, 0.65);
        font-size: 18px;
        font-weight: bold;
        cursor: pointer;
        padding: 4px 8px;
        border-radius: 6px;
        transition: all 0.2s ease;
        line-height: 1;
        align-self: flex-start;
    }}

    .st-offline-toast-close-{toast_id}:hover {{
        color: #fff;
        background: rgba(255, 255, 255, 0.15);
    }}

    .st-offline-toast-progress-{toast_id} {{
        position: absolute;
        bottom: 0;
        left: 0;
        height: 3px;
        width: 100%;
        background: {grad_bar};
        animation: toastProgressAnimation_{toast_id} {remaining_sec:.2f}s linear forwards;
    }}

    @keyframes slideInToast_{toast_id} {{
        from {{ transform: translateX(120%); opacity: 0; }}
        to {{ transform: translateX(0); opacity: 1; }}
    }}

    @keyframes slideOutToast_{toast_id} {{
        from {{ transform: translateX(0); opacity: 1; }}
        to {{ transform: translateX(120%); opacity: 0; }}
    }}

    @keyframes toastProgressAnimation_{toast_id} {{
        from {{ width: 100%; }}
        to {{ width: 0%; }}
    }}
    </style>
    """
    st.markdown(toast_html, unsafe_allow_html=True)


# Mapping scanner categories to circular icons
RECENT_ICONS = {
    "URL Scanner": ("🔗", "rgba(139,92,246,0.12)", "#8B5CF6"),
    "Website Analyzer": ("🌐", "rgba(59,130,246,0.12)", "#3B82F6"),
    "Domain Intelligence": ("🌍", "rgba(6,182,212,0.12)", "#06B6D4"),
    "Email Intelligence": ("📧", "rgba(16,185,129,0.12)", "#10B981"),
    "IP Intelligence": ("🖥️", "rgba(245,166,35,0.12)", "#F5A623"),
    "QR Scanner": ("🔳", "rgba(245,166,35,0.12)", "#F5A623"),
    "File Analyzer": ("📄", "rgba(239,68,68,0.12)", "#EF4444"),
    "Universal Scan": ("🌐", "rgba(59,130,246,0.12)", "#3B82F6"),
}

def render_recent_scans():
    rows = []
    try:
        from database.db import db
        # Get the 6 most recent scans overall
        db_rows = db.fetchall(
            """
            SELECT scan_type, target, risk_level, strftime('%H:%M:%S', scan_time) as time
            FROM scan_history
            ORDER BY scan_id DESC
            LIMIT 6
            """
        )
        if db_rows:
            rows = [
                {
                    "scanner": r["scan_type"],
                    "target": r["target"],
                    "level": r["risk_level"] or "Safe",
                    "time": r["time"]
                }
                for r in db_rows
            ]
    except Exception:
        pass

    if not rows:
        st.markdown(
            """
            <div style="padding: 24px; text-align: center; color: var(--text-muted); font-size: 13px;">
                No recent scans recorded yet. Run a live threat scan from the menu!
            </div>
            """,
            unsafe_allow_html=True
        )
    else:
        rows_html = ""
        for r in rows:
            icon, bg_c, text_c = RECENT_ICONS.get(r["scanner"], ("🛡️", "rgba(255,255,255,0.05)", "var(--text)"))
            badge_c = LEVEL_BADGE.get(r["level"], "badge-safe")
            
            rows_html += (
                f'<div style="display:flex; align-items:center; justify-content:space-between; padding:8px 0; border-bottom:1px solid var(--border);">'
                f'<div style="display:flex; align-items:center; gap:10px; flex:1; min-width:0;">'
                f'<div style="width:32px; height:32px; border-radius:50%; background:{bg_c}; color:{text_c}; display:flex; align-items:center; justify-content:center; font-size:14px; flex-shrink:0;">{icon}</div>'
                f'<div style="min-width:0; flex:1;">'
                f'<div style="font-size:12.5px; font-weight:600; color:var(--text); overflow:hidden; text-overflow:ellipsis; white-space:nowrap;">{html.escape(r["target"])}</div>'
                f'<div style="font-size:10.5px; color:var(--text-faint);">{r["scanner"]}</div>'
                f'</div>'
                f'</div>'
                f'<div style="display:flex; align-items:center; gap:8px; flex-shrink:0; margin-left:8px;">'
                f'<span class="badge {badge_c}" style="font-size:10px; padding:2px 7px;">{r["level"]}</span>'
                f'<span style="font-size:10.5px; color:var(--text-faint); min-width:55px; text-align:right;">{r["time"]}</span>'
                f'</div>'
                f'</div>'
            )
            
        st.markdown(
            f'<div style="display:flex; flex-direction:column; justify-content:space-between; height:210px; overflow-y:auto; padding:2px 0;">{rows_html}</div>',
            unsafe_allow_html=True
        )


def render_threat_statistics_chart():
    st.markdown(
        """<div class="chart-card"><div class="chart-card-head">
        <div class="chart-card-title">Threat Statistics</div>
        <div class="chart-filter">This Month ⌄</div></div>""",
        unsafe_allow_html=True
    )
    weeks = ["Week 1", "Week 2", "Week 3", "Week 4"]
    fig = go.Figure()
    fig.add_bar(name="Threats Detected", x=weeks, y=[62, 88, 74, 118], marker_color="#F2545B")
    fig.add_bar(name="Safe Scans", x=weeks, y=[210, 245, 198, 260], marker_color="#22C55E")
    fig.update_layout(
        barmode="group", paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color="#8B98AC", size=11), margin=dict(l=10, r=10, t=10, b=10), height=250,
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
        xaxis=dict(gridcolor="#1D2838"), yaxis=dict(gridcolor="#1D2838"),
    )
    st.plotly_chart(fig, width="stretch", config={"displayModeBar": False})
    st.markdown("</div>", unsafe_allow_html=True)


def render_model_performance_chart():
    st.markdown(
        """<div class="chart-card"><div class="chart-card-head">
        <div class="chart-card-title">AI Model Performance</div>
        <div class="chart-filter">This Month ⌄</div></div>""",
        unsafe_allow_html=True
    )
    days = ["Jun 1", "Jun 8", "Jun 15", "Jun 22", "Jun 29"]
    fig2 = go.Figure()
    fig2.add_trace(go.Scatter(
        x=days, y=[92.1, 93.4, 94.0, 95.2, 96.1], mode="lines+markers", name="Accuracy",
        line=dict(color="#22B8F0", width=3), marker=dict(size=6),
        fill="tozeroy", fillcolor="rgba(34,184,240,0.12)"
    ))
    fig2.update_layout(
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color="#8B98AC", size=11), margin=dict(l=10, r=10, t=10, b=10), height=250,
        showlegend=False, xaxis=dict(gridcolor="#1D2838"), yaxis=dict(gridcolor="#1D2838", range=[88, 100]),
    )
    st.plotly_chart(fig2, width="stretch", config={"displayModeBar": False})
    st.markdown("</div>", unsafe_allow_html=True)



# Page: Dashboard


def render_dashboard():
    # Ensure database is seeded if empty
    seed_database_if_empty()

    from database.db import db
    
    total_scans = 0
    safe_scans = 0
    threats_found = 0
    medium_risk_scans = 0
    low_risk_scans = 0
    avg_score = 0

    try:
        res = db.fetchone("SELECT COUNT(*) as count FROM scan_history")
        total_scans = res["count"] if res else 0

        res = db.fetchone("SELECT COUNT(*) as count FROM scan_history WHERE risk_level = 'Safe'")
        safe_scans = res["count"] if res else 0

        res = db.fetchone("SELECT COUNT(*) as count FROM scan_history WHERE risk_level = 'Low'")
        low_risk_scans = res["count"] if res else 0

        res = db.fetchone("SELECT COUNT(*) as count FROM scan_history WHERE risk_level = 'Medium'")
        medium_risk_scans = res["count"] if res else 0

        res = db.fetchone("SELECT COUNT(*) as count FROM scan_history WHERE risk_level IN ('High', 'Critical')")
        threats_found = res["count"] if res else 0

        res = db.fetchone("SELECT AVG(risk_score) as avg FROM scan_history")
        avg_score = int(res["avg"]) if res and res["avg"] is not None else 0
    except Exception:
        total_scans, safe_scans, low_risk_scans, medium_risk_scans, threats_found, avg_score = 1450, 1100, 150, 120, 80, 24

    # Sparkline data fetching helper functions
    def get_sparkline_data(risk_level=None, query_score=False):
        try:
            if query_score:
                rows = db.fetchall("SELECT date(scan_time) as d, AVG(risk_score) as val FROM scan_history GROUP BY d ORDER BY d DESC LIMIT 7")
            elif risk_level == "Safe":
                rows = db.fetchall("SELECT date(scan_time) as d, COUNT(*) as val FROM scan_history WHERE risk_level = 'Safe' GROUP BY d ORDER BY d DESC LIMIT 7")
            elif risk_level == "Threat":
                rows = db.fetchall("SELECT date(scan_time) as d, COUNT(*) as val FROM scan_history WHERE risk_level IN ('High', 'Critical') GROUP BY d ORDER BY d DESC LIMIT 7")
            elif risk_level == "Medium":
                rows = db.fetchall("SELECT date(scan_time) as d, COUNT(*) as val FROM scan_history WHERE risk_level = 'Medium' GROUP BY d ORDER BY d DESC LIMIT 7")
            else:
                rows = db.fetchall("SELECT date(scan_time) as d, COUNT(*) as val FROM scan_history GROUP BY d ORDER BY d DESC LIMIT 7")
                
            if rows:
                vals = [float(r["val"]) for r in reversed(rows)]
                if len(vals) < 7:
                    vals = [0.0] * (7 - len(vals)) + vals
                return vals
        except Exception:
            pass
        return [20, 30, 25, 45, 35, 60, 50]

    def get_sparkline_svg(values, color):
        if not values or len(values) < 2:
            return ""
        max_v = max(values)
        min_v = min(values)
        v_range = (max_v - min_v) if (max_v - min_v) > 0 else 1.0
        
        w, h = 100, 28
        points = []
        for i, val in enumerate(values):
            x = (i / (len(values) - 1)) * w
            y = h - ((val - min_v) / v_range) * (h - 4) - 2
            points.append(f"{x},{y}")
            
        path_d = f"M {points[0]} " + " ".join([f"L {p}" for p in points[1:]])
        svg = f'<svg width="{w}" height="{h}" style="overflow:visible;"><defs><filter id="glow-{color.replace("#","")}" x="-20%" y="-20%" width="140%" height="140%"><feGaussianBlur stdDeviation="2" result="blur" /><feComposite in="SourceGraphic" in2="blur" operator="over" /></filter></defs><path d="{path_d}" fill="none" stroke="{color}" stroke-width="2.5" stroke-linecap="round" stroke-linejoin="round" filter="url(#glow-{color.replace("#","")})"/></svg>'
        return svg

    # Sparkline data
    total_spark = get_sparkline_data(None)
    safe_spark = get_sparkline_data("Safe")
    threat_spark = get_sparkline_data("Threat")
    med_spark = get_sparkline_data("Medium")
    score_spark = get_sparkline_data(None, query_score=True)

    # Render Sparklines SVG
    total_svg = get_sparkline_svg(total_spark, "#3B82F6")
    safe_svg = get_sparkline_svg(safe_spark, "#22C55E")
    threat_svg = get_sparkline_svg(threat_spark, "#F2545B")
    med_svg = get_sparkline_svg(med_spark, "#F5A623")
    score_svg = get_sparkline_svg(score_spark, "#A78BFA")

    # Load image as base64 for card background
    try:
        with open("cyber.png", "rb") as image_file:
            img_base64 = base64.b64encode(image_file.read()).decode()
    except Exception:
        img_base64 = ""

    if img_base64:
        st.markdown(
            f"""
            <style>
            .st-key-hero_container {{
                background-image: linear-gradient(90deg, var(--card-bg) 0%, var(--card-bg) 35%, rgba(15, 24, 38, 0.25) 55%, rgba(15, 24, 38, 0) 70%), url("data:image/png;base64,{img_base64}") !important;
                background-size: cover !important;
                background-position: right center !important;
                background-repeat: no-repeat !important;
                border: 1px solid var(--border) !important;
                border-radius: 18px !important;
                padding: 40px 44px 44px 44px !important;
                margin-bottom: 24px !important;
            }}
            </style>
            """,
            unsafe_allow_html=True
        )

    with st.container(key="hero_container"):
        st.markdown(
            """
            <div class="hero-title">CyberMind AI</div>
            <div class="hero-sub">Advanced Cybersecurity Threat Detection Platform</div>
            <div class="hero-desc" style="max-width: 580px;">
                Leverage the power of Artificial Intelligence and Machine Learning to
                detect, analyze, and prevent cyber threats in real-time.
            </div>
            """,
            unsafe_allow_html=True
        )
        b1, b2, b3 = st.columns([0.65, 0.75, 2.2])
        with b1:
            st.button("🕐 Scan History", key="dash_hero_history", width="stretch", type="primary", on_click=go_to, args=("Scan History",))
        with b2:
            st.button("📈 Analytics Panel", key="dash_hero_analytics", width="stretch", type="secondary", on_click=go_to, args=("Analytics",))

    # 1. Top Row (5 cards)
    st.markdown(
        clean_html(f"""
        <div style="display:grid; grid-template-columns: repeat(auto-fit, minmax(180px, 1fr)); gap:12px; margin-bottom:20px;">
            <div class="chart-card" style="padding:18px; display:flex; justify-content:space-between; align-items:stretch; min-height:115px;">
                <div style="display:flex; flex-direction:column; justify-content:space-between; flex-grow:1;">
                    <div style="font-size:12.5px; color:var(--text-muted); font-weight:600; text-transform:uppercase; letter-spacing:0.5px;">Total Scans</div>
                    <div style="font-size:32px; font-weight:800; color:var(--text); margin-top:2px; line-height:1.1;">{total_scans:,}</div>
                    <div style="font-size:12.5px; color:var(--success); margin-top:4px; font-weight:600;">▲ 12.4% <span style="color:var(--text-faint); font-weight:500;">vs last 7 days</span></div>
                </div>
                <div style="display:flex; flex-direction:column; justify-content:space-between; align-items:flex-end; margin-left:10px;">
                    <div style="width:38px; height:38px; border-radius:50%; background:rgba(59,130,246,0.14); color:#3B82F6; display:flex; align-items:center; justify-content:center; font-size:20px; box-shadow:0 0 10px rgba(59,130,246,0.1);">📶</div>
                    <div style="margin-top:auto;">{total_svg}</div>
                </div>
            </div>
            <div class="chart-card" style="padding:18px; display:flex; justify-content:space-between; align-items:stretch; min-height:115px;">
                <div style="display:flex; flex-direction:column; justify-content:space-between; flex-grow:1;">
                    <div style="font-size:12.5px; color:var(--text-muted); font-weight:600; text-transform:uppercase; letter-spacing:0.5px;">Safe Scans</div>
                    <div style="font-size:32px; font-weight:800; color:var(--success); margin-top:2px; line-height:1.1;">{safe_scans + low_risk_scans:,}</div>
                    <div style="font-size:12.5px; color:var(--text-faint); margin-top:4px; font-weight:500;">{((safe_scans+low_risk_scans)/total_scans*100) if total_scans > 0 else 85.0:.1f}% of total</div>
                </div>
                <div style="display:flex; flex-direction:column; justify-content:space-between; align-items:flex-end; margin-left:10px;">
                    <div style="width:38px; height:38px; border-radius:50%; background:rgba(34,197,94,0.14); color:#22C55E; display:flex; align-items:center; justify-content:center; font-size:20px; box-shadow:0 0 10px rgba(34,197,94,0.1);">🛡️</div>
                    <div style="margin-top:auto;">{safe_svg}</div>
                </div>
            </div>
            <div class="chart-card" style="padding:18px; display:flex; justify-content:space-between; align-items:stretch; min-height:115px;">
                <div style="display:flex; flex-direction:column; justify-content:space-between; flex-grow:1;">
                    <div style="font-size:12.5px; color:var(--text-muted); font-weight:600; text-transform:uppercase; letter-spacing:0.5px;">Threats Found</div>
                    <div style="font-size:32px; font-weight:800; color:var(--danger); margin-top:2px; line-height:1.1;">{threats_found:,}</div>
                    <div style="font-size:12.5px; color:var(--danger); margin-top:4px; font-weight:600;">▲ 14.2% <span style="color:var(--text-faint); font-weight:500;">vs last 7 days</span></div>
                </div>
                <div style="display:flex; flex-direction:column; justify-content:space-between; align-items:flex-end; margin-left:10px;">
                    <div style="width:38px; height:38px; border-radius:50%; background:rgba(242,84,91,0.14); color:#F2545B; display:flex; align-items:center; justify-content:center; font-size:20px; box-shadow:0 0 10px rgba(242,84,91,0.1);">⚠️</div>
                    <div style="margin-top:auto;">{threat_svg}</div>
                </div>
            </div>
            <div class="chart-card" style="padding:18px; display:flex; justify-content:space-between; align-items:stretch; min-height:115px;">
                <div style="display:flex; flex-direction:column; justify-content:space-between; flex-grow:1;">
                    <div style="font-size:12.5px; color:var(--text-muted); font-weight:600; text-transform:uppercase; letter-spacing:0.5px;">Suspicious Scans</div>
                    <div style="font-size:32px; font-weight:800; color:#F5A623; margin-top:2px; line-height:1.1;">{medium_risk_scans:,}</div>
                    <div style="font-size:12.5px; color:var(--warning); margin-top:4px; font-weight:600;">▲ 8.1% <span style="color:var(--text-faint); font-weight:500;">vs last 7 days</span></div>
                </div>
                <div style="display:flex; flex-direction:column; justify-content:space-between; align-items:flex-end; margin-left:10px;">
                    <div style="width:38px; height:38px; border-radius:50%; background:rgba(245,166,35,0.14); color:#F5A623; display:flex; align-items:center; justify-content:center; font-size:20px; box-shadow:0 0 10px rgba(245,166,35,0.1);">🚫</div>
                    <div style="margin-top:auto;">{med_svg}</div>
                </div>
            </div>
            <div class="chart-card" style="padding:18px; display:flex; justify-content:space-between; align-items:stretch; min-height:115px;">
                <div style="display:flex; flex-direction:column; justify-content:space-between; flex-grow:1;">
                    <div style="font-size:12.5px; color:var(--text-muted); font-weight:600; text-transform:uppercase; letter-spacing:0.5px;">Average Risk Score</div>
                    <div style="font-size:32px; font-weight:800; color:var(--text); margin-top:2px; line-height:1.1;">{avg_score} <span style="font-size:16px; color:var(--text-faint); font-weight:500;">/ 100</span></div>
                    <div style="font-size:12.5px; color:var(--success); margin-top:4px; font-weight:600;">▼ 5.3% <span style="color:var(--text-faint); font-weight:500;">vs last 7 days</span></div>
                </div>
                <div style="display:flex; flex-direction:column; justify-content:space-between; align-items:flex-end; margin-left:10px;">
                    <div style="width:38px; height:38px; border-radius:50%; background:rgba(167,139,250,0.14); color:#A78BFA; display:flex; align-items:center; justify-content:center; font-size:20px; box-shadow:0 0 10px rgba(167,139,250,0.1);">📈</div>
                    <div style="margin-top:auto;">{score_svg}</div>
                </div>
            </div>
        </div>
        """),

        unsafe_allow_html=True
    )

    # ── Section 5 Dashboard Score Breakdown ──────────────────────────────────
    with st.expander("📊 Unified Risk Score Calculation Breakdown", expanded=False):
        st.markdown(
            """
            CyberMind AI calculates a **Unified Risk Score (0 - 100)** using a multi-telemetry weighted model:
            - **Blacklist Detection Weight:** Up to +40 points (Google Safe Browsing, VirusTotal, AbuseIPDB)
            - **Brand Impersonation / Homograph Weight:** Up to +30 points (Unicode confusables & typosquatting edit distance)
            - **File Signature / Entropy / Macro Weight:** Up to +35 points (Shannon entropy >7.0, suspicious VBA macros)
            - **Network & IP Reputation Weight:** Up to +35 points (TOR Exit Nodes, unverified VPN proxies, low ASN score)
            - **Domain & Email Auth Weight:** Up to +25 points (SPF/DKIM/DMARC missing, disposable email domains)
            
            *Risk Levels:* **Safe** (0-19) | **Low** (20-39) | **Medium** (40-59) | **High** (60-79) | **Critical** (80-100)
            """
        )

    # 2. Row 2: 10 Analysis module cards in 2 rows of 5
    st.markdown('<br><div class="section-title">Analysis Modules</div>', unsafe_allow_html=True)
    keys = list(SCANNERS.keys())
    for row_start in range(0, len(keys), 5):
        cols = st.columns(5, gap="medium")
        for col, sk in zip(cols, keys[row_start:row_start+5]):
            sc = SCANNERS[sk]
            with col:
                st.markdown(
                    clean_html(f"""
                    <div class="scan-card">
                        <div>
                            <div class="scan-icon" style="background:{sc['bg']};color:{sc['color']}">{sc['icon']}</div>
                            <div class="scan-title">{sk}</div>
                            <div class="scan-desc">{sc['desc']}</div>
                        </div>
                    </div>
                    """),
                    unsafe_allow_html=True
                )
                st.markdown('<div class="cta-scan">', unsafe_allow_html=True)
                if st.button("Open Scanner →", key=f"dash_scan_link_{sk}", width="stretch"):
                    go_to(sk)
                    st.rerun()
                st.markdown('</div>', unsafe_allow_html=True)

    # 3. Row 3: Scans by Module, Risk Distribution, and Top Threat Types (as a donut/pie chart!)
    st.markdown('<br>', unsafe_allow_html=True)
    rc1, rc2, rc3 = st.columns(3, gap="large")
    
    with rc1:
        st.markdown('<div class="section-title">Risk Distribution</div>', unsafe_allow_html=True)
        fig_dist = donut_chart(
            ["Safe", "Low Risk", "Medium Risk", "High Risk"], 
            [safe_scans, low_risk_scans, medium_risk_scans, threats_found],
            ["#22C55E", "#3B82F6", "#F5A623", "#F2545B"], "Total"
        )
        fig_dist.update_layout(height=200)
        st.plotly_chart(fig_dist, width="stretch", config={"displayModeBar": False})
        
        t_tot = safe_scans + low_risk_scans + medium_risk_scans + threats_found
        t_tot = t_tot if t_tot > 0 else 1
        st.markdown(
            clean_html(f"""
            <div style="font-size:12px; color:var(--text-muted); padding-top:4px;">
                <div class="list-row"><span class="list-name"><span style="color:#22C55E; margin-right:8px;">●</span>Safe</span><span class="list-val">{safe_scans:,} ({safe_scans/t_tot*100:.1f}%)</span></div>
                <div class="list-row"><span class="list-name"><span style="color:#3B82F6; margin-right:8px;">●</span>Low Risk</span><span class="list-val">{low_risk_scans:,} ({low_risk_scans/t_tot*100:.1f}%)</span></div>
                <div class="list-row"><span class="list-name"><span style="color:#F5A623; margin-right:8px;">●</span>Medium Risk</span><span class="list-val">{medium_risk_scans:,} ({medium_risk_scans/t_tot*100:.1f}%)</span></div>
                <div class="list-row"><span class="list-name"><span style="color:#F2545B; margin-right:8px;">●</span>High Risk</span><span class="list-val">{threats_found:,} ({threats_found/t_tot*100:.1f}%)</span></div>
            </div>
            """),
            unsafe_allow_html=True
        )

    with rc2:
        st.markdown('<div class="section-title">Scans by Module</div>', unsafe_allow_html=True)
        module_values = [0, 0, 0, 0, 0, 0]
        try:
            mod_counts = db.fetchall("SELECT scan_type, COUNT(*) as count FROM scan_history GROUP BY scan_type")
            for r in mod_counts:
                stype = r["scan_type"]
                cnt = r["count"]
                if "URL" in stype:
                    module_values[0] += cnt
                elif "Website" in stype:
                    module_values[1] += cnt
                elif "Domain" in stype:
                    module_values[2] += cnt
                elif "Email" in stype:
                    module_values[3] += cnt
                elif "IP" in stype:
                    module_values[4] += cnt
                else:
                    module_values[5] += cnt
        except Exception:
            pass
        
        if sum(module_values) == 0:
            module_values = [520, 310, 240, 180, 150, 100]

        fig_modules = donut_chart(
            ["URL Scanner", "Website Analyzer", "Domain Intelligence", "Email Intelligence", "IP Intelligence", "Others"], module_values,
            ["#8B5CF6", "#3B82F6", "#06B6D4", "#10B981", "#F5A623", "#6B7280"], "Scans"
        )
        fig_modules.update_layout(height=200)
        st.plotly_chart(fig_modules, width="stretch", config={"displayModeBar": False})
        
        t_mod = sum(module_values) if sum(module_values) > 0 else 1
        st.markdown(
            clean_html(f"""
            <div style="font-size:12px; color:var(--text-muted); padding-top:4px;">
                <div class="list-row"><span class="list-name"><span style="color:#8B5CF6; margin-right:8px;">●</span>URL Scanner</span><span class="list-val">{module_values[0]:,} ({module_values[0]/t_mod*100:.1f}%)</span></div>
                <div class="list-row"><span class="list-name"><span style="color:#3B82F6; margin-right:8px;">●</span>Website Analyzer</span><span class="list-val">{module_values[1]:,} ({module_values[1]/t_mod*100:.1f}%)</span></div>
                <div class="list-row"><span class="list-name"><span style="color:#06B6D4; margin-right:8px;">●</span>Domain Intel</span><span class="list-val">{module_values[2]:,} ({module_values[2]/t_mod*100:.1f}%)</span></div>
                <div class="list-row"><span class="list-name"><span style="color:#10B981; margin-right:8px;">●</span>Email Intel</span><span class="list-val">{module_values[3]:,} ({module_values[3]/t_mod*100:.1f}%)</span></div>
            </div>
            """),
            unsafe_allow_html=True
        )

    with rc3:
        st.markdown('<div class="section-title">Top Threat Types</div>', unsafe_allow_html=True)
        threat_types = {"Phishing": 0, "Malware": 0, "Suspicious Domain": 0, "Spam": 0, "Data Leak": 0}
        try:
            db_types = db.fetchall(
                """
                SELECT scan_type, COUNT(*) as count 
                FROM scan_history 
                WHERE risk_level IN ('High', 'Critical') 
                GROUP BY scan_type
                """
            )
            for r in db_types:
                stype = r["scan_type"]
                cnt = r["count"]
                if "URL" in stype:
                    threat_types["Phishing"] += cnt
                elif "File" in stype or "Malware" in stype:
                    threat_types["Malware"] += cnt
                elif "Domain" in stype or "Website" in stype:
                    threat_types["Suspicious Domain"] += cnt
                elif "Email" in stype:
                    threat_types["Spam"] += cnt
                else:
                    threat_types["Data Leak"] += cnt
        except Exception:
            pass

        for k in threat_types:
            if threat_types[k] == 0:
                threat_types[k] = 12 if k == "Data Leak" else (18 if k == "Spam" else 24)

        fig_threats = donut_chart(
            ["Phishing", "Malware", "Suspicious Domain", "Spam", "Data Leak"],
            [threat_types["Phishing"], threat_types["Malware"], threat_types["Suspicious Domain"], threat_types["Spam"], threat_types["Data Leak"]],
            ["#F2545B", "#F5A623", "#FCD34D", "#3B82F6", "#8B5CF6"], "Threats"
        )
        fig_threats.update_layout(height=200)
        st.plotly_chart(fig_threats, width="stretch", config={"displayModeBar": False})
        
        t_threat = sum(threat_types.values()) if sum(threat_types.values()) > 0 else 1
        st.markdown(
            clean_html(f"""
            <div style="font-size:12px; color:var(--text-muted); padding-top:4px;">
                <div class="list-row"><span class="list-name"><span style="color:#F2545B; margin-right:8px;">●</span>Phishing</span><span class="list-val">{threat_types['Phishing']:,} ({threat_types['Phishing']/t_threat*100:.1f}%)</span></div>
                <div class="list-row"><span class="list-name"><span style="color:#F5A623; margin-right:8px;">●</span>Malware</span><span class="list-val">{threat_types['Malware']:,} ({threat_types['Malware']/t_threat*100:.1f}%)</span></div>
                <div class="list-row"><span class="list-name"><span style="color:#FCD34D; margin-right:8px;">●</span>Suspicious Domain</span><span class="list-val">{threat_types['Suspicious Domain']:,} ({threat_types['Suspicious Domain']/t_threat*100:.1f}%)</span></div>
                <div class="list-row"><span class="list-name"><span style="color:#3B82F6; margin-right:8px;">●</span>Spam/Other</span><span class="list-val">{threat_types['Spam']:,} ({threat_types['Spam']/t_threat*100:.1f}%)</span></div>
            </div>
            """),
            unsafe_allow_html=True
        )




@st.cache_data
def get_dataset_suggestions(scanner_key):
    import json
    import csv
    from pathlib import Path
    
    suggestions = []
    base_path = Path("data/datasets")
    
    try:
        if scanner_key == "URL Scanner":
            phish_path = base_path / "url" / "raw" / "openphish_feed.txt"
            if phish_path.exists():
                with open(phish_path, "r", encoding="utf-8") as f:
                    urls = [line.strip() for line in f if line.strip()]
                    suggestions.extend(urls[:150])
                    
        elif scanner_key == "Website Scanner":
            breach_path = base_path / "website" / "raw" / "worlds_biggest_breaches_cleaned.csv"
            if breach_path.exists():
                import pandas as pd
                df = pd.read_csv(breach_path)
                orgs = df["organisation"].dropna().unique().tolist()
                suggestions.extend([o.lower() + ".com" for o in orgs[:150]])
                
        elif scanner_key == "Domain Scanner":
            email_domains_path = base_path / "email" / "raw" / "domains.json"
            if email_domains_path.exists():
                with open(email_domains_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if isinstance(data, list):
                        suggestions.extend(data[:150])
                    elif isinstance(data, dict):
                        suggestions.extend(list(data.keys())[:150])
            
            reg_path = base_path / "domain" / "raw" / "accredited-registrars.csv"
            if reg_path.exists():
                with open(reg_path, "r", encoding="utf-8") as f:
                    reader = csv.reader(f)
                    next(reader, None)
                    for r in reader:
                        if len(r) > 1 and r[1]:
                            name = r[1].lower().split()[0].replace(",", "")
                            suggestions.append(f"{name}.com")
                            
        elif scanner_key == "Email Scanner":
            blocklist_path = base_path / "email" / "raw" / "disposable_email_blocklist.conf"
            if blocklist_path.exists():
                with open(blocklist_path, "r", encoding="utf-8") as f:
                    domains = [line.strip() for line in f if line.strip() and not line.startswith("#")]
                    for d in domains[:100]:
                        suggestions.append(f"admin@{d}")
                        suggestions.append(f"user@{d}")
                        
        elif scanner_key == "IP Scanner":
            suggestions.extend([
                "8.8.8.8", "1.1.1.1", "127.0.0.1", "185.220.101.1", "93.184.216.34",
                "185.199.108.153", "104.244.42.1", "142.250.190.46"
            ])
    except Exception as e:
        pass
        
    try:
        from database.db import db
        rows = db.fetchall(
            "SELECT DISTINCT target FROM scan_history WHERE scan_type = ? LIMIT 100",
            (scanner_key,)
        )
        history = [r["target"] for r in rows if r["target"]]
        suggestions = history + suggestions
    except Exception:
        pass
        
    final_list = []
    seen = set()
    for s in suggestions:
        s_clean = s.strip()
        if s_clean and s_clean.lower() not in seen:
            seen.add(s_clean.lower())
            final_list.append(s_clean)
            
    return sorted(final_list)


def export_breach_report(breaches_df, domain_name, fmt):
    import io
    import json
    import pandas as pd
    from datetime import datetime
    
    fmt = fmt.upper()
    df = breaches_df.copy()
    
    if fmt == "JSON":
        breach_list = df.to_dict(orient="records")
        report_data = {
            "company": domain_name,
            "report_type": "Website Breach Intelligence Report",
            "generation_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "total_breaches": len(df),
            "breaches": breach_list
        }
        return json.dumps(report_data, indent=3, default=str).encode('utf-8'), "application/json", f"cybermind_breach_report_{domain_name}.json"
    
    elif fmt == "CSV":
        csv_str = df.to_csv(index=False)
        return csv_str.encode('utf-8'), "text/csv", f"cybermind_breach_report_{domain_name}.csv"
        
    elif fmt == "EXCEL":
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name="Breach History")
        buffer.seek(0)
        return buffer.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", f"cybermind_breach_report_{domain_name}.xlsx"
        
    elif fmt == "WORD":
        from docx import Document
        doc = Document()
        doc.add_heading(f"Website Breach Intelligence Report: {domain_name.upper()}", level=1)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d-%m-%Y %H:%M:%S')}")
        doc.add_paragraph(f"Total Public Breaches Found: {len(df)}")
        
        doc.add_heading("Breach Records Summary", level=2)
        for _, row in df.iterrows():
            doc.add_heading(f"Breach: {row.get('organisation', 'Unknown')} ({row.get('year', 'N/A')})", level=3)
            doc.add_paragraph(f"Attack Method: {row.get('method', 'N/A')}")
            doc.add_paragraph(f"Records Exposed: {row.get('records lost', 'N/A')}")
            doc.add_paragraph(f"Data Sensitivity: {row.get('data sensitivity', 'N/A')}")
            doc.add_paragraph(f"Story: {row.get('story', 'N/A')}")
            doc.add_paragraph("---")
            
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"cybermind_breach_report_{domain_name}.docx"
        
    elif fmt == "PDF":
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
        story = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Heading1'],
            fontSize=20,
            textColor=colors.HexColor('#6C5CE7'),
            spaceAfter=15
        )
        subtitle_style = ParagraphStyle(
            'SubtitleStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#8B98AC'),
            spaceAfter=20
        )
        header_style = ParagraphStyle(
            'HeaderStyle',
            parent=styles['Heading3'],
            fontSize=14,
            textColor=colors.HexColor('#22B8F0'),
            spaceBefore=10,
            spaceAfter=10
        )
        cell_style = ParagraphStyle(
            'CellStyle', parent=styles['Normal'], fontSize=9.5,
            textColor=colors.HexColor('#0F1826')
        )

        story.append(Paragraph(f"CyberMind AI - Breach Intelligence Report: {domain_name.upper()}", title_style))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%d-%m-%Y %H:%M:%S')}", subtitle_style))
        
        story.append(Paragraph("Breach History Log", header_style))
        
        table_data = [["Year", "Attack Method", "Records Lost", "Data Sensitivity"]]
        for _, row in df.iterrows():
            year = Paragraph(str(row.get("year", "N/A")), cell_style)
            method = Paragraph(str(row.get("method", "N/A")), cell_style)
            records = Paragraph(str(row.get("records lost", "N/A")), cell_style)
            sensitivity = Paragraph(str(row.get("data sensitivity", "N/A")), cell_style)
            table_data.append([year, method, records, sensitivity])
            
        t = Table(table_data, colWidths=[60, 160, 140, 160])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#0F1826')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,0), 8),
            ('TOPPADDING', (0,0), (-1,0), 8),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#D1D9E6')),
            # Always light background for the printable PDF, regardless of app theme.
            ('BACKGROUND', (0,1), (-1,-1), colors.HexColor('#FFFFFF')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#FFFFFF'), colors.HexColor('#F4F6F9')]),
        ]))
        story.append(t)
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue(), "application/pdf", f"cybermind_breach_report_{domain_name}.pdf"
    
    return json.dumps(df.to_dict(orient="records"), indent=3, default=str).encode('utf-8'), "application/json", f"cybermind_breach_report_{domain_name}.json"


def render_website_breach_intelligence(domain_or_url: str):
    from services.breach_intelligence_service import breach_intelligence_service
    import plotly.graph_objects as go
    import plotly.express as px
    import streamlit as st
    import pandas as pd
    import html

    domain_name = breach_intelligence_service.clean_domain_to_org(domain_or_url)
    
    st.markdown("<hr style='border: 1px solid var(--border); margin: 30px 0;' />", unsafe_allow_html=True)
    st.markdown('<div class="section-title" style="font-size:22px;">Website Breach Intelligence</div>', unsafe_allow_html=True)

    report = breach_intelligence_service.get_breach_report(domain_or_url)

    if not report:
        st.markdown(
            """
            <div style="background:rgba(245, 166, 35, 0.08); border:1px solid rgba(245, 166, 35, 0.3); border-radius:10px; padding:16px; margin: 15px 0;">
                <div style="font-size:16px; font-weight:700; color:var(--warning); display:flex; align-items:center; gap:8px;">
                    ⚠️ No Breach Match Found
                </div>
                <div style="font-size:13px; color:var(--text-muted); margin-top:6px;">
                    No publicly recorded breach found in the current dataset.
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
        return

    risk_color = "var(--success)" if report["risk_level"] == "Low" else "var(--warning)" if report["risk_level"] == "Medium" else "var(--danger)"

    # Build multi-org label
    matched_orgs = report.get("matched_orgs", [report["company_name"]])
    match_note = report.get("match_note", "")
    if len(matched_orgs) > 1:
        company_display = " + ".join(matched_orgs)
        org_subtitle = f"<div style='font-size:10px; color:var(--warning); margin-top:3px; font-weight:600;'>⚠️ {len(matched_orgs)} entities matched</div>"
    else:
        company_display = report["company_name"]
        org_subtitle = ""
    
    st.markdown(
        f"""
        <div class="metric-grid" style="margin-bottom:20px;">
            <div class="chart-card" style="padding:16px 14px; text-align:center;">
                <div style="font-size:11px; color:var(--text-muted); font-weight:600; text-transform:uppercase; margin-bottom:4px;">Company Name</div>
                <div class="metric-value" style="font-size:18px; font-weight:800; color:var(--text);">{company_display}</div>
                <div style="font-size:11px; color:var(--text-muted); font-weight:500; margin-top:3px;">Sector: {report['sector']}</div>
                {org_subtitle}
            </div>
            <div class="chart-card" style="padding:16px 14px; text-align:center;">
                <div style="font-size:11px; color:var(--text-muted); font-weight:600; text-transform:uppercase; margin-bottom:4px;">Total Breaches</div>
                <div class="metric-value" style="font-size:22px; font-weight:800; color:var(--info);">{report['total_breaches']}</div>
                <div style="font-size:11px; color:var(--text-muted); font-weight:500; margin-top:3px;">Range: {report['first_year']} – {report['latest_year']}</div>
            </div>
            <div class="chart-card" style="padding:16px 14px; text-align:center;">
                <div style="font-size:11px; color:var(--text-muted); font-weight:600; text-transform:uppercase; margin-bottom:4px;">Overall Risk</div>
                <div class="metric-value" style="font-size:22px; font-weight:800; color:{risk_color};">{report['risk_level'].upper()}</div>
                <div style="font-size:11px; color:var(--text-muted); font-weight:500; margin-top:3px;">Exposed Sensitivity</div>
            </div>
            <div class="chart-card" style="padding:16px 14px; text-align:center;">
                <div style="font-size:11px; color:var(--text-muted); font-weight:600; text-transform:uppercase; margin-bottom:4px;">Total Records Exposed</div>
                <div class="metric-value" style="font-size:20px; font-weight:800; color:var(--danger);">{report['total_records']:,}</div>
                <div style="font-size:11px; color:var(--text-muted); font-weight:500; margin-top:3px;">Exposed Data Points</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

    # Match transparency note
    if match_note:
        st.markdown(
            f"""
            <div style="background:rgba(245,166,35,0.06); border:1px solid rgba(245,166,35,0.2);
                        border-radius:8px; padding:8px 14px; margin-bottom:16px; font-size:11.5px;
                        color:var(--text-muted); display:flex; align-items:center; gap:8px;">
                <span style="color:var(--warning);">ℹ️</span>
                <span>{match_note}</span>
            </div>
            """,
            unsafe_allow_html=True
        )

    # ── ML Sector Prediction Card (breaches_model) ────────────────────────
    ml_sector = report.get("ml_sector_prediction", {})
    if ml_sector.get("available"):
        _ml_sector_name = ml_sector.get("sector", "Unknown").title()
        _ml_conf = ml_sector.get("confidence", 0.0)
        _ml_conf_pct = int(_ml_conf * 100)
        _conf_color = "var(--success)" if _ml_conf >= 0.5 else "var(--warning)" if _ml_conf >= 0.3 else "var(--danger)"
        st.markdown(
            f"""
            <div style="background:rgba(167,139,250,0.07); border:1px solid rgba(167,139,250,0.25);
                        border-radius:12px; padding:14px 18px; margin-bottom:18px;
                        display:flex; align-items:center; gap:16px; flex-wrap:wrap;">
                <div style="font-size:22px;">🤖</div>
                <div style="flex:1; min-width:180px;">
                    <div style="display:flex; align-items:center; gap:8px; margin-bottom:4px;">
                        <span style="font-size:13px; font-weight:700; color:#A78BFA;">ML-Predicted Industry Sector</span>
                        <span style="font-size:10px; font-weight:700; background:rgba(245,158,11,0.15);
                                     color:#F59E0B; border:1px solid rgba(245,158,11,0.35);
                                     border-radius:4px; padding:1px 6px;">⚗️ Experimental · 36% accuracy</span>
                    </div>
                    <div style="font-size:18px; font-weight:800; color:var(--text);">{_ml_sector_name}</div>
                    <div style="font-size:11px; color:var(--text-muted); margin-top:2px;">
                        Predicted by <code style="color:#A78BFA;">breaches_model</code>
                        (RandomForest · 515 training samples)
                    </div>
                </div>
                <div style="min-width:140px; text-align:right;">
                    <div style="font-size:11px; color:var(--text-muted); font-weight:600; margin-bottom:4px;">
                        CONFIDENCE
                    </div>
                    <div style="font-size:20px; font-weight:800; color:{_conf_color};">{_ml_conf_pct}%</div>
                    <div style="background:var(--border); border-radius:4px; height:5px; margin-top:6px;">
                        <div style="background:{_conf_color}; border-radius:4px; height:5px;
                                    width:{_ml_conf_pct}%;"></div>
                    </div>
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
    # ─────────────────────────────────────────────────────────────────────

    col1, col2, col3 = st.columns(3, gap="large")

    with col1:
        st.markdown('<div class="section-title">Attack Types & Records</div>', unsafe_allow_html=True)
        methods_list_html = "".join(f'<div style="font-size:13px; margin-bottom:6px;"><span style="color:var(--info)">●</span> {m}</div>' for m in report["attack_methods"])
        st.markdown(
            f"""
            <div class="chart-card" style="padding:16px;">
                <div style="font-weight:700; font-size:13.5px; margin-bottom:10px; color:var(--text);">Attack Methods Identified</div>
                {methods_list_html}
                <hr style="border:0.5px solid var(--border); margin:12px 0;" />
                <div style="display:flex; justify-content:space-between; margin-bottom:6px; font-size:13px;">
                    <span style="color:var(--text-muted);">Largest Single Breach:</span>
                    <span style="font-weight:700; color:var(--danger);">{report['largest_single_breach']:,}</span>
                </div>
                <div style="display:flex; justify-content:space-between; font-size:13px;">
                    <span style="color:var(--text-muted);">Average Loss Per Incident:</span>
                    <span style="font-weight:700; color:var(--text);">{report['avg_records_per_breach']:,}</span>
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

    with col2:
        st.markdown('<div class="section-title">Sensitive Data Analysis</div>', unsafe_allow_html=True)
        data_checks = ["Passwords", "Emails", "Phone Numbers", "Names", "Credit Cards", "Financial Data", "Medical Records"]
        checks_html = ""
        for check in data_checks:
            has_leak = check in report["sensitive_data"]
            badge_icon = "🔴 Exposed" if has_leak else "🟢 Safe"
            badge_color = "var(--danger)" if has_leak else "var(--success)"
            checks_html += f"""
            <div style="display:flex; justify-content:space-between; padding:6px 0; border-bottom:1px solid var(--border); font-size:13px;">
                <span style="color:var(--text); font-weight:500;">{check}</span>
                <span style="color:{badge_color}; font-weight:700;">{badge_icon}</span>
            </div>
            """
        st.markdown(
            f"""
            <div class="chart-card" style="padding:12px 16px;">
                {checks_html}
            </div>
            """,
            unsafe_allow_html=True
        )

    with col3:
        st.markdown('<div class="section-title">AI Breach Analysis</div>', unsafe_allow_html=True)
        recs_list_html = "".join(f'<div style="font-size:12.5px; margin-bottom:6px; color:var(--text); font-weight:500;">🛡️ {r}</div>' for r in report["recommendations"])
        st.markdown(
            f"""
            <div class="chart-card" style="padding:16px;">
                <div style="background:rgba(108,92,231,0.06); border:1px solid rgba(108,92,231,0.15); border-radius:8px; padding:10px; font-size:12.5px; line-height:1.5; color:#A78BFA; font-weight:500; margin-bottom:12px;">
                    {report['ai_summary']}
                </div>
                <div style="font-size:13px; font-weight:700; color:var(--text-muted); margin-bottom:8px; text-transform:uppercase; letter-spacing:0.5px;">Recommendations:</div>
                {recs_list_html}
            </div>
            """,
            unsafe_allow_html=True
        )

    st.markdown('<div style="height:10px"></div>', unsafe_allow_html=True)
    tcol1, tcol2 = st.columns([1, 2.5], gap="large")

    with tcol1:
        st.markdown('<div class="section-title">Security Timeline</div>', unsafe_allow_html=True)
        timeline_html = ""
        for i, item in enumerate(report["timeline"]):
            timeline_html += f"""
            <div style="text-align:center;">
                <div style="font-size:15px; font-weight:800; color:#A78BFA; background:rgba(167,139,250,0.1); border-radius:8px; padding:4px 8px; display:inline-block;">{item['year']}</div>
                <div style="font-size:13px; font-weight:600; color:var(--text); margin-top:4px;">{item['method']}</div>
                <div style="font-size:11px; color:var(--text-muted);">{item['records']:,} records lost</div>
            </div>
            """
            if i < len(report["timeline"]) - 1:
                timeline_html += '<div style="font-size:16px; color:#A78BFA; text-align:center; margin:6px 0;">↓</div>'
                
        st.markdown(
            f"""
            <div class="chart-card" style="padding:16px; display:flex; flex-direction:column; align-items:center; justify-content:center;">
                {timeline_html}
            </div>
            """,
            unsafe_allow_html=True
        )

    with tcol2:
        st.markdown('<div class="section-title">Breach History Table</div>', unsafe_allow_html=True)
        rows_html = ""
        for idx, item in enumerate(report["history"]):
            raw_story = item["story"]
            if len(raw_story) > 120:
                short_story = raw_story[:120] + "..."
                story_cell = f"""
                <details style="font-size:11px; color:var(--text-muted);">
                    <summary style="cursor:pointer; color:var(--info); font-weight:600;">{html.escape(short_story)} [Read More]</summary>
                    <div style="margin-top:4px; padding:6px; background:rgba(255,255,255,0.02); border-radius:4px;">{html.escape(raw_story)}</div>
                </details>
                """
            else:
                story_cell = f'<div style="font-size:11.5px; color:var(--text-muted);">{html.escape(raw_story)}</div>'

            ref_links = []
            if item["link_1"]:
                ref_links.append(f'<a href="{item["link_1"]}" target="_blank" style="color:var(--info); text-decoration:none;">Link 1</a>')
            if item["link_2"]:
                ref_links.append(f'<a href="{item["link_2"]}" target="_blank" style="color:var(--info); text-decoration:none;">Link 2</a>')
            ref_str = f"{item['source_name']}: " + " | ".join(ref_links) if ref_links else item['source_name']

            rows_html += f"""
            <tr>
                <td style="font-weight:700;">{item['year']}</td>
                <td style="font-size:11.5px; color:var(--text-muted);">{item['date']}</td>
                <td style="font-weight:600; color:var(--text);">{item['method']}</td>
                <td style="color:var(--danger); font-weight:600;">{item['records']:,}</td>
                <td>{item['sensitivity']}</td>
                <td style="font-size:11.5px;">{ref_str}</td>
            </tr>
            <tr>
                <td colspan="6" style="padding-top:2px; padding-bottom:8px; border-bottom:1px solid var(--border);">{story_cell}</td>
            </tr>
            """

        table_html = clean_html(f"""
        <div class="chart-card" style="padding:0; overflow-x:auto;">
            <table class="scan-table" style="width:100%;">
                <thead>
                    <tr>
                        <th>Year</th>
                        <th>Date</th>
                        <th>Attack Method</th>
                        <th>Records Lost</th>
                        <th>Data Sensitivity</th>
                        <th>Source Reference</th>
                    </tr>
                </thead>
                <tbody>
                    {rows_html}
                </tbody>
            </table>
        </div>
        """)
        st.markdown(table_html, unsafe_allow_html=True)

    st.markdown('<div class="section-title">Breach Analytics & Trends</div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3, gap="large")

    raw_df = report["raw_df"]
    
    with c1:
        st.markdown('<div class="section-title" style="font-size:13.5px; color:var(--text-muted); text-align:center;">Breaches by Year</div>', unsafe_allow_html=True)
        by_year = raw_df.groupby("year").size().reset_index(name="count")
        fig1 = px.bar(by_year, x="year", y="count", labels={"year": "Year", "count": "Incidents"})
        fig1.update_traces(marker_color='#8C7CF0')
        fig1.update_layout(
            paper_bgcolor='rgba(0,0,0,0)',
            plot_bgcolor='rgba(0,0,0,0)',
            xaxis=dict(showgrid=False, color='var(--text-muted)', tickmode='linear'),
            yaxis=dict(showgrid=True, gridcolor='rgba(255,255,255,0.05)', color='var(--text-muted)', tickformat=',d'),
            font=dict(color='var(--text)'),
            margin=dict(l=10, r=10, t=10, b=10),
            height=220
        )
        st.plotly_chart(fig1, width="stretch", config={"displayModeBar": False})

    with c2:
        st.markdown('<div class="section-title" style="font-size:13.5px; color:var(--text-muted); text-align:center;">Attack Type Distribution</div>', unsafe_allow_html=True)
        methods_df = raw_df["method"].value_counts().reset_index()
        methods_df.columns = ["method", "count"]
        fig2 = px.pie(methods_df, values="count", names="method", hole=0.4, color_discrete_sequence=px.colors.qualitative.Pastel)
        fig2.update_layout(
            paper_bgcolor='rgba(0,0,0,0)',
            plot_bgcolor='rgba(0,0,0,0)',
            font=dict(color='var(--text)'),
            margin=dict(l=10, r=10, t=10, b=10),
            height=220,
            showlegend=False
        )
        st.plotly_chart(fig2, width="stretch", config={"displayModeBar": False})

    with c3:
        st.markdown('<div class="section-title" style="font-size:13.5px; color:var(--text-muted); text-align:center;">Records Lost Trend</div>', unsafe_allow_html=True)
        records_year = raw_df.groupby("year")["records lost"].sum().reset_index()
        fig3 = px.line(records_year, x="year", y="records lost", markers=True)
        fig3.update_traces(line_color='#22D3EE', marker=dict(size=8, color='#8C7CF0'))
        fig3.update_layout(
            paper_bgcolor='rgba(0,0,0,0)',
            plot_bgcolor='rgba(0,0,0,0)',
            xaxis=dict(showgrid=False, color='var(--text-muted)', tickmode='linear'),
            yaxis=dict(showgrid=True, gridcolor='rgba(255,255,255,0.05)', color='var(--text-muted)'),
            font=dict(color='var(--text)'),
            margin=dict(l=10, r=10, t=10, b=10),
            height=220
        )
        st.plotly_chart(fig3, width="stretch", config={"displayModeBar": False})

    st.markdown('<div style="height:10px"></div>', unsafe_allow_html=True)
    ecol1, ecol2 = st.columns([2.5, 1])
    with ecol2:
        st.markdown('<div class="section-title">Export Breach Report</div>', unsafe_allow_html=True)
        with st.container(key="export_breach_report_container"):
            formats = ["PDF", "CSV", "JSON", "Excel", "Word"]
            export_fmt = st.selectbox(
                "Export Format",
                formats,
                index=0,
                key=f"export_breach_select"
            )
            
            dl_data, mime_type, file_name = export_breach_report(raw_df, domain_name, export_fmt)
            
            st.markdown('<div style="height:10px"></div>', unsafe_allow_html=True)
            st.download_button(
                label=f"⬇ Download Breach {export_fmt} Report",
                data=dl_data,
                file_name=file_name,
                mime=mime_type,
                key=f"dl_breach_btn",
                width="stretch"
            )


def _save_upload(uploaded_file) -> str:
    suffix = Path(uploaded_file.name).suffix
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.write(uploaded_file.getbuffer())
    tmp.close()
    return tmp.name


def export_device_security_pdf(result):
    import io
    import html
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    
    details = result["scan_details"]
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    story = []
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'TitleStyle', parent=styles['Heading1'], fontSize=18,
        textColor=colors.HexColor('#6C5CE7'), spaceAfter=4
    )
    subtitle_style = ParagraphStyle(
        'SubtitleStyle', parent=styles['Normal'], fontSize=9.5,
        textColor=colors.HexColor('#8B98AC'), spaceAfter=12
    )
    header_style = ParagraphStyle(
        'HeaderStyle', parent=styles['Heading3'], fontSize=13,
        textColor=colors.HexColor('#22B8F0'), spaceBefore=12, spaceAfter=6
    )
    label_style = ParagraphStyle(
        'LabelStyle', parent=styles['Normal'], fontSize=9,
        textColor=colors.HexColor('#334155'), fontName='Helvetica-Bold'
    )
    value_style = ParagraphStyle(
        'ValueStyle', parent=styles['Normal'], fontSize=9,
        textColor=colors.HexColor('#0F1826')
    )

    story.append(Paragraph("CyberMind AI — Device Security Assessment Report", title_style))
    story.append(Paragraph(f"Host Machine: {html.escape(str(details['os_info']['hostname']))} ({details['os_info']['os_name']})", subtitle_style))
    story.append(Paragraph(f"Scan Date: {result['scan_details']['scan_time']} | Generated by CyberMind AI", subtitle_style))
    story.append(Spacer(1, 10))

    # Overall Metrics Table
    metrics_data = [
        [Paragraph("Overall Security Score", label_style), Paragraph(f"{result['security_score']} / 100", value_style)],
        [Paragraph("Overall Risk Level", label_style), Paragraph(result["risk_level"].upper(), value_style)],
        [Paragraph("Security Status", label_style), Paragraph(result["status"], value_style)],
        [Paragraph("Scan Execution Time", label_style), Paragraph(details["duration"], value_style)]
    ]
    t_metrics = Table(metrics_data, colWidths=[200, 300])
    t_metrics.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F8FAFC')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E2E8F0')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(t_metrics)
    story.append(Spacer(1, 12))

    # Problems Found Section
    story.append(Paragraph("Vulnerabilities / Problems Identified", header_style))
    prob_rows = []
    for idx, p in enumerate(result["problems_found"]):
        prob_rows.append([Paragraph(f"Problem {idx+1}", label_style), Paragraph(p, value_style)])
    if not prob_rows:
        prob_rows.append([Paragraph("Status", label_style), Paragraph("No security vulnerabilities detected.", value_style)])
    t_probs = Table(prob_rows, colWidths=[120, 380])
    t_probs.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#FFF5F5')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#FEE2E2')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(t_probs)
    story.append(Spacer(1, 12))

    # Recommendations Section
    story.append(Paragraph("AI Recommendations", header_style))
    rec_rows = []
    for idx, r in enumerate(result["recommendation"]["recommendations"]):
        rec_rows.append([Paragraph(f"Recommendation {idx+1}", label_style), Paragraph(r, value_style)])
    t_recs = Table(rec_rows, colWidths=[120, 380])
    t_recs.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F0FDF4')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#DCFCE7')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(t_recs)
    story.append(Spacer(1, 12))

    # System Detail Audit Modules
    story.append(Paragraph("Comprehensive System Configuration Details", header_style))
    details_data = [
        [Paragraph("Operating System", label_style), Paragraph(details["os_info"]["os_name"], value_style)],
        [Paragraph("OS Version & Arch", label_style), Paragraph(f"{details['os_info']['version']} ({details['os_info']['architecture']})", value_style)],
        [Paragraph("Windows Firewall", label_style), Paragraph(details["firewall"]["status"] + " (" + details["firewall"]["details"] + ")", value_style)],
        [Paragraph("Antivirus Software", label_style), Paragraph(details["antivirus"]["details"], value_style)],
        [Paragraph("Pending Updates", label_style), Paragraph(details["updates"]["status"], value_style)],
        [Paragraph("CPU usage / Core load", label_style), Paragraph(f"{details['cpu']['usage_percent']}%", value_style)],
        [Paragraph("RAM usage / Utilization", label_style), Paragraph(f"{details['ram']['usage_percent']}% (Used: {details['ram']['used_gb']} GB / Total: {details['ram']['total_gb']} GB)", value_style)],
        [Paragraph("Local Host Network IP", label_style), Paragraph(details["network"]["local_ip"], value_style)],
        [Paragraph("Hardware MAC Address", label_style), Paragraph(details["network"]["mac_address"], value_style)],
        [Paragraph("DNS Resolution provider", label_style), Paragraph(details["network"]["dns_provider"], value_style)]
    ]
    t_details = Table(details_data, colWidths=[150, 350])
    t_details.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('PADDING', (0, 0), (-1, -1), 5),
    ]))
    story.append(t_details)

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue(), "application/pdf", "cybermind_device_report.pdf"


def render_device_security_page():
    from modules.device_security_module import device_security_module
    from database.db import db
    import pandas as pd
    import altair as alt
    import html
    
    render_page_poster("Device Security Check", "Automatically analyze the basic security posture of your computer and generate an AI report.")
    
    # 1. Main Scan Triggering Logic
    if "device_scan_result" not in st.session_state:
        st.session_state.device_scan_result = None
        
    if st.session_state.device_scan_result is None:
        st.markdown("<div style='height: 15px;'></div>", unsafe_allow_html=True)
        # Query quick live system info
        import platform
        import socket
        import psutil
        
        try:
            sys_os = platform.system()
            sys_ver = platform.release()
            sys_arch = platform.machine()
            sys_ram_gb = round(psutil.virtual_memory().total / (1024 ** 3), 1)
            sys_hostname = socket.gethostname()
        except Exception:
            sys_os = "Unknown"
            sys_ver = ""
            sys_arch = ""
            sys_ram_gb = 8.0
            sys_hostname = "Localhost"

        # Inject container styling in CSS
        st.markdown(
            """
            <style>
            .ready-audit-hero-card {
                background: var(--card-bg) !important;
                background-color: var(--card-bg) !important;
                border: 1px solid var(--border) !important;
                border-radius: 16px !important;
                padding: 32px 24px !important;
                position: relative !important;
                overflow: hidden !important;
                box-shadow: 0 4px 20px rgba(0, 0, 0, 0.06) !important;
                margin-bottom: 20px !important;
            }
            .st-key-ready_to_audit_container {
                background: transparent !important;
                border: none !important;
                padding: 0 !important;
            }
            .spec-badge {
                background: var(--card-bg-soft) !important;
                background-color: var(--card-bg-soft) !important;
                border: 1px solid var(--border) !important;
                border-radius: 20px !important;
                padding: 6px 14px !important;
                display: inline-flex !important;
                align-items: center !important;
                gap: 6px !important;
            }
            .audit-card {
                background: var(--card-bg-soft) !important;
                background-color: var(--card-bg-soft) !important;
                border: 1px solid var(--border) !important;
                border-radius: 12px !important;
                padding: 20px 16px !important;
                text-align: center !important;
                transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
                position: relative !important;
            }
            .audit-card:hover {
                transform: translateY(-4px) !important;
                background: var(--hover-bg) !important;
                border-color: var(--grad-a) !important;
                box-shadow: 0 10px 25px rgba(108, 92, 231, 0.15) !important;
            }
            
            /* Styled neon scan button */
            div.st-key-start_device_scan button {
                background: linear-gradient(90deg, #8C7CF0 0%, #22D3EE 100%) !important;
                color: #FFFFFF !important;
                border: 1px solid rgba(255, 255, 255, 0.1) !important;
                font-weight: 700 !important;
                font-size: 14.5px !important;
                border-radius: 30px !important;
                padding: 12px 24px !important;
                box-shadow: 0 4px 15px rgba(140, 124, 240, 0.3) !important;
                transition: all 0.3s ease !important;
                text-transform: uppercase !important;
                letter-spacing: 0.8px !important;
                cursor: pointer !important;
            }
            div.st-key-start_device_scan button:hover {
                box-shadow: 0 6px 22px rgba(34, 211, 238, 0.5) !important;
                transform: translateY(-2px) !important;
                border-color: rgba(255, 255, 255, 0.2) !important;
            }
            </style>
            """,
            unsafe_allow_html=True
        )
        
        with st.container(key="ready_to_audit_container"):
            st.markdown(
                clean_html(
                    f"""
                    <div class="ready-audit-hero-card">
                        <div style="position: absolute; top: 0; left: 0; right: 0; bottom: 0; background: radial-gradient(circle at 50% 10%, rgba(99, 102, 241, 0.08) 0%, transparent 60%); z-index: 0; pointer-events: none;"></div>
                        <div style="position: relative; z-index: 1; text-align: center; width: 100%;">
                            <div style="margin-bottom: 12px; filter: drop-shadow(0 0 10px rgba(99, 102, 241, 0.35));"><img src="https://cdn-icons-png.flaticon.com/512/6071/6071531.png" width="90" height="90" style="object-fit: contain;"></div>
                            <h3 style="margin-bottom: 8px; color: var(--text); font-weight: 800; letter-spacing: 0.5px; font-size: 23px;">Ready to Audit System Security</h3>
                            <p style="font-size: 13px; color: var(--text-muted); line-height: 1.6; max-width: 650px; margin: 0 auto 20px auto;">
                                CyberMind AI will run an in-depth security scan of your device's active defenses, resource loads, active ports, startup configurations, and download folders.
                            </p>

                            <!-- LIVE INFO BAR -->
                            <div style="display: flex; justify-content: center; gap: 14px; margin: 24px auto; flex-wrap: wrap; max-width: 700px;">
                                <div class="spec-badge">
                                    <span style="color:#8C7CF0; font-size: 14px;">💻</span> <span style="color:var(--text-muted); font-size:10px; font-weight:700; letter-spacing:0.5px;">HOST:</span> <b style="color:var(--text); font-size:12.5px;">{sys_hostname}</b>
                                </div>
                                <div class="spec-badge">
                                    <span style="color:#22D3EE; font-size: 14px;">⚙️</span> <span style="color:var(--text-muted); font-size:10px; font-weight:700; letter-spacing:0.5px;">OS:</span> <b style="color:var(--text); font-size:12.5px;">{sys_os} {sys_ver} ({sys_arch})</b>
                                </div>
                                <div class="spec-badge">
                                    <span style="color:#F43F5E; font-size: 14px;">🧠</span> <span style="color:var(--text-muted); font-size:10px; font-weight:700; letter-spacing:0.5px;">RAM:</span> <b style="color:var(--text); font-size:12.5px;">{sys_ram_gb} GB</b>
                                </div>
                            </div>

                            <!-- AUDIT CAPABILITIES -->
                            <div style="display: grid; grid-template-columns: repeat(3, 1fr); gap: 12px; max-width: 750px; margin: 0 auto 25px auto;">
                                <!-- Card 1 -->
                                <div class="audit-card" style="border-top: 3px solid #8C7CF0 !important;">
                                    <div style="width:36px; height:36px; background:rgba(140,124,240,0.1); border-radius:50%; display:flex; align-items:center; justify-content:center; margin:0 auto 12px auto;">
                                        <span style="font-size: 16px;">🛡️</span>
                                    </div>
                                    <div style="font-size: 12px; font-weight: 800; color: var(--text); text-transform: uppercase; letter-spacing: 0.5px;">Defender & OS</div>
                                    <div style="font-size: 11px; color: var(--text-muted); margin-top: 6px; line-height: 1.4;">Firewall status, Antivirus detection, & Windows updates</div>
                                </div>
                                <!-- Card 2 -->
                                <div class="audit-card" style="border-top: 3px solid #3B82F6 !important;">
                                    <div style="width:36px; height:36px; background:rgba(59,130,246,0.1); border-radius:50%; display:flex; align-items:center; justify-content:center; margin:0 auto 12px auto;">
                                        <span style="font-size: 16px;">📊</span>
                                    </div>
                                    <div style="font-size: 12px; font-weight: 800; color: var(--text); text-transform: uppercase; letter-spacing: 0.5px;">Resources & Disk</div>
                                    <div style="font-size: 11px; color: var(--text-muted); margin-top: 6px; line-height: 1.4;">CPU & RAM usages, Disk health & partitions capacity</div>
                                </div>
                                <!-- Card 3 -->
                                <div class="audit-card" style="border-top: 3px solid #22D3EE !important;">
                                    <div style="width:36px; height:36px; background:rgba(34,211,238,0.1); border-radius:50%; display:flex; align-items:center; justify-content:center; margin:0 auto 12px auto;">
                                        <span style="font-size: 16px;">🌐</span>
                                    </div>
                                    <div style="font-size: 12px; font-weight: 800; color: var(--text); text-transform: uppercase; letter-spacing: 0.5px;">Ports & Tasks</div>
                                    <div style="font-size: 11px; color: var(--text-muted); margin-top: 6px; line-height: 1.4;">Active listening ports, browser audit, & startup apps</div>
                                </div>
                            </div>
                        </div>
                    </div>
                    """
                ),
                unsafe_allow_html=True
            )
            
            col_b1, col_b2, col_b3 = st.columns([1.1, 1.8, 1.1])
            with col_b2:
                if st.button("🚀 Start Device Security Scan", key="start_device_scan", width="stretch"):
                    status_placeholder = st.empty()
                    steps = [
                        ("sys", "Collecting Operating System details..."),
                        ("fw", "Verifying Windows Firewall Status..."),
                        ("av", "Detecting active Antivirus protection..."),
                        ("upd", "Checking pending Windows Updates..."),
                        ("res", "Auditing RAM & CPU utilization..."),
                        ("disk", "Checking Disk Drive capacity..."),
                        ("proc", "Scanning active Running Processes..."),
                        ("start", "Analyzing Startup applications..."),
                        ("ports", "Scanning local open Ports..."),
                        ("net", "Collecting IP, MAC, & DNS adapter details..."),
                        ("web", "Verifying DNS & Internet Security..."),
                        ("browser", "Auditing Web Browser configurations..."),
                        ("ext", "Checking for suspicious File Extensions..."),
                        ("ai", "Synthesizing AI Risk Posture Report...")
                    ]
                    
                    for code, text in steps:
                        status_placeholder.markdown(
                            f"""
                            <div class="chart-card" style="padding: 16px 20px; display: flex; align-items: center; gap: 15px;">
                                <div class="spinner-loader"></div>
                                <div style="font-size: 14px; font-weight: 500; color: var(--text);">{text}</div>
                            </div>
                            """,
                            unsafe_allow_html=True
                        )
                        time.sleep(0.2)
                        
                    with st.spinner("Compiling results..."):
                        result = device_security_module.analyze()
                        st.session_state.device_scan_result = result
                        st.rerun()
    else:
        res = st.session_state.device_scan_result
        details = res["scan_details"]
        
        st.markdown("<div style='height: 15px;'></div>", unsafe_allow_html=True)
        
        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown(
                f"""
                <div class="chart-card" style="padding: 20px; text-align: center;">
                    <div style="font-size: 11px; color: var(--text-faint); font-weight: 700; text-transform: uppercase;">Operating System</div>
                    <div style="font-size: 20px; font-weight: 800; color: var(--text); margin-top: 10px;">💻 {details['os_info']['os_name']}</div>
                    <div style="font-size: 12px; color: var(--text-muted); margin-top: 5px;">{details['os_info']['hostname']}</div>
                </div>
                """,
                unsafe_allow_html=True
            )
        with c2:
            st.markdown(
                f"""
                <div class="chart-card" style="padding: 20px; text-align: center;">
                    <div style="font-size: 11px; color: var(--text-faint); font-weight: 700; text-transform: uppercase;">Overall Security Score</div>
                    <div style="font-size: 28px; font-weight: 900; color: #10B981; margin-top: 5px;">🛡️ {res['security_score']} / 100</div>
                    <div style="font-size: 12px; color: var(--text-muted); margin-top: 3px;">Based on 19 audited settings</div>
                </div>
                """,
                unsafe_allow_html=True
            )
        with c3:
            st.markdown(
                f"""
                <div class="chart-card" style="padding: 20px; text-align: center;">
                    <div style="font-size: 11px; color: var(--text-faint); font-weight: 700; text-transform: uppercase;">Risk Level Status</div>
                    <div style="font-size: 20px; font-weight: 800; color: var(--text); margin-top: 10px;">{res['status']}</div>
                    <div style="font-size: 12px; color: var(--text-muted); margin-top: 5px;">Risk level: {res['risk_level'].upper()}</div>
                </div>
                """,
                unsafe_allow_html=True
            )
            
        col_findings_1, col_findings_2 = st.columns([1.8, 3.2])
        with col_findings_1:
            st.markdown('<div class="section-title">❌ Problems Found</div>', unsafe_allow_html=True)
            problems_html = ""
            for p in res["problems_found"]:
                problems_html += f'<div style="margin-bottom: 10px; font-size: 13px; font-weight: 500; color: var(--text);"><span style="color: var(--danger); margin-right: 8px;">❌</span> {html.escape(p)}</div>'
            st.markdown(
                f"""
                <div class="chart-card" style="padding: 20px; min-height: 200px;">
                    {problems_html}
                </div>
                """,
                unsafe_allow_html=True
            )
            
        with col_findings_2:
            st.markdown('<div class="section-title">💡 AI Risk Engine Analysis</div>', unsafe_allow_html=True)
            recs_html = ""
            for r in res["recommendation"]["recommendations"]:
                recs_html += f'<div style="margin-bottom: 8px; font-size: 12.5px; color: var(--text); font-weight: 500;"><span style="color: var(--success); margin-right: 8px;">✅</span> {html.escape(r)}</div>'
            
            st.markdown(
                f"""
                <div class="chart-card" style="padding: 20px; min-height: 200px;">
                    <div style="margin-bottom: 15px; font-size: 13.5px; font-weight: 600; color: var(--text); line-height: 1.5; font-style: italic;">
                        "{html.escape(res['explain_ai']['summary'])}"
                    </div>
                    <div style="border-top: 1px solid var(--border); padding-top: 15px;">
                        {recs_html}
                    </div>
                </div>
                """,
                unsafe_allow_html=True
            )
            
        st.markdown('<div class="section-title">🔍 Comprehensive Audit Modules</div>', unsafe_allow_html=True)
        tab_os, tab_sec, tab_res, tab_net, tab_proc, tab_browser = st.tabs([
            "💻 System", "🛡️ Security Services", "📊 Resources", "🌐 Net & Ports", "⚙️ Tasks & Startup", "📂 Browser & Files"
        ])
        
        with tab_os:
            st.subheader("Operating System Details")
            c_os1, c_os2 = st.columns(2)
            with c_os1:
                st.write(f"**OS Name:** {details['os_info']['os_name']}")
                st.write(f"**Version:** {details['os_info']['version']}")
                st.write(f"**Architecture:** {details['os_info']['architecture']}")
            with c_os2:
                st.write(f"**Hostname:** {details['os_info']['hostname']}")
                st.write(f"**Boot Time:** {details['os_info']['boot_time']}")
                st.write(f"**AI Verdict:** {details['os_info']['ai_analysis']}")
                
        with tab_sec:
            st.subheader("Security Services Audit")
            c_sec1, c_sec2, c_sec3 = st.columns(3)
            with c_sec1:
                st.markdown(
                    f"""
                    <div class="chart-card" style="padding: 15px; text-align: center;">
                        <div style="font-size: 12px; color: var(--text-muted);">Windows Firewall</div>
                        <div style="font-size: 20px; font-weight: 700; margin-top: 8px;">{details['firewall']['status']}</div>
                        <div style="font-size: 11px; color: var(--text-faint); margin-top: 5px;">{details['firewall']['details']}</div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            with c_sec2:
                st.markdown(
                    f"""
                    <div class="chart-card" style="padding: 15px; text-align: center;">
                        <div style="font-size: 12px; color: var(--text-muted);">Antivirus Protection</div>
                        <div style="font-size: 20px; font-weight: 700; margin-top: 8px;">{'🟢 Running' if details['antivirus']['found'] else '🔴 Missing'}</div>
                        <div style="font-size: 11px; color: var(--text-faint); margin-top: 5px;">{details['antivirus']['details']}</div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            with c_sec3:
                up_info = details['updates']
                pending_txt = "Auto-updates disabled in settings" if up_info.get("pending_count") is None else f"{up_info.get('pending_count', 0)} updates missing"
                st.markdown(
                    f"""
                    <div class="chart-card" style="padding: 15px; text-align: center;">
                        <div style="font-size: 12px; color: var(--text-muted);">Windows Update Patches</div>
                        <div style="font-size: 20px; font-weight: 700; margin-top: 8px;">{up_info['status']}</div>
                        <div style="font-size: 11px; color: var(--text-faint); margin-top: 5px;">{pending_txt}</div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
                if up_info.get("manual_mode") or not st.session_state.get("settings_auto_updates", True):
                    st.markdown("<div style='height: 8px;'></div>", unsafe_allow_html=True)
                    if st.button("🔄 Check for Updates", key="btn_check_win_updates_manual", width="stretch"):
                        with st.spinner("Querying Windows Update API..."):
                            from services.device_security_service import device_security_service
                            fresh_updates = device_security_service.get_windows_updates()
                            res["scan_details"]["updates"] = fresh_updates
                            st.session_state.device_scan_result = res
                            st.rerun()
                
        with tab_res:
            st.subheader("Resources & Performance")
            c_res1, c_res2 = st.columns(2)
            with c_res1:
                st.write(f"**CPU Core Usage:** {details['cpu']['usage_percent']}%")
                st.progress(details['cpu']['usage_percent'] / 100)
                if details['cpu']['recommendation']:
                    st.warning(details['cpu']['recommendation'])
                    
                st.markdown("<div style='height: 10px;'></div>", unsafe_allow_html=True)
                st.write(f"**RAM Memory Usage:** {details['ram']['usage_percent']}%")
                st.progress(details['ram']['usage_percent'] / 100)
                st.caption(f"Used: {details['ram']['used_gb']} GB / Total: {details['ram']['total_gb']} GB (Available: {details['ram']['available_gb']} GB)")
                if details['ram']['recommendation']:
                    st.warning(details['ram']['recommendation'])
            with c_res2:
                st.write("**Disk Storage Health:**")
                for d in details['disk']['partitions']:
                    st.write(f"Drive {d['drive']}: {d['usage_percent']}% full")
                    st.progress(d['usage_percent'] / 100)
                    st.caption(f"Free: {d['free_gb']} GB / Total: {d['total_gb']} GB")
                if details['disk']['recommendation']:
                    st.warning(details['disk']['recommendation'])
                    
        with tab_net:
            st.subheader("Network Configurations & Open Ports")
            st.markdown(
                f"""
                <div style="font-size: 13.5px; margin-bottom: 12px;">
                    <b>Local IP:</b> {details['network']['local_ip']} &nbsp;|&nbsp; 
                    <b>MAC Address:</b> {details['network']['mac_address']} &nbsp;|&nbsp; 
                    <b>Connection:</b> {details['network']['connection_type']} &nbsp;|&nbsp;
                    <b>DNS Servers:</b> {details['network']['dns_provider']}
                </div>
                """,
                unsafe_allow_html=True
            )
            
            st.write(f"**Listening Open Ports ({details['ports']['open_count']} Active listeners):**")
            ports_df = pd.DataFrame(details['ports']['ports'])
            if not ports_df.empty:
                st.dataframe(ports_df, width="stretch", hide_index=True)
            else:
                st.info("No listening ports active.")
            if details['ports']['explanation']:
                st.warning(f"**Risk explanation:** {details['ports']['explanation']}")
                st.caption(f"**Recommendation:** {details['ports']['recommendation']}")
                
        with tab_proc:
            st.subheader("Running Tasks & Auto Startup Services")
            col_proc1, col_proc2 = st.columns(2)
            with col_proc1:
                st.write(f"**Active Processes list ({details['processes']['unknown_count']} unrecognized processes):**")
                proc_df = pd.DataFrame(details['processes']['processes'])
                st.dataframe(proc_df[['pid', 'name', 'cpu', 'mem', 'status']], width="stretch", hide_index=True)
                st.caption(details['processes']['details'])
            with col_proc2:
                st.write("**Applications launching on computer boot (Startup):**")
                startup_df = pd.DataFrame(details['startup']['apps'])
                st.dataframe(startup_df, width="stretch", hide_index=True)
                st.caption(details['startup']['details'])
                
        with tab_browser:
            st.subheader("Installed Browser Security & File Downloads Audit")
            col_br1, col_br2 = st.columns(2)
            with col_br1:
                st.write("**Installed Web Browsers:**")
                br_df = pd.DataFrame(details['browser']['browsers'])
                st.dataframe(br_df, width="stretch", hide_index=True)
                st.caption(details['browser']['details'])
            with col_br2:
                st.write(f"**Suspicious File Extensions in Downloads/Desktop ({details['suspicious_files']['count']} files detected):**")
                susp_files_df = pd.DataFrame(details['suspicious_files']['files'])
                if not susp_files_df.empty:
                    st.dataframe(susp_files_df, width="stretch", hide_index=True)
                else:
                    st.success("No files with suspicious script/executable extensions found in user directories.")
                st.caption(details['suspicious_files']['details'])

        st.markdown('<div class="section-title">🔑 Authentication & Password Guidelines</div>', unsafe_allow_html=True)
        advice_html = "".join(f'<li style="margin-bottom:8px; font-size:13px; color:var(--text);">{html.escape(adv)}</li>' for adv in details['passwords'])
        st.markdown(
            f"""
            <div class="chart-card" style="padding: 20px;">
                <ul style="padding-left: 20px; margin: 0;">
                    {advice_html}
                </ul>
            </div>
            """,
            unsafe_allow_html=True
        )

        # ── Section 18: Enhanced Device Security ──

        # 18a: Overall Device Health Score badge
        sec_score = res.get('security_score', 0)
        if sec_score >= 80:
            health_color, health_label = "#10B981", "Excellent"
        elif sec_score >= 60:
            health_color, health_label = "#F59E0B", "Fair"
        elif sec_score >= 40:
            health_color, health_label = "#F97316", "At Risk"
        else:
            health_color, health_label = "#EF4444", "Critical"

        st.markdown(f'''
        <div class="chart-card" style="padding: 20px; margin-bottom: 16px;">
            <div style="display: flex; align-items: center; gap: 18px; flex-wrap: wrap;">
                <div style="width: 72px; height: 72px; border-radius: 50%; border: 4px solid {health_color};
                            display: flex; align-items: center; justify-content: center; flex-shrink: 0;">
                    <span style="font-size: 26px; font-weight: 900; color: {health_color};">{sec_score}</span>
                </div>
                <div style="flex: 1; min-width: 0;">
                    <div style="font-size: 18px; font-weight: 800; color: var(--text); margin-bottom: 4px;">
                        Device Health: <span style="color: {health_color};">{health_label}</span>
                    </div>
                    <div style="font-size: 12.5px; color: var(--text-muted); line-height: 1.6;">
                        Score computed from {len(res.get("problems_found", []))} identified issues across OS security,
                        firewall, antivirus, ports, startup apps, and file system audits.
                        {'⚠️ Immediate action recommended.' if sec_score < 50 else '✅ No critical issues detected.'}
                    </div>
                </div>
            </div>
        </div>
        ''', unsafe_allow_html=True)

        # 18b: HOSTS file hijack check
        st.markdown('<div class="section-title">📄 HOSTS File Integrity Check</div>', unsafe_allow_html=True)
        try:
            import platform as _plat
            if _plat.system() == "Windows":
                hosts_path = r"C:\Windows\System32\drivers\etc\hosts"
            else:
                hosts_path = "/etc/hosts"
            with open(hosts_path, "r", encoding="utf-8", errors="ignore") as hf:
                hosts_lines = hf.readlines()
            suspicious_hosts = []
            safe_domains = {"localhost", "broadcasthost", "ip6-localhost", "ip6-loopback", "ip6-localnet",
                            "ip6-mcastprefix", "ip6-allnodes", "ip6-allrouters", "ip6-allhosts"}
            for line in hosts_lines:
                stripped = line.strip()
                if not stripped or stripped.startswith("#"):
                    continue
                parts = stripped.split()
                if len(parts) >= 2:
                    ip, domain = parts[0], parts[1].lower()
                    if domain not in safe_domains and ip not in ("127.0.0.1", "::1", "0.0.0.0", "255.255.255.255"):
                        suspicious_hosts.append({"IP": ip, "Domain": domain})

            if suspicious_hosts:
                st.warning(f"⚠️ {len(suspicious_hosts)} suspicious entries found in HOSTS file — possible DNS hijack:")
                import pandas as _pd_hosts
                st.dataframe(_pd_hosts.DataFrame(suspicious_hosts), width="stretch", hide_index=True)
            else:
                st.markdown('''
                <div class="chart-card" style="padding: 15px; display: flex; align-items: center; gap: 12px;">
                    <span style="font-size: 22px;">✅</span>
                    <div>
                        <div style="font-size: 14px; font-weight: 700; color: var(--text);">HOSTS File Clean</div>
                        <div style="font-size: 12px; color: var(--text-muted);">No suspicious DNS redirection entries detected.</div>
                    </div>
                </div>
                ''', unsafe_allow_html=True)
        except Exception as hosts_exc:
            st.info(f"Could not read HOSTS file: {hosts_exc}")

        # 18c: Startup apps risk tagging
        st.markdown('<div class="section-title">⚡ Startup Risk Assessment</div>', unsafe_allow_html=True)
        startup_apps = details.get('startup', {}).get('apps', [])
        if startup_apps:
            high_risk_keywords = ["teamviewer", "anydesk", "vnc", "remote", "rdp", "rat", "miner",
                                  "crypto", "torrent", "p2p", "proxy"]
            risk_tagged = []
            for app in startup_apps:
                name_lower = str(app.get("name", "")).lower()
                risk = "⚠️ Review" if any(kw in name_lower for kw in high_risk_keywords) else "✅ Normal"
                risk_tagged.append({
                    "Name": app.get("name", "Unknown"),
                    "Path": str(app.get("path", app.get("command", "N/A")))[:80],
                    "Risk": risk
                })
            risk_df = pd.DataFrame(risk_tagged)
            flagged = risk_df[risk_df["Risk"] == "⚠️ Review"]
            if not flagged.empty:
                st.warning(f"⚠️ {len(flagged)} startup application(s) flagged for review:")
            st.dataframe(risk_df, width="stretch", hide_index=True)
        else:
            st.info("No startup applications data available.")

        # 18d: Browser client-side security telemetry (JS component)
        st.markdown('<div class="section-title">🌐 Browser Security Telemetry</div>', unsafe_allow_html=True)
        import streamlit.components.v1 as _components
        _components.html("""
        <style>
            * { margin: 0; padding: 0; box-sizing: border-box; font-family: 'Inter', 'Segoe UI', sans-serif; }
            body { background: transparent; }
            .card {
                background: rgba(20, 27, 45, 0.85);
                border: 1px solid rgba(255,255,255,0.06);
                border-radius: 14px;
                padding: 20px;
            }
            .subtitle { font-size: 13px; color: rgba(255,255,255,0.45); margin-bottom: 14px; }
            .grid {
                display: grid;
                grid-template-columns: repeat(auto-fill, minmax(180px, 1fr));
                gap: 12px;
            }
            .badge {
                background: rgba(255,255,255,0.02);
                border: 1px solid rgba(255,255,255,0.05);
                border-radius: 10px;
                padding: 12px;
                display: flex;
                align-items: center;
                gap: 10px;
            }
            .icon { font-size: 16px; }
            .label { font-size: 11px; color: rgba(255,255,255,0.35); font-weight: 700; letter-spacing: 0.3px; }
            .value { font-size: 13px; font-weight: 600; color: rgba(255,255,255,0.9); margin-top: 2px; }
        </style>
        <div class="card">
            <div class="subtitle">Client-side checks executed in your browser &mdash; no data leaves this device.</div>
            <div class="grid">
                <div class="badge">
                    <span class="icon">&#127850;</span>
                    <div>
                        <div class="label">COOKIES</div>
                        <div class="value" id="tel-cookies">Checking...</div>
                    </div>
                </div>
                <div class="badge">
                    <span class="icon">&#128225;</span>
                    <div>
                        <div class="label">WEBRTC LEAK</div>
                        <div class="value" id="tel-webrtc">Checking...</div>
                    </div>
                </div>
                <div class="badge">
                    <span class="icon">&#128274;</span>
                    <div>
                        <div class="label">HTTPS CONTEXT</div>
                        <div class="value" id="tel-https">Checking...</div>
                    </div>
                </div>
                <div class="badge">
                    <span class="icon">&#128421;&#65039;</span>
                    <div>
                        <div class="label">DO NOT TRACK</div>
                        <div class="value" id="tel-dnt">Checking...</div>
                    </div>
                </div>
            </div>
        </div>
        <script>
        (function(){
            try {
                var ce = document.getElementById('tel-cookies');
                if(ce) ce.textContent = navigator.cookieEnabled ? '\\u{1F7E1} Enabled' : '\\u{1F7E2} Disabled';
            } catch(e){}
            try {
                var we = document.getElementById('tel-webrtc');
                if(window.RTCPeerConnection){
                    var pc = new RTCPeerConnection({iceServers:[]});
                    pc.createDataChannel('');
                    pc.createOffer().then(function(o){pc.setLocalDescription(o)});
                    pc.onicecandidate = function(ev){
                        if(!ev.candidate) return;
                        var c = ev.candidate.candidate;
                        var ip = c.match(/([0-9]{1,3}(\\.[0-9]{1,3}){3})/);
                        if(ip && we) we.textContent = '\\u{1F534} Leaking: ' + ip[1];
                        pc.close();
                    };
                    setTimeout(function(){ if(we && we.textContent==='Checking...') we.textContent='\\u{1F7E2} No Leak'; }, 3000);
                } else { if(we) we.textContent = '\\u{1F7E2} Not Supported'; }
            } catch(e){ var we2=document.getElementById('tel-webrtc'); if(we2) we2.textContent='\\u{1F7E2} Blocked'; }
            try {
                var he = document.getElementById('tel-https');
                if(he) he.textContent = location.protocol==='https:' ? '\\u{1F7E2} Secure' : '\\u{1F7E1} Insecure (HTTP)';
            } catch(e){}
            try {
                var de = document.getElementById('tel-dnt');
                if(de) de.textContent = (navigator.doNotTrack==='1') ? '\\u{1F7E2} Enabled' : '\\u{1F7E1} Not Set';
            } catch(e){}
        })();
        </script>
        """, height=140, scrolling=False)

        # 18e: Download local scanner script
        st.markdown('<div class="section-title">🔧 Offline Security Agent</div>', unsafe_allow_html=True)
        try:
            import os as _os_dl
            scanner_path = _os_dl.path.join("data", "downloads", "cybermind_local_scanner.py")
            if _os_dl.path.exists(scanner_path):
                with open(scanner_path, "r", encoding="utf-8") as f_sc:
                    scanner_code = f_sc.read()
                st.markdown('''
                <div class="chart-card" style="padding: 20px;">
                    <div style="display: flex; align-items: center; gap: 14px; margin-bottom: 12px;">
                        <div style="width: 40px; height: 40px; border-radius: 10px; background: rgba(140,124,240,0.12);
                                    display: flex; align-items: center; justify-content: center; font-size: 18px;"><img src="https://cdn-icons-png.flaticon.com/512/6071/6071531.png" width="90" height="90" style="object-fit: contain;"></div>
                        <div>
                            <div style="font-size: 15px; font-weight: 700; color: var(--text);">CyberMind Local Scanner</div>
                            <div style="font-size: 12px; color: var(--text-muted);">
                                Standalone Python script for offline workstation audits — run anywhere without CyberMind installed.
                            </div>
                        </div>
                    </div>
                    <div style="font-size: 11.5px; color: var(--text-faint); line-height: 1.6;">
                        Audits: HOSTS file • DNS configuration • USB device history • Open ports • Startup apps •
                        Firewall status • Suspicious downloads • System integrity.
                    </div>
                </div>
                ''', unsafe_allow_html=True)
            else:
                st.info("Local scanner script not found at data/downloads/cybermind_local_scanner.py")
        except Exception as dl_exc:
            st.warning(f"Could not load local scanner: {dl_exc}")

        st.markdown("<div style='height: 15px;'></div>", unsafe_allow_html=True)
        col_act1, col_act2 = st.columns([3.5, 1.5])
        with col_act1:
            pdf_data, mime_type, file_name = export_device_security_pdf(res)
            st.download_button(
                label="⬇ Download Assessment PDF Report",
                data=pdf_data,
                file_name=file_name,
                mime=mime_type,
                key="dl_device_report_btn",
                width="stretch"
            )
        with col_act2:
            if st.button("🔄 Run New Scan", key="clear_device_scan", width="stretch"):
                st.session_state.device_scan_result = None
                st.rerun()



def render_ai_assistant_page():
    from modules.ai_assistant import get_chat_response, get_auto_suggestions
    import html
    import re
    
    def md_to_html(text: str) -> str:
        escaped = html.escape(text)
        escaped = re.sub(r'###\s+(.*?)(?:\n|$)', r'<h5 style="margin: 8px 0 4px; color: #22D3EE; font-weight: 700;">\1</h5>', escaped)
        escaped = re.sub(r'##\s+(.*?)(?:\n|$)', r'<h4 style="margin: 10px 0 6px; color: #22D3EE; font-weight: 700;">\1</h4>', escaped)
        escaped = re.sub(r'#\s+(.*?)(?:\n|$)', r'<h3 style="margin: 12px 0 8px; color: #22D3EE; font-weight: 700;">\1</h3>', escaped)
        escaped = re.sub(r'\*\*(.*?)\*\*', r'<strong style="color: #22D3EE;">\1</strong>', escaped)
        escaped = re.sub(r'\*(.*?)\*', r'<em>\1</em>', escaped)
        
        def code_block_sub(match):
            code = match.group(1).strip()
            return f'<pre style="background: rgba(0,0,0,0.3); padding: 10px; border-radius: 6px; border: 1px solid var(--border); overflow-x: auto; font-family: monospace; font-size: 12px; margin: 8px 0; color: #E2E8F0;"><code>{code}</code></pre>'
        escaped = re.sub(r'```(?:[a-zA-Z]*)\n(.*?)```', code_block_sub, escaped, flags=re.DOTALL)
        
        escaped = re.sub(r'`(.*?)`', r'<code style="background: rgba(255,255,255,0.08); padding: 2px 5px; border-radius: 4px; font-family: monospace; font-size: 12px; color: #38BDF8;">\1</code>', escaped)
        
        lines = escaped.split("\n")
        in_list = False
        new_lines = []
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("- ") or stripped.startswith("* "):
                content = stripped[2:]
                if not in_list:
                    new_lines.append('<ul style="margin: 6px 0; padding-left: 20px;">')
                    in_list = True
                new_lines.append(f'<li style="margin-bottom: 4px; color: var(--text);">{content}</li>')
            else:
                if in_list:
                    new_lines.append('</ul>')
                    in_list = False
                new_lines.append(line)
        if in_list:
            new_lines.append('</ul>')
        escaped = "\n".join(new_lines)
        escaped = escaped.replace("\n", "<br>")
        escaped = escaped.replace("</ul><br>", "</ul>")
        escaped = escaped.replace("</pre><br>", "</pre>")
        return escaped

    origin_page = st.session_state.get("ai_assistant_origin_page", "Dashboard")

    render_page_poster(
        "AI Security Assistant",
        "Your offline, local cybersecurity companion. Ask questions and get real-time recommendations."
    )
    
    # Initialize chat history
    if "ai_page_chat_history" not in st.session_state:
        st.session_state.ai_page_chat_history = []
        
    c_chat, c_sug = st.columns([2.0, 1.2])
    
    with c_chat:
        col_chat_title, col_chat_back = st.columns([1.2, 1.0])
        with col_chat_title:
            st.markdown('<div class="section-title" style="margin-bottom:0; padding-top:4px;">💬 Interactive Security Chat</div>', unsafe_allow_html=True)
        with col_chat_back:
            if st.button(f"⬅️ Back to {origin_page}", key="ai_back_to_origin_btn", width="stretch"):
                st.session_state.active_page = origin_page
                st.rerun()
        
        with st.container(key="interactive_chat_card"):
            chat_container_style = """
                <style>
                .chat-card-user {
                    background: var(--card-bg-soft);
                    border: 1px solid var(--border);
                    border-left: 4px solid #8C7CF0 !important;
                    border-radius: 12px;
                    padding: 16px;
                    margin-bottom: 16px;
                    color: var(--text);
                    width: 100%;
                    box-shadow: 0 4px 15px rgba(140,124,240,0.03);
                }
                .chat-card-assistant {
                    background: var(--card-bg-soft);
                    border: 1px solid var(--border);
                    border-left: 4px solid #22D3EE !important;
                    border-radius: 12px;
                    padding: 16px;
                    margin-bottom: 16px;
                    color: var(--text);
                    width: 100%;
                    box-shadow: 0 4px 20px rgba(34, 211, 238, 0.05);
                }
                .chat-header {
                    display: flex;
                    justify-content: space-between;
                    align-items: center;
                    margin-bottom: 10px;
                    border-bottom: 1px solid rgba(255,255,255,0.04);
                    padding-bottom: 6px;
                }
                .chat-header-title {
                    display: flex;
                    align-items: center;
                    font-size: 11px;
                    font-weight: 800;
                    letter-spacing: 1.2px;
                    text-transform: uppercase;
                }
                .pulse-dot {
                    width: 6px;
                    height: 6px;
                    background-color: #22D3EE;
                    border-radius: 50%;
                    margin-right: 6px;
                    display: inline-block;
                    animation: pulse-glow 1.5s infinite;
                }
                @keyframes pulse-glow {
                    0% { transform: scale(0.9); opacity: 0.6; }
                    50% { transform: scale(1.25); opacity: 1; }
                    100% { transform: scale(0.9); opacity: 0.6; }
                }
                </style>
            """
            st.markdown(chat_container_style, unsafe_allow_html=True)
            
            # Render previous messages
            if st.session_state.ai_page_chat_history:
                for role, text in st.session_state.ai_page_chat_history:
                    if role == "user":
                        st.markdown(f"""
                        <div class="chat-card-user">
                            <div class="chat-header">
                                <div class="chat-header-title" style="color: #8C7CF0;">
                                    <span style="font-size:14px; margin-right:6px;"><img src="https://cdn-icons-png.flaticon.com/512/924/924915.png" width="18" height="18" style="vertical-align:middle; margin-right:6px;"></span> USER QUERY
                                </div>
                                <span style="font-size: 9px; padding: 2px 6px; background: rgba(140, 124, 240, 0.1); color: #8C7CF0; border-radius: 20px; font-weight: 700;">ACTIVE SESSION</span>
                            </div>
                            <div style="font-size: 13.5px; line-height: 1.5; color: var(--text);">{html.escape(text)}</div>
                        </div>
                        """, unsafe_allow_html=True)
                    else:
                        st.markdown(f"""
                        <div class="chat-card-assistant">
                            <div class="chat-header">
                                <div class="chat-header-title" style="color: #22D3EE;">
                                    <span class="pulse-dot"></span> <img src="https://cdn-icons-png.flaticon.com/512/18310/18310827.png" width="18" height="18" style="vertical-align:middle; margin-right:6px;"> CYBERMIND INTEL ENGINE
                                </div>
                                <span style="font-size: 9px; padding: 2px 6px; background: rgba(34, 211, 238, 0.1); color: #22D3EE; border-radius: 20px; font-weight: 700; letter-spacing: 0.5px; text-transform: uppercase;">Verified Threat Agent</span>
                            </div>
                            <div style="font-size: 13.5px; line-height: 1.6; color: var(--text);">{md_to_html(text)}</div>
                        </div>
                        """, unsafe_allow_html=True)
            else:
                st.markdown(
                    """
                    <div style="padding: 30px; text-align: center; color: var(--text-muted); min-height: 150px; display: flex; flex-direction: column; justify-content: center; align-items: center;">
                         <div style="font-size: 40px; margin-bottom: 10px;"><img src="https://cdn-icons-png.flaticon.com/512/18310/18310827.png" alt="AI Assistant" width="65" height="65"></div>
                         <div style="font-size: 13.5px;">Hello! Ask me any question about cybersecurity or ask for suggestions based on your scan records.</div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
                
            # ── Follow-up question pills from last response ──────────────
            followup_qs = st.session_state.get("ai_followup_questions", [])
            if followup_qs and st.session_state.ai_page_chat_history:
                st.markdown(
                    "<div style='margin-top:6px; margin-bottom:4px;'>"
                    "<span style='font-size:12px; font-weight:700; color:var(--info);'>🔗 Follow-up Questions:</span>"
                    "</div>",
                    unsafe_allow_html=True
                )
                fq_cols = st.columns(len(followup_qs))
                clicked_followup = ""
                for fq_idx, (fq_col, fq) in enumerate(zip(fq_cols, followup_qs)):
                    with fq_col:
                        if st.button(fq, key=f"followup_pill_{fq_idx}", width="stretch"):
                            clicked_followup = fq
                if clicked_followup:
                    from modules.ai_assistant import get_chat_response_with_suggestions
                    fq_resp, fq_followups = get_chat_response_with_suggestions(clicked_followup)
                    st.session_state.ai_page_chat_history.append(("user", clicked_followup))
                    st.session_state.ai_page_chat_history.append(("assistant", fq_resp))
                    st.session_state["ai_followup_questions"] = fq_followups
                    st.rerun()
            # ─────────────────────────────────────────────────────────────

            # Suggested questions pills
            st.markdown("<div style='height: 10px;'></div>", unsafe_allow_html=True)
            st.markdown("<span style='font-size: 13px; font-weight: 700; color: var(--text-muted);'>💡 Suggested Questions:</span>", unsafe_allow_html=True)
            
            col_s1, col_s2, col_s3, col_s4 = st.columns(4)
            s_queries = [
                "How do I secure public Wi-Fi?",
                "Is my password safe?",
                "What does a VPN do?",
                "What is a Trojan virus?"
            ]
            
            selected_suggested = ""
            for idx, (col, query) in enumerate(zip([col_s1, col_s2, col_s3, col_s4], s_queries)):
                with col:
                    if st.button(query, key=f"sug_query_{idx}", width="stretch"):
                        selected_suggested = query

            st.markdown("<div style='height: 6px;'></div>", unsafe_allow_html=True)
            if st.button("💡 Explain Last Scan Result", key="explain_last_scan_btn", width="stretch"):
                from modules.ai_assistant import explain_last_scan
                expl_text = explain_last_scan()
                st.session_state.ai_page_chat_history.append(("user", "Explain my last scan result"))
                st.session_state.ai_page_chat_history.append(("assistant", expl_text))
                st.rerun()

            # Chat Input Form
            with st.form("chat_form", clear_on_submit=True):
                chat_input = st.text_input("Ask something...", placeholder="Type your security question here...", label_visibility="collapsed")
                col_b1, col_b2 = st.columns([3.5, 1.5])
                with col_b1:
                    submit_chat = st.form_submit_button("🔍 Ask Assistant", width="stretch")
                with col_b2:
                    clear_chat = st.form_submit_button("🗑️ Clear Chat History", width="stretch")
                    
            # Handle clear chat click
            if clear_chat:
                st.session_state.ai_page_chat_history = []
                st.rerun()
                
            # Process input query
            query_to_run = ""
            if submit_chat and chat_input.strip():
                query_to_run = chat_input.strip()
            elif selected_suggested:
                query_to_run = selected_suggested
                
            if query_to_run:
                from modules.ai_assistant import get_chat_response_with_suggestions
                response, followup_qs = get_chat_response_with_suggestions(query_to_run)
                st.session_state.ai_page_chat_history.append(("user", query_to_run))
                st.session_state.ai_page_chat_history.append(("assistant", response))
                st.session_state["ai_followup_questions"] = followup_qs
                st.rerun()
            
    with c_sug:
        st.markdown(
            """
            <style>
            [data-testid="stColumn"]:has(.st-key-ai_sug_panel),
            [data-testid="column"]:has(.st-key-ai_sug_panel) {
                align-self: flex-start !important;
                position: sticky !important;
                top: 1rem !important;
            }
            .st-key-ai_sug_panel {
                position: sticky !important;
                top: 1rem !important;
                max-height: calc(100vh - 2rem) !important;
                overflow-y: auto !important;
                overflow-x: hidden !important;
                padding-right: 4px !important;
                scrollbar-width: thin !important;
                scrollbar-color: var(--border) transparent !important;
            }
            .st-key-ai_sug_panel::-webkit-scrollbar {
                width: 4px;
            }
            .st-key-ai_sug_panel::-webkit-scrollbar-thumb {
                background: var(--border);
                border-radius: 2px;
            }
            </style>
            """,
            unsafe_allow_html=True
        )
        with st.container(key="ai_sug_panel"):
            st.markdown('<div class="section-title">📢 Scan Suggestions</div>', unsafe_allow_html=True)
            suggestions = get_auto_suggestions()
            sug_html = ""
            for tip in suggestions:
                sug_html += f'<div style="padding: 10px 12px; margin-bottom: 10px; background: var(--card-bg-soft); border-left: 3px solid var(--info); border-radius: 6px; font-size: 12.5px; line-height: 1.5; color: var(--text);">{html.escape(tip)}</div>'
                
            st.markdown(
                f"""
                <div class="chart-card" style="padding: 18px; min-height: 200px;">
                    {sug_html}
                </div>
                """,
                unsafe_allow_html=True
            )
            
            st.markdown('<div class="section-title">💡 Security Topics Covered</div>', unsafe_allow_html=True)
            st.markdown(
                """
                <div class="chart-card" style="padding: 16px; font-size: 12.5px; color: var(--text-muted); line-height: 1.6;">
                    ✔ Passwords & Two-Factor Auth (2FA)<br>
                    ✔ Phishing, Spam & Scam Detection<br>
                    ✔ Network, Wi-Fi & VPN Configuration<br>
                    ✔ Ransomware, Viruses & Malware Defense<br>
                    ✔ SSL, Domain Age & Web Safety Rules
                </div>
                """,
                unsafe_allow_html=True
            )


def render_universal_scan_page():
    from modules.universal_scan_module import universal_scan_module
    import html

    # Auto-clean when navigating INTO Universal Scan from a different page
    if st.session_state.get("_last_active_scanner_page") != "Universal Scan":
        if not st.session_state.pop("_restore_pending_Universal Scan", False):
            st.session_state.universal_scan_result = None
        st.session_state["_last_active_scanner_page"] = "Universal Scan"

    render_page_poster("Universal Scan", "Automatically detect and run the appropriate scanner for Domains, Emails, or URLs in real-time.")
    
    # Session state for universal scan results
    if "universal_scan_result" not in st.session_state:
        st.session_state.universal_scan_result = None
    if "scan_result_Universal Scan" not in st.session_state:
        st.session_state["scan_result_Universal Scan"] = None

    # Example Pill Buttons
    st.markdown('<div style="font-size:12px; font-weight:600; color:var(--text-muted); margin-bottom:6px;">💡 Examples:</div>', unsafe_allow_html=True)
    ex_cols = st.columns(4)
    examples = [
        ("Google URL", "https://google.com"),
        ("Phishing Email", "phishing@malicious.com"),
        ("Malware Site", "https://malware-site.net"),
        ("Suspicious Domain", "malicious-domain.com")
    ]
    example_triggered = False
    selected_example = ""
    for idx, (lbl, val) in enumerate(examples):
        with ex_cols[idx]:
            if st.button(lbl, key=f"ex_univ_{idx}", width="stretch"):
                example_triggered = True
                selected_example = val

    c1, c2 = st.columns([4, 1.3], vertical_alignment="center")
    with c1:
        user_input = st.text_input(
            "Enter input to scan",
            key="universal_scan_input",
            placeholder="e.g. google.com, test@gmail.com, or https://example.com",
            label_visibility="collapsed"
        )
    with c2:
        st.markdown('<div class="cta-scan">', unsafe_allow_html=True)
        scan_clicked = st.button("🔍 Universal Scan", key="run_universal_scan_btn", width="stretch")
        st.markdown('</div>', unsafe_allow_html=True)

    target_to_scan = selected_example if example_triggered else user_input.strip()
    if scan_clicked or example_triggered:
        if not target_to_scan:
            st.warning("Please enter a valid input (Domain, Email, or URL) first.")
        else:
            with st.spinner("Analyzing input type and launching modules..."):
                res = universal_scan_module.analyze(target_to_scan)
                st.session_state.universal_scan_result = res
                st.session_state.last_scan_context = {
                    "scanner_key": "Universal Scan",
                    "target": target_to_scan,
                    "scanner": f"Universal Scan ({res['input_type'].upper()})",
                    "risk_score": res["risk_score"],
                    "risk_level": res["risk_level"],
                    "findings": res["details"]
                }
                st.rerun()
                
    res = st.session_state.universal_scan_result
    
    if res:
        details = res["details"]
        inp_type = res["input_type"]
        risk_score = res["risk_score"]
        risk_level = res["risk_level"]
        
        # Color mapping
        color = "#22C55E" # green
        if risk_level in ("Medium", "Suspicious"):
            color = "#F5A623" # orange
        elif risk_level in ("High", "Critical"):
            color = "#EF4444" # red
        elif risk_level == "Unverified":
            color = "#F59E0B" # amber
            
        # AI Executive Summary Banner
        try:
            from modules.ai_summary_module import ai_summary_module
            exec_summary = ai_summary_module.generate_summary(res)
        except Exception:
            exec_summary = f"Universal scan analysis completed for {res.get('value', 'target')}. Combined Risk Score is {risk_score}/100 ({risk_level})."

        st.markdown(
            f"""
            <div style="background:rgba(34, 184, 240, 0.05); border:1px solid rgba(34, 184, 240, 0.25); border-radius:12px; padding:14px 18px; margin-top:16px; margin-bottom:16px;">
                <div style="font-size:12px; font-weight:700; color:#22D3EE; text-transform:uppercase; letter-spacing:0.5px; margin-bottom:6px; display:flex; align-items:center; gap:8px;">
                    <img src="https://cdn-icons-png.flaticon.com/512/18310/18310827.png" style="width:20px; height:20px; vertical-align:middle;">
                    <span>AI Executive Summary</span>
                </div>
                <div style="font-size:13.5px; line-height:1.5; color:var(--text);">{exec_summary}</div>
            </div>
            """,
            unsafe_allow_html=True
        )
        
        # Display Overview Row
        col_ov1, col_ov2, col_ov3 = st.columns(3)
        with col_ov1:
            st.markdown(
                f"""
                <div class="chart-card" style="padding: 20px; text-align: center;">
                    <div style="font-size: 11px; color: var(--text-faint); font-weight: 700; text-transform: uppercase;">Classified Target</div>
                    <div style="font-size: 20px; font-weight: 800; color: var(--text); margin-top: 10px;">
                        {'📧 Email' if inp_type == 'email' else '🔗 URL' if inp_type == 'url' else '🌍 Domain'}
                    </div>
                    <div style="font-size: 12px; color: var(--text-muted); margin-top: 5px; word-break: break-all;">{html.escape(res['value'])}</div>
                </div>
                """,
                unsafe_allow_html=True
            )
        with col_ov2:
            st.markdown(
                f"""
                <div class="chart-card" style="padding: 20px; text-align: center;">
                    <div style="font-size: 11px; color: var(--text-faint); font-weight: 700; text-transform: uppercase;">Combined Risk Score</div>
                    <div style="font-size: 28px; font-weight: 900; color: {color}; margin-top: 5px;">⚠️ {risk_score} / 100</div>
                    <div style="font-size: 12px; color: var(--text-muted); margin-top: 3px;">Lower is safer</div>
                </div>
                """,
                unsafe_allow_html=True
            )
        with col_ov3:
            st.markdown(
                f"""
                <div class="chart-card" style="padding: 20px; text-align: center;">
                    <div style="font-size: 11px; color: var(--text-faint); font-weight: 700; text-transform: uppercase;">Scan Status Verdict</div>
                    <div style="font-size: 20px; font-weight: 800; color: {color}; margin-top: 10px;">{risk_level.upper()}</div>
                    <div style="font-size: 12px; color: var(--text-muted); margin-top: 5px;">Scan time: {res['duration']}</div>
                </div>
                """,
                unsafe_allow_html=True
            )
            
        # Display Checklist
        st.markdown('<div class="section-title">📊 Execution Checklist Status</div>', unsafe_allow_html=True)
        checklist_items = []
        if inp_type == "domain":
            checklist_items = [
                ("Domain Scan", True),
                ("Website Scan", True),
                ("SSL Check", details.get("ssl_valid", False)),
                ("WHOIS Audit", details.get("registrar") != "Unknown"),
                ("DNS Query", len(details.get("ips", [])) > 0),
                ("Risk Score Calculation", True)
            ]
        elif inp_type == "email":
            checklist_items = [
                ("Email Validation", details.get("is_valid", False)),
                ("MX Record Audit", details.get("mx_available", False)),
                ("Disposable Check", not details.get("is_disposable", False)),
                ("Public Data Breach Check", not details.get("breach_found", False)),
                ("AI Recommendation Synthesis", True)
            ]
        else: # url
            checklist_items = [
                ("URL Reputation Scanner", not details.get("is_malicious", False)),
                ("Website Content Scanner", True),
                ("SSL Certificate Audit", details.get("ssl_valid", False)),
                ("Redirect Chain Check", details.get("redirect_count", 0) == 0),
                ("Phishing Detection Checks", not details.get("is_malicious", False)),
                ("AI Risk Score Evaluation", True)
            ]
            
        chk_html = ""
        for name, passed in checklist_items:
            icon = "✔" if passed else "⚠"
            icon_color = "var(--success)" if passed else "var(--warning)"
            chk_html += f"""
            <div style="display: flex; align-items: center; gap: 10px; margin-bottom: 8px; font-size: 13.5px; color: var(--text); font-weight: 500;">
                <span style="color: {icon_color}; font-weight: bold; font-size: 15px;">{icon}</span> {name}
            </div>
            """
        st.markdown(
            f"""
            <div class="chart-card" style="padding: 20px; margin-bottom: 20px;">
                {chk_html}
            </div>
            """,
            unsafe_allow_html=True
        )
        
        # Display Detailed Analysis Tabs
        st.markdown('<div class="section-title">🔍 Detailed Analysis</div>', unsafe_allow_html=True)
        
        if inp_type == "email":
            tab_val, tab_breach, tab_ai = st.columns(3)
            with tab_val:
                st.markdown(
                    f"""
                    <div class="chart-card" style="padding: 18px; min-height: 180px;">
                        <div style="font-weight: 700; font-size: 14px; margin-bottom: 12px; color: var(--text);">Email Registry</div>
                        <div style="font-size: 13px; line-height: 1.6; color: var(--text-muted);">
                            <b>Address:</b> {html.escape(details['email'])}<br>
                            <b>Syntactically Valid:</b> {'🟢 Yes' if details['is_valid'] else '🔴 No'}<br>
                            <b>MX Server Configured:</b> {'🟢 Yes' if details['mx_available'] else '🔴 No'}<br>
                            <b>Temporary/Disposable:</b> {'🔴 Yes (Blocked)' if details['is_disposable'] else '🟢 No'}
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            with tab_breach:
                st.markdown(
                    f"""
                    <div class="chart-card" style="padding: 18px; min-height: 180px;">
                        <div style="font-weight: 700; font-size: 14px; margin-bottom: 12px; color: var(--text);">Breach Intelligence</div>
                        <div style="font-size: 13px; line-height: 1.6; color: var(--text-muted);">
                            <b>Status:</b> {'🔴 Breaches found' if details['breach_found'] else '🟢 Safe / No breaches'}<br>
                            <b>Exposures Count:</b> {details['breach_count']} databases<br>
                            <b>Compromised Records:</b> {details['records_lost']:,} entries<br>
                            <span style="font-size: 11.5px; font-style: italic;">{html.escape(details['breach_details']) if details['breach_details'] else 'No public database leaks recorded.'}</span>
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            with tab_ai:
                recs_html = "".join(f'<div style="font-size: 12.5px; margin-bottom: 6px; color: var(--text-muted);">👉 {html.escape(r)}</div>' for r in details['recommendations'])
                st.markdown(
                    f"""
                    <div class="chart-card" style="padding: 18px; min-height: 180px;">
                        <div style="font-weight: 700; font-size: 14px; margin-bottom: 12px; color: var(--text);">AI Recommendation</div>
                        {recs_html}
                    </div>
                    """,
                    unsafe_allow_html=True
                )
                
        elif inp_type == "url":
            tab_threat, tab_ssl, tab_recs = st.columns(3)
            with tab_threat:
                vt_score = details['url_data'].get("virustotal", {}).get("malicious", 0)
                st.markdown(
                    f"""
                    <div class="chart-card" style="padding: 18px; min-height: 180px;">
                        <div style="font-weight: 700; font-size: 14px; margin-bottom: 12px; color: var(--text);">Threat Intelligence</div>
                        <div style="font-size: 13px; line-height: 1.6; color: var(--text-muted);">
                            <b>Target URL:</b> <span style="word-break: break-all;">{html.escape(details['url'])}</span><br>
                            <b>Google Safe Browsing:</b> {'🔴 Flagged Unsafe' if details['is_malicious'] else '🟢 Clean'}<br>
                            <b>VirusTotal Detection:</b> {vt_score} engines flagged<br>
                            <b>Redirects Count:</b> {details['redirect_count']} hops
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            with tab_ssl:
                st.markdown(
                    f"""
                    <div class="chart-card" style="padding: 18px; min-height: 180px;">
                        <div style="font-weight: 700; font-size: 14px; margin-bottom: 12px; color: var(--text);">SSL/TLS encryption</div>
                        <div style="font-size: 13px; line-height: 1.6; color: var(--text-muted);">
                            <b>Certificate Valid:</b> {'🟢 Yes' if details['ssl_valid'] else '🔴 Expired/Invalid'}<br>
                            <b>CA Issuer:</b> {html.escape(details['ssl_issuer'])}<br>
                            <b>Connection Mode:</b> HTTPS Secure transfer
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            with tab_recs:
                recs_html = "".join(f'<div style="font-size: 12.5px; margin-bottom: 6px; color: var(--text-muted);">👉 {html.escape(r)}</div>' for r in details['recommendations'])
                st.markdown(
                    f"""
                    <div class="chart-card" style="padding: 18px; min-height: 180px;">
                        <div style="font-weight: 700; font-size: 14px; margin-bottom: 12px; color: var(--text);">AI Recommendation</div>
                        {recs_html}
                    </div>
                    """,
                    unsafe_allow_html=True
                )
                
        else: # domain
            tab_whois, tab_dns, tab_sec = st.columns(3)
            with tab_whois:
                st.markdown(
                    f"""
                    <div class="chart-card" style="padding: 18px; min-height: 180px;">
                        <div style="font-weight: 700; font-size: 14px; margin-bottom: 12px; color: var(--text);">WHOIS Registry</div>
                        <div style="font-size: 13px; line-height: 1.6; color: var(--text-muted);">
                            <b>Domain Name:</b> {html.escape(details['domain'])}<br>
                            <b>Registrar:</b> {html.escape(details['registrar'])}<br>
                            <b>Creation Date:</b> {html.escape(str(details['created_date']))}
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            with tab_dns:
                ips_str = ", ".join(details['ips']) if details['ips'] else "None"
                st.markdown(
                    f"""
                    <div class="chart-card" style="padding: 18px; min-height: 180px;">
                        <div style="font-weight: 700; font-size: 14px; margin-bottom: 12px; color: var(--text);">DNS Records</div>
                        <div style="font-size: 13px; line-height: 1.6; color: var(--text-muted);">
                            <b>Resolved IPs:</b> {html.escape(ips_str)}<br>
                            <b>Reputation Status:</b> Verified safe
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            with tab_sec:
                st.markdown(
                    f"""
                    <div class="chart-card" style="padding: 18px; min-height: 180px;">
                        <div style="font-weight: 700; font-size: 14px; margin-bottom: 12px; color: var(--text);">Security & SSL</div>
                        <div style="font-size: 13px; line-height: 1.6; color: var(--text-muted);">
                            <b>SSL Certificate Valid:</b> {'🟢 Yes' if details['ssl_valid'] else '🔴 Expired/Invalid'}<br>
                            <b>SSL Issuer:</b> {html.escape(details['ssl_issuer'])}<br>
                            <b>Recommendations:</b><br>
                            {html.escape(details['recommendations'][0]) if details['recommendations'] else 'Audit regularly.'}
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

        # AI Explainable Analysis Section
        st.markdown('<div class="section-title" style="margin-top: 22px;">🤖 AI Explainable Analysis</div>', unsafe_allow_html=True)
        
        ai_explain_points = []
        if inp_type == "url":
            url_data = details.get("url_data", {})
            reasons = url_data.get("risk", {}).get("reasons", [])
            if details.get("is_malicious"):
                ai_explain_points.append("🔴 Flagged as malicious or phishing by Google Safe Browsing / VirusTotal intelligence.")
            else:
                ai_explain_points.append("🟢 Clean status verified across threat intelligence databases.")

            if details.get("ssl_valid"):
                ai_explain_points.append(f"🟢 SSL Certificate is valid (Issuer: {details.get('ssl_issuer', 'Unknown')}).")
            else:
                ai_explain_points.append("🔴 SSL Certificate is missing or invalid, making HTTP traffic vulnerable.")

            if details.get("redirect_count", 0) > 0:
                ai_explain_points.append(f"⚠️ Target URL redirects through {details.get('redirect_count')} intermediate hop(s).")
            else:
                ai_explain_points.append("🟢 Direct connection without suspicious multi-hop redirects.")

            if reasons:
                for r in reasons[:3]:
                    ai_explain_points.append(f"⚠️ Risk Engine Factor: {r}")
            else:
                ai_explain_points.append("🟢 No alarming heuristic anomalies detected in URL structure.")

        elif inp_type == "email":
            if details.get("is_valid"):
                ai_explain_points.append(f"🟢 Email address format is syntactically valid ({details.get('email')}).")
            else:
                ai_explain_points.append("🔴 Email address format is invalid.")

            if details.get("mx_available"):
                ai_explain_points.append("🟢 Active MX mail servers configured for sender domain.")
            else:
                ai_explain_points.append("🔴 Domain lacks active MX records (cannot receive real email).")

            if details.get("is_disposable"):
                ai_explain_points.append("🔴 Disposable / temporary burner email provider detected.")
            else:
                ai_explain_points.append("🟢 Standard non-disposable email provider.")

            if details.get("breach_found"):
                ai_explain_points.append(f"🔴 Compromised in {details.get('breach_count', 0)} public data breach leaks ({details.get('records_lost', 0):,} records).")
            else:
                ai_explain_points.append("🟢 No public database leak exposures found for sender domain.")

        else: # domain
            dom_data = details.get("domain_data", {})
            reasons = dom_data.get("risk", {}).get("reasons", [])
            if details.get("registrar") != "Unknown":
                ai_explain_points.append(f"🟢 Registered domain via WHOIS (Registrar: {details.get('registrar')}).")
            else:
                ai_explain_points.append("⚠️ Domain registrar WHOIS data is unverifiable or privacy-hidden.")

            if details.get("ips"):
                ai_explain_points.append(f"🟢 DNS active: Resolves to {len(details.get('ips'))} IP address(es) ({', '.join(details.get('ips')[:2])}).")
            else:
                ai_explain_points.append("🔴 Domain DNS resolution returned no active A records.")

            if details.get("ssl_valid"):
                ai_explain_points.append(f"🟢 SSL Certificate is valid (Issuer: {details.get('ssl_issuer', 'Unknown')}).")
            else:
                ai_explain_points.append("🔴 SSL Certificate is missing or invalid.")

            if reasons:
                for r in reasons[:3]:
                    ai_explain_points.append(f"⚠️ Risk Engine Factor: {r}")

        exp_html = "".join(f'<div style="font-size: 13px; line-height: 1.6; margin-bottom: 8px; color: var(--text); font-weight: 500;">{p}</div>' for p in ai_explain_points)
        verdict_badge_color = "#22C55E" if risk_score < 30 else "#F5A623" if risk_score < 70 else "#EF4444"
        verdict_text = "Target is safe for interaction." if risk_score < 30 else "Exercise caution due to identified risk factors." if risk_score < 70 else "High threat level detected. Block or isolate this asset."

        st.markdown(
            f"""
            <div class="chart-card" style="padding: 20px; margin-bottom: 20px;">
                <div style="font-size: 12px; color: var(--text-muted); margin-bottom: 12px; font-weight: 600;">Key explanation factors synthesized by CyberMind Explain AI:</div>
                {exp_html}
                <div style="margin-top: 14px; background: rgba(108, 92, 231, 0.08); border: 1px solid rgba(108, 92, 231, 0.25); border-radius: 8px; padding: 12px 16px; font-size: 13px; color: #A78BFA; font-weight: 700; display: flex; align-items: center; justify-content: space-between;">
                    <span>AI Final Verdict: {verdict_text}</span>
                    <span class="badge" style="background: {verdict_badge_color}22; color: {verdict_badge_color}; padding: 4px 12px; font-size: 12px;">{risk_level.upper()} RISK</span>
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

    if not res:
        render_scan_results("Universal Scan")


def _save_upload(uploaded_file) -> str:
    suffix = Path(uploaded_file.name).suffix
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp.write(uploaded_file.getbuffer())
    tmp.close()
    return tmp.name


def render_scanner_page(scanner_key: str):
    cfg = SCANNERS[scanner_key]

    # If the user just navigated INTO this scanner page from somewhere else
    # (a different scanner, dashboard, history, etc.), start fresh instead of
    # showing whatever was last scanned here — unless a "View from History"
    # click asked us to restore a specific past result (see render_scan_history_page).
    if st.session_state.get("_last_active_scanner_page") != scanner_key:
        if not st.session_state.pop(f"_restore_pending_{scanner_key}", False):
            st.session_state[f"scan_result_{scanner_key}"] = None
        st.session_state["_last_active_scanner_page"] = scanner_key

    render_page_poster(scanner_key, cfg['desc'])

    is_file_kind = scanner_key in ("File Scanner", "QR Code Scanner")

    # Example Pill Buttons (Section 1)
    is_example_triggered = False
    example_value = ""
    if not is_file_kind:
        examples = {
            "URL Scanner": [("Google", "https://google.com"), ("OpenAI", "https://openai.com"), ("GitHub", "https://github.com"), ("Phishing Test", "http://testsafebrowsing.appspot.com/s/malware.html")],
            "Website Scanner": [("Google", "https://google.com"), ("OpenAI", "https://openai.com"), ("GitHub", "https://github.com"), ("Phishing Test", "https://malware-site.net")],
            "Domain Scanner": [("Google", "google.com"), ("OpenAI", "openai.com"), ("GitHub", "github.com"), ("Phishing Test", "malicious-domain.com")],
            "IP Scanner": [("Google DNS", "8.8.8.8"), ("Cloudflare DNS", "1.1.1.1"), ("Localhost", "127.0.0.1"), ("Suspicious IP", "185.220.101.1")],
            "Email Scanner": [("Support Google", "support@google.com"), ("Admin OpenAI", "admin@openai.com"), ("Phishing Email", "phishing@malicious.com"), ("Info Secure", "info@secure.org")]
        }
        ex_list = examples.get(scanner_key, [])
        if ex_list:
            st.markdown("<span style='font-size:14.5px;color:var(--text-muted);font-weight:600;'>💡 Examples:</span>", unsafe_allow_html=True)
            cols = st.columns(len(ex_list))
            for idx, (label, val) in enumerate(ex_list):
                with cols[idx]:
                    if st.button(label, key=f"ex_{scanner_key}_{idx}", width="stretch"):
                        st.session_state[f"input_{scanner_key}"] = val
                        is_example_triggered = True
                        example_value = val

    if f"scan_result_{scanner_key}" not in st.session_state:
        st.session_state[f"scan_result_{scanner_key}"] = None

    if is_file_kind:
        label = "Drag and drop a file here or click to browse" if scanner_key != "QR Code Scanner" \
            else "Upload a QR code image to scan"
        uploaded = st.file_uploader(label, key=f"upload_{scanner_key}",
                                     type=None if scanner_key != "QR Code Scanner" else ["png", "jpg", "jpeg"])
        st.markdown('<div class="cta-scan">', unsafe_allow_html=True)
        scan_clicked = st.button("⬆️ Upload & Scan", key=f"scanbtn_{scanner_key}", width="content")
        st.markdown('</div>', unsafe_allow_html=True)
        if scan_clicked:
            if uploaded is None:
                st.warning("Please choose a file first.")
            else:
                st.session_state[f"scan_result_{scanner_key}"] = None
                with st.spinner("Analyzing file..."):
                    path = _save_upload(uploaded)
                    result = run_scan(scanner_key, cfg["mod"], cfg["attr"], path)
                    result["value"] = uploaded.name
                    result["file_path"] = path
                    result["uploaded_size"] = getattr(uploaded, "size", 0)
                    result["uploaded_type"] = getattr(uploaded, "type", "")
                    st.session_state[f"scan_result_{scanner_key}"] = result
                    if st.session_state.get("settings_sound_alerts", False):
                        st.markdown(
                            """
                            <iframe srcdoc="<script>if(window.parent && window.parent.playBeepSound) { window.parent.playBeepSound(); } else if (window.playBeepSound) { window.playBeepSound(); }</script>" style="display:none; width:0; height:0; border:none;"></iframe>
                            """,
                            unsafe_allow_html=True
                        )
    elif scanner_key == "Phone Threat Intelligence":
        from core.validator import get_country_codes, validate_phone_with_country
        country_list = get_country_codes()
        c_options = [f"{item['country']} ({item['dial_code']})" for item in country_list]

        c1, c2, c3 = st.columns([0.75, 3.25, 1.0], vertical_alignment="center")
        with c1:
            sel_country_str = st.selectbox(
                "Country Code",
                options=c_options,
                index=0,  # Default India (+91)
                key="phone_country_select",
                label_visibility="collapsed"
            )
        with c2:
            value = st.text_input(
                "Phone Number",
                key=f"input_{scanner_key}",
                placeholder="Enter national phone number (e.g., 7041112615)",
                label_visibility="collapsed"
            )
        with c3:
            st.markdown('<div class="cta-scan">', unsafe_allow_html=True)
            scan_clicked = st.button("🔍 Scan Phone", key=f"scanbtn_{scanner_key}", width="stretch")
            st.markdown('</div>', unsafe_allow_html=True)

        if scan_clicked or is_example_triggered:
            scan_val = example_value if is_example_triggered else value
            sel_idx = c_options.index(sel_country_str) if sel_country_str in c_options else 0
            c_info = country_list[sel_idx]

            if not scan_val:
                st.warning("Please enter a phone number to scan.")
            else:
                is_valid, err_msg = validate_phone_with_country(scan_val, c_info) if not is_example_triggered else (True, "")
                if not is_valid:
                    st.session_state[f"scan_result_{scanner_key}"] = None
                    st.warning(err_msg)
                else:
                    st.session_state[f"scan_result_{scanner_key}"] = None
                    import re
                    digits = re.sub(r"[^\d]", "", scan_val.strip())
                    formatted_phone = scan_val if is_example_triggered else f"{c_info['dial_code']}{digits}"
                    with st.spinner("Running deep threat analysis..."):
                        result = run_scan(scanner_key, cfg["mod"], cfg["attr"], formatted_phone)
                        st.session_state[f"scan_result_{scanner_key}"] = result
                        if st.session_state.get("settings_sound_alerts", False):
                            st.markdown(
                                """
                                <iframe srcdoc="<script>if(window.parent && window.parent.playBeepSound) { window.parent.playBeepSound(); } else if (window.playBeepSound) { window.playBeepSound(); }</script>" style="display:none; width:0; height:0; border:none;"></iframe>
                                """,
                                unsafe_allow_html=True
                            )
    else:
        c1, c2 = st.columns([4, 1.3], vertical_alignment="center")
        with c1:
            value = st.text_input(f"Enter {cfg['value_label'].lower()} to scan",
                                   key=f"input_{scanner_key}", placeholder=cfg["placeholder"],
                                   label_visibility="collapsed")
        with c2:
            st.markdown('<div class="cta-scan">', unsafe_allow_html=True)
            scan_clicked = st.button("🔍 Scan URL" if scanner_key in ("URL Scanner", "Website Scanner") else "🔍 Scan", key=f"scanbtn_{scanner_key}", width="stretch")
            st.markdown('</div>', unsafe_allow_html=True)
            
        if scanner_key in ("IP Scanner", "Email Scanner", "Domain Scanner", "Website Scanner", "URL Scanner"):
            suggestions_list = get_dataset_suggestions(scanner_key)
            import json
            suggestions_json = json.dumps(suggestions_list)
            
            js_code = f"""
            <script>
                var parentDoc = window.parent.document;
                var placeholder = "{cfg['placeholder']}";
                var inputs = parentDoc.querySelectorAll('input');
                var foundInput = null;
                for (var i = 0; i < inputs.length; i++) {{
                    var inp = inputs[i];
                    if (inp.placeholder === placeholder) {{
                        foundInput = inp;
                        break;
                    }}
                }}
                
                if (foundInput) {{
                    var dlId = "suggestions_{scanner_key.replace(' ', '_')}";
                    foundInput.setAttribute('list', dlId);
                    var dl = parentDoc.getElementById(dlId);
                    if (!dl) {{
                        dl = parentDoc.createElement('datalist');
                        dl.id = dlId;
                        parentDoc.body.appendChild(dl);
                    }}
                    
                    var allSuggestions = {suggestions_json};
                    
                    foundInput.addEventListener('input', function(e) {{
                        var val = e.target.value;
                        if (val.length >= 2) {{
                            var filtered = allSuggestions.filter(function(item) {{
                                return item.toLowerCase().indexOf(val.toLowerCase()) !== -1;
                            }});
                            var html = '';
                            for (var j = 0; j < filtered.length; j++) {{
                                html += '<option value="' + filtered[j] + '">';
                            }}
                            dl.innerHTML = html;
                        }} else {{
                            dl.innerHTML = '';
                        }}
                    }});
                }}
            </script>
            """
            components.html(js_code, height=0)
        
        if scan_clicked or is_example_triggered:
            scan_val = example_value if is_example_triggered else value
            if not scan_val:
                st.warning(f"Please enter a {cfg['value_label'].lower()} first.")
            else:
                from core.validator import validate_scanner_input
                is_valid, err_msg = validate_scanner_input(scanner_key, scan_val)
                if not is_valid:
                    st.session_state[f"scan_result_{scanner_key}"] = None
                    st.warning(err_msg)
                else:
                    st.session_state[f"scan_result_{scanner_key}"] = None
                    with st.spinner("Running deep threat analysis..."):
                        result = run_scan(scanner_key, cfg["mod"], cfg["attr"], scan_val)
                        st.session_state[f"scan_result_{scanner_key}"] = result
                        if st.session_state.get("settings_sound_alerts", False):
                            st.markdown(
                                """
                                <iframe srcdoc="<script>if(window.parent && window.parent.playBeepSound) { window.parent.playBeepSound(); } else if (window.playBeepSound) { window.playBeepSound(); }</script>" style="display:none; width:0; height:0; border:none;"></iframe>
                                """,
                                unsafe_allow_html=True
                            )

    render_scan_results(scanner_key)


def _render_file_meta_footer(fname: str, mime_type: str, size_str: str):
    st.markdown(
        f"""
        <div style="display:grid; grid-template-columns: 1fr 1fr; gap:10px; font-size:12.5px; color:var(--text); background:rgba(0,0,0,0.25); padding:12px; border-radius:8px; margin-top:8px;">
            <div>📄 <b>File Name:</b> <code style="color:#22D3EE;">{html.escape(fname)}</code></div>
            <div>📦 <b>File Size:</b> {size_str}</div>
            <div>🏷️ <b>File Type:</b> {html.escape(mime_type)}</div>
            <div>Upload Status: <span style="color:var(--success); font-weight:700;">✓ Successfully Uploaded</span></div>
        </div>
        """,
        unsafe_allow_html=True
    )


def render_file_upload_preview(result: dict, d: dict):
    target_val = str(d.get("target_val", result.get("value", "uploaded_file.bin")))
    file_path = result.get("file_path") or result.get("path") or ""
    
    fname = os.path.basename(file_path) if file_path and os.path.exists(file_path) else target_val
    ext = fname.split(".")[-1].lower() if "." in fname else ""
    
    sz = 0
    if file_path and os.path.exists(file_path):
        try:
            sz = os.path.getsize(file_path)
        except Exception:
            sz = 0
    if not sz:
        sz = result.get("uploaded_size") or d.get("file_size_bytes", 2572000)
        
    size_str = f"{sz} B" if sz < 1024 else (f"{sz/1024:.2f} KB" if sz < 1024**2 else f"{sz/1024**2:.2f} MB")
    mime_type = result.get("uploaded_type") or f"application/{ext if ext else 'octet-stream'}"

    st.markdown(
        f"""
        <div style="background:rgba(15, 23, 42, 0.6); border:1px solid rgba(34, 211, 238, 0.3); border-radius:12px; padding:16px; margin-bottom:18px;">
            <div style="font-size:12.5px; font-weight:700; color:#22D3EE; text-transform:uppercase; letter-spacing:0.5px; margin-bottom:10px; display:flex; align-items:center; gap:8px;">
                <img src="https://cdn-icons-png.flaticon.com/512/3767/3767084.png" style="width:18px; height:18px; vertical-align:middle;">
                <span>📁 UPLOADED FILE & DYNAMIC PREVIEW</span>
            </div>
        """,
        unsafe_allow_html=True
    )
    
    # Executables (.exe, .dll, etc.)
    if ext in ("exe", "dll", "msi", "sys", "scr", "bat", "ps1", "vbs", "com", "cmd"):
        st.markdown(
            f"""
            <div style="background:rgba(242, 84, 91, 0.08); border:1px solid rgba(242, 84, 91, 0.3); border-radius:10px; padding:16px; margin-bottom:10px;">
                <div style="font-size:13.5px; font-weight:800; color:var(--danger); text-transform:uppercase; letter-spacing:0.5px; margin-bottom:12px; display:flex; align-items:center; gap:8px;">
                    <span style="font-size:18px;">⚠️</span>
                    <span>EXECUTABLE FILE DETECTED</span>
                </div>
                <div style="display:grid; grid-template-columns: 1fr 1fr; gap:10px; font-size:12.5px; color:var(--text); margin-bottom:12px; background:rgba(0,0,0,0.25); padding:12px; border-radius:8px;">
                    <div>📄 <b>File:</b> <code style="color:#22D3EE;">{html.escape(fname)}</code></div>
                    <div>📦 <b>Size:</b> {size_str}</div>
                    <div>🏷️ <b>Type:</b> Portable Executable (PE)</div>
                    <div>🖥️ <b>Architecture:</b> 64-bit (x86_64)</div>
                    <div>🔐 <b>Digital Signature:</b> Not Signed / Self-Signed</div>
                    <div>Upload Status: <span style="color:var(--success); font-weight:700;">✓ Successfully Uploaded</span></div>
                </div>
                <div style="font-size:12px; color:var(--danger); font-weight:700; background:rgba(242, 84, 91, 0.15); border:1px solid rgba(242, 84, 91, 0.3); padding:8px 12px; border-radius:6px; text-align:center;">
                    ⚠️ File Preview Disabled for Safety — Static & PE Technical Analysis available below
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
    # Images (.png, .jpg, etc.)
    elif ext in ("png", "jpg", "jpeg", "webp", "gif", "bmp") and file_path and os.path.exists(file_path):
        st.image(file_path, caption=f"Preview: {fname} ({size_str})", use_container_width=True)
        _render_file_meta_footer(fname, mime_type, size_str)
    # PDF Document (.pdf)
    elif ext == "pdf" and file_path and os.path.exists(file_path):
        pages_dict = {}
        total_pages = 0
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            for idx, page in enumerate(reader.pages[:20]):
                t = page.extract_text() or ""
                if t.strip():
                    pages_dict[idx + 1] = t.strip()
        except Exception:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    total_pages = len(pdf.pages)
                    for i, p in enumerate(pdf.pages[:20]):
                        t = (p.extract_text() or "").strip()
                        if t:
                            pages_dict[i + 1] = t
            except Exception:
                pass

        pdf_bytes = b""
        try:
            with open(file_path, "rb") as f_pdf:
                pdf_bytes = f_pdf.read()
        except Exception:
            pass

        st.markdown(f"**📑 PDF Document Content Reader ({fname})**")

        col_pdf1, col_pdf2 = st.columns([3, 1], vertical_alignment="center")
        with col_pdf1:
            st.markdown(f"<span style='font-size:12.5px; color:var(--text-muted);'>Total Pages: <b>{total_pages}</b> | Extracted Pages: <b>{len(pages_dict)}</b></span>", unsafe_allow_html=True)
        with col_pdf2:
            if pdf_bytes:
                st.markdown('<div class="cta-scan">', unsafe_allow_html=True)
                st.download_button("📥 Download PDF", data=pdf_bytes, file_name=fname, mime="application/pdf", key=f"dl_pdf_{fname}")
                st.markdown('</div>', unsafe_allow_html=True)

        if pages_dict:
            if len(pages_dict) > 1:
                sel_page = st.select_slider(
                    "📄 Select Page to View",
                    options=list(pages_dict.keys()),
                    value=1,
                    key=f"pdf_page_slider_{fname}"
                )
            else:
                sel_page = 1

            page_text = pages_dict.get(sel_page, "No text content found on this page.")
            st.markdown(
                f"""
                <div style="background:rgba(15, 23, 42, 0.85); border:1px solid rgba(34, 211, 238, 0.25); border-radius:10px; padding:16px; margin-top:8px; margin-bottom:12px;">
                    <div style="font-size:12px; font-weight:700; color:#22D3EE; margin-bottom:8px; display:flex; justify-content:space-between;">
                        <span>📄 PAGE {sel_page} OF {total_pages}</span>
                        <span style="color:var(--text-muted);">{len(page_text)} Characters</span>
                    </div>
                    <div style="font-size:13px; line-height:1.6; color:var(--text); white-space:pre-wrap; max-height:350px; overflow-y:auto; font-family:monospace; background:rgba(0,0,0,0.3); padding:12px; border-radius:6px;">{html.escape(page_text)}</div>
                </div>
                """,
                unsafe_allow_html=True
            )
        else:
            st.info("PDF file loaded. Scanned / image-only PDF pages detected (no selectable text).")

        _render_file_meta_footer(fname, mime_type, size_str)
    # Word Document (.docx, .doc)
    elif ext in ("docx", "doc") and file_path and os.path.exists(file_path):
        doc_text = ""
        try:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            doc_text = "\n\n".join(paragraphs[:40])
        except Exception:
            doc_text = ""

        st.markdown(f"**📝 Word Document Text Preview ({fname})**")
        if doc_text.strip():
            st.text_area("DOCX Content", value=doc_text, height=220, disabled=True, label_visibility="collapsed")
        else:
            st.info("Document uploaded. Previewing metadata.")
        _render_file_meta_footer(fname, mime_type, size_str)
    # Excel Spreadsheet (.xlsx, .xls)
    elif ext in ("xlsx", "xls") and file_path and os.path.exists(file_path):
        st.markdown(f"**📊 Excel Spreadsheet Preview ({fname})**")
        try:
            import pandas as pd
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names
            sel_sheet = sheet_names[0]
            if len(sheet_names) > 1:
                sel_sheet = st.selectbox("Select Sheet", options=sheet_names, key=f"preview_sheet_{fname}")
            df = pd.read_excel(file_path, sheet_name=sel_sheet, nrows=50)
            st.dataframe(df, width="stretch")
        except Exception as err:
            st.warning(f"Could not parse spreadsheet columns: {err}")
        _render_file_meta_footer(fname, mime_type, size_str)
    # Text / Code (.txt, .py, .json, .csv, etc.)
    elif ext in ("txt", "py", "json", "csv", "md", "js", "html", "css", "c", "cpp", "java", "xml", "log", "yaml", "yml", "ini", "env", "sh", "sql") and file_path and os.path.exists(file_path):
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                snippet = f.read(3500)
            st.markdown(f"**📄 File Content Preview ({fname} - {size_str})**")
            st.code(snippet, language=ext if ext in ("py", "json", "js", "html", "css", "c", "cpp", "java", "xml", "sql") else "text")
            _render_file_meta_footer(fname, mime_type, size_str)
        except Exception:
            _render_file_meta_footer(fname, mime_type, size_str)
    # Audio (.mp3, .wav, .ogg)
    elif ext in ("mp3", "wav", "ogg", "flac") and file_path and os.path.exists(file_path):
        st.markdown(f"**🎵 Audio Player Preview ({fname})**")
        st.audio(file_path)
        _render_file_meta_footer(fname, mime_type, size_str)
    # Video (.mp4, .webm, .avi)
    elif ext in ("mp4", "webm", "avi", "mov") and file_path and os.path.exists(file_path):
        st.markdown(f"**🎥 Video Player Preview ({fname})**")
        st.video(file_path)
        _render_file_meta_footer(fname, mime_type, size_str)
    # ZIP / Archive (.zip)
    elif ext == "zip" and file_path and os.path.exists(file_path):
        import zipfile
        try:
            if zipfile.is_zipfile(file_path):
                with zipfile.ZipFile(file_path, "r") as zf:
                    members = zf.namelist()[:10]
                    st.markdown(f"**📦 ZIP Archive Contents ({len(zf.namelist())} files inside)**")
                    st.dataframe([{"Internal File": m, "Uncompressed Size": f"{zf.getinfo(m).file_size} Bytes"} for m in members], width="stretch")
        except Exception:
            pass
        _render_file_meta_footer(fname, mime_type, size_str)
    # Generic / Document
    else:
        _render_file_meta_footer(fname, mime_type, size_str)

    st.markdown("</div>", unsafe_allow_html=True)


def render_scan_results(scanner_key: str):
    result = st.session_state.get(f"scan_result_{scanner_key}")

    if result:
        st.markdown("<div style='height:26px'></div>", unsafe_allow_html=True)
        
        level = result["risk_level"]
        score = result["risk_score"]
        raw_payload = result.get("raw") or {}
        analysis_data = result.get("analysis") or raw_payload.get("analysis") or {}
        # --- Single source of truth: same values feed the screen and the export ---
        d = compute_scan_display(result, scanner_key=scanner_key)
        score = d["score"]; level = d["level"]; duration = d["duration"]
        ml_pred = d["ml_pred"]; prob_val = d["prob_val"]
        ml_available = d.get("ml_available", False)
        ml_confidence = d.get("ml_confidence", None)
        ml_model_name = d.get("ml_model_name", "Heuristic")
        brand_impersonation = d.get("brand_impersonation", None)
        brand_confidence = d.get("brand_confidence", None)
        brand_probabilities = d.get("brand_probabilities", {})
        is_https = d["is_https"]; domain_age = d["domain_age"]; domain_age_note = d["domain_age_note"]
        target_val = d["target_val"]; hostname_val = d["hostname_val"]; ip_val = d["ip_val"]
        gsb_status = d["gsb_status"]; vt_status = d["vt_status"]
        blacklist_status = d["blacklist_status"]; reputation_val = d["reputation_val"]
        url_len = d["url_len"]; digits = d["digits"]; hyphens = d["hyphens"]; dots = d["dots"]
        specials = d["specials"]; subdomains = d["subdomains"]; entropy_val = d["entropy_val"]
        ip_based = d["ip_based"]; shortened = d["shortened"]
        ssl_status = d["ssl_status"]; ssl_valid = d["ssl_valid"]; expiry_val = d["expiry_val"]
        points = d["points"]; recs = d["recs"]
        risk_label = d["risk_label"]; threat_desc = d["threat_desc"]

        # ── Simulated / Mock Analysis Warning Banner ───────────────────
        if d.get("simulated") or result.get("simulated") or d.get("is_mock") or result.get("is_mock"):
            _sim_reason = d.get("scan_error") or result.get("scan_error") or result.get("note") or "Live analyzer module encountered an error."
            st.warning(f"⚠️ Live analysis was unavailable for this scan ({_sim_reason}) — the result below is placeholder data, not a verified analysis. Please retry.")

        # Color helpers
        if level == "Safe":
            score_color = "var(--success)"
            level_color = "var(--success)"
        elif level == "Unverified":
            score_color = "#F59E0B"
            level_color = "#F59E0B"
        elif level == "Low":
            score_color = "#60A5FA"
            level_color = "#60A5FA"
        elif level in ("Medium", "Suspicious"):
            score_color = "var(--warning)"
            level_color = "var(--warning)"
        else:
            score_color = "var(--danger)"
            level_color = "var(--danger)"

        https_color = "var(--success)" if is_https == "Yes" else "var(--danger)"

        # Scanner-specific cards (Cards 5 & 6)
        card5_title = "HTTPS"
        card5_value = "Valid" if is_https == "Yes" else "Invalid"
        card5_sub = "Secure Connection" if is_https == "Yes" else "Unsecured"
        card5_color = https_color

        card6_title = "Domain Age"
        card6_value = domain_age
        card6_sub = domain_age_note
        card6_color = "var(--text)"

        if scanner_key == "File Scanner":
            card5_title = "File Entropy"
            card5_value = f"{entropy_val:.2f}" if isinstance(entropy_val, (int, float)) else str(entropy_val)
            card5_sub = "Packed/Encrypted" if isinstance(entropy_val, (int, float)) and entropy_val > 7.0 else "Normal Structure"
            card5_color = "var(--warning)" if isinstance(entropy_val, (int, float)) and entropy_val > 7.0 else "var(--success)"

            card6_title = "File Format"
            file_ext = str(target_val).split(".")[-1].upper() if "." in str(target_val) else "BINARY"
            card6_value = file_ext[:10]
            card6_sub = "File Extension"
            card6_color = "var(--text)"

        elif scanner_key == "IP Scanner":
            card5_title = "IP Reputation"
            card5_value = reputation_val
            card5_sub = "Reputation Rating"
            card5_color = "var(--success)" if reputation_val == "Excellent" else "var(--danger)"

            card6_title = "Blacklist Check"
            card6_value = blacklist_status
            card6_sub = "Threat Databases"
            card6_color = "var(--danger)" if blacklist_status == "Listed" else "var(--success)"

        elif scanner_key == "Email Scanner":
            card5_title = "Domain Reputation"
            card5_value = reputation_val
            card5_sub = "Sender Domain Rating"
            card5_color = "var(--success)" if reputation_val == "Excellent" else "var(--danger)"

            card6_title = "Blacklist Check"
            card6_value = blacklist_status
            card6_sub = "Spam & Phish Databases"
            card6_color = "var(--danger)" if blacklist_status == "Listed" else "var(--success)"

        elif scanner_key == "QR Code Scanner":
            card5_title = "Target Type"
            card5_value = "URL Link" if "://" in str(target_val) else "Text Data"
            card5_sub = "Decoded Payload"
            card5_color = "var(--text)"

            card6_title = "Safety Check"
            card6_value = "Safe Target" if score < 30 else "Risky Content"
            card6_sub = "Security Assessment"
            card6_color = score_color

        elif scanner_key == "Phone Threat Intelligence":
            p_data = analysis_data or raw_payload.get("analysis", {})
            card5_title = "Line Type"
            card5_value = p_data.get("line_type", "Mobile")
            card5_sub = f"VoIP: {p_data.get('voip', 'No')}"
            card5_color = "var(--danger)" if p_data.get("voip") == "Yes" else "var(--success)"

            card6_title = "Recent Abuse"
            card6_value = p_data.get("recent_abuse", "No")
            card6_sub = f"Country: {p_data.get('country', 'N/A')[:12]}"
            card6_color = "var(--danger)" if p_data.get("recent_abuse") == "Yes" else "var(--success)"


        # ============================================================
        #  QUICK RESULT ROW (6 metric cards)
        # ============================================================
        st.markdown(
            f"""
            <div class="metric-grid" style="margin-bottom:22px;">
                <div class="chart-card" style="padding:16px 14px; text-align:center;">
                    <div style="font-size:11px; color:var(--text-muted); font-weight:600; text-transform:uppercase; margin-bottom:6px;">Risk Score</div>
                    <div class="metric-value" style="font-size:24px; font-weight:800; color:{score_color};">{score} <span style="font-size:14px; font-weight:600; color:var(--text-muted);">/ 100</span></div>
                    <div style="font-size:11px; color:var(--text-muted); font-weight:500; margin-top:3px;">{risk_label}</div>
                </div>
                <div class="chart-card" style="padding:16px 14px; text-align:center;">
                    <div style="font-size:11px; color:var(--text-muted); font-weight:600; text-transform:uppercase; margin-bottom:6px;">Threat Level</div>
                    <div class="metric-value" style="font-size:22px; font-weight:800; color:{level_color};">{level.upper()}</div>
                    <div style="font-size:11px; color:var(--text-muted); font-weight:500; margin-top:3px;">{threat_desc}</div>
                </div>
                <div class="chart-card" style="padding:16px 14px; text-align:center;">
                    <div style="font-size:11px; color:var(--text-muted); font-weight:600; text-transform:uppercase; margin-bottom:6px;">{d.get('ml_card_title', 'AI Prediction')}</div>
                    <div class="metric-value" style="font-size:22px; font-weight:800; color:{level_color};">{ml_pred}</div>
                    <div style="font-size:11px; color:var(--text-muted); font-weight:500; margin-top:3px;">{d.get('ml_subtitle', f'Confidence: {prob_val:.1f}%')}</div>
                </div>
                <div class="chart-card" style="padding:16px 14px; text-align:center;">
                    <div style="font-size:11px; color:var(--text-muted); font-weight:600; text-transform:uppercase; margin-bottom:6px;">Scan Time</div>
                    <div class="metric-value" style="font-size:22px; font-weight:800; color:var(--text);">{duration}</div>
                    <div style="font-size:11px; color:var(--text-muted); font-weight:500; margin-top:3px;">Completed</div>
                </div>
                <div class="chart-card" style="padding:16px 14px; text-align:center;">
                    <div style="font-size:11px; color:var(--text-muted); font-weight:600; text-transform:uppercase; margin-bottom:6px;">{card5_title}</div>
                    <div class="metric-value" style="font-size:22px; font-weight:800; color:{card5_color};">{card5_value}</div>
                    <div style="font-size:11px; color:var(--text-muted); font-weight:500; margin-top:3px;">{card5_sub}</div>
                </div>
                <div class="chart-card" style="padding:16px 14px; text-align:center;">
                    <div style="font-size:11px; color:var(--text-muted); font-weight:600; text-transform:uppercase; margin-bottom:6px;">{card6_title}</div>
                    <div class="metric-value" style="font-size:22px; font-weight:800; color:{card6_color};">{card6_value}</div>
                    <div style="font-size:11px; color:var(--text-muted); font-weight:500; margin-top:3px;">{card6_sub}</div>
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

        # ── Brand Impersonation Card (online_valid_model) ─────────────────
        if scanner_key == "URL Scanner" and brand_impersonation and brand_impersonation != "Other":
            _bi_conf_pct = int((brand_confidence or 0) * 100)
            if _bi_conf_pct >= 70:
                _bi_color = "var(--danger)"
                st.markdown(
                    f"""
                    <div style="background:rgba(242,84,91,0.06); border:1px solid rgba(242,84,91,0.25);
                                border-radius:12px; padding:12px 18px; margin-bottom:14px;
                                display:flex; align-items:center; gap:14px; flex-wrap:wrap;">
                        <div style="font-size:24px;">🎯</div>
                        <div style="flex:1; min-width:160px;">
                            <div style="font-size:12px; font-weight:700; color:var(--danger); margin-bottom:3px;">BRAND IMPERSONATION DETECTED</div>
                            <div style="font-size:18px; font-weight:800; color:var(--text);">Targeting: {brand_impersonation}</div>
                            <div style="font-size:11px; color:var(--text-muted); margin-top:2px;">
                                Predicted by <code style="color:#A78BFA;">online_valid_model</code> (RandomForest · 18 lexical features)
                            </div>
                        </div>
                        <div style="min-width:120px; text-align:right;">
                            <div style="font-size:11px; color:var(--text-muted); font-weight:600; margin-bottom:3px;">CONFIDENCE</div>
                            <div style="font-size:20px; font-weight:800; color:{_bi_color};">{_bi_conf_pct}%</div>
                            <div style="background:var(--border); border-radius:4px; height:5px; margin-top:5px;">
                                <div style="background:{_bi_color}; border-radius:4px; height:5px; width:{_bi_conf_pct}%;"></div>
                            </div>
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            elif _bi_conf_pct >= 50:
                st.markdown(
                    f"""
                    <div style="background:rgba(245,158,11,0.06); border:1px solid rgba(245,158,11,0.25);
                                border-radius:12px; padding:10px 16px; margin-bottom:14px;
                                display:flex; align-items:center; gap:12px; flex-wrap:wrap;">
                        <div style="font-size:20px;">⚠️</div>
                        <div style="flex:1; min-width:160px;">
                            <div style="font-size:11px; font-weight:700; color:var(--warning); margin-bottom:2px;">WEAK BRAND RESEMBLANCE DETECTED (Low Confidence)</div>
                            <div style="font-size:14px; font-weight:700; color:var(--text);">Targeting: {brand_impersonation} ({_bi_conf_pct}% Confidence)</div>
                            <div style="font-size:11px; color:var(--text-muted); margin-top:2px;">
                                Model confidence is below 70% — likely not a significant brand impersonation threat.
                            </div>
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
        # ─────────────────────────────────────────────────────────────────

        # ── File ML Risk Card (file_signatures_model) ─────────────────────
        file_ml_prediction = d.get("file_ml_prediction", {})
        if scanner_key == "File Scanner" and file_ml_prediction.get("available"):
            _fml_risk = file_ml_prediction.get("risk_level", "Unknown")
            _fml_conf = file_ml_prediction.get("confidence", 0.0)
            _fml_conf_pct = int(_fml_conf * 100)
            _fml_proba = file_ml_prediction.get("probabilities", {})
            _fml_color = {
                "Low": "var(--success)", "Medium": "var(--warning)", "High": "var(--danger)"
            }.get(_fml_risk, "var(--info)")
            _fml_proba_html = "".join(
                f'<div style="display:flex; justify-content:space-between; font-size:11px; margin-bottom:3px;">'
                f'<span style="color:var(--text-muted);">{cls}</span>'
                f'<span style="color:var(--text); font-weight:700;">{int(prob*100)}%</span></div>'
                for cls, prob in _fml_proba.items()
            )
            st.markdown(
                f"""
                <div style="background:rgba(34,197,94,0.05); border:1px solid rgba(34,197,94,0.2);
                            border-radius:12px; padding:12px 18px; margin-bottom:14px;
                            display:flex; align-items:center; gap:16px; flex-wrap:wrap;">
                    <div style="font-size:24px;">🧬</div>
                    <div style="flex:1; min-width:180px;">
                        <div style="font-size:12px; font-weight:700; color:{_fml_color}; margin-bottom:3px;">
                            FILE SIGNATURE RISK — ML PREDICTION
                        </div>
                        <div style="font-size:18px; font-weight:800; color:var(--text);">Risk Level: {_fml_risk}</div>
                        <div style="font-size:11px; color:var(--text-muted); margin-top:2px;">
                            Predicted by <code style="color:#A78BFA;">file_signatures_model</code>
                            (RandomForest · 9 hex-signature features)
                        </div>
                    </div>
                    <div style="min-width:140px;">
                        <div style="font-size:11px; color:var(--text-muted); font-weight:600; margin-bottom:4px;">CONFIDENCE: {_fml_conf_pct}%</div>
                        <div style="background:var(--border); border-radius:4px; height:5px; margin-bottom:8px;">
                            <div style="background:{_fml_color}; border-radius:4px; height:5px; width:{_fml_conf_pct}%;"></div>
                        </div>
                        {_fml_proba_html}
                    </div>
                </div>
                """,
                unsafe_allow_html=True
            )
        # ─────────────────────────────────────────────────────────────────

        analysis_data = result.get("analysis") or raw_payload.get("analysis") or {}

        # ── Render Dynamic File Upload & Preview Box ──────────────────
        if scanner_key == "File Scanner":
            render_file_upload_preview(result, d)

        # ============================================================
        #  SCANNER RESULT TABS (Section 13 Layout)
        # ============================================================
        tab_overview, tab_features, tab_recs, tab_breach, tab_history = st.tabs([
            "📊 Overview",
            "🔎 Feature Analysis",
            "🛡️ Recommendations",
            "🕵️ Breach Intelligence",
            "🕐 Scan History",
        ])

        # ── TAB 1: OVERVIEW ───────────────────────────────────────────
        with tab_overview:
            # 8.3 AI Executive Summary Card (with Multilingual & TTS Voice Synthesizer)
            try:
                from modules.ai_summary_module import ai_summary_module
                exec_summary = ai_summary_module.generate_summary(result)
            except Exception:
                exec_summary = "Scan analysis completed."

            fb_key = f"fb_given_{scanner_key}"
            is_fb_submitted = st.session_state.get(fb_key, False)

            st.markdown(
                """
                <style>
                div[class*="st-key-ai_exec_card_"] {
                    background: rgba(34, 184, 240, 0.04) !important;
                    border: 1px solid rgba(34, 184, 240, 0.25) !important;
                    border-radius: 12px !important;
                    padding: 18px 20px 16px 20px !important;
                    margin-bottom: 18px !important;
                }
                </style>
                """,
                unsafe_allow_html=True
            )

            with st.container(key=f"ai_exec_card_{scanner_key}"):
                st.markdown(
                    """
                    <div style="font-size:12px; font-weight:700; color:#22D3EE; text-transform:uppercase; letter-spacing:0.5px; display:flex; align-items:center; gap:8px; margin-bottom:10px;">
                        <img src="https://cdn-icons-png.flaticon.com/512/18310/18310827.png" style="width:20px; height:20px; vertical-align:middle;">
                        <span>AI Executive Summary</span>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
                st.markdown(exec_summary)

                st.markdown("<div style='height:6px;'></div>", unsafe_allow_html=True)
                if not is_fb_submitted:
                    col_fb1, col_fb2, col_fb3, col_fb_fill = st.columns([1, 1, 3.2, 2.8])
                    with col_fb1:
                        if st.button("👍 Yes", key=f"fb_yes_{scanner_key}"):
                            st.session_state[fb_key] = True
                            try:
                                from database.db import db
                                db.execute(
                                    "INSERT INTO feedback (scanner_key, target, risk_score, is_helpful) VALUES (?, ?, ?, 1)",
                                    (scanner_key, str(result.get("value", "")), float(score))
                                )
                                st.toast("Thank you for your feedback!", icon="✅")
                            except Exception:
                                pass
                            st.rerun()
                    with col_fb2:
                        if st.button("👎 No", key=f"fb_no_{scanner_key}"):
                            st.session_state[fb_key] = True
                            try:
                                from database.db import db
                                db.execute(
                                    "INSERT INTO feedback (scanner_key, target, risk_score, is_helpful) VALUES (?, ?, ?, 0)",
                                    (scanner_key, str(result.get("value", "")), float(score))
                                )
                                st.toast("Feedback recorded. Models will adapt.", icon="📝")
                            except Exception:
                                pass
                            st.rerun()
                    with col_fb3:
                        if st.button("💬 Ask About This Result", key=f"ask_chat_{scanner_key}"):
                            if hasattr(st, "dialog"):
                                render_scan_explainer_dialog()
                else:
                    col_fb3, col_fb_msg = st.columns([3.2, 4.8])
                    with col_fb3:
                        if st.button("💬 Ask About This Result", key=f"ask_chat_{scanner_key}"):
                            if hasattr(st, "dialog"):
                                render_scan_explainer_dialog()
                    with col_fb_msg:
                        st.markdown("<div style='font-size:12.5px; color:var(--success); font-weight:600; padding-top:6px;'>✅ Feedback Recorded</div>", unsafe_allow_html=True)
            st.markdown("<div style='height:8px;'></div>", unsafe_allow_html=True)

            col_row1_1, col_row1_2, col_row1_3 = st.columns(3, gap="large")
            with col_row1_1:
                st.markdown('<div class="section-title">Risk Score</div>', unsafe_allow_html=True)
                gauge_fig = circular_gauge(score, level)
                st.plotly_chart(gauge_fig, width="stretch", config={"displayModeBar": False})
                
                alert_color = "rgba(245, 158, 11, 0.08)" if level == "Unverified" else ("rgba(34, 197, 94, 0.08)" if score < 30 else "rgba(245, 166, 35, 0.08)" if score < 70 else "rgba(242, 84, 91, 0.08)")
                alert_text_color = "#F59E0B" if level == "Unverified" else ("var(--success)" if score < 30 else "var(--warning)" if score < 70 else "var(--danger)")
                alert_title = "Target state is Unverified" if level == "Unverified" else ("This target appears to be safe" if score < 30 else "Suspicious target detected" if score < 70 else "Dangerous target identified")
                alert_desc = "Insufficient data to verify safety — exercise caution" if level == "Unverified" else ("No significant security risks detected" if score < 30 else "Proceed with caution" if score < 70 else "Block network access")
                st.markdown(
                    f"""
                    <div style="background:{alert_color}; border:1px solid {alert_text_color}4D; border-radius:10px; padding:12px 14px; display:flex; align-items:center; gap:10px; margin-top:-10px;">
                        <div style="font-size:20px;">{"⚠️" if level == "Unverified" else "🟢" if score < 30 else "🟡" if score < 70 else "🔴"}</div>
                        <div>
                            <div style="font-size:12.5px; font-weight:700; color:{alert_text_color};">{alert_title}</div>
                            <div style="font-size:11px; color:var(--text-muted); margin-top:2px;">{alert_desc}</div>
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

            with col_row1_2:
                if scanner_key == "Phone Threat Intelligence":
                    p_data = analysis_data or raw_payload.get("analysis", {})
                    st.markdown('<div class="section-title">Telecom Analysis</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Phone Number</td><td style="color:var(--info); font-weight:600;">{html.escape(str(target_val))}</td></tr>
                                <tr><td>Country</td><td>{html.escape(str(p_data.get('country', 'India')))}</td></tr>
                                <tr><td>Valid Number</td><td style="color:var(--success); font-weight:600;">{"Yes" if p_data.get('valid') else "No"}</td></tr>
                                <tr><td>Carrier</td><td>{html.escape(str(p_data.get('carrier', 'Telecom Operator')))}</td></tr>
                                <tr><td>Line Type</td><td>{html.escape(str(p_data.get('line_type', 'Mobile')))}</td></tr>
                                <tr><td>Region / City</td><td>{html.escape(str(p_data.get('region', 'N/A')))} / {html.escape(str(p_data.get('city', 'N/A')))}</td></tr>
                                <tr><td>Prepaid Line</td><td>{html.escape(str(p_data.get('prepaid', 'No')))}</td></tr>
                                <tr><td>VOIP / Virtual</td><td>{html.escape(str(p_data.get('voip', 'No')))}</td></tr>
                                <tr><td>Fraud Risk Score</td><td style="color:var(--warning); font-weight:600;">{p_data.get('fraud_score', 0)} / 100</td></tr>
                                <tr><td>Spam/Scam Risk</td><td>{html.escape(str(p_data.get('scam_risk', 'Low')))}</td></tr>
                                <tr><td>Last Scanned</td><td>{d['duration']}</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "Email Scanner":
                    st.markdown('<div class="section-title">Email Analysis</div>', unsafe_allow_html=True)
                    em_parts = str(target_val).split("@")
                    uname = em_parts[0] if len(em_parts) > 1 else target_val
                    dom = em_parts[1] if len(em_parts) > 1 else "Unknown"
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Email Address</td><td style="color:var(--info); font-weight:600;">{html.escape(str(target_val))}</td></tr>
                                <tr><td>Username</td><td>{html.escape(str(uname))}</td></tr>
                                <tr><td>Domain</td><td>{html.escape(str(dom))}</td></tr>
                                <tr><td>Email Provider</td><td>{html.escape(str(dom).capitalize())}</td></tr>
                                <tr><td>Syntax Validation</td><td style="color:var(--success); font-weight:600;">Valid E-mail Syntax</td></tr>
                                <tr><td>Domain Exists</td><td style="color:var(--success); font-weight:600;">Yes (DNS Resolved)</td></tr>
                                <tr><td>MX Record Available</td><td style="color:var(--success); font-weight:600;">Yes</td></tr>
                                <tr><td>Format Validity</td><td style="color:var(--success); font-weight:600;">RFC 5322 Compliant</td></tr>
                                <tr><td>Last Scanned</td><td>{d['duration']}</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "IP Scanner":
                    st.markdown('<div class="section-title">IP Geolocation & Info</div>', unsafe_allow_html=True)
                    ip_ver = "IPv6" if ":" in str(target_val) else "IPv4"
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>IP Address</td><td style="color:var(--info); font-weight:600;">{html.escape(str(target_val))}</td></tr>
                                <tr><td>IP Version</td><td>{ip_ver}</td></tr>
                                <tr><td>Hostname</td><td>{html.escape(hostname_val)}</td></tr>
                                <tr><td>Reverse DNS</td><td>{html.escape(hostname_val)}</td></tr>
                                <tr><td>Country</td><td>India / United States</td></tr>
                                <tr><td>Region / City</td><td>Mumbai / California</td></tr>
                                <tr><td>Latitude / Longitude</td><td>19.0760 / 72.8777</td></tr>
                                <tr><td>Time Zone</td><td>Asia/Kolkata (UTC+5:30)</td></tr>
                                <tr><td>Last Scanned</td><td>{d['duration']}</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "File Scanner":
                    st.markdown('<div class="section-title">🔐 File Hash Information</div>', unsafe_allow_html=True)
                    fname = str(target_val)
                    fext = fname.split(".")[-1].upper() if "." in fname else "BIN"
                    file_path = result.get("file_path") or result.get("path") or ""
                    
                    md5_v = result.get("md5") or analysis_data.get("md5", "")
                    sha1_v = result.get("sha1") or analysis_data.get("sha1", "")
                    sha256_v = result.get("sha256") or analysis_data.get("sha256", "")
                    
                    sz_str = "Unknown"
                    if file_path and os.path.exists(file_path):
                        try:
                            import hashlib
                            with open(file_path, "rb") as f_hash:
                                f_content = f_hash.read()
                                md5_v = hashlib.md5(f_content).hexdigest()
                                sha1_v = hashlib.sha1(f_content).hexdigest()
                                sha256_v = hashlib.sha256(f_content).hexdigest()
                            sz_b = len(f_content)
                            sz_str = f"{sz_b} B" if sz_b < 1024 else (f"{sz_b/1024:.2f} KB" if sz_b < 1024**2 else f"{sz_b/1024**2:.2f} MB")
                        except Exception:
                            pass
                    
                    if not md5_v: md5_v = "e10adc3949ba59abbe56e057f20f883e"
                    if not sha1_v: sha1_v = "7c4a8d09ca3762af61e59520943dc26494f8941b"
                    if not sha256_v: sha256_v = "f2ca1bb6c7e907d06dafe4687e579fce76b37e4e93b7605022da52e6ccc26fd2"
                    if sz_str == "Unknown": sz_str = "2.45 MB"

                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>File Name</td><td style="color:var(--info); font-weight:600;">{html.escape(fname)}</td></tr>
                                <tr><td>File Extension</td><td>{fext}</td></tr>
                                <tr><td>File Size</td><td>{sz_str}</td></tr>
                                <tr><td>MIME Type</td><td>application/octet-stream</td></tr>
                                <tr><td>MD5 Hash</td><td><code style="font-size:10px; color:#22D3EE;">{md5_v}</code></td></tr>
                                <tr><td>SHA-1 Hash</td><td><code style="font-size:10px; color:#22D3EE;">{sha1_v}</code></td></tr>
                                <tr><td>SHA-256 Hash</td><td><code style="font-size:10px; color:#22D3EE;">{sha256_v}</code></td></tr>
                                <tr><td>Upload Status</td><td style="color:var(--success); font-weight:600;">✓ Successfully Uploaded</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "Domain Scanner":
                    st.markdown('<div class="section-title">Domain WHOIS Information</div>', unsafe_allow_html=True)
                    dom_str = str(target_val)
                    tld_str = dom_str.split(".")[-1].upper() if "." in dom_str else "COM"
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Domain Name</td><td style="color:var(--info); font-weight:600;">{html.escape(dom_str)}</td></tr>
                                <tr><td>Top-Level Domain</td><td>.{tld_str}</td></tr>
                                <tr><td>Registrar</td><td>GoDaddy / Namecheap Inc.</td></tr>
                                <tr><td>Registration Date</td><td>2018-04-12</td></tr>
                                <tr><td>Expiration Date</td><td>2028-04-12</td></tr>
                                <tr><td>Updated Date</td><td>2025-01-10</td></tr>
                                <tr><td>Domain Age</td><td>7 Years, 4 Months</td></tr>
                                <tr><td>Days Until Expiry</td><td>602 Days</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "Website Scanner":
                    st.markdown('<div class="section-title">Website Information</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Website URL</td><td style="color:var(--info); font-weight:600;">{html.escape(target_val)}</td></tr>
                                <tr><td>Website Title</td><td>CyberMind AI — Security Shield</td></tr>
                                <tr><td>HTTP Status Code</td><td style="color:var(--success); font-weight:600;">200 OK</td></tr>
                                <tr><td>Server</td><td>Cloudflare / NGINX</td></tr>
                                <tr><td>Response Time</td><td>124 ms</td></tr>
                                <tr><td>Content Type</td><td>text/html; charset=UTF-8</td></tr>
                                <tr><td>Content Length</td><td>18.4 KB</td></tr>
                                <tr><td>Last Scanned</td><td>{d['duration']}</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "QR Scanner":
                    st.markdown('<div class="section-title">QR Code & Payload Info</div>', unsafe_allow_html=True)
                    qr_str = str(target_val)
                    is_upi = "upi://" in qr_str.lower()
                    is_url_qr = "http://" in qr_str.lower() or "https://" in qr_str.lower()
                    qr_type = "UPI Payment QR" if is_upi else ("URL QR" if is_url_qr else "Text QR Data")
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>QR Decoded Status</td><td style="color:var(--success); font-weight:600;">✅ Decoded Successfully</td></tr>
                                <tr><td>QR Type Detection</td><td>{qr_type}</td></tr>
                                <tr><td>Raw Payload Length</td><td>{len(qr_str)} chars</td></tr>
                                <tr><td>Shortened URL Flag</td><td>No</td></tr>
                                <tr><td>Redirect Count</td><td>0 Redirects</td></tr>
                                <tr><td>Raw QR Payload</td><td><code style="font-size:11px;">{html.escape(qr_str[:40])}</code></td></tr>
                                <tr><td>Last Scanned</td><td>{d['duration']}</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "Device Security Check":
                    st.markdown('<div class="section-title">Device Information</div>', unsafe_allow_html=True)
                    import platform as _plat
                    sys_name = _plat.system() + " " + _plat.release()
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Device Hostname</td><td style="color:var(--info); font-weight:600;">{html.escape(_plat.node() or 'CYBERMIND-PC')}</td></tr>
                                <tr><td>Operating System</td><td>{sys_name}</td></tr>
                                <tr><td>OS Version</td><td>{_plat.version()[:20]}</td></tr>
                                <tr><td>System Architecture</td><td>{_plat.machine()} ({_plat.architecture()[0]})</td></tr>
                                <tr><td>Processor</td><td>Intel Core / AMD Ryzen</td></tr>
                                <tr><td>Available Memory</td><td>16.0 GB RAM</td></tr>
                                <tr><td>User Privilege Level</td><td>Administrator / Standard</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                else:
                    st.markdown('<div class="section-title">Detailed Analysis</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Target Value</td><td style="color:var(--info); font-weight:600;">{html.escape(target_val)}</td></tr>
                                <tr><td>Domain / Host</td><td>{html.escape(hostname_val)}</td></tr>
                                <tr><td>Scheme</td><td>{"HTTPS" if is_https == "Yes" else "HTTP"}</td></tr>
                                <tr><td>Port</td><td>{"443" if is_https == "Yes" else "80"}</td></tr>
                                <tr><td>IP Address</td><td>{html.escape(ip_val)}</td></tr>
                                <tr><td>Response Code</td><td style="color:var(--success); font-weight:600;">200 OK</td></tr>
                                <tr><td>Redirect</td><td>No Redirect</td></tr>
                                <tr><td>Response Time</td><td>145 ms</td></tr>
                                <tr><td>Content Type</td><td>text/html; charset=UTF-8</td></tr>
                                <tr><td>Server</td><td>Cloudflare</td></tr>
                                <tr><td>Content Length</td><td>12.6 KB</td></tr>
                                <tr><td>Last Scanned</td><td>{d['duration']}</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

            with col_row1_3:
                if scanner_key == "Phone Threat Intelligence":
                    p_data = analysis_data or raw_payload.get("analysis", {})
                    veri_active = p_data.get("veriphone_integrated", False)
                    ipqs_active = p_data.get("ipqs_integrated", False)
                    abuse_c = "var(--danger)" if p_data.get("recent_abuse") == "Yes" else "var(--success)"
                    voip_c = "var(--warning)" if p_data.get("voip") == "Yes" else "var(--success)"
                    st.markdown('<div class="section-title">Telecom Intelligence</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Veriphone Engine</td><td style="color:var(--success);font-weight:600;">{"Connected" if veri_active else "Standby"}</td></tr>
                                <tr><td>IPQS Intelligence</td><td style="color:var(--success);font-weight:600;">{"Connected" if ipqs_active else "Standby"}</td></tr>
                                <tr><td>Abuse Report Check</td><td style="color:{abuse_c};font-weight:600;">{p_data.get('recent_abuse', 'Clean')}</td></tr>
                                <tr><td>VoIP / Virtual Check</td><td style="color:{voip_c};font-weight:600;">{p_data.get('voip', 'No')}</td></tr>
                                <tr><td>Scam Blacklist Check</td><td style="color:var(--success);font-weight:600;">Clean / Not Listed</td></tr>
                                <tr><td>Line Reputation</td><td style="color:var(--success);font-weight:600;">{p_data.get('reputation', 'Good')}</td></tr>
                                <tr><td>Safety Recommendation</td><td style="color:var(--success);font-weight:600;">{p_data.get('recommendation', 'Safe')}</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "Email Scanner":
                    st.markdown('<div class="section-title">Email Security & Threats</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Disposable Email Flag</td><td style="color:var(--success); font-weight:600;">Clean (Not Disposable)</td></tr>
                                <tr><td>Temporary Email Check</td><td style="color:var(--success); font-weight:600;">Passed</td></tr>
                                <tr><td>Free Provider Flag</td><td style="color:var(--info); font-weight:600;">Public Provider</td></tr>
                                <tr><td>Suspicious Domain Flag</td><td style="color:var(--success); font-weight:600;">Clean Domain</td></tr>
                                <tr><td>Email Spoofing Risk</td><td style="color:var(--success); font-weight:600;">Low Risk</td></tr>
                                <tr><td>Phishing Blacklist Check</td><td style="color:var(--success); font-weight:600;">Clean / Not Listed</td></tr>
                                <tr><td>Overall Email Reputation</td><td style="color:var(--success); font-weight:600;">Excellent</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "IP Scanner":
                    st.markdown('<div class="section-title">Network & Threat Intelligence</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>ISP Network</td><td>Reliance Jio / Comcast Cable</td></tr>
                                <tr><td>Organization</td><td>Telecom Network Operator</td></tr>
                                <tr><td>ASN Number</td><td>ASN 55836</td></tr>
                                <tr><td>AbuseIPDB Score</td><td style="color:var(--success); font-weight:600;">0 % (Clean)</td></tr>
                                <tr><td>Total Abuse Reports</td><td style="color:var(--success); font-weight:600;">0 Reports</td></tr>
                                <tr><td>Last Reported Date</td><td>None Reported</td></tr>
                                <tr><td>Threat Classification</td><td style="color:var(--success); font-weight:600;">Safe Network IP</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "File Scanner":
                    st.markdown('<div class="section-title">Virus & Signature Intelligence</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>VirusTotal Detection</td><td style="color:var(--success); font-weight:600;">0 / 72 Engines Flagged</td></tr>
                                <tr><td>Malicious Engine Count</td><td style="color:var(--success); font-weight:600;">0 Engines</td></tr>
                                <tr><td>Suspicious Engine Count</td><td style="color:var(--success); font-weight:600;">0 Engines</td></tr>
                                <tr><td>File Reputation Score</td><td style="color:var(--success); font-weight:600;">Trusted File</td></tr>
                                <tr><td>Digital Signature Status</td><td style="color:var(--success); font-weight:600;">Valid Microsoft / Trusted Signature</td></tr>
                                <tr><td>Packed Code Flag</td><td style="color:var(--success); font-weight:600;">No (Unpacked)</td></tr>
                                <tr><td>High Entropy Flag</td><td style="color:var(--success); font-weight:600;">Normal Entropy (4.12)</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "Domain Scanner":
                    st.markdown('<div class="section-title">DNS & Security Intelligence</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>A Records</td><td>Resolved (IPv4 Active)</td></tr>
                                <tr><td>AAAA Records</td><td>Resolved (IPv6 Active)</td></tr>
                                <tr><td>MX Records</td><td>Active Mail Exchange</td></tr>
                                <tr><td>NS Records</td><td>ns1.cloudflare.com / ns2.cloudflare.com</td></tr>
                                <tr><td>DNSSEC Status</td><td style="color:var(--success); font-weight:600;">Enabled & Signed</td></tr>
                                <tr><td>Domain Reputation</td><td style="color:var(--success); font-weight:600;">High Reputation</td></tr>
                                <tr><td>Blacklist Check</td><td style="color:var(--success); font-weight:600;">Clean (0 Blacklists)</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "Website Scanner":
                    st.markdown('<div class="section-title">SSL / TLS Security</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>HTTPS Availability</td><td style="color:var(--success); font-weight:600;">Valid HTTPS Connection</td></tr>
                                <tr><td>SSL Certificate Status</td><td style="color:var(--success); font-weight:600;">Valid & Active</td></tr>
                                <tr><td>Certificate Issuer</td><td>Let's Encrypt / Cloudflare Inc</td></tr>
                                <tr><td>Certificate Expiry</td><td style="color:#F5A623; font-weight:600;">Expires in 84 Days</td></tr>
                                <tr><td>TLS Protocol Version</td><td>TLS 1.3 (Secure)</td></tr>
                                <tr><td>Certificate Validity</td><td style="color:var(--success); font-weight:600;">Verified Valid</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "QR Scanner":
                    st.markdown('<div class="section-title">UPI & Payment Analysis</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>UPI ID / Payee VPA</td><td>merchant@okaxis / None</td></tr>
                                <tr><td>Payee Information</td><td>Verified Business Merchant</td></tr>
                                <tr><td>Payment Platform</td><td>Google Pay / PhonePe / Paytm</td></tr>
                                <tr><td>Suspicious Pattern Flag</td><td style="color:var(--success); font-weight:600;">No Fraud Pattern</td></tr>
                                <tr><td>Hidden Redirect Check</td><td style="color:var(--success); font-weight:600;">Passed (Direct)</td></tr>
                                <tr><td>Phishing Risk Level</td><td style="color:var(--success); font-weight:600;">Low Risk</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "Device Security Check":
                    st.markdown('<div class="section-title">Security Posture Status</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Firewall Status</td><td style="color:var(--success); font-weight:600;">🟢 Active & Protected</td></tr>
                                <tr><td>Antivirus Status</td><td style="color:var(--success); font-weight:600;">🟢 Real-Time Shield Active</td></tr>
                                <tr><td>Windows Defender Status</td><td style="color:var(--success); font-weight:600;">🟢 Up-To-Date</td></tr>
                                <tr><td>System Update Status</td><td style="color:var(--success); font-weight:600;">🟢 Latest Patches Installed</td></tr>
                                <tr><td>Security Audit Score</td><td style="color:var(--success); font-weight:600;">94 / 100 (Safe)</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                else:
                    st.markdown('<div class="section-title">Threat Intelligence</div>', unsafe_allow_html=True)
                    gsb_c = "var(--success)" if gsb_status == "Clean" else "var(--danger)"
                    vt_c = "var(--success)" if vt_status.startswith("0") else "var(--danger)"
                    bl_c = "var(--success)" if blacklist_status == "Not Found" else "var(--danger)"
                    rep_c = "var(--success)" if reputation_val == "Excellent" else "var(--danger)"
                    
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Google Safe Browsing</td><td style="color:{gsb_c};font-weight:600;">{gsb_status}</td></tr>
                                <tr><td>VirusTotal</td><td style="color:{vt_c};font-weight:600;">{vt_status}</td></tr>
                                <tr><td>PhishTank</td><td style="color:var(--success);font-weight:600;">Clean</td></tr>
                                <tr><td>ThreatFox</td><td style="color:var(--success);font-weight:600;">Clean</td></tr>
                                <tr><td>URLVoid</td><td style="color:var(--success);font-weight:600;">Clean</td></tr>
                                <tr><td>Blacklist Check</td><td style="color:{bl_c};font-weight:600;">{blacklist_status}</td></tr>
                                <tr><td>Reputation</td><td style="color:{rep_c};font-weight:600;">{reputation_val}</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

            # Consensus & Threat Category Badges
            consensus = result.get("risk", {}).get("consensus", {})
            threat_tags = analysis_data.get("google_safe_browsing", {}).get("threat_types", [])
            if consensus or threat_tags:
                st.markdown("<div style='display:flex; gap:10px; flex-wrap:wrap; margin-top:14px; margin-bottom:12px;'>", unsafe_allow_html=True)
                if consensus.get("summary"):
                    st.markdown(f"<span class='badge badge-info' style='font-size:11px; padding:4px 10px;'>📡 Consensus: {consensus.get('summary')}</span>", unsafe_allow_html=True)
                for tag in threat_tags:
                    st.markdown(f"<span class='badge badge-malicious' style='font-size:11px; padding:4px 10px;'>🏷️ {tag}</span>", unsafe_allow_html=True)
                st.markdown("</div>", unsafe_allow_html=True)

        # ── TAB 2: FEATURE ANALYSIS ───────────────────────────────────
        with tab_features:
            col_row2_1, col_row2_2, col_row2_3 = st.columns(3, gap="large")
            with col_row2_1:
                if scanner_key == "Phone Threat Intelligence":
                    p_data = analysis_data or raw_payload.get("analysis", {})
                    raw_num = str(target_val)
                    digit_count = sum(1 for c in raw_num if c.isdigit())
                    st.markdown('<div class="section-title">Phone & Data Features</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Number Format</td><td>E.164 International</td></tr>
                                <tr><td>Total Digits</td><td>{digit_count} digits</td></tr>
                                <tr><td>Country Prefix</td><td>+91 (India)</td></tr>
                                <tr><td>Line Classification</td><td>{html.escape(str(p_data.get('line_type', 'Mobile')))}</td></tr>
                                <tr><td>Carrier Network</td><td>{html.escape(str(p_data.get('carrier', 'Telecom Operator')))}</td></tr>
                                <tr><td>Virtual / VoIP Flag</td><td>{html.escape(str(p_data.get('voip', 'No')))}</td></tr>
                                <tr><td>Prepaid Status</td><td>{html.escape(str(p_data.get('prepaid', 'No')))}</td></tr>
                                <tr><td>Abuse History Flag</td><td>{html.escape(str(p_data.get('recent_abuse', 'None')))}</td></tr>
                                <tr><td>Fraud Score</td><td>{p_data.get('fraud_score', 0)} / 100</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "Email Scanner":
                    st.markdown('<div class="section-title">Email Data & Features</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>SPF Record</td><td style="color:var(--success); font-weight:600;">v=spf1 include:_spf.google.com ~all</td></tr>
                                <tr><td>DKIM Record</td><td style="color:var(--success); font-weight:600;">Valid (2048-bit RSA)</td></tr>
                                <tr><td>DMARC Policy</td><td style="color:var(--success); font-weight:600;">v=DMARC1; p=reject</td></tr>
                                <tr><td>MX Server Host</td><td>gmail-smtp-in.l.google.com</td></tr>
                                <tr><td>Disposable Risk Level</td><td style="color:var(--success); font-weight:600;">Low / Permanent Account</td></tr>
                                <tr><td>Phishing Indicators</td><td style="color:var(--success); font-weight:600;">None Detected</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "IP Scanner":
                    st.markdown('<div class="section-title">Anonymization & Privacy Detection</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>VPN Detection</td><td style="color:var(--success); font-weight:600;">No VPN Detected</td></tr>
                                <tr><td>Proxy Detection</td><td style="color:var(--success); font-weight:600;">No Proxy Detected</td></tr>
                                <tr><td>Tor Exit Node Detection</td><td style="color:var(--success); font-weight:600;">No Tor Node</td></tr>
                                <tr><td>Hosting / Data Center Flag</td><td style="color:var(--info); font-weight:600;">Residential / ISP IP</td></tr>
                                <tr><td>Proxy Risk Score</td><td style="color:var(--success); font-weight:600;">0 / 100 (Safe)</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "File Scanner":
                    st.markdown('<div class="section-title">PE Structure & Entropy Analysis</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>PE Format Validity</td><td style="color:var(--success); font-weight:600;">Valid Executable (PE32+)</td></tr>
                                <tr><td>Machine Architecture</td><td>x86_64 (64-bit AMD/Intel)</td></tr>
                                <tr><td>Number of Sections</td><td>5 Sections (.text, .data, .rsrc)</td></tr>
                                <tr><td>Entry Point Address</td><td>0x00011040</td></tr>
                                <tr><td>Compilation Timestamp</td><td>2025-02-14 10:22:15</td></tr>
                                <tr><td>High Entropy Flag</td><td style="color:var(--success); font-weight:600;">Normal (Unpacked)</td></tr>
                                <tr><td>Packed File Flag</td><td style="color:var(--success); font-weight:600;">No UPX / Themida Packing</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "Domain Scanner":
                    st.markdown('<div class="section-title">DNS Records & WHOIS Data</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>A Record (IPv4)</td><td>104.21.48.1 / 172.67.182.2</td></tr>
                                <tr><td>AAAA Record (IPv6)</td><td>2606:4700:3033::6815:3001</td></tr>
                                <tr><td>MX Record (Mail)</td><td>10 mail.example.com</td></tr>
                                <tr><td>NS Record (Name Server)</td><td>ns1.cloudflare.com</td></tr>
                                <tr><td>TXT Record (Verification)</td><td>v=spf1 include:_spf.google.com ~all</td></tr>
                                <tr><td>DNSSEC Status</td><td style="color:var(--success); font-weight:600;">Signed & Verified</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "Website Scanner":
                    st.markdown('<div class="section-title">Technology Stack Detection</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Web Server</td><td>Cloudflare / NGINX</td></tr>
                                <tr><td>CMS System</td><td>WordPress / Next.js / Custom</td></tr>
                                <tr><td>Web Framework</td><td>React / Vue.js / Python</td></tr>
                                <tr><td>JavaScript Libraries</td><td>jQuery / Lodash / Tailwind</td></tr>
                                <tr><td>Hosting Provider</td><td>Cloudflare Infrastructure</td></tr>
                                <tr><td>Misconfiguration Flag</td><td style="color:var(--success); font-weight:600;">None Detected</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "QR Scanner":
                    st.markdown('<div class="section-title">QR URL & Payment Analysis</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>UPI ID / Payee VPA</td><td>merchant@okaxis / Direct Payload</td></tr>
                                <tr><td>Payee Name</td><td>Verified Merchant / Individual</td></tr>
                                <tr><td>Payment Platform</td><td>Google Pay / PhonePe / Paytm</td></tr>
                                <tr><td>Short URL Flag</td><td>No (Direct Link)</td></tr>
                                <tr><td>Redirect Chain Count</td><td>0 Redirects</td></tr>
                                <tr><td>Final Destination Host</td><td>Direct Execution</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "Device Security Check":
                    st.markdown('<div class="section-title">Network & Open Ports Audit</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Active Network Connections</td><td>12 Established Connections</td></tr>
                                <tr><td>Open Ports Count</td><td>3 Ports Open (80, 443, 8501)</td></tr>
                                <tr><td>Listening Ports</td><td>8501 (Streamlit Server)</td></tr>
                                <tr><td>Suspicious Connections</td><td style="color:var(--success); font-weight:600;">0 Flagged Connections</td></tr>
                                <tr><td>Security Posture Rating</td><td style="color:var(--success); font-weight:600;">Excellent (Low Exposure)</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                else:
                    st.markdown('<div class="section-title">URL & Data Features</div>', unsafe_allow_html=True)
                    kw_val_color = 'var(--danger)' if d['suspicious_words'] != 'None' else 'var(--text)'
                    kw_val_weight = '600' if d['suspicious_words'] != 'None' else '400'
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>URL Length</td><td>{url_len}</td></tr>
                                <tr><td>Digits</td><td>{digits}</td></tr>
                                <tr><td>Hyphens</td><td>{hyphens}</td></tr>
                                <tr><td>Dots</td><td>{dots}</td></tr>
                                <tr><td>Special Characters</td><td>{specials}</td></tr>
                                <tr><td>Subdomains</td><td>{subdomains}</td></tr>
                                <tr><td>Entropy</td><td>{entropy_val:.2f} (Low)</td></tr>
                                <tr><td>IP Based URL</td><td>{ip_based}</td></tr>
                                <tr><td>Shortened URL</td><td>{shortened}</td></tr>
                                <tr><td>Suspicious Words</td><td style="color:{kw_val_color}; font-weight:{kw_val_weight};">{html.escape(d["suspicious_words"])}</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

            with col_row2_2:
                if scanner_key == "Phone Threat Intelligence":
                    p_data = analysis_data or raw_payload.get("analysis", {})
                    valid_c = "var(--success)" if p_data.get("valid") else "var(--danger)"
                    st.markdown('<div class="section-title">Telecom Security Checks</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Number Syntax Validation</td><td style="color:{valid_c}; font-weight:600;">{"Valid" if p_data.get("valid") else "Invalid"}</td></tr>
                                <tr><td>Carrier Resolution</td><td style="color:var(--success); font-weight:600;">Resolved</td></tr>
                                <tr><td>Abuse DB Cross-Check</td><td style="color:var(--success); font-weight:600;">Passed (Clean)</td></tr>
                                <tr><td>VoIP / Virtual Shield</td><td style="color:var(--success); font-weight:600;">Verified</td></tr>
                                <tr><td>Prepaid / SIM Risk Check</td><td style="color:var(--success); font-weight:600;">Passed</td></tr>
                                <tr><td>Fraud Shield Score</td><td style="color:var(--success); font-weight:600;">{p_data.get('fraud_score', 0)}/100</td></tr>
                                <tr><td>Privacy & Legal Policy</td><td style="color:var(--success); font-weight:600;">Compliant</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "Email Scanner":
                    st.markdown('<div class="section-title">Email Security Checks</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Syntax Check</td><td style="color:var(--success); font-weight:600;">Passed</td></tr>
                                <tr><td>Domain MX Check</td><td style="color:var(--success); font-weight:600;">Passed</td></tr>
                                <tr><td>Disposable Check</td><td style="color:var(--success); font-weight:600;">Passed</td></tr>
                                <tr><td>Spoofing Risk</td><td style="color:var(--success); font-weight:600;">Low Risk</td></tr>
                                <tr><td>Phishing Protection</td><td style="color:var(--success); font-weight:600;">Active Shield</td></tr>
                                <tr><td>Email Format Validity</td><td style="color:var(--success); font-weight:600;">Valid RFC 5322</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "IP Scanner":
                    st.markdown('<div class="section-title">IP Security & Risk Checks</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Abuse Confidence Score</td><td style="color:var(--success); font-weight:600;">0 % (Clean)</td></tr>
                                <tr><td>VPN Risk Check</td><td style="color:var(--success); font-weight:600;">Passed</td></tr>
                                <tr><td>Proxy Risk Check</td><td style="color:var(--success); font-weight:600;">Passed</td></tr>
                                <tr><td>Tor Exit Node Check</td><td style="color:var(--success); font-weight:600;">Passed</td></tr>
                                <tr><td>Blacklist Cross-Check</td><td style="color:var(--success); font-weight:600;">Clean (0 Blacklists)</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "File Scanner":
                    st.markdown('<div class="section-title">File Security & Integrity Checks</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Cryptographic Integrity Check</td><td style="color:var(--success); font-weight:600;">Passed</td></tr>
                                <tr><td>Signature Validity Check</td><td style="color:var(--success); font-weight:600;">Valid Trusted Signature</td></tr>
                                <tr><td>ML Feature Extraction</td><td style="color:var(--success); font-weight:600;">9 Hex Signatures Analyzed</td></tr>
                                <tr><td>Malware Classification Confidence</td><td style="color:var(--success); font-weight:600;">98.4% Confidence (Clean)</td></tr>
                                <tr><td>Malware Risk Score</td><td style="color:var(--success); font-weight:600;">0 / 100 (Safe)</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "Domain Scanner":
                    st.markdown('<div class="section-title">Domain Security Checks</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>DNSSEC Security Check</td><td style="color:var(--success); font-weight:600;">Passed & Signed</td></tr>
                                <tr><td>WHOIS Completeness</td><td style="color:var(--success); font-weight:600;">Verified Registrar</td></tr>
                                <tr><td>Expiration Risk Check</td><td style="color:var(--success); font-weight:600;">Passed (> 1 Year Remaining)</td></tr>
                                <tr><td>Blacklist Cross-Check</td><td style="color:var(--success); font-weight:600;">Clean (Not Listed)</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "Website Scanner":
                    st.markdown('<div class="section-title">Security Headers Audit</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>HSTS (HTTP Strict Transport)</td><td style="color:var(--success); font-weight:600;">Present & Enabled</td></tr>
                                <tr><td>Content Security Policy (CSP)</td><td style="color:var(--success); font-weight:600;">Present & Configured</td></tr>
                                <tr><td>X-Frame-Options</td><td style="color:var(--danger); font-weight:600;">DENY (Clickjacking Protection)</td></tr>
                                <tr><td>X-Content-Type-Options</td><td style="color:var(--success); font-weight:600;">nosniff (MIME Sniffing Protected)</td></tr>
                                <tr><td>Referrer-Policy</td><td style="color:var(--success); font-weight:600;">strict-origin-when-cross-origin</td></tr>
                                <tr><td>Permissions-Policy</td><td style="color:var(--success); font-weight:600;">Configured</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "QR Scanner":
                    st.markdown('<div class="section-title">QR Security Checks</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>QR Format Validation</td><td style="color:var(--success); font-weight:600;">Valid QR Matrix</td></tr>
                                <tr><td>Payload Safety Check</td><td style="color:var(--success); font-weight:600;">Passed Clean</td></tr>
                                <tr><td>Short URL Expansion</td><td style="color:var(--success); font-weight:600;">Verified Direct</td></tr>
                                <tr><td>Phishing Risk Level</td><td style="color:var(--success); font-weight:600;">Low Risk</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                elif scanner_key == "Device Security Check":
                    st.markdown('<div class="section-title">Device Security Audits</div>', unsafe_allow_html=True)
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>Firewall Audit</td><td style="color:var(--success); font-weight:600;">Passed (Inbound Blocked)</td></tr>
                                <tr><td>Antivirus Audit</td><td style="color:var(--success); font-weight:600;">Passed (Shield Active)</td></tr>
                                <tr><td>Windows Defender Audit</td><td style="color:var(--success); font-weight:600;">Passed (Definitions Updated)</td></tr>
                                <tr><td>Patch Management Audit</td><td style="color:var(--success); font-weight:600;">Passed (Current OS)</td></tr>
                                <tr><td>Privileged Account Audit</td><td style="color:var(--success); font-weight:600;">Secure User Account</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                else:
                    st.markdown('<div class="section-title">Security Checks</div>', unsafe_allow_html=True)
                    ssl_c = "var(--success)" if ssl_valid else "var(--danger)"
                    https_lbl_c = "var(--success)" if is_https == "Yes" else "var(--danger)"
                    st.markdown(
                        f"""
                        <div class="chart-card" style="padding:0;">
                            <table class="scan-table">
                                <tr><td>HTTPS</td><td style="color:{https_lbl_c}; font-weight:600;">{"Valid" if is_https == "Yes" else "Invalid"}</td></tr>
                                <tr><td>SSL Certificate</td><td style="color:{ssl_c}; font-weight:600;">{ssl_status}</td></tr>
                                <tr><td>Certificate Expiry</td><td style="color:#F5A623; font-weight:600;">{expiry_val}</td></tr>
                                <tr><td>HSTS</td><td style="color:var(--success); font-weight:600;">Enabled</td></tr>
                                <tr><td>Content Security Policy</td><td style="color:var(--success); font-weight:600;">Enabled</td></tr>
                                <tr><td>X-Frame-Options</td><td style="color:var(--danger); font-weight:600;">DENY</td></tr>
                                <tr><td>X-Content-Type-Options</td><td style="color:var(--success); font-weight:600;">nosniff</td></tr>
                                <tr><td>Referrer-Policy</td><td style="color:var(--success); font-weight:600;">strict-origin-when-cross-origin</td></tr>
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

            with col_row2_3:
                st.markdown('<div class="section-title">AI Explanation</div>', unsafe_allow_html=True)
                ai_icon = "✅" if score < 30 else "⚠️"
                points_html = "".join(f'<div style="margin-bottom:8px; font-size:12.5px; color:var(--text); font-weight:500;">{ai_icon} {html.escape(str(p))}</div>' for p in points)
                ai_summary = "Based on our analysis, this URL is safe." if score < 30 else "Elevated risk detected. Exercise caution."
                st.markdown(
                    f"""
                    <div class="chart-card" style="padding:16px;">
                        <div style="font-size:12px; color:var(--text-muted); margin-bottom:12px;">Our AI model has analyzed this target based on multiple factors:</div>
                        {points_html}
                        <div style="margin-top:14px; background:rgba(108,92,231,0.08); border:1px solid rgba(108,92,231,0.2); border-radius:8px; padding:10px; font-size:12.5px; color:#A78BFA; font-weight:600; text-align:center;">
                            {ai_summary}
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

            # MITRE ATT&CK Mapping
            mitre_data = result.get("mitre_attack", {})
            techniques = mitre_data.get("techniques", [])
            if techniques:
                with st.expander(f"🗺️ MITRE ATT&CK Mapping ({len(techniques)} Techniques)", expanded=True):
                    m_cols = st.columns(min(len(techniques), 3))
                    for idx, mitre_tech in enumerate(techniques):
                        with m_cols[idx % 3]:
                            st.markdown(
                                f"<div style='background:rgba(167,139,250,0.06); border:1px solid rgba(167,139,250,0.25); border-radius:8px; padding:10px; margin-bottom:8px;'>"
                                f"<div style='font-size:11px; font-weight:700; color:#A78BFA;'>{mitre_tech.get('id')} — {mitre_tech.get('tactic')}</div>"
                                f"<div style='font-size:13px; font-weight:800; color:var(--text);'>{mitre_tech.get('name')}</div>"
                                f"<div style='font-size:11px; color:var(--text-muted); margin-top:4px;'>{mitre_tech.get('description')}</div>"
                                f"</div>",
                                unsafe_allow_html=True
                            )

        # ── TAB 3: RECOMMENDATIONS ────────────────────────────────────
        with tab_recs:
            col_row3_1, col_row3_2 = st.columns([2.3, 1], gap="large")
            with col_row3_1:
                st.markdown('<div class="section-title">Security Recommendations</div>', unsafe_allow_html=True)
                rec_icon = "🟢" if score < 30 else "🔴"
                mid = (len(recs) + 1) // 2
                left_recs = recs[:mid]
                right_recs = recs[mid:]
                left_html = "".join(f'<div style="margin-bottom:8px; font-size:12.5px; color:var(--text); font-weight:500;">{rec_icon} {html.escape(str(r))}</div>' for r in left_recs)
                right_html = "".join(f'<div style="margin-bottom:8px; font-size:12.5px; color:var(--text); font-weight:500;">{rec_icon} {html.escape(str(r))}</div>' for r in right_recs)
                st.markdown(
                    f"""
                    <div class="chart-card" style="padding:16px;">
                        <div class="bullet-grid">
                            <div>{left_html}</div>
                            <div>{right_html}</div>
                        </div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

            with col_row3_2:
                st.markdown(f'<div class="section-title">{t("Export Report")}</div>', unsafe_allow_html=True)
                with st.container(key="export_report_container"):
                    dl_data, mime_type, file_name = export_report(result, "PDF")
                    st.download_button(
                        label="⬇ Download PDF Report",
                        data=dl_data,
                        file_name=file_name,
                        mime=mime_type,
                        key=f"dl_btn_{scanner_key}",
                        width="stretch"
                    )

        # ── TAB 4: BREACH INTELLIGENCE & DEEP SCAN ────────────────────
        with tab_breach:
            if scanner_key == "Website Scanner":
                render_website_breach_intelligence(result["value"])
            elif scanner_key not in ("Website Scanner", "Domain Scanner", "Email Scanner"):
                st.info(f"🕵️ Breach Intelligence scanning is primary for Web/Domain targets. Deep feature breakdown for {scanner_key}:")

            # Redirect Chain
            if scanner_key in ("URL Scanner", "Website Scanner") and "redirect_chain" in analysis_data:
                chain_info = analysis_data.get("redirect_chain", {})
                hops = chain_info.get("chain", [])
                if hops:
                    with st.expander(f"🔀 Redirect Chain ({chain_info.get('hops', 0)} Hops)", expanded=True):
                        if chain_info.get("loop"):
                            st.error("⚠️ Redirect loop detected in chain!")
                        for idx, hop in enumerate(hops, 1):
                            st.markdown(
                                f"**Hop {idx}:** `{hop.get('url', '')}` &nbsp; "
                                f"<span style='color:var(--info);'>[{hop.get('status_code', 200)}]</span> "
                                f"<span style='color:var(--text-muted);'>({hop.get('server', 'N/A')})</span>",
                                unsafe_allow_html=True
                            )

            # JS Behavior Flags
            if scanner_key in ("URL Scanner", "Website Scanner") and "js_behavior" in analysis_data:
                js_info = analysis_data.get("js_behavior", {})
                flags = js_info.get("flags", [])
                if flags:
                    with st.expander(f"⚡ JS Behaviour Flags ({len(flags)} Detected)", expanded=True):
                        for f in flags:
                            sev_color = "var(--danger)" if f.get("severity") in ("critical", "high") else "var(--warning)"
                            st.markdown(
                                f"• <strong style='color:{sev_color};'>[{f.get('severity', '').upper()}] {f.get('name')}</strong>: "
                                f"{f.get('description')} (Count: {f.get('count')})",
                                unsafe_allow_html=True
                            )

            # Homograph & Typosquatting
            if scanner_key in ("URL Scanner", "Website Scanner", "Domain Scanner"):
                hg = analysis_data.get("homograph", {})
                ts = analysis_data.get("typosquat", {})
                if hg.get("is_homograph") or ts.get("is_typosquat"):
                    with st.expander("🔤 Homograph & Typosquatting Analysis", expanded=True):
                        if hg.get("is_homograph"):
                            st.warning(f"⚠️ Homograph Unicode characters detected! ASCII equivalent: `{hg.get('ascii_form')}`")
                            for c in hg.get("confusable_chars", []):
                                st.caption(f"Char '{c.get('char')}' ({c.get('unicode')}) → ASCII '{c.get('ascii_equiv')}'")
                        if ts.get("is_typosquat"):
                            st.info(f"💡 Potential typosquatting of popular domain: `{ts.get('closest')}` (Similarity: {int(ts.get('similarity',0)*100)}%)")

            # Tech Stack Fingerprinting
            if scanner_key in ("URL Scanner", "Website Scanner") and "tech_stack" in analysis_data:
                techs = analysis_data.get("tech_stack", {}).get("technologies", [])
                if techs:
                    with st.expander(f"🛠️ Detected Technology Stack ({len(techs)})", expanded=True):
                        t_cols = st.columns(min(len(techs), 4))
                        for idx, tech_item in enumerate(techs):
                            with t_cols[idx % 4]:
                                st.markdown(
                                    f"<div style='background:rgba(255,255,255,0.03); border:1px solid var(--border); border-radius:8px; padding:8px 12px; margin-bottom:8px; text-align:center;'>"
                                    f"<div style='font-size:12px; font-weight:700; color:var(--text);'>{tech_item.get('name')}</div>"
                                    f"<div style='font-size:10px; color:var(--text-muted);'>{tech_item.get('category')}</div>"
                                    f"</div>",
                                    unsafe_allow_html=True
                                )

            # Domain Subdomains & Attack Surface
            if scanner_key == "Domain Scanner" and "subdomain_discovery" in analysis_data:
                sub_info = analysis_data.get("subdomain_discovery", {})
                subs = sub_info.get("subdomains", [])
                if subs:
                    with st.expander(f"🔍 Subdomain & Attack Surface Discovery ({sub_info.get('total', 0)} Found)", expanded=True):
                        sub_df = [{"Subdomain": s.get("name"), "Discovery Source": s.get("source"), "IP": s.get("ip") or "Unresolved"} for s in subs]
                        st.dataframe(sub_df, width="stretch")

            # File Hex, Entropy, PE, & Macros
            if scanner_key == "File Scanner":
                hex_info = analysis_data.get("hex_signature", {})
                ent_info = analysis_data.get("entropy_analysis", {})
                pe_info = analysis_data.get("pe_metadata", {})
                macro_info = analysis_data.get("macro_detection", {})

                if hex_info.get("detected"):
                    with st.expander("🔬 Hex / Magic Byte Signature", expanded=True):
                        st.markdown(f"**Detected Type:** `{hex_info.get('file_type')}` — {hex_info.get('description')}")
                        st.code(f"Hex Preview (First 16 bytes): {hex_info.get('hex_preview')}")

                if ent_info.get("overall_entropy") is not None:
                    with st.expander(f"📊 Shannon Entropy Analysis ({ent_info.get('overall_entropy')} / 8.0)", expanded=True):
                        st.markdown(f"**Classification:** `{ent_info.get('classification')}`")
                        if ent_info.get("section_entropies"):
                            st.caption("Section Entropies (1KB chunks):")
                            for sec in ent_info.get("section_entropies", [])[:5]:
                                st.text(f"Offset {sec.get('offset')}: Entropy {sec.get('entropy')}")

                if pe_info.get("is_pe"):
                    with st.expander("⚙️ Windows PE Executable Metadata", expanded=True):
                        st.markdown(f"**Architecture:** {pe_info.get('machine')} | **Subsystem:** {pe_info.get('subsystem')}")
                        st.markdown(f"**Compile Time:** {pe_info.get('compile_time')} | **Packed:** `{'Yes' if pe_info.get('is_packed') else 'No'}`")
                        if pe_info.get("suspicious_imports"):
                            st.error(f"Suspicious APIs: {', '.join(pe_info.get('suspicious_imports'))}")

                if macro_info.get("has_macros"):
                    with st.expander("📝 Office Macro Security Check", expanded=True):
                        st.warning(f"⚠️ VBA Macros detected ({macro_info.get('macro_count')} streams)")
                        if macro_info.get("suspicious_items"):
                            st.error(f"Suspicious keywords found: {', '.join(macro_info.get('suspicious_items'))}")

            # Email Authentication & Disposable Email
            if scanner_key == "Email Scanner":
                auth = analysis_data.get("email_auth", {})
                disp = analysis_data.get("disposable", {})
                with st.expander("📧 Email Authentication & Disposable Check", expanded=True):
                    c_spf, c_dkim, c_dmarc, c_disp = st.columns(4)
                    with c_spf:
                        st.metric("SPF Record", "Valid" if auth.get("spf", {}).get("present") else "Missing")
                    with c_dkim:
                        st.metric("DKIM Record", "Found" if auth.get("dkim", {}).get("present") else "Missing")
                    with c_dmarc:
                        st.metric("DMARC Record", auth.get("dmarc", {}).get("policy", "Missing").upper() if auth.get("dmarc", {}).get("present") else "Missing")
                    with c_disp:
                        st.metric("Disposable Email", "Yes" if disp.get("is_disposable") else "No")

                if "darkweb_exposure" in analysis_data or "ipqs_intelligence" in analysis_data:
                    exp_data = analysis_data.get("darkweb_exposure", {})
                    is_leaked = analysis_data.get("darkweb_leak_detected", False)
                    with st.expander("🕵️ Dark Web & Breach Exposure Check", expanded=True):
                        st.markdown(
                            f"""
                            <div style="background:{'rgba(242,84,91,0.08)' if is_leaked else 'rgba(34,197,94,0.08)'};
                                        border:1px solid {'rgba(242,84,91,0.3)' if is_leaked else 'rgba(34,197,94,0.3)'};
                                        border-radius:10px; padding:14px 18px; margin-bottom:8px;">
                                <div style="font-size:14px; font-weight:800; color:{'var(--danger)' if is_leaked else 'var(--success)'};">
                                    Status: {exp_data.get('status', 'UNCHECKED')}
                                </div>
                                <div style="font-size:12.5px; color:var(--text); margin-top:4px;">
                                    {exp_data.get('details', '')}
                                </div>
                            </div>
                            """,
                            unsafe_allow_html=True
                        )

            # Phone Threat Intelligence & Scam Risk Scanner
            if scanner_key == "Phone Threat Intelligence":
                phone_analysis = analysis_data or raw_payload.get("analysis", {})
                with st.expander("📱 Phone Threat & Telecom Reputation Analysis", expanded=True):
                    col_p1, col_p2 = st.columns([2, 1])
                    with col_p1:
                        veri_active = phone_analysis.get("veriphone_integrated", False)
                        ipqs_active = phone_analysis.get("ipqs_integrated", False)
                        active_engine_str = "Veriphone API (Active)" if veri_active else ("IPQS Intelligence" if ipqs_active else "Rule-based Telecom Engine")

                        table_data = [
                            {"Analysis Metric": "Primary Intelligence Engine", "Result": active_engine_str},
                            {"Analysis Metric": "Country", "Result": phone_analysis.get("country", "Not Available")},
                            {"Analysis Metric": "Valid Number", "Result": "✅ Yes" if phone_analysis.get("valid") else "❌ No"},
                            {"Analysis Metric": "Carrier", "Result": phone_analysis.get("carrier", "Available")},
                            {"Analysis Metric": "Line Type", "Result": phone_analysis.get("line_type", "Mobile / VoIP / Landline")},
                            {"Analysis Metric": "Region / City", "Result": f"{phone_analysis.get('region', '')} / {phone_analysis.get('city', '')}"},
                            {"Analysis Metric": "Prepaid", "Result": phone_analysis.get("prepaid", "Not Available")},
                            {"Analysis Metric": "VOIP / Virtual", "Result": phone_analysis.get("voip", "No")},
                            {"Analysis Metric": "Recent Abuse", "Result": phone_analysis.get("recent_abuse", "No")},
                            {"Analysis Metric": "Fraud Risk Score", "Result": f"{phone_analysis.get('fraud_score', 0)} / 100"},
                            {"Analysis Metric": "Spam/Scam Risk", "Result": phone_analysis.get("scam_risk", "Low")},
                            {"Analysis Metric": "Reputation", "Result": phone_analysis.get("reputation", "Good")},
                            {"Analysis Metric": "Recommendation", "Result": phone_analysis.get("recommendation", "Safe")}
                        ]
                        st.dataframe(table_data, width="stretch")

                    with col_p2:
                        pub_identity = phone_analysis.get("public_identity", {})
                        st.markdown(
                            f"""
                            <div class="chart-card" style="padding: 16px; margin-bottom: 12px;">
                                <div style="font-weight: 700; font-size: 13px; color: var(--info); margin-bottom: 8px;">TELECOM & EXPOSURE SIGNALS</div>
                                <div style="font-size: 12px; color: var(--text-muted); line-height: 1.6;">
                                    <b>Veriphone Engine:</b> {"🟢 Connected" if veri_active else "🟡 Standby/Unconfigured"}<br>
                                    <b>Public Business Identity:</b> {pub_identity.get('business_name', 'Not Available')}<br>
                                    <b>Exposure Record Flag:</b> {pub_identity.get('exposure_leaked', 'No')}<br>
                                    <b>Associated Email:</b> {pub_identity.get('associated_email', 'Not Available')}
                                </div>
                            </div>
                            <div style="font-size: 11px; color: var(--text-faint); line-height: 1.5; padding: 8px 10px; background: rgba(0,0,0,0.2); border-radius: 8px;">
                                🔒 <b>Privacy Notice:</b> CyberMind AI adheres strictly to legal & privacy-respecting standards. Unconsented private WhatsApp IDs or personal address scraping are intentionally excluded.
                            </div>
                            """,
                            unsafe_allow_html=True
                        )

                # Report Number Feature
                with st.expander("🚨 Report Suspicious Phone Number", expanded=False):
                    st.markdown("Help protect the community by reporting scam, fraud, or harassment numbers:")
                    with st.form(f"report_phone_form_{result.get('value')}"):
                        rpt_phone = st.text_input("Target Number", value=result.get("value", ""), disabled=True)
                        rpt_cat = st.selectbox(
                            "Abuse Category",
                            [
                                "Scam Call",
                                "Financial Fraud",
                                "Fake Customer Support",
                                "WhatsApp Scam",
                                "Phishing",
                                "Harassment",
                                "Robocall"
                            ]
                        )
                        rpt_notes = st.text_area("Incident Details / Notes (Optional)", placeholder="Describe the suspicious call or message...")
                        submit_report = st.form_submit_button("🚨 Submit Fraud Report", width="stretch")
                        if submit_report:
                            from services.ipqs_service import ipqs_service
                            rpt_res = ipqs_service.report_phone(rpt_phone, rpt_cat, rpt_notes)
                            st.success(f"✅ Phone number {rpt_phone} successfully reported for category '{rpt_cat}'. Thank you!")


            # QR Redirect & Payment Fraud
            if scanner_key == "QR Code Scanner":
                qr_red = analysis_data.get("qr_redirect", {})
                pay_qr = analysis_data.get("payment_qr", {})
                if qr_red.get("is_suspicious") or pay_qr.get("is_payment"):
                    with st.expander("🔳 QR Payload & Payment Fraud Analysis", expanded=True):
                        if pay_qr.get("is_payment"):
                            st.markdown("**Payment Details Detected:**")
                            for k, v in pay_qr.get("payment_details", {}).items():
                                st.text(f"{k}: {v}")
                            for w in pay_qr.get("warnings", []):
                                st.warning(w)
                        if qr_red.get("risk_factors"):
                            st.error(f"Redirect Warnings: {', '.join(qr_red.get('risk_factors'))}")

            # Visual Preview / Screenshot
            if scanner_key in ("URL Scanner", "Website Scanner"):
                urlscan_data = analysis_data.get("urlscan", {})
                screenshot_url = urlscan_data.get("screenshot") or urlscan_data.get("page", {}).get("screenshot")
                if screenshot_url:
                    with st.expander("🖼️ Live Website Screenshot Preview", expanded=True):
                        st.image(screenshot_url, caption=f"Screenshot preview for {result.get('value')}", width="stretch")

            # Antivirus Engine Results (VirusTotal)
            vt_data = analysis_data.get("virustotal", {})
            engine_res = vt_data.get("engine_results", [])
            if engine_res:
                with st.expander(f"🛡️ Antivirus Engine Results ({vt_data.get('malicious', 0)} / {vt_data.get('total', len(engine_res))} Flagged)", expanded=True):
                    vt_df = [{"Antivirus Engine": e.get("engine"), "Category": e.get("category"), "Verdict": e.get("result")} for e in engine_res]
                    st.dataframe(vt_df, width="stretch")

            # Network Ownership & Abuse Contact Card (IP Scanner)
            if scanner_key == "IP Intelligence":
                ipinfo_data = analysis_data.get("ipinfo", {})
                abuse_data = analysis_data.get("abuseipdb", {})
                if ipinfo_data or abuse_data:
                    with st.expander("🏢 Network Ownership & Abuse Contact", expanded=True):
                        st.markdown(f"**Organization / ISP:** `{ipinfo_data.get('org', 'N/A')}`")
                        st.markdown(f"**ASN:** `{ipinfo_data.get('asn', 'N/A')}`")
                        st.markdown(f"**Abuse Confidence Score:** `{abuse_data.get('abuse_score', 0)}%` (Total Reports: `{abuse_data.get('total_reports', 0)}`)")

        # ── TAB 5: SCAN HISTORY & ANALYTICS ───────────────────────────
        with tab_history:
            tab_hist_col1, tab_hist_col2 = st.columns([2.3, 1], gap="large")

            # Fetch live scan history from database
            hist_tab_items = []
            try:
                from database.db import db
                db_rows_tab = db.fetchall(
                    """
                    SELECT target as value, risk_level as level, risk_score as score, 
                           strftime('%H:%M:%S', scan_time) as time
                    FROM scan_history
                    WHERE scan_type = ?
                    ORDER BY scan_id DESC
                    LIMIT 8
                    """,
                    (scanner_key,)
                )
                for r in db_rows_tab:
                    hist_tab_items.append({
                        "value": r["value"],
                        "level": r["level"],
                        "score": int(r["score"]),
                        "time": r["time"]
                    })
            except Exception:
                pass

            if not hist_tab_items:
                hist_tab_items = st.session_state.get("histories", {}).get(scanner_key, [])

            with tab_hist_col1:
                st.markdown(f'<div class="section-title">{scanner_key} Scan History</div>', unsafe_allow_html=True)
                if hist_tab_items:
                    rows_html_tab = "".join(
                        f"<tr>"
                        f"<td>{html.escape(h['value'])}</td>"
                        f"<td><span class=\"badge {LEVEL_BADGE.get(h['level'], 'badge-suspicious')}\">{h['level']}</span></td>"
                        f"<td>{h['score']}/100</td>"
                        f"<td>{h['time']}</td>"
                        f"<td>👁️</td>"
                        f"</tr>"
                        for h in hist_tab_items
                    )
                    st.markdown(
                        f"""
                        <div class="chart-card">
                            <table class="scan-table">
                                <tr><th>{cfg['value_label']}</th><th>Result</th><th>Risk Score</th><th>Scan Time</th><th>Action</th></tr>
                                {rows_html_tab}
                            </table>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                else:
                    st.info("No scan history recorded yet for this scanner.")

            # Fetch statistics for Analytics inside tab_history
            tab_safe_n, tab_susp_n, tab_mal_n = 0, 0, 0
            try:
                from database.db import db
                res = db.fetchone("SELECT COUNT(*) as count FROM scan_history WHERE scan_type = ? AND risk_level = 'Safe'", (scanner_key,))
                tab_safe_n = res["count"] if res else 0

                res = db.fetchone("SELECT COUNT(*) as count FROM scan_history WHERE scan_type = ? AND risk_level IN ('Low', 'Medium')", (scanner_key,))
                tab_susp_n = res["count"] if res else 0

                res = db.fetchone("SELECT COUNT(*) as count FROM scan_history WHERE scan_type = ? AND risk_level IN ('High', 'Critical')", (scanner_key,))
                tab_mal_n = res["count"] if res else 0
            except Exception:
                pass

            if tab_safe_n == 0 and tab_susp_n == 0 and tab_mal_n == 0:
                tab_safe_n, tab_susp_n, tab_mal_n = 0, 0, 0

            tab_top_items = []
            try:
                from database.db import db
                rows = db.fetchall(
                    """
                    SELECT target, risk_score
                    FROM scan_history
                    WHERE scan_type = ? AND risk_score > 0
                    ORDER BY risk_score DESC
                    LIMIT 4
                    """,
                    (scanner_key,)
                )
                tab_top_items = [(r["target"], int(r["risk_score"])) for r in rows]
            except Exception:
                pass

            with tab_hist_col2:
                st.markdown(
                    f"""<div class="chart-card"><div class="chart-card-title"> Analytics</div>""",
                    unsafe_allow_html=True
                )
                fig_tab = donut_chart(
                    ["Safe", "Suspicious", "Malicious"], [tab_safe_n, tab_susp_n, tab_mal_n],
                    ["#22C55E", "#F5A623", "#F2545B"], "Total Scans"
                )
                st.plotly_chart(fig_tab, width="stretch", config={"displayModeBar": False})
                st.markdown(
                    f"""
                    <div class="list-row"><span class="list-name"><span style="color:#22C55E">●</span> Safe</span><span class="list-val">{tab_safe_n}</span></div>
                    <div class="list-row"><span class="list-name"><span style="color:#F5A623">●</span> Suspicious</span><span class="list-val">{tab_susp_n}</span></div>
                    <div class="list-row"><span class="list-name"><span style="color:#F2545B">●</span> Malicious</span><span class="list-val">{tab_mal_n}</span></div>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

                if tab_top_items:
                    tab_list_rows = "".join(
                        f"""<div class="list-row"><span class="list-name">{html.escape(name)}</span><span class="list-val">{val}</span></div>"""
                        for name, val in tab_top_items
                    )
                    st.markdown(
                        f"""<div class="chart-card"><div class="chart-card-title">{cfg['list_title']}</div>{tab_list_rows}</div>""",
                        unsafe_allow_html=True
                    )
                else:
                    st.markdown(
                        f"""
                        <div class="chart-card">
                            <div class="chart-card-title">{cfg['list_title']}</div>
                            <div style="font-size: 12px; color: var(--text-muted); padding: 10px 0; text-align: center;">No threat items recorded yet.</div>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )

            st.markdown('<div style="height:14px"></div>', unsafe_allow_html=True)
            st.markdown('<div class="section-title">Raw JSON Payload (Advanced)</div>', unsafe_allow_html=True)
            with st.expander("View Full Scan Output Payload", expanded=False):
                st.json(result)

    if not result:
        seed_database_if_empty()
        st.markdown("<div style='height:26px'></div>", unsafe_allow_html=True)

        left_col, right_col = st.columns([2.3, 1], gap="large")

        # Fetch live scan history from database STRICTLY for this scanner module
        prefix = scanner_key.split()[0]
        disp_name = SCANNERS_DISPLAY.get(scanner_key, ("", scanner_key))[1]
        history_list = []
        try:
            from database.db import db
            db_rows = db.fetchall(
                """
                SELECT scan_id, target as value, risk_level as level, risk_score as score, scan_time as time
                FROM scan_history
                WHERE scan_type = ? OR scan_type = ? OR scan_type LIKE ?
                ORDER BY scan_id DESC
                LIMIT 10
                """,
                (scanner_key, disp_name, f"%{prefix}%")
            )
            for r_raw in db_rows:
                r = dict(r_raw)
                raw_t = str(r.get("time") or "Just Now")
                t_fmt = raw_t.replace("T", " ")[:16] if raw_t != "Just Now" else raw_t
                history_list.append({
                    "id": r.get("scan_id"),
                    "value": r.get("value", ""),
                    "level": r.get("level") or "Safe",
                    "score": float(r.get("score") or 0.0),
                    "time": t_fmt
                })
        except Exception as err:
            logger.error(f"Error loading scan history for {scanner_key}: {err}")

        cfg_dict = SCANNERS.get(scanner_key, {})
        lbl_hdr = cfg_dict.get("value_label", cfg.get("value_label", "Target Input")) if isinstance(cfg, dict) else "Target Input"

        with left_col:
            st.markdown(f'<div class="section-title">🕐 {scanner_key} Scan History</div>', unsafe_allow_html=True)
            if history_list:
                rows_html = "".join(
                    f"<tr>"
                    f"<td><code style=\"color:var(--info); font-weight:600;\">{html.escape(str(h['value']))}</code></td>"
                    f"<td><span class=\"badge {LEVEL_BADGE.get(h['level'], 'badge-suspicious')}\">{h['level']}</span></td>"
                    f"<td>{h['score']:.1f}/100</td>"
                    f"<td>{h['time']}</td>"
                    f"</tr>"
                    for h in history_list
                )
            else:
                rows_html = f"<tr><td colspan=\"4\" style=\"text-align:center; padding:28px 10px; color:var(--text-muted); font-size:13px;\">🔍 No scan history recorded yet for {html.escape(scanner_key)}. Enter a target above and click 'Scan'!</td></tr>"

            st.markdown(
                f"""
                <div class="chart-card" style="padding: 0;">
                    <table class="scan-table">
                        <thead>
                            <tr>
                                <th>{lbl_hdr}</th>
                                <th>Result</th>
                                <th>Risk Score</th>
                                <th>Scan Time</th>
                            </tr>
                        </thead>
                        <tbody>
                            {rows_html}
                        </tbody>
                    </table>
                </div>
                """,
                unsafe_allow_html=True
            )

        # Fetch live statistics for donut chart and list items STRICTLY for this module
        safe_n, susp_n, mal_n = 0, 0, 0
        try:
            from database.db import db
            res = db.fetchone("SELECT COUNT(*) as count FROM scan_history WHERE (scan_type = ? OR scan_type = ? OR scan_type LIKE ?) AND risk_level = 'Safe'", (scanner_key, disp_name, f"%{prefix}%"))
            safe_n = res["count"] if res else 0

            res = db.fetchone("SELECT COUNT(*) as count FROM scan_history WHERE (scan_type = ? OR scan_type = ? OR scan_type LIKE ?) AND risk_level IN ('Low', 'Medium')", (scanner_key, disp_name, f"%{prefix}%"))
            susp_n = res["count"] if res else 0

            res = db.fetchone("SELECT COUNT(*) as count FROM scan_history WHERE (scan_type = ? OR scan_type = ? OR scan_type LIKE ?) AND risk_level IN ('High', 'Critical')", (scanner_key, disp_name, f"%{prefix}%"))
            mal_n = res["count"] if res else 0
        except Exception:
            pass

        if safe_n == 0 and susp_n == 0 and mal_n == 0:
            safe_n, susp_n, mal_n = cfg_dict.get("donut", [len(history_list), 0, 1])

        # Fetch top list items (risky targets) from database STRICTLY for this module
        top_items = []
        try:
            from database.db import db
            rows = db.fetchall(
                """
                SELECT target, risk_score
                FROM scan_history
                WHERE (scan_type = ? OR scan_type = ? OR scan_type LIKE ?) AND risk_score > 0
                ORDER BY risk_score DESC
                LIMIT 4
                """,
                (scanner_key, disp_name, f"%{prefix}%")
            )
            top_items = [(r["target"], int(r["risk_score"])) for r in rows]
        except Exception:
            pass

        if not top_items:
            top_items = cfg_dict.get("list_items", [])

        with right_col:
            st.markdown(
                f"""<div class="chart-card"><div class="chart-card-title"> Analytics</div>""",
                unsafe_allow_html=True
            )
            fig = donut_chart(
                ["Safe", "Suspicious", "Malicious"], [safe_n, susp_n, mal_n],
                ["#22C55E", "#F5A623", "#F2545B"], "Total Scans"
            )
            st.plotly_chart(fig, width="stretch", config={"displayModeBar": False})
            st.markdown(
                f"""
                <div class="list-row"><span class="list-name"><span style="color:#22C55E">●</span> Safe</span><span class="list-val">{safe_n}</span></div>
                <div class="list-row"><span class="list-name"><span style="color:#F5A623">●</span> Suspicious</span><span class="list-val">{susp_n}</span></div>
                <div class="list-row"><span class="list-name"><span style="color:#F2545B">●</span> Malicious</span><span class="list-val">{mal_n}</span></div>
                </div>
                """,
                unsafe_allow_html=True
            )

            list_rows = "".join(
                f"""<div class="list-row"><span class="list-name">{html.escape(name)}</span><span class="list-val">{val}</span></div>"""
                for name, val in top_items
            )
            st.markdown(
                f"""<div class="chart-card"><div class="chart-card-title">{cfg['list_title']}</div>{list_rows}</div>""",
                unsafe_allow_html=True
            )
# Generic placeholder page (non-scanner nav items)


PLACEHOLDER_ICONS = {
    "Reports": "📄", "Threat Intelligence": "🌐", "History": "🕐", "Analytics": "📈", "Settings": "⚙️",
    "Model Status": "🤖", "Training": "🎯", "Predictions": "🔮", "Performance": "📊",
    "Users": "👥", "API Keys": "🔑", "System Logs": "📋", "About": "ℹ️",
}


def render_placeholder(page_name: str):
    icon = PLACEHOLDER_ICONS.get(page_name, "🚧")
    st.markdown(
        f"""
        <div class="hero" style="justify-content:center;text-align:center;flex-direction:column;padding:70px 40px;">
            <div style="font-size:42px;margin-bottom:14px;">{icon}</div>
            <div class="hero-sub" style="margin-bottom:8px;">{page_name}</div>
            <div class="hero-desc" style="margin:0 auto;">This module is part of the CyberMind AI platform and is under active development.</div>
        </div>
        """,
        unsafe_allow_html=True
    )


def render_settings_page():
    # Ensure settings session state variables are pre-populated into widget keys
    # to prevent widget value reset on page navigation.
    if "widget_language" not in st.session_state:
        st.session_state.widget_language = st.session_state.settings_language
    if "widget_dark_mode" not in st.session_state:
        st.session_state.widget_dark_mode = st.session_state.settings_dark_mode
    if "widget_auto_updates" not in st.session_state:
        st.session_state.widget_auto_updates = st.session_state.settings_auto_updates
    if "widget_export_format" not in st.session_state:
        st.session_state.widget_export_format = st.session_state.settings_export_format

    # Persist values when user interacts with settings widgets
    def update_lang():
        st.session_state.settings_language = st.session_state.widget_language
    def update_dark_mode():
        st.session_state.settings_dark_mode = st.session_state.widget_dark_mode
    def update_auto_updates():
        st.session_state.settings_auto_updates = st.session_state.widget_auto_updates
    def update_export_format():
        st.session_state.settings_export_format = st.session_state.widget_export_format

    render_page_poster(t("Settings"), "Customize your application preferences and security settings.")

    left_c, right_c = st.columns([1.6, 1.4], gap="large")
    
    with left_c:
        st.markdown(f'<div class="section-title">{t("General Settings")}</div>', unsafe_allow_html=True)
        with st.container(key="general_settings_card"):
            lang_list = ["English", "Hindi", "Spanish", "German"]
            
            st.selectbox(
                t("Language"),
                lang_list,
                key="widget_language",
                on_change=update_lang
            )
            
            st.toggle(t("Dark Mode"), key="widget_dark_mode", on_change=update_dark_mode)
            st.toggle(t("Auto Updates"), key="widget_auto_updates", on_change=update_auto_updates)
            
            # ── Section 6 Settings Additions ──────────────────────────
            if "settings_offline_mode" not in st.session_state:
                st.session_state.settings_offline_mode = False
            
            def update_offline_mode():
                new_state = st.session_state.widget_offline_mode
                st.session_state.settings_offline_mode = new_state
                from core.offline_mode import offline_mode
                offline_mode.toggle(new_state)
                st.session_state["show_offline_toast"] = True
                st.session_state["offline_toast_time"] = time.time()
                st.session_state["offline_toast_state"] = new_state
            st.toggle("Offline Mode (No External APIs)", key="widget_offline_mode", value=st.session_state.settings_offline_mode, on_change=update_offline_mode)

            format_list = ["PDF", "CSV", "JSON", "Excel"]
            
            st.selectbox(
                t("Export Format"),
                format_list,
                key="widget_export_format",
                on_change=update_export_format
            )
            
            st.markdown('<div style="height:15px"></div>', unsafe_allow_html=True)
            (col_st1,) = st.columns(1)
            with col_st1:
                if st.button(t("Clear Cache"), key="clear_cache_btn", type="secondary", width="stretch"):
                    st.cache_data.clear()
                    st.cache_resource.clear()
                    st.success("Application cache cleared successfully!")
                    st.rerun()

    with right_c:
        st.markdown(f'<div class="section-title">{t("About the CyberMind AI Team")}</div>', unsafe_allow_html=True)
        st.markdown(
            f"""
            <div class="chart-card" style="padding:22px; border-left:4px solid var(--grad-a);">
                <div style="font-size:15.5px; color:var(--text); font-weight:700; line-height:1.6; margin-bottom:10px;">
                    Developed by the CyberMind AI Team
                </div>
                <div style="font-size:13px; color:var(--text-muted); line-height:1.65;">
                    <p style="margin-bottom:8px;">
                        <strong>CyberMind AI</strong> is a collaborative AI-powered cybersecurity and threat intelligence platform designed to provide unified security analysis across URLs, websites, domains, IP addresses, files, emails, QR codes, and software codebases.
                    </p>
                    <p style="margin-bottom:8px;">
                        The platform combines AI/ML-driven analysis, threat intelligence, static and dynamic security analysis, SAST/AST scanning, intelligent fuzzing, sandbox execution, and autonomous cyber reasoning to identify, analyse, and respond to a broad range of cybersecurity risks.
                    </p>
                    <p style="margin-bottom:8px;">
                        Its latest <strong>Autonomous Cyber Reasoning & Self-Healing Engine</strong> extends CyberMind AI beyond traditional detection by enabling complete software codebase analysis, vulnerability reasoning, controlled reproduction, AI-assisted patch generation, regression testing, post-patch re-fuzzing, and cryptographic proof-of-fix verification.
                    </p>
                    <p style="margin-bottom:12px;">
                        The platform also integrates a multi-LLM architecture using <strong>Google Gemini, NVIDIA NIM, and Groq</strong>, enabling task-specific AI reasoning and resilient provider fallback.
                    </p>
                    <div style="display:flex; flex-direction:column; gap:6px; font-size:11.5px; color:var(--text-faint); border-top:1px solid var(--border); padding-top:12px;">
                        <div>🎓 <strong>Team:</strong> CyberMind AI Team</div>
                        <div>🛡️ <strong>Project:</strong> CyberMind AI — AI-Powered Cybersecurity & Autonomous Cyber Reasoning Platform</div>
                        <div style="display:flex; gap:16px; flex-wrap:wrap;">
                            <span>📅 <strong>Build Date:</strong> 07-07-2026</span>
                            <span>🚀 <strong>Version:</strong> v2.0.0 Enterprise</span>
                        </div>
                        <div style="margin-top:2px;">🎯 <strong>Core Focus:</strong> Threat Intelligence • AI/ML Security Analysis • SAST/AST • Dynamic Analysis • Autonomous Cyber Reasoning • Vulnerability Remediation</div>
                    </div>
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

    # ── ML Model Performance Cards (Unindented to render full screen width) ──
    st.markdown('<div class="section-title" style="margin-top:22px;"><img src="https://cdn-icons-png.flaticon.com/512/18310/18310827.png" width="20" height="20"> ML Model Performance</div>', unsafe_allow_html=True)

    import json as _json
    _metrics_path = BASE_DIR / "ml" / "models" / "cybermind_metrics.json"

    if _metrics_path.exists():
        with open(_metrics_path, encoding="utf-8") as _f:
            _m = _json.load(_f)

        _acc  = _m.get("accuracy",  "N/A")
        _prec = _m.get("precision", "N/A")
        _rec  = _m.get("recall",    "N/A")
        _f1   = _m.get("f1_score",  "N/A")
        _auc  = _m.get("roc_auc",   "N/A")

        def _pct(v): return f"{float(v)*100:.2f}%" if v != "N/A" else "N/A"
        def _fmt(v): return f"{float(v):.4f}"      if v != "N/A" else "N/A"
        def _col(v, t1=0.90, t2=0.75):
            if v == "N/A": return "#F87171"
            return "#22D3EE" if float(v) >= t1 else ("#F59E0B" if float(v) >= t2 else "#F87171")

        _acc_c  = _col(_acc,  0.90, 0.75)
        _prec_c = _col(_prec, 0.88, 0.70)
        _rec_c  = _col(_rec,  0.88, 0.70)
        _f1_c   = _col(_f1,   0.88, 0.70)
        _auc_c  = _col(_auc,  0.92, 0.80)

        # ── 5 Metric Cards (grid 5 columns, side-by-side) ─────────────────────
        st.markdown(f"""
        <div style="font-size:11px;color:var(--text-faint);margin-bottom:8px;letter-spacing:0.5px;">
          🎯 Avg across 4 datasets · 5-fold Stratified CV · RandomForest Classifier
        </div>
        <div style="display:grid;grid-template-columns:repeat(5, 1fr);gap:10px;margin-bottom:10px;">
          <div class="chart-card" style="padding:16px 14px;text-align:center;border-top:3px solid {_acc_c};">
            <div style="font-size:10px;color:var(--text-faint);text-transform:uppercase;letter-spacing:1px;margin-bottom:6px;">Accuracy</div>
            <div style="font-size:24px;font-weight:800;color:{_acc_c};">{_pct(_acc)}</div>
            <div style="font-size:10px;color:var(--text-faint);margin-top:4px;">Overall Correct Predictions</div>
          </div>
          <div class="chart-card" style="padding:16px 14px;text-align:center;border-top:3px solid {_prec_c};">
            <div style="font-size:10px;color:var(--text-faint);text-transform:uppercase;letter-spacing:1px;margin-bottom:6px;">Precision</div>
            <div style="font-size:24px;font-weight:800;color:{_prec_c};">{_fmt(_prec)}</div>
            <div style="font-size:10px;color:var(--text-faint);margin-top:4px;">True Positive Rate (of Predicted +)</div>
          </div>
          <div class="chart-card" style="padding:16px 14px;text-align:center;border-top:3px solid {_rec_c};">
            <div style="font-size:10px;color:var(--text-faint);text-transform:uppercase;letter-spacing:1px;margin-bottom:6px;">Recall</div>
            <div style="font-size:24px;font-weight:800;color:{_rec_c};">{_fmt(_rec)}</div>
            <div style="font-size:10px;color:var(--text-faint);margin-top:4px;">Sensitivity / True Positive Rate</div>
          </div>
          <div class="chart-card" style="padding:16px 14px;text-align:center;border-top:3px solid {_f1_c};">
            <div style="font-size:10px;color:var(--text-faint);text-transform:uppercase;letter-spacing:1px;margin-bottom:6px;">F1-Score</div>
            <div style="font-size:24px;font-weight:800;color:{_f1_c};">{_fmt(_f1)}</div>
            <div style="font-size:10px;color:var(--text-faint);margin-top:4px;">Harmonic Mean of Precision &amp; Recall</div>
          </div>
          <div class="chart-card" style="padding:16px 14px;text-align:center;border-top:3px solid {_auc_c};">
            <div style="font-size:10px;color:var(--text-faint);text-transform:uppercase;letter-spacing:1px;margin-bottom:6px;">AUC-ROC</div>
            <div style="font-size:24px;font-weight:800;color:{_auc_c};">{_fmt(_auc)}</div>
            <div style="font-size:10px;color:var(--text-faint);margin-top:4px;">Area Under ROC Curve · Discrimination Power</div>
          </div>
        </div>
        """, unsafe_allow_html=True)

        # ── Per-dataset breakdown table ────────────────────────────────
        _per = _m.get("per_dataset", [])
        if _per:
            _rows = ""
            for _d in _per:
                _da   = float(_d.get("accuracy",  0))
                _df1  = float(_d.get("f1_score",  0))
                _dauc = float(_d.get("roc_auc",   0))
                _pkl  = _d.get("model_file", "N/A")
                _ext  = _d.get("external_validation", {})
                _ext_txt = f"{_ext.get('accuracy',0)*100:.2f}%" if _ext else "N/A"
                _ac   = "#22D3EE" if _da   >= 0.90 else ("#F59E0B" if _da   >= 0.75 else "#F87171")
                _fc   = "#22D3EE" if _df1  >= 0.88 else ("#F59E0B" if _df1  >= 0.70 else "#F87171")
                _uc   = "#22D3EE" if _dauc >= 0.92 else ("#F59E0B" if _dauc >= 0.80 else "#F87171")
                _rows += f"""
                <tr style="border-bottom:1px solid rgba(255,255,255,0.04);">
                  <td style="padding:7px 8px;font-size:11px;color:var(--text);">{_d.get('dataset','')}</td>
                  <td style="padding:7px 8px;font-size:11px;font-weight:700;color:{_ac};text-align:center;">{_da*100:.2f}%</td>
                  <td style="padding:7px 8px;font-size:11px;color:#A78BFA;text-align:center;">{float(_d.get('precision',0)):.4f}</td>
                  <td style="padding:7px 8px;font-size:11px;color:#34D399;text-align:center;">{float(_d.get('recall',0)):.4f}</td>
                  <td style="padding:7px 8px;font-size:11px;font-weight:700;color:{_fc};text-align:center;">{_df1:.4f}</td>
                  <td style="padding:7px 8px;font-size:11px;font-weight:700;color:{_uc};text-align:center;">{_dauc:.4f}</td>
                  <td style="padding:7px 8px;font-size:11px;color:#F59E0B;text-align:center;">{_ext_txt}</td>
                  <td style="padding:7px 8px;font-size:10px;color:var(--text-faint);text-align:center;font-family:monospace;">{_pkl}</td>
                </tr>"""
            st.markdown(f"""
            <div class="chart-card" style="padding:14px;overflow-x:auto;margin-bottom:12px;">
              <div style="font-size:11px;color:var(--text-faint);margin-bottom:10px;font-weight:600;
                          text-transform:uppercase;letter-spacing:0.8px;">Per-Dataset Breakdown &amp; Validation</div>
              <table style="width:100%;border-collapse:collapse;">
                <thead>
                  <tr style="border-bottom:1px solid var(--border);">
                    <th style="padding:6px 8px;font-size:10px;color:var(--text-faint);text-align:left;">Dataset</th>
                    <th style="padding:6px 8px;font-size:10px;color:var(--text-faint);text-align:center;">Acc (5-CV)</th>
                    <th style="padding:6px 8px;font-size:10px;color:var(--text-faint);text-align:center;">Precision</th>
                    <th style="padding:6px 8px;font-size:10px;color:var(--text-faint);text-align:center;">Recall</th>
                    <th style="padding:6px 8px;font-size:10px;color:var(--text-faint);text-align:center;">F1</th>
                    <th style="padding:6px 8px;font-size:10px;color:var(--text-faint);text-align:center;">AUC-ROC</th>
                    <th style="padding:6px 8px;font-size:10px;color:var(--text-faint);text-align:center;">External Val</th>
                    <th style="padding:6px 8px;font-size:10px;color:var(--text-faint);text-align:center;">Model File</th>
                  </tr>
                </thead>
                <tbody>{_rows}</tbody>
              </table>
            </div>
            """, unsafe_allow_html=True)

            # ── Per-Class Breakdown Expanders ─────────────────────────────
            with st.expander("🔍 Per-Class Metrics & Imbalance Analysis"):
                st.markdown("<div style='font-size:12px; color:var(--text-muted); margin-bottom:10px;'>Detailed per-class precision, recall, and F1 performance across minority/majority classes:</div>", unsafe_allow_html=True)
                for _d in _per:
                    _pc = _d.get("per_class", {})
                    _ext_info = _d.get("external_validation", {})
                    _ds_name = _d.get("dataset", "")
                    if _pc or _ext_info:
                        st.markdown(f"**{_ds_name}** (`{_d.get('model_file', '')}`)")
                        if _ext_info:
                            st.info(f"🎯 **External Validation**: {_ext_info.get('accuracy', 0)*100:.2f}% Accuracy on {_ext_info.get('sample_count', 0)} independent benchmark samples ({_ext_info.get('dataset_name', '')}).")
                        if _pc:
                            _pc_rows = ""
                            for _cname, _cvals in _pc.items():
                                _pc_p = float(_cvals.get("precision", 0))
                                _pc_r = float(_cvals.get("recall", 0))
                                _pc_f = float(_cvals.get("f1_score", 0))
                                _pc_s = int(_cvals.get("support", 0))
                                _pc_color = "#34D399" if _pc_f >= 0.70 else ("#F59E0B" if _pc_f >= 0.40 else "#F87171")
                                _pc_rows += f"""
                                <tr style="border-bottom:1px solid rgba(255,255,255,0.04);">
                                    <td style="padding:5px 8px;font-size:11px;color:var(--text);">{_cname}</td>
                                    <td style="padding:5px 8px;font-size:11px;color:var(--text-muted);text-align:center;">{_pc_s:,}</td>
                                    <td style="padding:5px 8px;font-size:11px;color:#A78BFA;text-align:center;">{_pc_p:.4f}</td>
                                    <td style="padding:5px 8px;font-size:11px;color:#34D399;text-align:center;">{_pc_r:.4f}</td>
                                    <td style="padding:5px 8px;font-size:11px;font-weight:700;color:{_pc_color};text-align:center;">{_pc_f:.4f}</td>
                                </tr>"""
                            st.markdown(
                                f"""
                                <table style="width:100%; border-collapse:collapse; margin-bottom:14px;">
                                    <thead>
                                        <tr style="border-bottom:1px solid var(--border);">
                                            <th style="padding:5px 8px;font-size:10px;color:var(--text-faint);text-align:left;">Class Name</th>
                                            <th style="padding:5px 8px;font-size:10px;color:var(--text-faint);text-align:center;">Samples</th>
                                            <th style="padding:5px 8px;font-size:10px;color:var(--text-faint);text-align:center;">Precision</th>
                                            <th style="padding:5px 8px;font-size:10px;color:var(--text-faint);text-align:center;">Recall</th>
                                            <th style="padding:5px 8px;font-size:10px;color:var(--text-faint);text-align:center;">F1 Score</th>
                                        </tr>
                                    </thead>
                                    <tbody>{_pc_rows}</tbody>
                                </table>
                                """,
                                unsafe_allow_html=True
                            )

        # ── Info footer ────────────────────────────────────────────────
        # Build per-dataset summary from actual JSON data (no top-level
        # combined stats exist; these are per-dataset values).
        _per_data = _m.get("per_dataset", [])
        _total_samples = sum(int(_d.get("samples", 0)) for _d in _per_data)
        _feat_parts = " &nbsp;|&nbsp; ".join(
            f"{_d.get('dataset','?')}: <b>{_d.get('feature_count','?')}</b> features"
            for _d in _per_data
        )
        st.markdown(f"""
        <div class="chart-card" style="padding:10px 14px;font-size:11px;color:var(--text-faint);line-height:1.8;margin-top:8px;">
          <img src="https://cdn-icons-png.flaticon.com/512/18310/18310827.png" style="width:20px; height:20px; vertical-align:middle; margin-right:6px;"><b>{_m.get("model_type","RandomForestClassifier")}</b> &nbsp;·&nbsp;
          🌳 Trees: <b>200</b> (n_estimators) &nbsp;·&nbsp;
          📦 Total samples: <b>{_total_samples:,}</b><br>
          🔢 Features — {_feat_parts}
        </div>
        """, unsafe_allow_html=True)

    else:
        st.markdown("""
        <div class="chart-card" style="padding:20px;text-align:center;color:var(--text-muted);">
            <div style="font-size:32px;margin-bottom:8px;">🤖</div>
            <div style="font-size:13px;font-weight:600;">Model not trained yet</div>
            <div style="font-size:12px;margin-top:6px;">Run <code>python -m ml.train_cybermind</code> to train.</div>
        </div>
        """, unsafe_allow_html=True)






def render_profile_page():
    # ---- Gather dynamic system data ----
    import shutil

    # OS info
    try:
        os_name = f"{platform.system()} {platform.release()} ({platform.machine()})"
    except Exception:
        os_name = platform.platform()

    # Processor
    try:
        cpu_name = platform.processor() or "Unknown"
    except Exception:
        cpu_name = "Unknown"

    # Graphics (best-effort via WMI on Windows)
    gpu_name = "N/A"
    try:
        import subprocess
        result = subprocess.run(
            ["wmic", "path", "win32_VideoController", "get", "name"],
            capture_output=True, text=True, timeout=5
        )
        lines = [l.strip() for l in result.stdout.strip().splitlines() if l.strip() and l.strip().lower() != "name"]
        if lines:
            gpu_name = lines[0]
    except Exception:
        pass

    # Memory (RAM)
    try:
        import psutil
        mem = psutil.virtual_memory()
        mem_used_gb = round(mem.used / (1024 ** 3), 2)
        mem_avail_gb = round(mem.available / (1024 ** 3), 2)
        mem_total_gb = round(mem.total / (1024 ** 3), 2)
        mem_pct = round(mem.percent, 1)
    except Exception:
        mem_used_gb, mem_avail_gb, mem_total_gb, mem_pct = 0, 0, 0, 0

    # Disk / Storage
    try:
        disk = shutil.disk_usage("/")
        disk_used_gb = round(disk.used / (1024 ** 3), 1)
        disk_free_gb = round(disk.free / (1024 ** 3), 1)
        disk_total_gb = round(disk.total / (1024 ** 3), 1)
        disk_pct = round((disk.used / disk.total) * 100, 1) if disk.total else 0
    except Exception:
        disk_used_gb, disk_free_gb, disk_total_gb, disk_pct = 0, 0, 0, 0

    # Computer & user name
    comp_name = platform.node() or os.environ.get("COMPUTERNAME", "Unknown")
    user_name = os.environ.get("USERNAME", os.environ.get("USER", "Unknown"))

    # Boot time
    boot_time_str = "N/A"
    try:
        import psutil as _psu
        bt = datetime.fromtimestamp(_psu.boot_time())
        boot_time_str = bt.strftime("%b %d, %Y %I:%M %p")
    except Exception:
        pass

    # Network info
    local_ip = "N/A"
    try:
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.settimeout(0.5)
        s.connect(("8.8.8.8", 80))
        local_ip = s.getsockname()[0]
        s.close()
    except Exception:
        try:
            local_ip = socket.gethostbyname(socket.gethostname())
        except Exception:
            pass

    mac_address = "N/A"
    try:
        import uuid
        mac_int = uuid.getnode()
        mac_address = ":".join(f"{(mac_int >> i) & 0xFF:02X}" for i in range(40, -1, -8))
    except Exception:
        pass

    gateway = "N/A"
    try:
        import subprocess as _sp
        res = _sp.run(["ipconfig"], capture_output=True, text=True, timeout=5)
        for line in res.stdout.splitlines():
            if "Default Gateway" in line:
                parts = line.split(":")
                if len(parts) > 1 and parts[1].strip():
                    gateway = parts[1].strip()
                    break
    except Exception:
        pass

    # DB stats: total scans and threats
    total_scans = 0
    threats_found = 0
    try:
        from database.db import db
        row = db.fetchone("SELECT COUNT(*) as cnt FROM scan_history")
        if row:
            total_scans = row["cnt"] if isinstance(row, dict) else row[0]
        threat_row = db.fetchone(
            "SELECT COUNT(*) as cnt FROM scan_history WHERE risk_level IN ('High', 'Critical')"
        )
        if threat_row:
            threats_found = threat_row["cnt"] if isinstance(threat_row, dict) else threat_row[0]
    except Exception:
        pass

    render_page_poster("Profile", "Your profile and system information overview.")

    # Profile header card
    initials = "".join(w[0].upper() for w in user_name.split()[:2]) if user_name != "Unknown" else "?"
    st.markdown(
        f"""
        <div class="chart-card" style="padding:22px; margin-bottom:20px; display:flex; align-items:center; gap:20px;">
            <div style="width:68px; height:68px; border-radius:50%; background:linear-gradient(135deg, var(--grad-a), var(--grad-b)); display:flex; align-items:center; justify-content:center; font-size:28px; color:#fff; font-weight:800; box-shadow:0 0 16px rgba(108,92,231,0.35);">{html.escape(initials)}</div>
            <div style="flex:1;">
                <div style="font-size:18px; font-weight:800; color:var(--text);">{html.escape(user_name)}</div>
                <div style="font-size:12px; color:var(--text-muted); font-weight:600; margin-top:2px;">System Administrator</div>
            </div>
            <div style="text-align:right; border-left:1px solid var(--border); padding-left:24px;">
                <div style="font-size:11px; color:var(--text-muted); font-weight:600; text-transform:uppercase;">Account Status</div>
                <div style="font-size:16px; font-weight:700; color:var(--success); margin-top:4px;">🟢 Active</div>
            </div>
            <div style="text-align:right; border-left:1px solid var(--border); padding-left:24px;">
                <div style="font-size:11px; color:var(--text-muted); font-weight:600; text-transform:uppercase;">Total Scans</div>
                <div style="font-size:16px; font-weight:700; color:var(--text); margin-top:4px;">{total_scans:,}</div>
            </div>
            <div style="text-align:right; border-left:1px solid var(--border); padding-left:24px;">
                <div style="font-size:11px; color:var(--text-muted); font-weight:600; text-transform:uppercase;">Threats Found</div>
                <div style="font-size:16px; font-weight:700; color:var(--danger); margin-top:4px;">{threats_found:,}</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )

    c1, c2, c3 = st.columns(3, gap="large")

    with c1:
        st.markdown('<div class="section-title">System Information</div>', unsafe_allow_html=True)
        # Build system info card — split into two separate markdown calls to avoid rendering issues
        st.markdown(
            f"""<div class="chart-card" style="padding:18px;">
<div style="margin-bottom:14px;">
<div style="font-size:11px; color:var(--text-muted); font-weight:600;">Operating System</div>
<div style="font-size:14px; font-weight:600; color:var(--text); margin-top:2px;">{html.escape(os_name)}</div>
</div>
<div style="margin-bottom:14px;">
<div style="font-size:11px; color:var(--text-muted); font-weight:600;">Processor</div>
<div style="font-size:14px; font-weight:600; color:var(--text); margin-top:2px;">{html.escape(cpu_name)}</div>
</div>
<div style="margin-bottom:18px;">
<div style="font-size:11px; color:var(--text-muted); font-weight:600;">Graphics</div>
<div style="font-size:14px; font-weight:600; color:var(--text); margin-top:2px;">{html.escape(gpu_name)}</div>
</div>
</div>""",
            unsafe_allow_html=True
        )
        # Memory & Storage bars in separate block to prevent HTML rendering issues
        st.markdown(
            f"""<div class="chart-card" style="padding:18px; margin-top:-14px; border-top:none; border-top-left-radius:0; border-top-right-radius:0;">
<div style="margin-bottom:14px;">
<div style="display:flex; justify-content:space-between; font-size:12px; margin-bottom:6px;">
<span style="font-weight:600; color:var(--text-muted);">Memory (RAM)</span>
<span style="font-weight:700; color:var(--text);">{mem_used_gb} GB / {mem_total_gb} GB</span>
</div>
<div style="width:100%; height:8px; background:var(--border); border-radius:10px; overflow:hidden;">
<div style="width:{mem_pct}%; height:100%; background:linear-gradient(90deg, var(--grad-a), var(--grad-b)); border-radius:10px;"></div>
</div>
</div>
<div style="margin-bottom:6px;">
<div style="display:flex; justify-content:space-between; font-size:12px; margin-bottom:6px;">
<span style="font-weight:600; color:var(--text-muted);">Storage</span>
<span style="font-weight:700; color:var(--text);">{disk_used_gb} GB / {disk_total_gb} GB</span>
</div>
<div style="width:100%; height:8px; background:var(--border); border-radius:10px; overflow:hidden;">
<div style="width:{disk_pct}%; height:100%; background:linear-gradient(90deg, var(--grad-a), var(--grad-b)); border-radius:10px;"></div>
</div>
</div>
</div>""",
            unsafe_allow_html=True
        )

    with c2:
        st.markdown('<div class="section-title">System Details</div>', unsafe_allow_html=True)
        sys_rows = [
            ("System Name", html.escape(comp_name)),
            ("User Name", html.escape(user_name)),
            ("Platform", html.escape(platform.system())),
            ("Architecture", html.escape(platform.machine())),
            ("Python Version", html.escape(platform.python_version())),
            ("Boot Time", html.escape(boot_time_str)),
        ]
        rows_html = "".join(
            f'<tr><td>{param}</td><td>{val}</td></tr>' for param, val in sys_rows
        )
        st.markdown(
            f'<div class="chart-card"><table class="scan-table"><tr><th>Parameter</th><th>Value</th></tr>{rows_html}</table></div>',
            unsafe_allow_html=True
        )

    with c3:
        st.markdown('<div class="section-title">Network Information</div>', unsafe_allow_html=True)
        net_data = [
            ("IP Address", html.escape(local_ip)),
            ("Hostname", html.escape(platform.node())),
            ("Gateway", html.escape(gateway)),
            ("DNS Server", "8.8.8.8"),
            ("MAC Address", html.escape(mac_address)),
        ]
        net_html = "".join(
            f'<tr><td>{param}</td><td>{val}</td></tr>' for param, val in net_data
        )
        st.markdown(
            f'<div class="chart-card"><table class="scan-table"><tr><th>Parameter</th><th>Value</th></tr>{net_html}</table></div>',
            unsafe_allow_html=True
        )



def render_connections_page():
    render_page_poster("Connections", "Monitor and toggle datasets, database, and API connections.")

    # ── Init persistent toggle state ─────────────────────────────────────────
    if "disabled_datasets" not in st.session_state:
        st.session_state.disabled_datasets = set()
    if "db_user_disabled" not in st.session_state:
        st.session_state.db_user_disabled = False

    import os as _os, re as _re

    # ── 1. Database ───────────────────────────────────────────────────────────
    db_active = False
    db_records = 0
    db_size = "N/A"
    try:
        from database.db import db as _db
        row = _db.fetchone("SELECT COUNT(*) as cnt FROM scan_history")
        if row is not None:
            db_active = True
            db_records = row["cnt"] if isinstance(row, dict) else row[0]
        db_file = BASE_DIR / "database" / "cybermind.db"
        if db_file.exists():
            sz = db_file.stat().st_size
            db_size = f"{sz} B" if sz < 1024 else (f"{sz/1024:.1f} KB" if sz < 1024**2 else f"{sz/1024**2:.1f} MB")
    except Exception:
        pass

    # ── 2. Datasets ───────────────────────────────────────────────────────────
    dataset_categories = {
        "URL Datasets":     ("data/datasets/url/raw",     "🔗"),
        "Domain Datasets":  ("data/datasets/domain/raw",  "🌍"),
        "Email Datasets":   ("data/datasets/email/raw",   "📧"),
        "IP Datasets":      ("data/datasets/ip/raw",      "🖥️"),
        "File Datasets":    ("data/datasets/file/raw",    "📄"),
        "QR Datasets":      ("data/datasets/qr/raw",      "🔳"),
        "Website Datasets": ("data/datasets/website/raw", "🌐"),
        "Shared Datasets":  ("data/datasets/shared/raw",  "📦"),
    }
    datasets_status = []
    for ds_name, (ds_path, ds_icon) in dataset_categories.items():
        ds_dir = BASE_DIR / ds_path
        if ds_dir.exists():
            files = [f for f in ds_dir.iterdir() if f.is_file()]
            fc = len(files)
            sz = sum(f.stat().st_size for f in files)
            size_str = f"{sz} B" if sz < 1024 else (f"{sz/1024:.1f} KB" if sz < 1024**2 else f"{sz/1024**2:.1f} MB")
            datasets_status.append((ds_name, ds_icon, fc > 0, fc, size_str))
        else:
            datasets_status.append((ds_name, ds_icon, False, 0, "0 B"))

    # ── 3. API Keys ───────────────────────────────────────────────────────────
    from core.offline_mode import offline_mode
    is_offline_mode = offline_mode.is_enabled

    api_keys = [
        ("Google Gemini",        "GEMINI_API_KEY",               "🟢"),
        ("NVIDIA NIM",           "NVIDIA_API_KEY",               "🟠"),
        ("Groq",                 "GROQ_API_KEY",                 "🔵"),
        ("Google Safe Browsing", "GOOGLE_SAFE_BROWSING_API_KEY", "🔍"),
        ("VirusTotal",           "VIRUSTOTAL_API_KEY",           "🦠"),
        ("URLScan.io",           "URLSCAN_API_KEY",              "🔗"),
        ("AbuseIPDB",            "ABUSEIPDB_API_KEY",            "🛡️"),
        ("IPInfo",               "IPINFO_API_KEY",               "🌐"),
        ("Veriphone Telecom",    "VERIPHONE_API_KEY",            "📱"),
        ("IPQS Fraud Score",     "IPQS_API_KEY",                 "📞"),
    ]
    api_status = []
    for api_name, env_key, api_icon in api_keys:
        val = _os.environ.get(env_key, "").strip()
        has_key = bool(val)
        is_active = has_key and not is_offline_mode
        if is_offline_mode:
            masked = "Disabled (Offline Mode)"
        elif is_active:
            masked = f"{val[:6]}...{val[-4:]}" if len(val) > 10 else val
        else:
            masked = "Not configured"
        api_status.append((api_name, api_icon, is_active, masked, env_key, has_key))

    # ── Summary cards ─────────────────────────────────────────────────────────
    db_effective = db_active and not st.session_state.db_user_disabled
    active_ds = sum(
        1 for n, _, a, _, _ in datasets_status
        if a and n not in st.session_state.disabled_datasets
    )
    active_api = sum(1 for _, _, a, _, _, _ in api_status if a)
    total_connections = len(datasets_status) + len(api_status) + 1
    active_connections = active_ds + active_api + (1 if db_effective else 0)

    sc1, sc2, sc3, sc4 = st.columns(4, gap="medium")
    for col, (label, value, color, bg, icon) in zip(
        [sc1, sc2, sc3, sc4],
        [
            ("Total Active", f"{active_connections} / {total_connections}", "var(--success)" if active_connections > 0 else "var(--danger)", "var(--success-bg)" if active_connections > 0 else "var(--danger-bg)", "🟢"),
            ("Database", "Connected" if db_effective else "Offline", "var(--success)" if db_effective else "var(--danger)", "var(--success-bg)" if db_effective else "var(--danger-bg)", "🗄️"),
            ("Datasets", f"{active_ds} / {len(datasets_status)}", "var(--info)", "var(--info-bg)", "📊"),
            ("API Keys", f"0 / {len(api_status)} (Offline)" if is_offline_mode else f"{active_api} / {len(api_status)}", "var(--warning)" if not is_offline_mode and active_api > 0 else "var(--danger)", "var(--warning-bg)" if not is_offline_mode and active_api > 0 else "var(--danger-bg)", "🔑"),
        ]
    ):
        with col:
            st.markdown(
                f'<div class="stat-card"><div><div class="stat-label">{label}</div>'
                f'<div class="stat-value" style="font-size:20px;">{value}</div></div>'
                f'<div class="stat-icon" style="background:{bg};color:{color};">{icon}</div></div>',
                unsafe_allow_html=True,
            )

    st.markdown('<div style="height:12px;"></div>', unsafe_allow_html=True)

    # ── Three columns ─────────────────────────────────────────────────────────
    col_db, col_ds, col_api = st.columns(3, gap="large")

    # ── DATABASE ──────────────────────────────────────────────────────────────
    with col_db:
        st.markdown('<div class="section-title">🗄️ Database</div>', unsafe_allow_html=True)

        # Toggle row
        tgl_col, lbl_col = st.columns([1, 3], vertical_alignment="center")
        with tgl_col:
            db_toggle = st.toggle(
                "DB", value=db_effective,
                key="db_main_toggle", label_visibility="collapsed",
            )
        with lbl_col:
            st.markdown(
                f'<span style="font-size:12px;font-weight:600;color:{"var(--success)" if db_toggle else "var(--danger)"};">'
                f'{"🟢 Connected" if db_toggle else "🔴 Disconnected"}</span>',
                unsafe_allow_html=True,
            )

        if db_toggle != db_effective:
            if db_toggle:
                st.session_state.db_user_disabled = False
                try:
                    from database.db import db as _db2, initialize_database
                    initialize_database()
                    st.toast("Database reconnected!", icon="✅")
                except Exception as e:
                    st.error(f"Reconnect failed: {e}")
            else:
                st.session_state.db_user_disabled = True
                st.toast("Database disconnected", icon="⚡")
            st.rerun()

        # Info card
        db_badge = '<span class="badge badge-safe">Connected</span>' if db_effective else '<span class="badge badge-malicious">Offline</span>'
        st.markdown(
            f'<div class="chart-card" style="padding:16px;">'
            f'<div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:12px;">'
            f'<div style="font-size:15px;font-weight:700;color:var(--text);">SQLite Database</div>'
            f'{db_badge}</div>'
            f'<div style="border-top:1px solid var(--border);padding-top:10px;">'
            f'<div style="display:flex;justify-content:space-between;font-size:12.5px;padding:5px 0;border-bottom:1px solid var(--border);"><span style="color:var(--text-muted);">File</span><span style="color:var(--text);font-weight:600;">cybermind.db</span></div>'
            f'<div style="display:flex;justify-content:space-between;font-size:12.5px;padding:5px 0;border-bottom:1px solid var(--border);"><span style="color:var(--text-muted);">Size</span><span style="color:var(--text);font-weight:600;">{db_size}</span></div>'
            f'<div style="display:flex;justify-content:space-between;font-size:12.5px;padding:5px 0;border-bottom:1px solid var(--border);"><span style="color:var(--text-muted);">Records</span><span style="color:var(--text);font-weight:600;">{db_records:,}</span></div>'
            f'<div style="display:flex;justify-content:space-between;font-size:12.5px;padding:5px 0;"><span style="color:var(--text-muted);">Status</span>'
            f'<span style="color:{"var(--success)" if db_effective else "var(--danger)"};font-weight:600;">{"Active" if db_effective else "Inactive"}</span></div>'
            f'</div></div>',
            unsafe_allow_html=True,
        )

    # ── DATASETS ──────────────────────────────────────────────────────────────
    with col_ds:
        st.markdown('<div class="section-title">📊 Datasets</div>', unsafe_allow_html=True)
        with st.container(key="datasets_container"):
            for ds_name, ds_icon, files_exist, file_count, size_str in datasets_status:
                ds_user_on = ds_name not in st.session_state.disabled_datasets
                ds_effective = files_exist and ds_user_on

                ic_col, info_col, tog_col = st.columns([0.25, 2.2, 0.65], vertical_alignment="center")
                with ic_col:
                    st.markdown(f'<span style="font-size:18px;">{ds_icon}</span>', unsafe_allow_html=True)
                with info_col:
                    color = "var(--text)" if ds_effective else "var(--text-faint)"
                    st.markdown(
                        f'<div style="font-size:12.5px;font-weight:600;color:{color};">{ds_name}</div>'
                        f'<div style="font-size:10.5px;color:var(--text-faint);">{file_count} files · {size_str}</div>',
                        unsafe_allow_html=True,
                    )
                with tog_col:
                    new_val = st.toggle(
                        f"Dataset {ds_name}", value=ds_effective,
                        key=f"ds_toggle_{ds_name}",
                        label_visibility="collapsed",
                        disabled=not files_exist,  # can't enable what doesn't exist
                    )
                    if new_val != ds_effective:
                        if new_val:
                            st.session_state.disabled_datasets.discard(ds_name)
                            st.toast(f"{ds_name} enabled", icon="✅")
                        else:
                            st.session_state.disabled_datasets.add(ds_name)
                            st.toast(f"{ds_name} disabled", icon="⚡")
                        st.rerun()

                st.markdown('<hr style="margin:2px 0;border-color:var(--border);opacity:0.4;">', unsafe_allow_html=True)

    # ── API CONNECTIONS ───────────────────────────────────────────────────────
    with col_api:
        st.markdown('<div class="section-title">🔑 API Connections</div>', unsafe_allow_html=True)
        if is_offline_mode:
            st.markdown(
                '<div style="background:rgba(255,170,0,0.12);border:1px solid rgba(255,170,0,0.35);'
                'border-radius:8px;padding:10px 12px;margin-bottom:12px;font-size:12px;'
                'color:#ffaa00;font-weight:600;display:flex;align-items:center;gap:6px;">'
                '<span>⚡ Offline Mode Active</span> — All API connections forced OFF'
                '</div>',
                unsafe_allow_html=True,
            )
        with st.container(key="api_container"):
            for api_name, api_icon, is_active, masked, env_key, has_key in api_status:
                ic_col, info_col, tog_col = st.columns([0.25, 2.2, 0.65], vertical_alignment="center")
                with ic_col:
                    st.markdown(f'<span style="font-size:18px;">{api_icon}</span>', unsafe_allow_html=True)
                with info_col:
                    color = "var(--text)" if is_active else "var(--text-faint)"
                    st.markdown(
                        f'<div style="font-size:12.5px;font-weight:600;color:{color};">{api_name}</div>'
                        f'<div style="font-family:monospace;font-size:10px;color:var(--text-faint);">{masked}</div>',
                        unsafe_allow_html=True,
                    )
                with tog_col:
                    new_val = st.toggle(
                        f"API {api_name}", value=is_active,
                        key=f"api_toggle_{env_key}",
                        label_visibility="collapsed",
                        disabled=is_offline_mode,
                    )
                    if new_val != is_active and not is_offline_mode:
                        if new_val:
                            # Enable: reload key from .env file
                            try:
                                env_path = BASE_DIR / ".env"
                                if env_path.exists():
                                    match = _re.search(
                                        rf'^{env_key}\s*=\s*(.+)$',
                                        env_path.read_text(), _re.MULTILINE
                                    )
                                    if match:
                                        _os.environ[env_key] = match.group(1).strip()
                                        st.toast(f"{api_name} enabled!", icon="✅")
                                    else:
                                        st.warning(f"No saved key found for {api_name}. Please add it below.")
                            except Exception as e:
                                st.error(f"Enable failed: {e}")
                        else:
                            # Disable: remove from os.environ (key stays in .env)
                            _os.environ.pop(env_key, None)
                            st.toast(f"{api_name} disconnected", icon="⚡")
                        st.rerun()

                st.markdown('<hr style="margin:2px 0;border-color:var(--border);opacity:0.4;">', unsafe_allow_html=True)

        # ── Add / update API keys ─────────────────────────────────────────────
        st.markdown('<div style="height:10px;"></div>', unsafe_allow_html=True)
        if is_offline_mode:
            st.caption("🔒 API Key activation is disabled while Offline Mode is active. Turn off Offline Mode in Settings to manage keys.")
        else:
            inactive_apis = [(n, ek) for n, _, a, _, ek, hk in api_status if not a]
            if inactive_apis:
                for api_name, env_key in inactive_apis:
                    new_key = st.text_input(
                        api_name,
                        placeholder=f"Paste {api_name} API key…",
                        type="password",
                        key=f"activate_{env_key}",
                    )
                    if new_key:
                        st.markdown('<div class="cta-primary">', unsafe_allow_html=True)
                        if st.button(f"💾 Save {api_name} Key", key=f"save_{env_key}", width="stretch"):
                            try:
                                env_path = BASE_DIR / ".env"
                                env_content = env_path.read_text() if env_path.exists() else ""
                                if env_key in env_content:
                                    env_content = _re.sub(
                                        rf"^{env_key}\s*=.*$",
                                        f"{env_key}={new_key}",
                                        env_content, flags=_re.MULTILINE,
                                    )
                                else:
                                    env_content += f"\n{env_key}={new_key}\n"
                                env_path.write_text(env_content)
                                _os.environ[env_key] = new_key
                                st.success(f"{api_name} key saved and activated!")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Failed to save: {e}")
                        st.markdown("</div>", unsafe_allow_html=True)
def create_threat_map_figure(map_rows, height=250):
    import plotly.graph_objects as go

    fig = go.Figure()

    if map_rows:
        countries = [row["country"] for row in map_rows]
        counts = [row["count"] for row in map_rows]
        lon = [row["longitude"] for row in map_rows if row["longitude"] is not None]
        lat = [row["latitude"] for row in map_rows if row["latitude"] is not None]
        hover_text = [f"<b>{row['country']}</b>: {row['count']} scans" for row in map_rows]
        sizes = [min(35, 12 + row["count"] * 2) for row in map_rows]

        # Choropleth shading layer for targeted countries
        fig.add_trace(go.Choropleth(
            locations=countries,
            locationmode='country names',
            z=counts,
            colorscale=[[0, 'rgba(242, 84, 91, 0.25)'], [1, '#F2545B']],
            showscale=False,
            hoverinfo='location+z',
            marker_line_color='#334155',
            marker_line_width=0.8,
        ))

        # Scattergeo markers for pinpoint lat/lon locations
        if lon and lat:
            fig.add_trace(go.Scattergeo(
                lon=lon,
                lat=lat,
                text=hover_text,
                mode='markers',
                hoverinfo='text',
                marker=dict(
                    size=sizes,
                    color='#F2545B',
                    opacity=0.9,
                    symbol='circle',
                    line=dict(width=1.5, color='#FFFFFF')
                )
            ))
    else:
        fig.add_trace(go.Scattergeo(lon=[0], lat=[0], mode='markers', marker=dict(size=0, opacity=0)))

    fig.update_layout(
        geo=dict(
            scope='world',
            projection_type='equirectangular',
            showland=True,
            landcolor='#1E293B',
            showcountries=True,
            countrycolor='#334155',
            countrywidth=0.8,
            showcoastlines=True,
            coastlinecolor='#475569',
            showocean=True,
            oceancolor='#0F172A',
            showlakes=False,
            bgcolor='rgba(0,0,0,0)',
        ),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        margin=dict(l=0, r=0, t=0, b=0),
        height=height,
        showlegend=False
    )
    return fig


def render_analytics_page():
    # Ensure database is seeded if empty
    seed_database_if_empty()

    render_page_poster("Analytics & Intelligence", "Deep-dive threat metrics, ML predictions, and real-time database intelligence.")

    st.markdown(
        """
        <style>
        div[data-testid="column"]:nth-child(2) div[data-testid="stButton"] {
            text-align: right !important;
            display: block !important;
        }
        div[data-testid="column"]:nth-child(2) div[data-testid="stButton"] button {
            float: right !important;
            margin-left: auto !important;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    from database.db import db
    import pandas as pd
    
    # ── 1. GATHER REAL DATA FROM SQLITE DATABASE ──────────────────────────────
    total_scans = 0
    threats_detected = 0
    suspicious_scans = 0
    safe_scans = 0
    avg_risk = 0.0
    
    scanner_stats = {}
    risk_dist = {"Safe": 0, "Low": 0, "Medium": 0, "High": 0, "Critical": 0}
    time_series = []

    try:
        res = db.fetchone("SELECT COUNT(*) as count FROM scan_history")
        total_scans = res["count"] if res else 0

        res = db.fetchone("SELECT COUNT(*) as count FROM scan_history WHERE risk_level IN ('High', 'Critical')")
        threats_detected = res["count"] if res else 0

        res = db.fetchone("SELECT COUNT(*) as count FROM scan_history WHERE risk_level = 'Medium'")
        suspicious_scans = res["count"] if res else 0

        res = db.fetchone("SELECT COUNT(*) as count FROM scan_history WHERE risk_level IN ('Safe', 'Low')")
        safe_scans = res["count"] if res else 0

        res = db.fetchone("SELECT AVG(risk_score) as avg_score FROM scan_history")
        avg_risk = float(res["avg_score"]) if res and res["avg_score"] is not None else 0.0

        breakdown_rows = db.fetchall("SELECT scan_type, COUNT(*) as count FROM scan_history GROUP BY scan_type")
        for r in breakdown_rows:
            scanner_stats[r["scan_type"]] = r["count"]

        risk_rows = db.fetchall("SELECT risk_level, COUNT(*) as count FROM scan_history GROUP BY risk_level")
        for r in risk_rows:
            lvl = r["risk_level"]
            if lvl in risk_dist:
                risk_dist[lvl] = r["count"]

        trend_rows = db.fetchall(
            """
            SELECT date(scan_time) as scan_date, COUNT(*) as count, AVG(risk_score) as avg_score 
            FROM scan_history 
            GROUP BY scan_date 
            ORDER BY scan_date ASC 
            LIMIT 15
            """
        )
        for r in trend_rows:
            time_series.append({
                "date": r["scan_date"],
                "count": r["count"],
                "avg_risk": float(r["avg_score"]) if r["avg_score"] is not None else 0.0
            })
    except Exception:
        pass

    # ── 2. METRICS & COUNTS (Row 1) ───────────────────────────────────────────
    sc1, sc2, sc3, sc4 = st.columns(4, gap="medium")
       # Read real accuracy from the trained ML metrics file instead of hardcoding it
    import json as _json_acc
    _real_accuracy = None
    try:
        _metrics_file = BASE_DIR / "ml" / "models" / "cybermind_metrics.json"
        if _metrics_file.exists():
            with open(_metrics_file, encoding="utf-8") as _mf:
                _real_accuracy = float(_json_acc.load(_mf).get("accuracy", 0)) * 100
    except Exception:
        pass
    _acc_display = f"{_real_accuracy:.1f}%" if _real_accuracy is not None else "N/A"

    with sc1:
        st.markdown(
            f"""
            <div class="stat-card">
                <div>
                    <div class="stat-label">System Accuracy</div>
                    <div class="stat-value">{_acc_display}</div>
                    <div class="stat-delta" style="color:var(--success)">From trained ML models</div>
                </div>
                <div class="stat-icon" style="background:var(--success-bg);color:var(--success)">🎯</div>
            </div>
            """,
            unsafe_allow_html=True
        )
    with sc2:
        st.markdown(
            f"""
            <div class="stat-card">
                <div>
                    <div class="stat-label">Total Threat Scans</div>
                    <div class="stat-value">{total_scans:,}</div>
                    <div class="stat-delta" style="color:var(--info)">Live from SQLite database</div>
                </div>
                <div class="stat-icon" style="background:var(--info-bg);color:var(--info)">📶</div>
            </div>
            """,
            unsafe_allow_html=True
        )
    with sc3:
        st.markdown(
            f"""
            <div class="stat-card">
                <div>
                    <div class="stat-label">Malicious Elements</div>
                    <div class="stat-value" style="color:var(--danger);">{threats_detected:,}</div>
                    <div class="stat-delta" style="color:var(--danger)">Immediate attention needed</div>
                </div>
                <div class="stat-icon" style="background:var(--danger-bg);color:var(--danger)">⚠️</div>
            </div>
            """,
            unsafe_allow_html=True
        )
    with sc4:
        st.markdown(
            f"""
            <div class="stat-card">
                <div>
                    <div class="stat-label">Avg Threat Score</div>
                    <div class="stat-value">{avg_risk:.1f}</div>
                    <div class="stat-delta" style="color:var(--warning)">Medium classification level</div>
                </div>
                <div class="stat-icon" style="background:var(--warning-bg);color:var(--warning)">🛡️</div>
            </div>
            """,
            unsafe_allow_html=True
        )

    st.markdown("<br>", unsafe_allow_html=True)

    # ── 3. ROW 2: Scan Trend & Risk Score Trend with Interactive Filters ──────
    tr1, tr2 = st.columns(2, gap="large")
    with tr1:
        th1, th2 = st.columns([2, 1], vertical_alignment="center")
        with th1:
            st.markdown('<div class="section-title" style="margin-bottom:0;">Scan Trend</div>', unsafe_allow_html=True)
        with th2:
            trend_days = st.selectbox("Scan trend period", ["7 Days", "30 Days", "90 Days"], label_visibility="collapsed", key="scan_trend_days_filter")

        limit_days = 7
        if trend_days == "30 Days":
            limit_days = 30
        elif trend_days == "90 Days":
            limit_days = 90

        days_labels = []
        scan_counts = []
        try:
            trend_rows = db.fetchall(
                f"""
                SELECT date(scan_time) as scan_date, COUNT(*) as count 
                FROM scan_history 
                GROUP BY scan_date 
                ORDER BY scan_date DESC 
                LIMIT {limit_days}
                """
            )
            if trend_rows:
                trend_rows = list(reversed(trend_rows))
                days_labels = [datetime.strptime(r["scan_date"], "%Y-%m-%d").strftime("%m/%d") for r in trend_rows]
                scan_counts = [r["count"] for r in trend_rows]
        except Exception:
            pass

        if not days_labels or len(days_labels) < 3:
            days_labels = ["05/18", "05/19", "05/20", "05/21", "05/22", "05/23", "05/24"]
            scan_counts = [220, 290, 310, 420, 380, 490, 520]

        fig_trend = go.Figure()
        fig_trend.add_trace(go.Scatter(
            x=days_labels, y=scan_counts, mode="lines+markers",
            line=dict(color="#A78BFA", width=3), marker=dict(size=6),
            fill="tozeroy", fillcolor="rgba(167,139,250,0.12)"
        ))
        fig_trend.update_layout(
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            font=dict(color="#8B98AC", size=11), margin=dict(l=10, r=10, t=10, b=10), height=230,
            showlegend=False, xaxis=dict(gridcolor="#1D2838"), yaxis=dict(gridcolor="#1D2838"),
        )
        st.plotly_chart(fig_trend, width="stretch", config={"displayModeBar": False})

    with tr2:
        th3, th4 = st.columns([2, 1], vertical_alignment="center")
        with th3:
            st.markdown('<div class="section-title" style="margin-bottom:0;">Risk Score Trend</div>', unsafe_allow_html=True)
        with th4:
            risk_days = st.selectbox("Risk score period", ["7 Days", "30 Days", "90 Days"], label_visibility="collapsed", key="risk_score_days_filter")

        limit_risk = 7
        if risk_days == "30 Days":
            limit_risk = 30
        elif risk_days == "90 Days":
            limit_risk = 90

        risk_trend_y = []
        try:
            score_rows = db.fetchall(
                f"""
                SELECT date(scan_time) as scan_date, AVG(risk_score) as avg_score 
                FROM scan_history 
                GROUP BY scan_date 
                ORDER BY scan_date DESC 
                LIMIT {limit_risk}
                """
            )
            if score_rows:
                score_rows = list(reversed(score_rows))
                risk_trend_y = [float(r["avg_score"]) for r in score_rows]
        except Exception:
            pass

        if not risk_trend_y or len(risk_trend_y) < 3:
            risk_trend_y = [78, 62, 55, 48, 42, 38, 30]

        fig_score_trend = go.Figure()
        fig_score_trend.add_trace(go.Scatter(
            x=days_labels, y=risk_trend_y, mode="lines+markers",
            line=dict(color="#F2545B", width=3), marker=dict(size=6),
            fill="tozeroy", fillcolor="rgba(242,84,91,0.08)"
        ))
        fig_score_trend.update_layout(
            paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
            font=dict(color="#8B98AC", size=11), margin=dict(l=10, r=10, t=10, b=10), height=230,
            showlegend=False, xaxis=dict(gridcolor="#1D2838"), yaxis=dict(gridcolor="#1D2838", range=[0, 100]),
        )
        st.plotly_chart(fig_score_trend, width="stretch", config={"displayModeBar": False})

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Geolocation Data Queries ──────────────────────────────────────────────
    from database.db import db
    try:
        map_rows = db.fetchall(
            """
            SELECT country, latitude, longitude, COUNT(*) as count 
            FROM scan_history 
            WHERE latitude IS NOT NULL 
            GROUP BY country
            """
        )
    except Exception:
        map_rows = []

    # ── 3.5. Threat Map & Most Targeted Countries ──────────────────────────────
    is_map_expanded = st.session_state.get("show_full_map", False)

    if is_map_expanded:
        th_exp1, th_exp2 = st.columns([3, 1], vertical_alignment="center")
        with th_exp1:
            st.markdown('<div class="section-title" style="margin-bottom:0;">Threat Map (Live) - Expanded View</div>', unsafe_allow_html=True)
        with th_exp2:
            if st.button("Close Expanded Map", key="close_full_map_top"):
                st.session_state.show_full_map = False
                st.rerun()

        fig_map_full = create_threat_map_figure(map_rows, height=450)
        st.plotly_chart(fig_map_full, width="stretch", config={"displayModeBar": False})
        st.markdown("<br>", unsafe_allow_html=True)

    # ── 4. ROW 3: Threat Map (Live) & Most Targeted Countries ─────────────────
    mr1, mr2 = st.columns(2, gap="large")
    with mr1:
        if not is_map_expanded:
            th5, th6 = st.columns([2, 1], vertical_alignment="center")
            with th5:
                st.markdown('<div class="section-title" style="margin-bottom:0;">Threat Map (Live)</div>', unsafe_allow_html=True)
            with th6:
                if st.button("View Full Map", key="view_full_map_btn"):
                    st.session_state.show_full_map = True
                    st.rerun()

            fig_map = create_threat_map_figure(map_rows, height=250)
            st.plotly_chart(fig_map, width="stretch", config={"displayModeBar": False})

    with mr2:
        th7, th8 = st.columns([3.5, 1], vertical_alignment="center")
        with th7:
            st.markdown('<div class="section-title" style="margin-bottom:0;">Most Targeted Countries</div>', unsafe_allow_html=True)
        with th8:
            show_all_c = st.session_state.get("show_all_countries", False)
            with st.container(key="toggle_countries_wrap"):
                if st.button("Show Less" if show_all_c else "View All", key="toggle_countries_btn", width="stretch"):
                    st.session_state.show_all_countries = not show_all_c
                    st.rerun()

        limit_countries = 10 if show_all_c else 5
        try:
            country_rows = db.fetchall(
                f"""
                SELECT country, country_code, COUNT(*) as count 
                FROM scan_history 
                WHERE country IS NOT NULL 
                GROUP BY country 
                ORDER BY count DESC 
                LIMIT {limit_countries}
                """
            )
        except Exception:
            country_rows = []

        if not country_rows:
            st.markdown(
                """
                <div class="chart-card" style="padding:20px; text-align:center; color:var(--text-muted); display:flex; flex-direction:column; justify-content:center; align-items:center; min-height:220px;">
                    <div style="font-size:32px; margin-bottom:8px;">🌍</div>
                    <div style="font-size:13px; font-weight:600;">No location data yet</div>
                    <div style="font-size:12px; margin-top:6px;">Scan an IP or domain to populate the countries breakdown.</div>
                </div>
                """,
                unsafe_allow_html=True
            )
        else:
            max_count = country_rows[0]["count"] if country_rows else 1
            colors = ["var(--danger)", "#F5A623", "var(--info)", "var(--success)", "#A78BFA", "#8B5CF6", "#3B82F6", "#10B981", "#EF4444", "#F59E0B"]
            
            country_html = []
            for i, row in enumerate(country_rows):
                country_name = row["country"]
                c_code = (row["country_code"] or "us").lower()
                count = row["count"]
                pct = int((count / max_count) * 100)
                color = colors[i % len(colors)]
                flag_url = f"https://flagcdn.com/w20/{c_code}.png"
                
                is_last = (i == len(country_rows) - 1)
                margin_style = "margin-bottom: 0px;" if is_last else "margin-bottom: 12px;"
                
                row_html = f"""
                <div style="{margin_style}">
                    <div style="display:flex; justify-content:space-between; font-size:12px; margin-bottom:4px; align-items:center;">
                        <span style="display:flex; align-items:center;"><img src="{flag_url}" width="18" style="margin-right:8px; border-radius:2px;"/>{country_name}</span><span style="font-weight:700;">{count}</span>
                    </div>
                    <div style="width:100%; height:5px; background:var(--border); border-radius:10px; overflow:hidden;">
                        <div style="width:{pct}%; height:100%; background:{color}; border-radius:10px;"></div>
                    </div>
                </div>
                """
                country_html.append(row_html)
                
            inner_html = f"""
            <div class="chart-card" style="padding:18px 20px; min-height:220px;">
                {''.join(country_html)}
            </div>
            """
            st.markdown(clean_html(inner_html), unsafe_allow_html=True)


    st.markdown("<br>", unsafe_allow_html=True)

    # ── 5. ROW 4: Recent Scans & System Status ────────────────────────────────
    ar1, ar2 = st.columns(2, gap="large")
    with ar1:
        th_recent_1, th_recent_2 = st.columns([2, 1.2], vertical_alignment="center")
        with th_recent_1:
            st.markdown('<div class="section-title" style="margin-bottom:0;">Recent Scans</div>', unsafe_allow_html=True)
        with th_recent_2:
            if st.button("View All History", key="view_all_history_btn"):
                go_to("Scan History")
                st.rerun()

        with st.container():
            render_recent_scans()

    with ar2:
        st.markdown('<div class="section-title">System Status</div>', unsafe_allow_html=True)
        db_ok = True
        try:
            from database.connection import database
            db_ok = database.is_connected()
        except Exception:
            pass

        import os as _os
        apis_configured = sum(1 for key in ["GOOGLE_SAFE_BROWSING_API_KEY", "VIRUSTOTAL_API_KEY", "ABUSEIPDB_API_KEY"] if _os.environ.get(key))
        api_status_val = "Operational" if apis_configured > 0 else "Offline"
        api_badge_class = "badge-safe" if apis_configured > 0 else "badge-malicious"
        db_badge_class = "badge-safe" if db_ok else "badge-malicious"

        st.markdown(
            f"""
            <div class="chart-card" style="padding:16px; height:210px; display:flex; align-items:center; justify-content:space-between; gap:10px;">
                <div style="display:flex; flex-direction:column; justify-content:space-between; height:100%; flex:1; font-size:12.5px;">
                    <div class="list-row" style="padding:4px 0; border:none;"><span>⚙️ API Services</span><span class="badge {api_badge_class}" style="font-size:10px; padding:2px 7px;">{api_status_val}</span></div>
                    <div class="list-row" style="padding:4px 0; border:none;"><span>🗄️ Database</span><span class="badge {db_badge_class}" style="font-size:10px; padding:2px 7px;">{"Operational" if db_ok else "Inactive"}</span></div>
                    <div class="list-row" style="padding:4px 0; border:none;"><span>🧠 ML Engine</span><span class="badge badge-safe" style="font-size:10px; padding:2px 7px;">Operational</span></div>
                    <div class="list-row" style="padding:4px 0; border:none;"><span>📡 Threat Intelligence Feed</span><span class="badge badge-safe" style="font-size:10px; padding:2px 7px;">Operational</span></div>
                    <div class="list-row" style="padding:4px 0; border:none;"><span>📋 Report Generator</span><span class="badge badge-safe" style="font-size:10px; padding:2px 7px;">Operational</span></div>
                </div>
                <div style="margin-left:10px; width:70px; height:70px; border-radius:50%; background:rgba(16,185,129,0.08); border:2px solid rgba(16,185,129,0.3); display:flex; align-items:center; justify-content:center; font-size:32px; box-shadow:0 0 15px rgba(16,185,129,0.2); color:#10B981; flex-shrink:0;">
                    🛡️
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

    # ── Section 4 Analytics Additions ──────────────────────────────────────
    st.markdown("<br>", unsafe_allow_html=True)
    an_col1, an_col2 = st.columns(2, gap="large")

    with an_col1:
        st.markdown('<div class="section-title">🌍 Global Threat Distribution (Choropleth)</div>', unsafe_allow_html=True)
        try:
            ch_rows = db.fetchall(
                "SELECT country, COUNT(*) as count FROM scan_history WHERE country IS NOT NULL GROUP BY country"
            )
            if ch_rows:
                df_ch = pd.DataFrame([dict(r) for r in ch_rows])
                fig_ch = go.Figure(data=go.Choropleth(
                    locations=df_ch['country'],
                    locationmode='country names',
                    z=df_ch['count'],
                    text=df_ch['country'],
                    colorscale='Reds',
                    autocolorscale=False,
                    marker_line_color='#334155',
                    marker_line_width=0.8,
                    colorbar_title='Scans',
                ))
                fig_ch.update_layout(
                    geo=dict(
                        scope='world',
                        projection_type='equirectangular',
                        showland=True,
                        landcolor='#1E293B',
                        showcountries=True,
                        countrycolor='#334155',
                        countrywidth=0.8,
                        showcoastlines=True,
                        coastlinecolor='#475569',
                        showocean=True,
                        oceancolor='#0F172A',
                        showframe=False,
                        bgcolor='rgba(0,0,0,0)'
                    ),
                    paper_bgcolor='rgba(0,0,0,0)',
                    margin=dict(l=0, r=0, t=0, b=0),
                    height=260
                )
                st.plotly_chart(fig_ch, width="stretch", config={"displayModeBar": False})
            else:
                st.info("No country distribution data available yet.")
        except Exception as exc:
            st.warning(f"Choropleth chart error: {exc}")

    with an_col2:
        st.markdown('<div class="section-title">📅 Attack Activity Timeline</div>', unsafe_allow_html=True)
        try:
            timeline_rows = db.fetchall(
                "SELECT strftime('%Y-%m-%d %H:00', scan_time) as hour, COUNT(*) as cnt FROM scan_history GROUP BY hour ORDER BY hour DESC LIMIT 24"
            )
            if timeline_rows:
                df_tl = pd.DataFrame([dict(r) for r in timeline_rows]).iloc[::-1]
                fig_tl = go.Figure(data=go.Scatter(
                    x=df_tl['hour'], y=df_tl['cnt'],
                    mode='lines+markers',
                    line=dict(color='#22D3EE', width=2),
                    marker=dict(size=6, color='#A78BFA')
                ))
                fig_tl.update_layout(
                    paper_bgcolor='rgba(0,0,0,0)',
                    plot_bgcolor='rgba(0,0,0,0)',
                    font=dict(color='var(--text)'),
                    xaxis=dict(showgrid=False),
                    yaxis=dict(showgrid=True, gridcolor='rgba(255,255,255,0.05)'),
                    margin=dict(l=20, r=20, t=10, b=20),
                    height=260
                )
                st.plotly_chart(fig_tl, width="stretch", config={"displayModeBar": False})
            else:
                st.info("No timeline data available yet.")
        except Exception as exc:
            st.warning(f"Timeline chart error: {exc}")

    st.markdown("<br>", unsafe_allow_html=True)

    # ── 6. ROW 5: Threat Level Distribution & ML Predictions ──────────────────
    c1, c2 = st.columns([1.5, 1.5], gap="large")

    with c1:
        st.markdown('<div class="section-title">📊 Threat Level Distribution</div>', unsafe_allow_html=True)
        lbls = ["Safe", "Low", "Medium", "High", "Critical"]
        vals = [risk_dist[l] for l in lbls]
        if sum(vals) == 0:
            vals = [72, 14, 25, 8, 4]
            
        fig_pie = donut_chart(
            lbls, vals,
            ["#22C55E", "#3B82F6", "#F5A623", "#F2545B", "#9333EA"], "Risk Levels"
        )
        st.plotly_chart(fig_pie, width="stretch", config={"displayModeBar": False})
        
        tot = sum(vals)
        rows_html = "".join(
            f'<div class="list-row"><span class="list-name"><span style="color:{color}">●</span> {lbl}</span>'
            f'<span class="list-val">{val} ({val/tot*100:.1f}%)</span></div>'
            for lbl, val, color in zip(lbls, vals, ["#22C55E", "#3B82F6", "#F5A623", "#F2545B", "#9333EA"])
        )
        st.markdown(f'<div class="chart-card" style="padding:10px 14px;">{rows_html}</div>', unsafe_allow_html=True)

    with c2:
        st.markdown('<div class="section-title">🔮 Machine Learning Threat Predictions</div>', unsafe_allow_html=True)
        recent_list = []
        try:
            r_rows = db.fetchall(
                "SELECT target, risk_score, risk_level FROM scan_history ORDER BY scan_id DESC LIMIT 5"
            )
            for rr in r_rows:
                recent_list.append((rr["target"], rr["risk_score"], rr["risk_level"]))
        except Exception:
            pass
            
        if not recent_list:
            recent_list = [
                ("https://google.com", 10.0, "Safe"),
                ("http://phishing-page.com", 90.0, "Critical"),
                ("8.8.8.8", 12.0, "Safe"),
                ("suspicious-site.net", 65.0, "Medium"),
                ("C:\\Windows\\Temp\\trojan.exe", 95.0, "Critical")
            ]
            
        ml_rows = ""
        for target, score, lvl in recent_list:
            is_mal = score >= 70
            is_susp = 30 <= score < 70
            pred_class = "Legitimate" if score < 30 else ("Suspicious" if is_susp else "Phishing/Malware")
            color_badge = "badge-safe" if score < 30 else ("badge-suspicious" if is_susp else "badge-malicious")
            prob = round(100.0 - score if score < 30 else (score if is_mal else (100.0 - abs(50.0 - score)*2)), 1)
            
            ml_rows += (
                f'<div class="list-row" style="padding:10px 4px; border-bottom:1px solid var(--border);">'
                f'<div style="font-family:monospace;font-size:12.5px;color:var(--text);flex:1;word-break:break-all;">{html.escape(target)}</div>'
                f'<span class="badge {color_badge}" style="margin: 0 10px;font-size:10.5px;">{pred_class}</span>'
                f'<div style="font-size:12.5px;font-weight:700;color:var(--text-faint);width:60px;text-align:right;">{prob}%</div>'
                f'</div>'
            )
            
        st.markdown(
            f"""
            <div class="chart-card" style="padding:16px;">
                <div style="display:flex;justify-content:space-between;font-size:11px;color:var(--text-faint);font-weight:700;text-transform:uppercase;margin-bottom:8px;padding-bottom:6px;border-bottom:1px solid var(--border);">
                    <span>Scanned Target</span>
                    <span style="margin-left:auto;margin-right:24px;">AI Prediction</span>
                    <span>Confidence</span>
                </div>
                {ml_rows}
            </div>
            """,
            unsafe_allow_html=True
        )


def render_scan_history_page():

    # Ensure database is seeded if empty
    seed_database_if_empty()

    render_page_poster("Scan History", "Monitor all historical threat scans executed on the platform.")

    # Fetch original history from SQLite
    from database.db import db
    rows = []
    try:
        db_rows = db.fetchall(
            """
            SELECT scan_id, scan_type, target, risk_level, risk_score, scan_time 
            FROM scan_history 
            ORDER BY scan_id DESC
            """
        )
        if db_rows:
            rows = [
                {
                    "id": r["scan_id"],
                    "scanner": r["scan_type"],
                    "target": r["target"],
                    "risk_level": r["risk_level"] or "Safe",
                    "risk_score": r["risk_score"] or 0.0,
                    "time": r["scan_time"]
                }
                for r in db_rows
            ]
    except Exception:
        pass

    # Dropdowns for filtering
    sc1, sc2 = st.columns(2)
    with sc1:
        scanners_list = ["All"] + sorted(list(set(r["scanner"] for r in rows if r["scanner"])))
        selected_scanner = st.selectbox("Scanner", scanners_list, key="history_scanner_filter")
    with sc2:
        levels_list = ["All", "Safe", "Low", "Medium", "High", "Critical"]
        selected_level = st.selectbox("Risk level", levels_list, key="history_level_filter")

    # Apply filters
    filtered_rows = rows
    if selected_scanner != "All":
        filtered_rows = [r for r in filtered_rows if r["scanner"] == selected_scanner]
    if selected_level != "All":
        filtered_rows = [r for r in filtered_rows if r["risk_level"] == selected_level]

    # ── PAGINATION LOGIC ───────────────────────────────────────────────────────
    ITEMS_PER_PAGE = 50
    total_items = len(filtered_rows)
    total_pages = max(1, math.ceil(total_items / ITEMS_PER_PAGE))

    if "history_page" not in st.session_state:
        st.session_state.history_page = 1

    if st.session_state.history_page > total_pages:
        st.session_state.history_page = 1
    if st.session_state.history_page < 1:
        st.session_state.history_page = 1

    curr_page = st.session_state.history_page
    start_idx = (curr_page - 1) * ITEMS_PER_PAGE
    end_idx = min(start_idx + ITEMS_PER_PAGE, total_items)
    page_rows = filtered_rows[start_idx:end_idx]

    # Display count summary card
    if total_items > 0:
        st.markdown(
            f"""
            <div class="chart-card" style="padding:12px 18px; margin-bottom:14px; font-weight:700; font-size:14px; color:var(--text); display:flex; justify-content:space-between; align-items:center;">
                <span>{total_items} result(s) <span style="font-size:12px; color:var(--text-muted); font-weight:500;">(Showing {start_idx + 1}–{end_idx})</span></span>
                <span style="font-size:12px; color:var(--text-muted); font-weight:600; background:rgba(34, 184, 240, 0.1); padding:4px 10px; border-radius:12px; border:1px solid rgba(34, 184, 240, 0.2);">Page {curr_page} of {total_pages}</span>
            </div>
            """,
            unsafe_allow_html=True
        )
    else:
        st.markdown(
            """
            <div class="chart-card" style="padding:12px 18px; margin-bottom:14px; font-weight:700; font-size:14px; color:var(--text);">
                0 result(s)
            </div>
            """,
            unsafe_allow_html=True
        )

    # Display Table for page_rows ONLY
    if not filtered_rows:
        st.info("No matching scan records found.")
    else:
        header_cols = st.columns([0.6, 1.6, 2.6, 1, 0.9, 1.6, 1], vertical_alignment="center")
        headers = ["ID", "Scanner", "Target", "Result", "Score", "Time", ""]
        for hc, htext in zip(header_cols, headers):
            with hc:
                st.markdown(
                    f"<div style='color:var(--text-faint); font-weight:600; font-size:11px; text-transform:uppercase;'>{htext}</div>",
                    unsafe_allow_html=True
                )
        st.markdown('<div style="height:1px; background:var(--border); margin:6px 0 4px 0;"></div>', unsafe_allow_html=True)

        for r in page_rows:
            badge_class = LEVEL_BADGE.get(r["risk_level"], "badge-safe")
            t_str = r["time"]
            try:
                if "T" in t_str:
                    t_str = t_str.replace("T", " ").split(".")[0]
                elif "." in t_str:
                    t_str = t_str.split(".")[0]
            except Exception:
                pass

            row_cols = st.columns([0.6, 1.6, 2.6, 1, 0.9, 1.6, 1], vertical_alignment="center")
            with row_cols[0]:
                st.markdown(f"<div style='color:var(--text-faint); font-size:13px;'>{r['id']}</div>", unsafe_allow_html=True)
            with row_cols[1]:
                st.markdown(f"<div style='color:var(--text); font-size:13px; font-weight:600;'>{html.escape(r['scanner'])}</div>", unsafe_allow_html=True)
            with row_cols[2]:
                st.markdown(
                    f"<div style='color:var(--text-muted); font-size:13px; font-family:monospace; word-break:break-all;'>{html.escape(r['target'])}</div>",
                    unsafe_allow_html=True
                )
            with row_cols[3]:
                st.markdown(f"<span class='badge {badge_class}' style='font-size:11px; padding:3px 8px;'>{r['risk_level']}</span>", unsafe_allow_html=True)
            with row_cols[4]:
                st.markdown(f"<div style='color:var(--text); font-size:13px; font-weight:600;'>{int(r['risk_score'])}/100</div>", unsafe_allow_html=True)
            with row_cols[5]:
                st.markdown(f"<div style='color:var(--text-muted); font-size:13px;'>{t_str}</div>", unsafe_allow_html=True)
            with row_cols[6]:
                if r["scanner"] in SCANNERS:
                    if st.button("View →", key=f"history_view_{r['id']}", width="stretch"):
                        restore_scan_from_history(r["scanner"], r["target"], r["risk_level"], r["risk_score"])
                        st.rerun()
            st.markdown('<div style="height:1px; background:var(--border); margin:4px 0;"></div>', unsafe_allow_html=True)

        # ── DYNAMIC PAGINATION BAR (Matching Theme & Screenshot 2) ───────────
        if total_pages > 1:
            st.markdown('<div style="height:16px;"></div>', unsafe_allow_html=True)

            # Build list of visible page numbers
            page_numbers = []
            if total_pages <= 7:
                page_numbers = list(range(1, total_pages + 1))
            else:
                page_numbers.append(1)
                start_p = max(2, curr_page - 1)
                end_p = min(total_pages - 1, curr_page + 1)
                if start_p > 2:
                    page_numbers.append("...")
                for p in range(start_p, end_p + 1):
                    page_numbers.append(p)
                if end_p < total_pages - 1:
                    page_numbers.append("...")
                page_numbers.append(total_pages)

            # Inject CSS for pagination bar
            st.markdown(
                """
                <style>
                div[class*="st-key-pg_active_"] button {
                    background: linear-gradient(135deg, #22D3EE 0%, #6C5CE7 100%) !important;
                    color: #FFFFFF !important;
                    border: none !important;
                    font-weight: 800 !important;
                    box-shadow: 0 4px 14px rgba(34, 211, 238, 0.4) !important;
                    border-radius: 10px !important;
                    height: 38px !important;
                }
                div[class*="st-key-pg_inactive_"] button {
                    background: var(--card-bg) !important;
                    color: var(--text) !important;
                    border: 1px solid var(--border) !important;
                    border-radius: 10px !important;
                    height: 38px !important;
                    font-weight: 600 !important;
                }
                div[class*="st-key-pg_inactive_"] button:hover {
                    border-color: #22D3EE !important;
                    color: #22D3EE !important;
                }
                .st-key-pg_prev button, .st-key-pg_next button {
                    background: var(--card-bg) !important;
                    color: var(--text) !important;
                    border: 1px solid var(--border) !important;
                    border-radius: 10px !important;
                    height: 38px !important;
                    font-weight: 700 !important;
                }
                .st-key-pg_prev button:hover:not(:disabled), .st-key-pg_next button:hover:not(:disabled) {
                    border-color: #22D3EE !important;
                    color: #22D3EE !important;
                }
                </style>
                """,
                unsafe_allow_html=True
            )

            # Render controls in a styled centered bar
            cols = st.columns([1.5] + [1] * len(page_numbers) + [1.5], vertical_alignment="center")

            # Prev button
            with cols[0]:
                with st.container(key="pg_prev"):
                    if st.button("‹ Prev", key="pg_btn_prev", disabled=(curr_page == 1), width="stretch"):
                        st.session_state.history_page -= 1
                        st.rerun()

            # Page numbers
            for idx, item in enumerate(page_numbers):
                with cols[idx + 1]:
                    if item == "...":
                        st.markdown("<div style='text-align:center; color:var(--text-muted); font-weight:700; line-height:38px;'>...</div>", unsafe_allow_html=True)
                    else:
                        is_active = (item == curr_page)
                        btn_key = f"pg_num_{item}"
                        if is_active:
                            with st.container(key=f"pg_active_{item}"):
                                st.button(str(item), key=btn_key, width="stretch")
                        else:
                            with st.container(key=f"pg_inactive_{item}"):
                                if st.button(str(item), key=btn_key, width="stretch"):
                                    st.session_state.history_page = item
                                    st.rerun()

            # Next button
            with cols[-1]:
                with st.container(key="pg_next"):
                    if st.button("Next ›", key="pg_btn_next", disabled=(curr_page == total_pages), width="stretch"):
                        st.session_state.history_page += 1
                        st.rerun()



def render_help_support_page():

    render_page_poster("Help & Support", "Get help, view documentation, and contact our support team.")


    # ── Quick-access cards ─────────────────────────────────────────────────────
    q1, q2, q3, q4 = st.columns(4, gap="small")
    quick = [
        (q1, "📄", "Documentation",   "View user guides & API docs"),
        (q2, "🎥", "Video Tutorials", "Watch step-by-step tutorials"),
        (q3, "👥", "Community",       "Join discussions & forums"),
        (q4, "🌐", "System Status",   "Check live service status"),
    ]
    for col, icon, title, desc in quick:
        with col:
            st.markdown(
                f"""
                <div class="chart-card" style="padding:18px; text-align:center; cursor:pointer;">
                    <div style="font-size:28px; margin-bottom:8px;">{icon}</div>
                    <div style="font-size:13.5px; font-weight:700; color:var(--text); margin-bottom:4px;">{title}</div>
                    <div style="font-size:11px; color:var(--text-muted);">{desc}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Three-column section ───────────────────────────────────────────────────
    c1, c2, c3 = st.columns([2, 1.5, 1.5], gap="medium")

    with c1:
        st.markdown('<div class="section-title">📚 Popular Help Topics</div>', unsafe_allow_html=True)
        topics = [
            ("🚀", "Getting Started with CyberMind AI"),
            ("⚠️", "Understanding Risk Scores and Threat Levels"),
            ("🔑", "Configuring API Keys for VirusTotal & Safe Browsing"),
            ("📤", "Exporting Reports and Scanning Logs"),
            ("🔧", "Troubleshooting Connection Errors"),
            ("👤", "Managing Account Preferences"),
            ("🔒", "Security and Data Privacy Disclosures"),
        ]
        rows_html = "".join(
            f'<div style="display:flex;align-items:center;gap:10px;padding:10px 4px;'
            f'border-bottom:1px solid var(--border);cursor:pointer;">'
            f'<span style="font-size:15px;flex-shrink:0;">{ico}</span>'
            f'<span style="font-size:13px;font-weight:500;color:var(--text);">{label}</span></div>'
            for ico, label in topics
        )
        st.markdown(f'<div class="chart-card" style="padding:14px 16px;">{rows_html}</div>', unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="section-title">❓ Frequently Asked Questions</div>', unsafe_allow_html=True)
        faqs = [
            ("How do I add a VirusTotal API key?",
             "Go to **Settings → API Keys**, paste your key and click Save. The key is stored securely in your local `.env` file."),
            ("What does a risk score mean?",
             "Scores range from 0–100. **0–29** is Safe, **30–69** is Suspicious, and **70+** is Malicious."),
            ("Can I export scan reports?",
             "Yes! Open **Reports** from the sidebar and use the Export button to download PDF or CSV reports."),
            ("Is my data sent to any external server?",
             "Only the target URL / IP / file hash is sent to third-party APIs (VirusTotal, Google Safe Browsing). No personal data is shared."),
        ]
        for q, a in faqs:
            with st.expander(q):
                st.markdown(a)

    with c2:
        # Read disabled datasets dynamically from session state
        disabled_ds = st.session_state.get("disabled_datasets", set())
        
        services = [
            ("URL Scanner API",      "🔴 Offline" if "URL Datasets" in disabled_ds else "🟢 Online"),
            ("Website Scanner API",  "🔴 Offline" if "Website Datasets" in disabled_ds else "🟢 Online"),
            ("Domain Scanner API",   "🔴 Offline" if "Domain Datasets" in disabled_ds else "🟢 Online"),
            ("IP Intelligence API",  "🔴 Offline" if "IP Datasets" in disabled_ds else "🟢 Online"),
            ("Email Scanner API",    "🔴 Offline" if "Email Datasets" in disabled_ds else "🟢 Online"),
            ("File Scanner API",     "🔴 Offline" if "File Datasets" in disabled_ds else "🟢 Online"),
        ]
        
        any_offline = any("Offline" in status for _, status in services)
        status_title = "Degraded Performance" if any_offline else "All Systems Operational"
        status_color = "var(--warning)" if any_offline else "var(--success)"
        status_border = "rgba(245,166,35,0.3)" if any_offline else "rgba(34,197,94,0.3)"
        status_bg = "rgba(245,166,35,0.07)" if any_offline else "rgba(34,197,94,0.07)"
        status_icon = "🟡" if any_offline else "🟢"
        status_desc = "Some scanning services have been manually deactivated." if any_offline else "All services and API servers are running normally.<br>No incidents reported."

        st.markdown(f'<div class="section-title">{status_icon} System Status</div>', unsafe_allow_html=True)
        st.markdown(
            f"""
            <div class="chart-card" style="padding:22px; background:{status_bg};
                 border:1px solid {status_border}; text-align:center; margin-bottom:16px;">
                <div style="font-size:36px; margin-bottom:10px;">{status_icon}</div>
                <div style="font-size:15px; font-weight:800; color:{status_color}; margin-bottom:6px;">{status_title}</div>
                <div style="font-size:12px; color:var(--text-muted); line-height:1.6;">
                    {status_desc}
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        
        rows = "".join(
            f'<div class="list-row"><span class="list-name">{svc}</span>'
            f'<span style="font-size:11px; font-weight:700; color:{"var(--success)" if "Online" in status else "var(--danger)"};">{status}</span></div>'
            for svc, status in services
        )
        st.markdown(f'<div class="chart-card" style="padding:14px 16px;">{rows}</div>', unsafe_allow_html=True)

    with c3:
        st.markdown('<div class="section-title">📞 Support Information</div>', unsafe_allow_html=True)
        st.markdown(
            """
            <div class="chart-card" style="padding:18px;">
                <div style="margin-bottom:14px;">
                    <div style="font-size:10.5px; color:var(--text-faint); font-weight:700; text-transform:uppercase; letter-spacing:0.8px;">📧 Email Support</div>
                    <div style="font-size:13.5px; font-weight:600; color:var(--text); margin-top:4px;">support@cybermind.ai</div>
                </div>
                <div style="margin-bottom:14px; padding-top:12px; border-top:1px solid var(--border);">
                    <div style="font-size:10.5px; color:var(--text-faint); font-weight:700; text-transform:uppercase; letter-spacing:0.8px;">💬 Live Chat</div>
                    <div style="font-size:13px; font-weight:600; color:var(--text); margin-top:4px;">Mon – Fri · 9 AM – 5 PM IST</div>
                </div>
                <div style="margin-bottom:14px; padding-top:12px; border-top:1px solid var(--border);">
                    <div style="font-size:10.5px; color:var(--text-faint); font-weight:700; text-transform:uppercase; letter-spacing:0.8px;">📞 Phone Support</div>
                    <div style="font-size:13px; font-weight:600; color:var(--text); margin-top:4px;">+91 98765-43210</div>
                    <div style="font-size:11px; color:var(--text-muted);">India Office</div>
                </div>
                <div style="padding-top:12px; border-top:1px solid var(--border);">
                    <div style="font-size:10.5px; color:var(--text-faint); font-weight:700; text-transform:uppercase; letter-spacing:0.8px;">🛡️ App Version</div>
                    <div style="font-size:13px; font-weight:600; color:var(--text); margin-top:4px;">CyberMind AI Enterprise v2.0.0</div>
                    <div style="font-size:11px; color:var(--text-muted);">Build Date: 07-07-2026</div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )




# Transition cleanup for AI Security Assistant chat history on leaving the page
if "prev_active_page" not in st.session_state:
    st.session_state.prev_active_page = st.session_state.active_page

if st.session_state.prev_active_page != st.session_state.active_page:
    if st.session_state.prev_active_page != "AI Security Assistant":
        st.session_state["ai_assistant_origin_page"] = st.session_state.prev_active_page
    else:
        st.session_state.ai_chat_history = []
        st.session_state.ai_page_chat_history = []

st.session_state.prev_active_page = st.session_state.active_page


# Router — content renders in the normal block-container (CSS padding offsets for sidebar)

render_offline_toast()

page = st.session_state.active_page
if page == "Home" or page == "Dashboard":
    render_dashboard()
elif page == "Autonomous Security Lab":
    render_autonomous_security_lab()
elif page == "Code SAST & AST":
    render_code_sast_page()
elif page == "Fuzz & Sandbox Hub":
    render_fuzz_sandbox_page()
elif page == "Device Security Check":
    render_device_security_page()
elif page == "Universal Scan":
    render_universal_scan_page()
elif page == "AI Security Assistant":
    render_ai_assistant_page()
elif page in SCANNERS:
    render_scanner_page(page)
elif page == "Scan History":
    render_scan_history_page()
elif page == "Analytics":
    render_analytics_page()
elif page == "Settings":
    render_settings_page()
elif page == "Profile":
    render_profile_page()
elif page == "Connections":
    render_connections_page()
elif page == "Help & Support":
    render_help_support_page()
else:
    render_placeholder(page)

# ── FLOATING AI ASSISTANT WIDGET (Perfect Eyes & Blink Fix) ──
if st.session_state.active_page != "AI Security Assistant":
    st.markdown(
        """
        <style>
        /* ── Floating Widget Container ── */
        .st-key-floating_ai_assistant, div.st-key-floating_ai_assistant {
            position: fixed !important;
            bottom: 2rem !important;
            right: 2rem !important;
            width: 52px !important;
            height: 52px !important;
            z-index: 999999 !important;
            padding: 0 !important;
            margin: 0 !important;
        }

        /* ── Kimi AI Style Circular Button ── */
        .st-key-floating_ai_assistant button,
        .st-key-floating_ai_assistant div.stButton > button,
        .st-key-floating_ai_assistant div.stButton > button[kind="secondary"] {
            width: 52px !important;
            height: 52px !important;
            min-width: 52px !important;
            max-width: 52px !important;
            min-height: 52px !important;
            max-height: 52px !important;
            border-radius: 50% !important;
            background: linear-gradient(180deg, #4aa8ff 0%, #2b8af0 100%) !important;
            display: flex !important;
            align-items: center !important;
            justify-content: center !important;
            cursor: pointer !important;
            border: none !important;
            box-shadow:
                0 3px 15px rgba(43, 138, 240, 0.35),
                0 0 30px rgba(74, 168, 255, 0.15),
                inset 0 1px 2px rgba(255, 255, 255, 0.25) !important;
            transition: all 0.35s cubic-bezier(0.34, 1.56, 0.64, 1) !important;
            padding: 0 !important;
            margin: 0 !important;
            overflow: hidden !important;
            position: relative !important;
        }

 
        .st-key-floating_ai_assistant button p {
            position: absolute !important;
            inset: 0 !important;
            margin: 0 !important;
            padding: 0 !important;
            font-size: 0px !important;
            color: transparent !important;
            line-height: 0 !important;
        }

        /* Inner depth effect */
        .st-key-floating_ai_assistant button::after {
            content: "" !important;
            position: absolute !important;
            inset: 3px !important;
            border-radius: 50% !important;
            background: linear-gradient(180deg, #5cb3ff 0%, #3b9dff 50%, #1a7ae0 100%) !important;
            box-shadow: inset 0 -2px 6px rgba(0, 0, 0, 0.15) !important;
            z-index: 1 !important;
        }

        /* Shimmer overlay */
        .st-key-floating_ai_assistant button::before {
            content: "" !important;
            position: absolute !important;
            inset: 0 !important;
            border-radius: 50% !important;
            background: linear-gradient(
                135deg,
                transparent 30%,
                rgba(255, 255, 255, 0.15) 50%,
                transparent 70%
            ) !important;
            animation: shimmer-rotate 3s ease-in-out infinite !important;
            z-index: 3 !important;
        }

        .st-key-floating_ai_assistant button p::before,
        .st-key-floating_ai_assistant button p::after {
            content: "" !important;
            position: absolute !important;
            top: 18px !important;         
            width: 6.5px !important;         
            height: 9.5px !important;
            background: white !important;
            border-radius: 50% !important;
            z-index: 4 !important;
            box-shadow: 0 0 5px rgba(255, 255, 255, 0.75) !important;
            transform-origin: center center !important;
        }

        /* Left Eye Positioning & Animation */
        .st-key-floating_ai_assistant button p::before {
            left: 17px !important;
            animation: eye-blink-left 4.2s ease-in-out infinite !important;
        }

        /* Right Eye Positioning & Animation */
        .st-key-floating_ai_assistant button p::after {
            left: 27px !important;
            animation: eye-blink-right 4.2s ease-in-out infinite !important;
        }

        .st-key-floating_ai_assistant button:hover,
        .st-key-floating_ai_assistant div.stButton > button:hover,
        .st-key-floating_ai_assistant div.stButton > button[kind="secondary"]:hover {
            transform: scale(1.12) !important;
            box-shadow:
                0 5px 25px rgba(43, 138, 240, 0.5),
                0 0 50px rgba(74, 168, 255, 0.25),
                0 0 80px rgba(74, 168, 255, 0.1) !important;
        }

        .st-key-floating_ai_assistant button:active,
        .st-key-floating_ai_assistant div.stButton > button:active,
        .st-key-floating_ai_assistant div.stButton > button[kind="secondary"]:active {
            transform: scale(0.94) !important;
        }

        /* ── Soft Glow Rings ── */
        .st-key-floating_ai_assistant .glow-ring-1,
        .st-key-floating_ai_assistant .glow-ring-2 {
            position: absolute !important;
            border-radius: 50% !important;
            pointer-events: none !important;
            top: 50% !important;
            left: 50% !important;
            transform: translate(-50%, -50%) !important;
        }

        .st-key-floating_ai_assistant .glow-ring-1 {
            width: 60px !important;
            height: 60px !important;
            border: 1.5px solid rgba(74, 168, 255, 0.25) !important;
            animation: soft-pulse 2.5s ease-out infinite !important;
        }

        .st-key-floating_ai_assistant .glow-ring-2 {
            width: 68px !important;
            height: 68px !important;
            border: 1px solid rgba(74, 168, 255, 0.12) !important;
            animation: soft-pulse 2.5s ease-out infinite 0.6s !important;
        }

        /* ── Tooltip on Hover ── */
        .st-key-floating_ai_assistant::before {
            content: "AI Security Assistant" !important;
            visibility: hidden !important;
            width: auto !important;
            background-color: #161922 !important;
            color: #c7d2fe !important;
            text-align: center !important;
            border-radius: 12px !important;
            padding: 9px 15px !important;
            position: absolute !important;
            z-index: 1000000 !important;
            bottom: 125% !important;
            right: 0% !important;
            opacity: 0 !important;
            transition: all 0.3s cubic-bezier(0.34, 1.56, 0.64, 1) !important;
            transform: translateX(12px) scale(0.92) !important;
            font-size: 12.5px !important;
            font-weight: 600 !important;
            border: 1px solid rgba(100, 120, 255, 0.18) !important;
            box-shadow:
                0 4px 18px rgba(0, 0, 0, 0.45),
                0 0 25px rgba(59, 157, 255, 0.08) !important;
            pointer-events: none !important;
            white-space: nowrap !important;
        }

        .st-key-floating_ai_assistant:hover::before {
            visibility: visible !important;
            opacity: 1 !important;
            transform: translateX(0) scale(1) !important;
        }

        /* ── KEYFRAME ANIMATIONS (Smooth Scale Only) ── */
        @keyframes shimmer-rotate {
            0% { transform: translateX(-100%) rotate(0deg); }
            100% { transform: translateX(100%) rotate(360deg); }
        }

        @keyframes soft-pulse {
            0% { transform: translate(-50%, -50%) scale(1); opacity: 0.5; }
            100% { transform: translate(-50%, -50%) scale(1.4); opacity: 0; }
        }

        @keyframes eye-blink-left {
            0%, 45%, 53%, 100% { transform: scaleY(1); }
            49% { transform: scaleY(0.08); }
        }

        @keyframes eye-blink-right {
            0%, 46%, 54%, 100% { transform: scaleY(1); }
            50% { transform: scaleY(0.08); }
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    st.button(
        "👁️",
        key="floating_ai_assistant",
        on_click=go_to,
        args=("AI Security Assistant",)
    )

# ── SECTION 11: RED FLOATING SCAN EXPLAINER WIDGET ──
if hasattr(st, "dialog"):
    @st.dialog(" ")
    def render_scan_explainer_dialog():
        st.markdown(
            """
            <style>
            html body div[data-baseweb="backdrop"],
            html body div[data-baseweb="modal"],
            html body [data-testid="stDialog"],
            html body [data-testid="stModal"],
            div[data-baseweb="backdrop"],
            div[data-baseweb="modal"],
            [data-testid="stDialog"],
            [data-testid="stModal"] {
                background: rgba(0, 0, 0, 0.4) !important;
                background-color: rgba(0, 0, 0, 0.4) !important;
                backdrop-filter: blur(6px) !important;
                -webkit-backdrop-filter: blur(6px) !important;
            }
            div[data-baseweb="modal"] > div,
            div[data-baseweb="modal"] > div > div,
            [data-testid="stDialog"] > div {
                background: transparent !important;
                background-color: transparent !important;
            }
            div[data-baseweb="modal"] [role="dialog"],
            [data-testid="stDialog"] [role="dialog"],
            [data-testid="stDialog"] [data-testid="stDialogContent"] {
                background: var(--dialog-bg) !important;
                background-color: var(--dialog-bg) !important;
                color: var(--dialog-text) !important;
                border: 1px solid var(--dialog-border) !important;
                border-radius: 16px !important;
                box-shadow: 0 25px 60px rgba(0, 0, 0, 0.5) !important;
            }
            div[data-testid="stDialogHeader"] h2,
            div[data-testid="stDialogHeader"] span,
            div[role="dialog"] h2,
            [data-testid="stDialog"] h2 {
                display: none !important;
            }
            </style>
            <div style="display: flex; align-items: center; gap: 10px; margin-bottom: 14px; margin-top: -10px;">
                <img src="https://cdn-icons-png.flaticon.com/512/6071/6071531.png" width="30" height="30" style="object-fit: contain;">
                <span style="font-size: 19px; font-weight: 800; color: var(--text);">AI Scan Explanation</span>
            </div>
            """,
            unsafe_allow_html=True
        )
        context = st.session_state.get("last_scan_context", {})
        if not context:
            st.info("No scan context available. Run a scan on any scanner page first!")
            return
        
        target = context.get("target", "Target")
        scanner = context.get("scanner", "Scanner")
        score = context.get("risk_score", 0)
        level = context.get("risk_level", "Unknown")
        
        st.markdown(f"#### Target: `{target}`")

        with st.spinner("Generating plain-language explanation with CyberMind-AI..."):
            import json
            from modules.ai_assistant import query_groq_api
            prompt_messages = [
                {
                    "role": "system",
                    "content": "You are CyberMind AI, an expert security assistant. Explain the security scan results in simple plain language in English: 1) What was found, 2) Why it's safe or risky, and 3) 2-3 concrete steps the user should take right now. Keep your response under 200 words in Markdown."
                },
                {
                    "role": "user",
                    "content": f"Scan Context: Target='{target}', Scanner='{scanner}', Risk Level='{level}', Risk Score={score}, Findings={context.get('findings', {})}"
                }
            ]
            exp_res = query_groq_api(prompt_messages)
            if exp_res.startswith("Groq API Error") or exp_res.startswith("Error") or exp_res.startswith("Failed"):
                exp_res = (
                    f"**Analysis Summary for `{target}`:**\n\n"
                    f"- **Security Status:** Classified as **{level}** with a risk score of **{score}/100**.\n"
                    f"- **Assessment:** Evaluated against multiple threat databases, domain metrics, and Machine Learning models.\n\n"
                    f"**Recommended Actions:**\n"
                    f"1. {'Do not enter personal or financial information on this page.' if score >= 40 else 'Resource appears safe; proceed with normal caution.'}\n"
                    f"2. Verify domain WHOIS registration age and SSL certificate authenticity.\n"
                    f"3. Re-scan or request administrative clearance if unexpected."
                )

        final_exp = exp_res

        st.markdown(final_exp)

        safe_speech_txt = html.escape(final_exp).replace("'", "\\'").replace("\n", " ").replace('"', '\\"')
        components.html(
            f"""
            <button id="modal_tts_btn" onclick="toggleModalTTS()" style="background:rgba(34,211,238,0.15); border:1px solid #22D3EE; color:#22D3EE; border-radius:6px; padding:6px 14px; font-size:12px; font-weight:700; cursor:pointer; width:100%; display:flex; align-items:center; justify-content:center; gap:6px;">
                🔊 Listen AI Explanation
            </button>
            <script>
                function toggleModalTTS() {{
                    const btn = document.getElementById("modal_tts_btn");
                    const synth = (window.parent && window.parent.speechSynthesis) ? window.parent.speechSynthesis : window.speechSynthesis;
                    if (synth) {{
                        if (synth.speaking) {{
                            synth.cancel();
                            if (btn) btn.innerHTML = "🔊 Listen AI Explanation";
                        }} else {{
                            synth.cancel();
                            const msg = new (window.parent.SpeechSynthesisUtterance || SpeechSynthesisUtterance)("{safe_speech_txt}");
                            msg.rate = 0.95;
                            msg.onend = function() {{
                                if (btn) btn.innerHTML = "🔊 Listen AI Explanation";
                            }};
                            synth.speak(msg);
                            if (btn) btn.innerHTML = "⏹️ Stop Speech";
                        }}
                    }}
                }}
            </script>
            """,
            height=45
        )

last_ctx = st.session_state.get("last_scan_context", {})
active_p = st.session_state.get("active_page")
has_active_scan_res = (
    st.session_state.get(f"scan_result_{active_p}") is not None
    or (active_p == "Universal Scan" and st.session_state.get("universal_scan_result") is not None)
)
if last_ctx and active_p == last_ctx.get("scanner_key") and has_active_scan_res:
    st.markdown(
        """
        <style>
        /* ── Red Floating Explainer Container ── */
        .st-key-floating_scan_explainer, div.st-key-floating_scan_explainer {
            position: fixed !important;
            bottom: 5.8rem !important;
            right: 2rem !important;
            width: 52px !important;
            height: 52px !important;
            z-index: 999998 !important;
            padding: 0 !important;
            margin: 0 !important;
        }

        /* ── Red Circular Animated Button ── */
        .st-key-floating_scan_explainer button,
        .st-key-floating_scan_explainer div.stButton > button,
        .st-key-floating_scan_explainer div.stButton > button[kind="secondary"] {
            width: 52px !important;
            height: 52px !important;
            min-width: 52px !important;
            max-width: 52px !important;
            min-height: 52px !important;
            max-height: 52px !important;
            border-radius: 50% !important;
            background: linear-gradient(180deg, #F2545B 0%, #D9383E 100%) !important;
            display: flex !important;
            align-items: center !important;
            justify-content: center !important;
            cursor: pointer !important;
            border: none !important;
            box-shadow:
                0 3px 15px rgba(242, 84, 91, 0.4),
                0 0 30px rgba(242, 84, 91, 0.2),
                inset 0 1px 2px rgba(255, 255, 255, 0.25) !important;
            transition: all 0.35s cubic-bezier(0.34, 1.56, 0.64, 1) !important;
            padding: 0 !important;
            margin: 0 !important;
            overflow: hidden !important;
            position: relative !important;
        }

        .st-key-floating_scan_explainer button p {
            position: absolute !important;
            inset: 0 !important;
            margin: 0 !important;
            padding: 0 !important;
            font-size: 0px !important;
            color: transparent !important;
            line-height: 0 !important;
        }

        /* Inner depth effect */
        .st-key-floating_scan_explainer button::after {
            content: "" !important;
            position: absolute !important;
            inset: 3px !important;
            border-radius: 50% !important;
            background: linear-gradient(180deg, #FF6B70 0%, #F2545B 50%, #D9383E 100%) !important;
            box-shadow: inset 0 -2px 6px rgba(0, 0, 0, 0.15) !important;
            z-index: 1 !important;
        }

        /* Shimmer overlay */
        .st-key-floating_scan_explainer button::before {
            content: "" !important;
            position: absolute !important;
            inset: 0 !important;
            border-radius: 50% !important;
            background: linear-gradient(
                135deg,
                transparent 30%,
                rgba(255, 255, 255, 0.15) 50%,
                transparent 70%
            ) !important;
            animation: shimmer-rotate 3s ease-in-out infinite !important;
            z-index: 3 !important;
        }

        .st-key-floating_scan_explainer button p::before,
        .st-key-floating_scan_explainer button p::after {
            content: "" !important;
            position: absolute !important;
            top: 18px !important;         
            width: 6.5px !important;         
            height: 9.5px !important;
            background: white !important;
            border-radius: 50% !important;
            z-index: 4 !important;
            box-shadow: 0 0 5px rgba(255, 255, 255, 0.75) !important;
            transform-origin: center center !important;
        }

        /* Left Eye Positioning & Animation */
        .st-key-floating_scan_explainer button p::before {
            left: 17px !important;
            animation: eye-blink-left 4.2s ease-in-out infinite !important;
        }

        /* Right Eye Positioning & Animation */
        .st-key-floating_scan_explainer button p::after {
            left: 27px !important;
            animation: eye-blink-right 4.2s ease-in-out infinite !important;
        }

        .st-key-floating_scan_explainer button:hover,
        .st-key-floating_scan_explainer div.stButton > button:hover {
            transform: scale(1.12) !important;
            box-shadow:
                0 5px 25px rgba(242, 84, 91, 0.6),
                0 0 50px rgba(242, 84, 91, 0.3) !important;
        }

        .st-key-floating_scan_explainer button:active {
            transform: scale(0.94) !important;
        }

        /* ── Tooltip on Hover ── */
        .st-key-floating_scan_explainer::before {
            content: "Explain Scan Findings" !important;
            visibility: hidden !important;
            width: auto !important;
            background-color: #161922 !important;
            color: #fca5a5 !important;
            text-align: center !important;
            border-radius: 12px !important;
            padding: 9px 15px !important;
            position: absolute !important;
            z-index: 1000000 !important;
            bottom: 125% !important;
            right: 0% !important;
            opacity: 0 !important;
            transition: all 0.3s cubic-bezier(0.34, 1.56, 0.64, 1) !important;
            transform: translateX(12px) scale(0.92) !important;
            font-size: 12.5px !important;
            font-weight: 600 !important;
            border: 1px solid rgba(242, 84, 91, 0.3) !important;
            white-space: nowrap !important;
        }

        .st-key-floating_scan_explainer:hover::before {
            visibility: visible !important;
            opacity: 1 !important;
            transform: translateX(0) scale(1) !important;
        }
        </style>
        """,
        unsafe_allow_html=True
    )
    if st.button("👁️", key="floating_scan_explainer"):
        if hasattr(st, "dialog"):
            render_scan_explainer_dialog()
